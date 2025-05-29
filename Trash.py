import os
import docx
from docx.oxml.ns import nsmap
from docx.oxml import OxmlElement
from Formated_Base_PEP8 import configurar_directorio_trabajo

# Define the WordprocessingML namespace for easier XML element checking
W_NS = nsmap['w']
configurar_directorio_trabajo()

def iter_elements(parent_element):
    """
    Recursively yields XML elements from a parent element in document order.
    Helps traverse complex structures like paragraphs, runs, tables, SDTs, etc.
    """
    if parent_element is None:
        return

    # Iterate directly over the children of the parent element
    for child in parent_element.getchildren():
        # Yield the child element itself FIRST
        yield child

        # Then, if the child is a container type that can hold other elements
        # we care about (like runs, hyperlinks, fields, etc.), recurse into it.
        # This list of tags covers common inline containers and fields.
        if child.tag in (
                OxmlElement('{%s}r' % W_NS).tag,  # Run
                OxmlElement('{%s}hyperlink' % W_NS).tag,  # Hyperlink
                OxmlElement('{%s}fldSimple' % W_NS).tag,  # Simple field
                OxmlElement('{%s}smartTag' % W_NS).tag,  # Deprecated Smart Tag
                # OxmlElement('{%s}sdt' % W_NS).tag,     # SDT handled below
        ):
            yield from iter_elements(child)

        # Handle Structured Document Tags (SDTs) - need to find the content part
        elif child.tag == OxmlElement('{%s}sdt' % W_NS).tag:
            sdtContent = child.find(OxmlElement('{%s}sdtContent' % W_NS).tag)
            if sdtContent is not None:
                yield from iter_elements(sdtContent)  # Recurse into content

        # Handle tables - need to recurse into cells (within rows)
        elif child.tag == OxmlElement('{%s}tbl' % W_NS).tag:
            # Iterate through rows and cells and recurse into cell content
            # Use xpath to find all tr and tc within the table element
            for row in child.xpath('.//w:tr', namespaces={'w': W_NS}):
                for cell in row.xpath('.//w:tc', namespaces={'w': W_NS}):
                    yield from iter_elements(cell)


def get_bookmark_text_data(doc):
    """
    Extracts text content associated with each bookmark in the document.

    Args:
        doc (docx.document.Document): The python-docx document object.

    Returns:
        dict: A dictionary where keys are bookmark names (str) and values
              are the extracted text (str) between the bookmark start and end.
              Note: Whitespace/newlines within the bookmark might not be
              perfectly preserved depending on the document's XML structure
              and how w:t elements handle spacing. This function concatenates
              text from all w:t and w:instrText elements found within an
              active bookmark range.
    """
    bookmark_data = {}
    active_bookmarks = {}  # {bookmark_id: bookmark_name}

    # Iterate through all XML elements in the document body in order
    # doc.element.body is the root of the main document content XML
    for element in iter_elements(doc.element.body):

        # Check if it's a bookmarkStart element
        if element.tag == OxmlElement('{%s}bookmarkStart' % W_NS).tag:
            bkm_name = element.get(OxmlElement('{%s}name' % W_NS).tag)
            bkm_id = element.get(OxmlElement('{%s}id' % W_NS).tag)
            if bkm_name and bkm_id:
                # Activate this bookmark for text collection
                active_bookmarks[bkm_id] = bkm_name
                # Initialize the text entry if it doesn't exist
                bookmark_data[bkm_name] = bookmark_data.get(bkm_name, "")

        # Check if it's a bookmarkEnd element
        elif element.tag == OxmlElement('{%s}bookmarkEnd' % W_NS).tag:
            bkm_id = element.get(OxmlElement('{%s}id' % W_NS).tag)
            # Deactivate this bookmark by removing its ID
            # This assumes typical bookmarks; complex nested cases might need more logic
            if bkm_id in active_bookmarks:
                del active_bookmarks[bkm_id]

        # Check if it's a text element (w:t) or field instruction text (w:instrText)
        # Field instruction text is common within bookmarks used for merge fields
        elif element.tag == OxmlElement('{%s}t' % W_NS).tag or element.tag == OxmlElement('{%s}instrText' % W_NS).tag:
            text = element.text
            if text is not None and active_bookmarks:  # Only add text if any bookmark is currently active
                # Append this text to the accumulated text for ALL currently active bookmarks
                # This handles nested bookmarks by duplicating text, which is a common
                # way text inside nested bookmarks is treated for extraction purposes.
                for bkm_name in active_bookmarks.values():
                    bookmark_data[bkm_name] += text

        # Note: This simple text concatenation approach doesn't perfectly handle
        # complex whitespace, tabs (w:tab), breaks (w:br), or complex field codes.
        # It focuses on extracting raw text between the start and end tags.

    return bookmark_data


# --- Your existing script starts here ---

doc_filename_base = "contrato_automatizado_con_marcadores"
doc_path = doc_filename_base + ".docx"

# Check if the document exists
if not os.path.exists(doc_path):
    print(f"Error: El archivo '{doc_path}' no fue encontrado en el directorio de trabajo: {os.getcwd()}")
    # The commented-out dummy document creation code is a good way to test
    # if you don't have a real document ready. Uncomment if needed.
    # print("Creando un documento de prueba...")
    # try:
    #     dummy_doc = docx.Document()
    #     p = dummy_doc.add_paragraph()
    #     # Add bookmarkStart for BM_Prueba
    #     bm_start_el = docx.oxml.OxmlElement('w:bookmarkStart')
    #     bm_start_el.set(docx.oxml.ns.qn('w:name'), 'BM_Prueba')
    #     bm_start_el.set(docx.oxml.ns.qn('w:id'), '0')
    #     p._p.append(bm_start_el)
    #     # Add a run with text
    #     p.add_run("Este es el texto para ")
    #     # Add bookmarkEnd for BM_Prueba
    #     bm_end_el = docx.oxml.OxmlElement('w:bookmarkEnd')
    #     bm_end_el.set(docx.oxml.ns.qn('w:id'), '0')
    #     p._p.append(bm_end_el)
    #     # Add more text in another run after the bookmark
    #     p.add_run("BM_Prueba.")

    #     # Add a second paragraph with another bookmark
    #     p2 = dummy_doc.add_paragraph()
    #     bm_start_el2 = docx.oxml.OxmlElement('w:bookmarkStart')
    #     bm_start_el2.set(docx.oxml.ns.qn('w:name'), 'OtroMarcador')
    #     bm_start_el2.set(docx.oxml.ns.qn('w:id'), '1')
    #     p2._p.append(bm_start_el2)
    #     p2.add_run("Texto del segundo marcador.")
    #     bm_end_el2 = docx.oxml.OxmlElement('w:bookmarkEnd')
    #     bm_end_el2.set(docx.oxml.ns.qn('w:id'), '1')
    #     p2._p.append(bm_end_el2)

    #     dummy_doc.save(doc_path)
    #     print(f"Documento de prueba '{doc_path}' creado con marcadores 'BM_Prueba' y 'OtroMarcador'.")
    # except Exception as create_e:
    #     print(f"Error creating dummy document: {create_e}")
    exit()  # Exit if the file doesn't exist and we aren't creating a dummy

try:
    doc = docx.Document(doc_path)
except Exception as e:
    print(f"Error al abrir el documento '{doc_path}': {e}")
    exit()

# Call the function to extract all bookmark names and their texts
all_bookmarks_data = get_bookmark_text_data(doc)

# Now, 'all_bookmarks_data' is a dictionary like {'BookmarkName': 'Text found inside'}.
# You can iterate through it to access each bookmark's name and text.
print("\n--- Datos de los Marcadores Extraídos ---")
if all_bookmarks_data:
    # Sort bookmarks by name for consistent output (optional)
    sorted_bookmark_names = sorted(all_bookmarks_data.keys())

    for bookmark_name in sorted_bookmark_names:
        bookmark_text = all_bookmarks_data[bookmark_name]

        # Format the text by wrapping it in double curly braces
        formatted_text = "{{" + bookmark_text + "}}"

        # Print the extracted data
        print(f"Nombre del Marcador: {bookmark_name}")
        print(f"Texto Asociado     : '{formatted_text}'")  # Using '' to show if text is empty
        print("-" * 30)

        # You can now use these variables as needed. For example, store them,
        # print them, or use them in further processing.

        # Example: Assigning specific bookmark texts to unique variables
        # You would typically do this by checking the bookmark_name:
        #
        # if bookmark_name == "NombreCliente":
        #     nombre_cliente_variable = bookmark_text # Use the original text if you don't need {{}} yet
        #     print(f"(Asignado a 'nombre_cliente_variable': '{nombre_cliente_variable}')")
        # elif bookmark_name == "FechaDocumento":
        #     fecha_documento_variable = bookmark_text
        #     print(f"(Asignado a 'fecha_documento_variable': '{fecha_documento_variable}')")
        # elif bookmark_name == "MontoContrato":
        #      monto_contrato_variable = bookmark_text
        #      print(f"(Asignado a 'monto_contrato_variable': '{monto_contrato_variable}')")

    # The dictionary 'all_bookmarks_data' itself stores all the data.
    # You can access specific bookmark texts directly like this:
    # (Using .get() is safer as it returns None if the key doesn't exist, avoiding errors)
    texto_cliente = all_bookmarks_data.get("NombreCliente", "Marcador 'NombreCliente' no encontrado")
    print(f"\nAcceso directo a 'NombreCliente': '{texto_cliente}'")

    texto_fecha = all_bookmarks_data.get("FechaDocumento", "Marcador 'FechaDocumento' no encontrado")
    print(f"Acceso directo a 'FechaDocumento': '{texto_fecha}'")

    # You could also create a new dictionary with the formatted text if needed
    formatted_bookmarks_data = {name: "{{" + text + "}}" for name, text in all_bookmarks_data.items()}
    # print("\nFormatted Dictionary:", formatted_bookmarks_data)

    # The original document object 'doc' is not modified by get_bookmark_text_data,
    # so saving it here will save the original document as is.
    # If you were *replacing* bookmark text, you would save *after* replacement.
    # doc.save(doc_filename_base + "_testing_v1.docx") # Saving the original doc
    # If you needed to save a *new* doc with the formatted text *replacing* bookmarks,
    # that would be a separate step involving finding bookmarks and replacing runs.


else:
    print("No se encontraron marcadores en el documento, o no se pudo extraer texto de ellos.")

print("\nProceso de extracción completado.")