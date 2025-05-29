import os
import re
import docx
from docx.shared import Inches, Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT

def configurar_directorio_trabajo():
    """Configura el directorio de trabajo en la subcarpeta 'Files'."""
    cwd = os.getcwd()
    target_dir_name = "Files"
    wd = os.path.join(cwd, target_dir_name)
    pattern = r"Files\\Files"
    if re.search(pattern, wd):
        wd = wd.replace(r"\Files\Files", r"\Files")
    if os.path.isdir(wd):
        print(f"Directorio de trabajo configurado a: {wd}")
    else:
        print(f"Advertencia: El directorio '{wd}' no existe o no es válido. Usando directorio actual: {cwd}")
        wd = cwd
    return wd

def create_melipilla_document(archivo="base", wd=None):
    """Crea el documento de portada para Melipilla con la configuración especificada.

    Args:
        archivo (str): Tipo de documento, puede ser 'base' o 'contrato'. Por defecto es 'base'.
        wd (str, optional): Directorio de trabajo donde se guardará el documento.
                           Si no se especifica, usa el directorio actual o la subcarpeta 'Files'.
    """
    # Configurar el directorio de trabajo
    if wd is None:
        wd = configurar_directorio_trabajo()
    else:
        wd = os.path.abspath(wd)  # Normalizar la ruta a absoluta
        if not os.path.isdir(wd):
            print(f"Advertencia: El directorio '{wd}' no existe. No se puede crear el documento.")
            return

    print(f"Usando directorio de trabajo: {wd}")

    # Define image paths (usar rutas fijas de red)
    logo_melipilla_name = r"\\10.5.130.24\Abastecimiento\Compartido Abastecimiento\Otros\NO_MODIFICAR\logo_melipilla.png"  # Imagen a la derecha
    ssmo_alta_name = r"\\10.5.130.24\Abastecimiento\Compartido Abastecimiento\Otros\NO_MODIFICAR\SSMOalta.png"  # Imagen a la izquierda

    # Verificar si las imágenes existen antes de intentar usarlas
    # Create a new Document
    doc = docx.Document()

    # Add images side-by-side using a table
    table = doc.add_table(rows=1, cols=3)
    table.autofit = False  # Disable autofit to control column widths manually

    # Set column widths
    table.columns[0].width = Inches(2.5)  # Width for the left logo cell (SSMOalta)
    table.columns[1].width = Inches(3)  # Width for spacing
    table.columns[2].width = Inches(2.5)  # Width for the right logo cell (logo_melipilla)

    left_logo_cell = table.cell(0, 0)  # This cell holds the left logo (SSMOalta)
    right_logo_cell = table.cell(0, 2)  # This cell holds the right logo (logo_melipilla)

    left_logo_paragraph = left_logo_cell.paragraphs[0]
    run = left_logo_paragraph.add_run()
    run.add_picture(ssmo_alta_name, height=Cm(2.87))
    left_logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    left_logo_paragraph.paragraph_format.space_after = Pt(0)
    left_logo_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

    right_logo_paragraph = right_logo_cell.paragraphs[0]
    run = right_logo_paragraph.add_run()
    run.add_picture(logo_melipilla_name, width=Cm(3.74))
    right_logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    right_logo_paragraph.paragraph_format.space_after = Pt(0)
    right_logo_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


    # Add text below the table
    if archivo == "base":
        text_lines = (
            "SERVICIO SALUD OCCIDENTE",
            "HOSPITAL DE MELIPILLA",
            "UNIDAD DE ABASTECIMIENTO",
            "BASE N° {{ numero_base }}"
        )
    else:
        text_lines = [
            "SERVICIO SALUD OCCIDENTE",
            "HOSPITAL SAN JOSE DE MELIPILLA",
            "UNIDAD DE ABASTECIMIENTO",
            "Convenios",
            "N° {{ numero_contrato }}",
            "{{ involucrados }}"
        ]

    for line in text_lines:
        paragraph = doc.add_paragraph(line)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)

        if paragraph.runs:
            font = paragraph.runs[0].font
            font.size = Pt(8)
            if line == "BASE N° {{ numero_base }}":
                font.bold = True

    # Save the document in the specified working directory
    if archivo == "base":
        output_filename = "portada_melipilla_base.docx"
    else:
        output_filename = "portada_melipilla_contrato.docx"

    output_path = os.path.join(wd, output_filename)
    try:
        doc.save(output_path)
        print(f"Documento guardado en: {output_path}")
    except Exception as e:
        print(f"Error al guardar el documento en {output_path}: {e}")
        return

    # Verify the saved file
    try:
        if os.path.exists(output_path):
            tamaño_archivo = os.path.getsize(output_path) / 1024  # Tamaño en KB
            print(f"✅ El archivo '{output_filename}' se guardó correctamente.")
            print(f"   Ruta completa: {output_path}")
            print(f"   Tamaño del archivo: {tamaño_archivo:.2f} KB")
        else:
            print(f"❌ Error: No se encontró el archivo '{output_filename}' después de intentar guardarlo.")
    except Exception as e:
        print(f"❌ Error al verificar el archivo guardado: {str(e)}")

if __name__ == "__main__":
    # Ejecutar la función principal si se corre directamente
    create_melipilla_document(archivo="base")
