from openai import OpenAI

from PIL import Image
import base64
import io
import docx
import fitz

# from Formated_Base_PEP8 import configurar_directorio_trabajo
#configurar_directorio_trabajo()

# Initialize OpenAI API
openai_api_key = "AIzaSyAgXYnDNJ6gQmMVUk88G6wGPdXz-qG2Gbw"
client = OpenAI(
    api_key=openai_api_key,
    base_url="https://generativelanguage.googleapis.com/v1beta/openai/"
)

model = "gemini-2.5-flash-preview-04-17"

# Cargar y convertir la imagen a base64

image_name = "CERTIFICADO DE FIEL CUMPLIMIENTO ID 1057480-11-LE25. ELITEC SPA.pdf"



import re
regex = r".*(.pdf)"
formato = re.match(regex, image_name)


# Si es PDF
if formato != None:
    try:
        pdf_document = fitz.open(image_name)
        primera_pagina = pdf_document.load_page(0)
        pix = primera_pagina.get_pixmap(alpha=False)
        img_bytes = pix.tobytes("jpeg")
        image_base64 = base64.b64encode(img_bytes).decode('utf-8')
        pdf_document.close()
    except Exception as e:
        print(f"Error al procesar el PDF: {e}")
        exit(1)
# Si es Imagen
else:
    image = Image.open(image_name)
# Convertir la imagen a bytes y luego a base64
    buffer = io.BytesIO()
    image.save(buffer, format="JPEG")
    image_bytes = buffer.getvalue()
    image_base64 = base64.b64encode(image_bytes).decode('utf-8')

# Crear el mensaje con la imagen
response = client.chat.completions.create(
    model=model,
    messages=[
        {  "role": "system",
    "content": "You are a transcriber in a company specializing in standardizing documents for Garntía de Fiel Cumplimiento from various institutions. Your task is to extract and interpret key values into a single, consistent format. Note that fields may vary; for example, use 'Tomador' if not explicit, but interpret 'Afianzado' as equivalent. Similarly, interpret 'Acreedor' as 'Beneficiario' if needed. Output the results in a dictionary format based on the provided example, asegurado suele ser el mismo beneficiado el rut del tomador no puede ser el mismo que el rut del beneficiado, tipicamente 61.602.123-0 el rut del beneficiado será . The files are in Spanish and you must work in spanish"},
        {
            "role": "user",
            "content": [
                {"type": "text", "text": """
Your role is to read the file an complete the fields an generate a dictionary, follow the same formating as this example, use the keys from the example, and the values for those keys should be the corresponding data found in the image. If not found, put NULL. If numero Dias is not found it gotta be calculated as the difference of days betwen the second value of vigencia del seguro and the first one, if "Valor a pagar del Contrato" is not on the file, it should be written in letters takin as a base the ammuont in numbers. Vigencia del seguro is since the document is emited until it expires, and Numero de dias is the difference betwen those two elements  
{
    "Tomador": "Servicios de Bioingenieria Limitada",
    "RUT_Tomador": "76.644.150-5",
    "Asegurado": "Hospital San José de Melipilla",
    "RUT_Asegurado": "66.644.150-5",
    "Beneficiario": "Hospital San José de Melipilla",
    "RUT_Beneficiario": "66.644.150-5",
    "Dirección del Tomador": "Avenida Condell 1680",
    "Ciudad": "Nuñoa",
    "Cobertura": "Fiel Cumplimiento",
    "Vigencia del Seguro": "01/01/2025–02/01-2025",
    "Numero de Días": "1",
    "Valor Asegurado": "699,50",
    "Prima Neta": "1633,92",
    "IVA": "3,21",
    "Total a Pagar": "20,13",
    "Valor a pagar en Letra": "Veinte coma trece UF"
    "Ciudad y fecha de emisión": "Santiago, 18 de marzo de 2025",
    "Poliza N°/ID: 17306798"
}

                """},
                {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{image_base64}"}}
            ]
        }
    ]
)


print(response.choices[0].message.content)


# Obtener la respuesta del modelo Gemini
respuesta = response.choices[0].message.content
import json


# esto es para que la respuesta sea legible, porque viene e n un formato json que no es valido
try:
    # Intentar parsear directamente como JSON
    datos = json.loads(response.choices[0].message.content)
except json.JSONDecodeError:
    # Si falla, intenta limpiar el texto antes de parsear
    contenido_limpio = response.choices[0].message.content
    # Buscar donde comienza y termina el JSON
    inicio = contenido_limpio.find('{')
    fin = contenido_limpio.rfind('}') + 1
    if inicio >= 0 and fin > 0:
        json_texto = contenido_limpio[inicio:fin]
        datos = json.loads(json_texto)
    else:
        datos = {}  # Si no se puede parsear, usar diccionario vacío
# Tengo que
# datos

# Importar las bibliotecas necesarias
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches

# Crear un nuevo documento
documento = Document()

# Añadir un título
documento.add_heading('Certificado de Fianza', 0)

# Crear una tabla con 2 columnas
tabla = documento.add_table(rows=1, cols=2)
tabla.style = 'Table Grid'

# Configurar encabezados

# Definir los campos que queremos extraer
# Definir los campos que queremos extraer
campos = [
    "Tomador",
    "RUT",
    "Asegurado",
    "Beneficiario",
    "Dirección del Tomador",
    "Ciudad",
    "Cobertura",
    "Vigencia del Seguro",
    "Numero de Días",
    "Valor Asegurado",
    "Prima Neta",
    "IVA",
    "Total a Pagar",
    "Valor a pagar en Letra",
    "Ciudad y fecha de emisión",
    "Poliza N°/ID"
]

# Procesar la respuesta para extraer pares clave-valor
lineas = respuesta.split('\n')
valores = {}

# Intentar diferentes patrones de formato en la respuesta
for linea in lineas:
    for campo in campos:
        # Patrón 1: "campo", "text_content": "valor"}
        patron1 = f'"{campo}", "text_content": "(.+?)"'
        import re

        match1 = re.search(patron1, linea)
        if match1:
            valores[campo] = match1.group(1)
            break

        # Patrón 2: campo: valor
        if f"{campo}:" in linea:
            valor = linea.split(f"{campo}:", 1)[1].strip()
            valores[campo] = valor
            break

        # Patrón 3: "campo": "valor"
        if f'"{campo}":' in linea:
            try:
                valor = linea.split(f'"{campo}":', 1)[1].strip()
                # Eliminar comillas y comas al final si existen
                valor = valor.strip('",')
                valores[campo] = valor
                break
            except:
                pass

# Añadir los campos y valores a la tabla
for campo in campos:
    fila = tabla.add_row().cells
    fila[0].text = campo
    fila[1].text = valores.get(campo, "")

# Guardar el documento
documento.save('tabla_fea_sin_rellenar.docx')

print("Documento Word creado con éxito.")

# Rellenado de la Tabla

# Recorrer tablas y filas
# Cargar el documento
doc = docx.Document("prototipo_tabla.docx")


# Rellena la tabla con los valores encontrados
for table in doc.tables:
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            # Obtener el texto completo de la celda
            texto_completo = cell.text

            # Buscar cuál de los campos está en esta celda
            campo_encontrado = None
            for campo in campos:
                if campo in texto_completo:
                    campo_encontrado = campo
                    break

            if campo_encontrado and campo_encontrado in valores:
                # Creamos un nuevo texto con el formato "Campo\nNuevo valor"
                nuevo_texto = f"{campo_encontrado}\n{valores[campo_encontrado]}"

                # Limpiamos la celda y añadimos el nuevo contenido
                for paragraph in cell.paragraphs:
                    if paragraph.text:
                        paragraph.clear()

                # Añadir el nuevo texto
                cell.paragraphs[0].add_run(nuevo_texto)
                print(f"Actualizada celda con {campo_encontrado}: {valores[campo_encontrado]}")

# Guardar el documento actualizado

doc.save("prototipo_tabla_rellenado.docx")

# Now I have a se
for i in doc.tables:
    for j in i.rows:
        for k in j.cells:
            print(k.text)

