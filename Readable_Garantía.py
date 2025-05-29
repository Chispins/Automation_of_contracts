import os
import re
import json
import ast
from PIL import Image
import google.generativeai as genai
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches



# Set up the working directory and load the image
def setup_and_load_image(file_name):
    # Assuming configurar_directorio_trabajo() is defined elsewhere and needed
    from Formated_Base_PEP8 import configurar_directorio_trabajo
    configurar_directorio_trabajo()

    # Load the image file
    image = Image.open(file_name)
    return image


file_name = "CERTIFICADO GARANTIA HOSPITAL MELIPILLA.png"  # Updated to the specified file
image = setup_and_load_image(file_name)


# Configure Google Generative AI and generate content
google_api_key = "AIzaSyAgXYnDNJ6gQmMVUk88G6wGPdXz-qG2Gbw"
genai.configure(api_key=google_api_key)  # Replace with actual API key if not set

model = genai.GenerativeModel('models/gemini-2.5-flash-preview-04-17')

prompt = """
Markdown

Extrae los siguientes datos del documento proporcionado y devuelve la información estrictamente en formato JSON, sin texto adicional antes ni después.

Los datos a extraer son:

Certificado de fianza Web
Fecha (día, mes y año)
Nombre de afianzado
R.U.T. N° del afianzado
Domicilio del afianzado
Nombre del mandante
R.U.T. N° del mandante
Domicilio del mandante
Obligación caucionada
Monto
Glosa
El formato de salida debe ser un array JSON que contenga exactamente dos elementos:

Un objeto JSON donde las claves sean los nombres exactos de los datos solicitados (tal como se listan arriba) y los valores correspondan a una lista y sean los datos extraídos junto con su nivel de confianza.
Asegúrate de que la respuesta sea solo el array JSON, comenzando con [ y terminando con ]. No incluyas prefijos (como "Confianza:"), explicaciones o cualquier otro texto. El objetivo es que la salida sea directamente parseable como un array donde el primer elemento sea el diccionario de datos.

Ejemplo de estructura esperada (los valores y confianzas serán los extraídos del documento):
[
{
"Certificado de fianza Web": {"valor extraído": Confianza},
"Fecha": "valor extraído": Confianza,
"Nombre de afianzado": "valor extraído": Confianza,
"R.U.T. N° del afianzado": "valor extraído": Confianza,
"Domicilio del afianzado": "valor extraído": Confianza,
"Nombre del mandante": "valor extraído": Confianza,
"R.U.T. N° del mandante": "valor extraído": Confianza,
"Domicilio del mandante": "valor extraído": Confianza,
"Obligación caucionada": "valor extraído": Confianza,
"Monto": "valor extraído": Confianza,
"Glosa": "valor extraído": Confianza
},
    """


# Generate content from the AI model
response = model.generate_content([prompt, image])
response_text = response.text

# Parse the AI response
# Extract and parse the JSON-like structure
m = re.search(r'(\[.*\])', response_text, re.S)
if m:
    json_text = m.group(1)
else:
    json_text = response_text.strip()

parsed_data = ast.literal_eval(json_text)  # Safely parse the string as a Python literal

if isinstance(parsed_data, list) and len(parsed_data) == 2:
    dict_data = parsed_data[0]  # First element: dictionary of extracted data
confidence_list = parsed_data[1]  # Second element: list of confidence values



# First document creation: Based on extracted data
def create_first_document(dict_data):
    document = Document()
    table = document.add_table(rows=1, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Campo'
    hdr_cells[1].text = 'Valor'

    for campo, (valor, _) in dict_data.items():  # Assuming values are tuples like (valor, confidence)
        row_cells = table.add_row().cells
        row_cells[0].text = campo
        row_cells[1].text = str(valor) or ''  # Use str() for safety and handle empty values

    document.save('resultado.docx')


create_first_document(dict_data)


# Second document creation: Hardcoded tables (as per your repeated section)
def create_second_document():
    document = Document()

    # First table for initial data
    table1 = document.add_table(rows=0, cols=4)
    table1.autofit = False
    table1.columns[0].width = Inches(1.5)  # Label 1
    table1.columns[1].width = Inches(2.0)  # Value 1
    table1.columns[2].width = Inches(1.5)  # Label 2
    table1.columns[3].width = Inches(2.5)  # Value 2

    row = table1.add_row().cells
    row[0].text = 'Certificado de Fianza:'
    row[1].text = "FJEH124U2U33"
    row[2].text = 'Fecha de Emisión:'
    row[3].text = '27 de junio de 2024'

    row = table1.add_row().cells
    row[0].text = 'Tomador:'
    merged_cell = row[1].merge(row[2]).merge(row[3])
    merged_cell.text = 'DISEÑO Y MANTENCIÓN DE JARDINES PABLO NAVARRETE EIRL'

    row = table1.add_row().cells
    row[0].text = 'R.U.T. N°:'
    merged_cell = row[1].merge(row[2]).merge(row[3])
    merged_cell.text = '76.594.422-8'

    row = table1.add_row().cells
    row[0].text = 'Dirección del Tomador:'
    merged_cell = row[1].merge(row[2]).merge(row[3])
    merged_cell.text = 'Condominio almendros norte, casa 47, comuna de Calera de Tango'

    row = table1.add_row().cells
    row[0].text = 'Beneficiario:'
    row[1].text = 'Hospital San José de Melipilla'
    row[2].text = 'R.U.T. N°'
    row[3].text = '61.602.123-0'

    # Add COBERTURA heading
    p = document.add_paragraph()
    run = p.add_run('COBERTURA')
    run.bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Second table for coverage details
    table2 = document.add_table(rows=0, cols=4)
    table2.autofit = False
    table2.columns[0].width = Inches(1.5)
    table2.columns[1].width = Inches(2.0)
    table2.columns[2].width = Inches(1.5)
    table2.columns[3].width = Inches(2.5)

    row = table2.add_row().cells
    row[0].text = 'Fecha de inicio:'
    row[1].text = '21-06-2024'
    row[2].text = 'Fecha de termino:'
    row[3].text = '23-10-2025'

    row = table2.add_row().cells
    row[0].text = 'Valor Asegurado:'
    merged_cell = row[1].merge(row[2]).merge(row[3])
    merged_cell.text = 'UF 56,50'

    row = table2.add_row().cells
    row[0].text = 'Finalidad:'
    merged_cell = row[1].merge(row[2]).merge(row[3])
    merged_cell.text = 'Fiel cumplimiento'

    document.save('insurance_policy_snippet.docx')


create_second_document()
