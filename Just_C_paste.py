import os
import docx
from docx.enum.section import WD_SECTION_START
from docx.enum.style import WD_STYLE_TYPE
from Bases import configurar_directorio_trabajo # Línea añadida

configurar_directorio_trabajo() # Línea añadida

doc = docx.Document()

# Crear estilos de lista numerada
try:
    # Asegúrate de que el estilo base 'List_Number_1' existe
    # Si no existe en tu plantilla por defecto, podrías necesitar crearlo o usar uno existente
    if 'List Number' not in doc.styles:
        # Intenta usar un estilo de lista base común si 'List_Number_1' no está
        # O maneja el error como prefieras
        print("Advertencia: El estilo 'List_Number' no se encontró. Usando 'ListParagraph' como base si existe.")
        base_style_name = 'List Paragraph' # O cualquier otro estilo de lista disponible
        if base_style_name not in doc.styles:
             raise KeyError("No se encontró un estilo de lista base adecuado ('List_Number_1' o 'List Paragraph').")
    else:
        base_style_name = 'List_Number_1'

    base_style = doc.styles[base_style_name]

    for i in range(1, 21):
        style_name = f'List_Number_{i}'
        # Verificar si el estilo ya existe antes de intentar añadirlo
        if style_name not in doc.styles:
            # Crear el estilo sin el argumento base_style
            new_style = doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
            # Asignar el estilo base después de la creación
            new_style.base_style = base_style
            # Opcionalmente, copiar propiedades específicas si es necesario
            # new_style.font.name = base_style.font.name
            # new_style.font.size = base_style.font.size
            # ... etc. para otras propiedades que quieras heredar explícitamente
        else:
            print(f"El estilo '{style_name}' ya existe.")
except KeyError as e:
    print(f"Error: {e}")
    # Considera manejar este error de for


# Titulos Nivel 0
doc.add_heading("RESOLUCIÓN EXENTA Nº1", level=0)

# Primeros parrafos
doc.add_heading("VISTOS", level=2)
doc.add_paragraph("Visto: La Ley N° 19.880, de 2003, que establece normas sobre los actos administrativos; la Ley N° 20.285, de 2008, sobre acceso a la información pública; la Ley N° 21.000, de 2017, que establece un sistema de compras públicas; el Decreto Exento N° 1.000, de 2020, del Ministerio de Salud; y la Resolución Exenta N° 1.000, de 2020, del Ministerio de Salud.", style="List_Number_1")
doc.add_paragraph("Que, el Hospital de San José de Melipilla perteneciente a la red de salud del Servicio de Salud Metropolitano Occidente, tiene como misión otorgar una atención integral, oportuna y resolutiva a las personas y familias de la provincia de Melipilla y sus alrededores, con un equipo de salud competente, comprometido y solidario, entregando un servicio de calidad y seguridad, en coordinación con la red asistencial;", style="List_Number_1")
doc.add_paragraph("Que, dada la naturaleza del Establecimiento, la atención de los beneficiarios requiere una oportuna e inmediata resolución, que no puede en caso alguno diferirse en el tiempo, lo que nos compromete a disponer en forma constante, continua y permanente de los servicios necesarios para responder adecuadamente a la demanda asistencial y administrativa a su población beneficiaria.", style="List_Number_1")

run_1 = "Que, existe la necesidad"
run_2 = "suministro de insumos y accesorios para terapia de presión negativa con equipos en comodato"
run_3 = ", a fin de entregar una prestación de salud integral y oportuna a los usuarios del Hospital de San José de Melipilla, y de esta manera dar cumplimiento con el tratamiento de los pacientes."

# -- Crear el párrafo y añadir los runs con formato --
vistos_parrafo_4 = doc.add_paragraph()
vistos_parrafo_4.add_run(run_1 + " ") # -- Añadir espacio al final si es necesario --
run_bold = vistos_parrafo_4.add_run(run_2)
run_bold.bold = True
vistos_parrafo_4.add_run(run_3)

vistos_parrafo_4.style = "List_Number_1"

doc.add_paragraph("Que corresponde asegurar la transparencia en este proceso y conocer las condiciones de oferta imperantes en el mercado bajo la modalidad de la licitación pública en el sistema de compras y contratación públicas establecido en la Ley Nº 19.886 y su Reglamento.", style="List_Number_1")
doc.add_paragraph("Que, considerando los montos de la contratación y en virtud de lo establecido en las resoluciones N°7/2019 y 16/2020 de la Contraloría General de la República, la presenta contratación no está sometida al trámite de toma de razón.", style="List_Number_1")
doc.add_paragraph("Que revisado el catálogo de bienes y servicios ofrecidos en el sistema de información Mercado Público, se ha verificado la ausencia de contratos marcos vigentes para el servicio antes mencionado.", style="List_Number_1")
doc.add_paragraph("Que, en consecuencia y en mérito de lo expuesto, para esta contratación se requiere llamar a licitación pública, debiendo esta regularse por la Bases Administrativas, Técnicas, Formularios y Anexos que se aprueban a través del presente acto administrativo.", style="List_Number_1")
doc.add_paragraph("Que, en razón de lo expuesto y la normativa vigente;", style="List_Number_1")



doc.add_heading("CONSIDERANDO", level=2)
doc.add_paragraph("Que dada la alta complejidad que caracteriza al Hospital San José de Melipilla, obliga a efectuar mejoras constantes y permanentes a fin de brindar a toda nuestra comunidad el desarrollo de diversas funciones con alta calidad que el sistema público puede brindar.")


# -- Ahora vamos con la parte de resolición --
doc.add_heading("RESOLUCIÓN", level=2)
run_1 = "LLÁMASE "
run_2 = "a Licitación Pública Nacional a través del Portal Mercado Público, para la compra de"
run_3 = "Suministro de Insumos y Accesorios para Terapia de Presión Negativa con Equipos en Comodato"
run_4 = "para el Hospital San José de Melipilla."

resolucion_parrafo_1 = doc.add_paragraph()
resolucion_parrafo_1.add_run(run_1+" ").bold = True
resolucion_parrafo_1.add_run(run_2+"")
resolucion_parrafo_1.add_run(run_3+" ").bold = True
resolucion_parrafo_1.add_run(run_4)

run_1 = "ACOGIENDOSE"
run_2 = "al Art.º 25 del decreto 250 que aprueba el reglamento de la ley Nº 19.886 de las Bases sobre Contratos Administrativos de Suministros y Prestación de Servicios, se reduce el tiempo de publicación de las bases en el portal de Mercado Público de 20 a 10 días, ya que se trata de la contratación de bienes o servicios de simple y objetiva especificación, y que conlleva un esfuerzo menor en la preparación de ofertas."

resolocion_parrafo_2 = doc.add_paragraph()
resolocion_parrafo_2.add_run(run_1+" ").bold = True
resolocion_parrafo_2.add_run(run_2)

run_1 = "APRUÉBENSE las bases administrativas, técnicas y anexos N.º 1, 2, 3, 4, 5, 6, 7, 8 y 9"
run_2 = "desarrollados para efectuar el llamado a licitación, que se transcriben a continuación:"
resolucion_parrafo_3 = doc.add_paragraph()
resolucion_parrafo_3.add_run(run_1+" ").bold = True
resolucion_parrafo_3.add_run(run_2)

# Nueva Sección
new_section = doc.add_section(WD_SECTION_START.NEW_PAGE)
doc.add_heading("BASES ADMINISTRATIVAS PARA EL SUMINISTRO DE INSUMOS Y ACCESORIOS PARA TERAPIA DE PRESIÓN NEGATIVA CON EQUIPOS EN COMODATO PARA EL HOSPITAL SAN JOSÉ DE MELIPILLA", level=1)
doc.add_heading("RESOLUCIÓN", level=2)
doc.add_paragraph("En Santiago, a 1 de enero de 2023, se resuelve lo siguiente:")

run_1 = "Antecedentes Básicos de la ENTIDAD LICITANTE"
doc.add_paragraph(run_1, style="List_Number_1").bold = True

tabla_entidad_licitante = doc.add_table(6, 2, style="Table Grid")

# Fila 1
tabla_entidad_licitante.cell(0, 0).text = "Razón Social del organismo"
tabla_entidad_licitante.cell(0, 1).text = "Hospital San José de Melipilla"

# Fila 2
tabla_entidad_licitante.cell(1, 0).text = "Unidad de Compra"
tabla_entidad_licitante.cell(1, 1).text = "Unidad de Abastecimiento de Bienes y Servicios"

# Fila 3
tabla_entidad_licitante.cell(2, 0).text = "R.U.T. del organismo"
tabla_entidad_licitante.cell(2, 1).text = "61.602.123-0"

# Fila 4
tabla_entidad_licitante.cell(3, 0).text = "Dirección"
tabla_entidad_licitante.cell(3, 1).text = "O’Higgins #551"

# Fila 5
tabla_entidad_licitante.cell(4, 0).text = "Comuna"
tabla_entidad_licitante.cell(4, 1).text = "Melipilla"

# Fila 6
tabla_entidad_licitante.cell(5, 0).text = "Región en que se genera la Adquisición"
tabla_entidad_licitante.cell(5, 1).text = "Región Metropolitana"

doc_path = 'resolucion_estilos.docx'
doc.save(doc_path)
# --- Fin de la creación y guardado del documento original ---


