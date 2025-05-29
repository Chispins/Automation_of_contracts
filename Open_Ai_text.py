import os

from openai import OpenAI
from PIL import Image
import base64
import io
import docx
import fitz


os.chdir(r"\\10.5.130.24\Abastecimiento\Compartido Abastecimiento\Otros\Licitaciones_Testing\Nueva carpeta (10)\Nueva carpeta\Nueva carpeta\Nueva carpeta (2)")


# Initialize OpenAI API
openai_api_key = "AIzaSyAgXYnDNJ6gQmMVUk88G6wGPdXz-qG2Gbw"
client = OpenAI(
    api_key=openai_api_key,
    base_url="https://generativelanguage.googleapis.com/v1beta/openai/"
)

model = "gemini-2.5-flash-preview-04-17"

# Cargar y convertir la imagen a base64

image_name = "FIEL CUMPLIMIENTO.pdf"

import re
regex = r".*(.pdf)"
formato = re.match(regex, image_name)


# Si es PDF
if formato != None:
    try:
        pdf_document = fitz.open(image_name)
        primera_pagina = pdf_document.load_page(0)
        pix = primera_pagina.get_pixmap(alpha=False)
        img_bytes = pix.tobytes("jpeg")
        image_base64 = base64.b64encode(img_bytes).decode('utf-8')
        pdf_document.close()
    except Exception as e:
        print(f"Error al procesar el PDF: {e}")

# Si es Imagen
else:
    image = Image.open(image_name)
# Convertir la imagen a bytes y luego a base64
    buffer = io.BytesIO()
    image.save(buffer, format="JPEG")
    image_bytes = buffer.getvalue()
    image_base64 = base64.b64encode(image_bytes).decode('utf-8')

# Crear el mensaje con la imagen
response = client.chat.completions.create(
    model=model,
    messages=[
        {"role": "system", "content": "You are a helpful assistant."},
        {
            "role": "user",
            "content": [
                {"type": "text", "text": """
                Certificado de fianza Web",
"Fecha",
"Nombre de afianzado",
"R.U.T. N° del afianzado",
"Domicilio del afianzado",
"Nombre del mandante":,
"R.U.T. N° del mandante":,
"Domicilio del mandante":,
"Obligación caucionada":,
"Monto":,
"Glosa"
Those are the fields, but the output should follow the same formating as this example 
```json
[
  {"label": "Certificado de fianza Web", "text_content": "W6498-017917"},
  {"label": "Fecha", "text_content": "Santiago, 18 de MARZO de 2025"},
  {"label": "Nombre de afianzado", "text_content": "SERVICIOS DE BIOINGENIERIA LIMITADA"},
  {"label": "R.U.T. N° del afianzado", "text_content": "76.644.150-5"},
  {"label": "Domicilio del afianzado", "text_content": "AVDA. CONDELL 1680, ÑUÑOA, REGIÓN METROPOLITANA DE SANTIAGO"},
  {"label": "Nombre del mandante", "text_content": "HOSPITAL DE MELIPILLA"},
  {"label": "R.U.T. N° del mandante", "text_content": "61.602.123-0"},
  {"label": "Domicilio del mandante", "text_content": "O'HIGGINS 551, MELIPILLA, REGIÓN METROPOLITANA DE SANTIAGO"},
  {"label": "Obligación caucionada", "text_content": "Fiel cumplimiento"},
  {"label": "Monto", "text_content": "$2.000.000"},
  {"label": "Glosa", "text_content": "Para garantizar el fiel cumplimiento del contrato denominado: COMPRA DEL SERVICIO DE MANTENCION PREVENTIVA DE EQUIPOS DE MONITOREO PARA EL HOSPITAL SAN JOSE DE MELIPILLA\" ID 1057480-12-LE25 y/o de las obligaciones laborales y sociales del adjudicatario"}
]
```
                """},
                {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{image_base64}"}}
            ]
        }
    ]
)


print(response.choices[0].message.content)

# Obtener la respuesta del modelo Gemini
respuesta = response.choices[0].message.content

# Importar las bibliotecas necesarias
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches

# Crear un nuevo documento
documento = Document()

# Añadir un título
documento.add_heading('Certificado de Fianza', 0)

# Crear una tabla con 2 columnas
tabla = documento.add_table(rows=1, cols=2)
tabla.style = 'Table Grid'

# Configurar encabezados
encabezados = tabla.rows[0].cells
encabezados[0].text = 'Campo'
encabezados[1].text = 'Valor'

# Definir los campos que queremos extraer
# Definir los campos que queremos extraer
campos = [
    "Tomador",
    "RUT",
    "Asegurado",
    "Beneficiario",
    "Dirección del Tomador",
    "Ciudad",
    "Cobertura",
    "Vigencia del Seguro",
    "Numero de Días",
    "Valor Asegurado",
    "Prima Neta",
    "IVA",
    "Total a Pagar",
    "Valor a pagar en Letra"
]

# Procesar la respuesta para extraer pares clave-valor
lineas = respuesta.split('\n')
valores = {}

# Intentar diferentes patrones de formato en la respuesta
for linea in lineas:
    for campo in campos:
        # Patrón 1: "campo", "text_content": "valor"}
        patron1 = f'"{campo}", "text_content": "(.+?)"'
        import re

        match1 = re.search(patron1, linea)
        if match1:
            valores[campo] = match1.group(1)
            break

        # Patrón 2: campo: valor
        if f"{campo}:" in linea:
            valor = linea.split(f"{campo}:", 1)[1].strip()
            valores[campo] = valor
            break

        # Patrón 3: "campo": "valor"
        if f'"{campo}":' in linea:
            try:
                valor = linea.split(f'"{campo}":', 1)[1].strip()
                # Eliminar comillas y comas al final si existen
                valor = valor.strip('",')
                valores[campo] = valor
                break
            except:
                pass

# Añadir los campos y valores a la tabla
for campo in campos:
    fila = tabla.add_row().cells
    fila[0].text = campo
    fila[1].text = valores.get(campo, "")


# Guardar el documento
documento.save("tabla_fea.docx")

print("Documento Word creado con éxito.")

# Rellenado de la Tabla

# Cargar el documento
doc = docx.Document("prototipo_tabla.docx")


# Diccionario con los datos a rellenar
data = {
    "Tomador": "Servicios de Bioingenieria Limitada",
    "RUT": "76.644.150-5",
    "Asegurado": "Hospital San José de Melipilla",
    "Beneficiario": "Hospital San José de Melipilla",
    "Dirección del Tomador": "Avenida Condell 1680",
    "Ciudad": "Nuñoa",
    "Cobertura": "Fiel Cumplimiento",
    "Vigencia del Seguro": "01/01/2025–02/01-2025",
    "Numero de Días": "100",
    "Valor Asegurado": "699,50",
    "Prima Neta": "1633,92",
    "IVA": "3,21",
    "Total a Pagar": "20,13",
    "Valor a pagar en Letra": "Veinte coma trece UF"
}

# Recorrer tablas y filas
for table in doc.tables:
    for row in table.rows:
        cells = row.cells
        for i in range(len(cells) - 1):  # Evitar último índice para no salir del rango
            cell_text = cells[i].text.strip()
            if cell_text in data:
                cells[i + 1].text = data[cell_text]

# Guardar el documento actualizado
doc.save("prototipo_tabla_rellenado.docx")
