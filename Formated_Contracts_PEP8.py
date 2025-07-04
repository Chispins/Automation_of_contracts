import os
import re
import random
import copy
import shutil
import tempfile
from datetime import datetime
from docx import Document
from Formated_Base_PEP8 import aplicar_formato_global
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_BREAK
from win32com import client
import pythoncom
from docxtpl import DocxTemplate
from win32com.client import constants



def configurar_directorio_trabajo():
    """Configura el directorio de trabajo en la subcarpeta 'Files'."""
    cwd = os.getcwd()
    target_dir_name = "Files"
    wd = os.path.join(cwd, target_dir_name)
    pattern = r"Files\\Files"
    if re.search(pattern, wd):
        wd = wd.replace(r"\Files\Files", r"\Files")
    if os.path.isdir(wd):
        os.chdir(wd)
        print(f"Directorio de trabajo cambiado a: {wd}")
    else:
        print(f"Advertencia: El directorio '{wd}' no existe o no es válido.")
# Set up the working directory
configurar_directorio_trabajo()


def crear_numeracion(doc):
    """Crea un formato de numeración y devuelve su ID único."""
    part = doc._part
    if not hasattr(part, 'numbering_part'):
        part._add_numbering_part()
    num_id = random.randint(1000, 9999)
    return num_id


def aplicar_numeracion(parrafo, num_id, nivel=0):
    """
    Aplica numeración a un párrafo con un ID y nivel específicos.
    Elimina numeración previa si existe y ajusta la sangría.
    """
    p = parrafo._p
    pPr = p.get_or_add_pPr()

    # Eliminar numeración previa
    for child in pPr.iterchildren():
        if child.tag.endswith('numPr'):
            pPr.remove(child)

    # Crear nueva numeración
    numPr = OxmlElement('w:numPr')
    ilvl = OxmlElement('w:ilvl')
    ilvl.set(qn('w:val'), str(nivel))
    numPr.append(ilvl)
    numId = OxmlElement('w:numId')
    numId.set(qn('w:val'), str(num_id))
    numPr.append(numId)
    pPr.append(numPr)

    # Configurar sangría
    ind = OxmlElement('w:ind')
    ind.set(qn('w:left'), '720')
    ind.set(qn('w:hanging'), '360')
    pPr.append(ind)

    return parrafo


# --- Funciones para Extracción y Copia de Secciones ---
def extraer_seccion_completa(doc_cargado, titulo_seccion):
    """
    Extrae una sección completa (párrafos y tablas) hasta el siguiente encabezado
    del mismo nivel o superior.
    Retorna el encabezado, los elementos de la sección y el nivel del encabezado.
    """
    seccion_heading = None
    elementos_seccion = []
    nivel_seccion = None
    indice_inicio = -1
    indice_fin = None

    # Buscar el encabezado de la sección
    for i, element in enumerate(doc_cargado.paragraphs):
        if element.style.name.startswith('Heading') and element.text.strip() == titulo_seccion:
            seccion_heading = element
            indice_inicio = i
            try:
                nivel_seccion = int(element.style.name.split()[-1])
            except (ValueError, IndexError):
                nivel_seccion = 2  # Nivel por defecto
            break

    if indice_inicio == -1:
        print(f"No se encontró la sección '{titulo_seccion}'")
        return None

    # Encontrar el fin de la sección
    for i in range(indice_inicio + 1, len(doc_cargado.paragraphs)):
        parrafo = doc_cargado.paragraphs[i]
        if parrafo.style.name.startswith('Heading'):
            try:
                nivel_actual = int(parrafo.style.name.split()[-1])
                if nivel_actual <= nivel_seccion:
                    indice_fin = i
                    break
            except (ValueError, IndexError):
                pass

    if indice_fin is None:
        indice_fin = len(doc_cargado.paragraphs)

    # Agregar párrafos entre inicio y fin (excepto encabezados)
    for i in range(indice_inicio + 1, indice_fin):
        parrafo = doc_cargado.paragraphs[i]
        elementos_seccion.append(('parrafo', parrafo, i))

    # Mapa de posiciones de párrafos en XML
    parrafos_posicion_xml = {p._p: i for i, p in enumerate(doc_cargado.paragraphs)}

    # Extraer posición de tablas en XML
    tablas_posicion = []
    for i, tabla in enumerate(doc_cargado.tables):
        tabla_p = tabla._tbl
        elemento_anterior = tabla_p.getprevious()
        pos_anterior = -1
        while elemento_anterior is not None:
            if elemento_anterior in parrafos_posicion_xml:
                pos_anterior = parrafos_posicion_xml[elemento_anterior]
                break
            elemento_anterior = elemento_anterior.getprevious()
        if indice_inicio <= pos_anterior < indice_fin:
            tablas_posicion.append((i, tabla, pos_anterior + 0.5))

    # Agregar tablas a la sección
    for i, tabla, pos in tablas_posicion:
        elementos_seccion.append(('tabla', tabla, pos))

    # Ordenar elementos por posición
    elementos_seccion.sort(key=lambda x: x[2])
    elementos_seccion = [(tipo, elem) for tipo, elem, _ in elementos_seccion]

    return seccion_heading, elementos_seccion, nivel_seccion if seccion_heading else None


def copiar_seccion_completa(doc_destino, seccion_heading, elementos_seccion, nivel_seccion):
    # Copiar encabezado
    nuevo_encabezado = doc_destino.add_heading(level=nivel_seccion)
    if seccion_heading._p.pPr is not None:
        nuevo_encabezado._p.append(copy.deepcopy(seccion_heading._p.pPr))
    # Copiar contenido y marcadores del encabezado
    for child in seccion_heading._p:
        if child.tag in (qn('w:r'), qn('w:bookmarkStart'), qn('w:bookmarkEnd')):
            nuevo_encabezado._p.append(copy.deepcopy(child))
        elif child.tag == qn('w:bookmarkStart'):
            new_bm_start = OxmlElement('w:bookmarkStart')
            bm_id, bm_name = child.get(qn('w:id')), child.get(qn('w:name'))
            if bm_id:
                new_bm_start.set(qn('w:id'), bm_id)
            if bm_name:
                new_bm_start.set(qn('w:name'), bm_name)
            nuevo_encabezado._p.append(new_bm_start)
        elif child.tag == qn('w:bookmarkEnd'):
            new_bm_end = OxmlElement('w:bookmarkEnd')
            bm_id = child.get(qn('w:id'))
            if bm_id:
                new_bm_end.set(qn('w:id'), bm_id)
            nuevo_encabezado._p.append(new_bm_end)

    # ID único de numeración para la sección
    seccion_num_id = None
    # Copiar elementos de la sección ( resto del código sin cambios )
    for tipo, elemento in elementos_seccion:
        if tipo == 'parrafo':
            origen = elemento
            # Encabezados
            if origen.style.name.startswith('Heading'):
                nivel = int(origen.style.name.split()[-1])
                destino = doc_destino.add_heading(level=nivel)
                if origen._p.pPr is not None:
                    destino._p.append(copy.deepcopy(origen._p.pPr))
                for child in origen._p:
                    if child.tag in (qn('w:r'), qn('w:bookmarkStart'), qn('w:bookmarkEnd')):
                        destino._p.append(copy.deepcopy(child))
                continue

            # Párrafos normales
            nuevo = doc_destino.add_paragraph(style=origen.style.name)
            if origen._p.pPr is not None:
                nuevo._p.append(copy.deepcopy(origen._p.pPr))
            for child in origen._p:
                if child.tag in (qn('w:r'), qn('w:bookmarkStart'), qn('w:bookmarkEnd')):
                    nuevo._p.append(copy.deepcopy(child))

            # Lista numerada
            if origen.style.name == 'List Number':
                if seccion_num_id is None:
                    seccion_num_id = crear_numeracion(doc_destino)
                aplicar_numeracion(nuevo, seccion_num_id)

        elif tipo == 'tabla':
            # No copiar tablas en esta etapa, solo agregar un marcador de posición
            doc_destino.add_paragraph(f"[[TABLE_PLACEHOLDER]]")


def copiar_tablas_con_win32(source_path, intermediate_path, prototipo_path, output_path):
    pythoncom.CoInitialize()
    try:
        word = client.Dispatch("Word.Application")
        word.Visible = False

        src = word.Documents.Open(os.path.abspath(source_path))
        dst = word.Documents.Open(os.path.abspath(intermediate_path))

        # 1. Pegar tablas 8, 9 y 10 en los primeros 3 marcadores
        table_indices = [8, 9, 10, 11]
        placeholders = [p for p in dst.Paragraphs if '[[TABLE_PLACEHOLDER]]' in p.Range.Text]

        for i, idx in enumerate(table_indices):
            if i < len(placeholders) and idx <= src.Tables.Count:
                src.Tables.Item(idx).Range.Copy()
                placeholders[i].Range.PasteAndFormat(constants.wdFormatOriginalFormatting)

        # 2. Pegar las dos tablas de `prototipo_tabla_rellenado.docx` en el último marcador
        if placeholders:
            last_ph = placeholders[-1]
            prot = word.Documents.Open(os.path.abspath(prototipo_path))
            for j in range(1, prot.Tables.Count + 1):
                prot.Tables.Item(j).Range.Copy()
                last_ph.Range.PasteAndFormat(constants.wdFormatOriginalFormatting)
            prot.Close(False)

        dst.SaveAs(os.path.abspath(output_path))
        dst.Close(False)
        src.Close(False)
        word.Quit()
    finally:
        pythoncom.CoUninitialize()

# --- Flujo Principal del Programa ---
def main():
    """Función principal que ejecuta el proceso de copia de secciones de un documento."""
    # Configurar directorio de trabajo
    configurar_directorio_trabajo()

    # Obtener directorio actual
    current_dir = os.getcwd()
    print(f"Directorio de trabajo actual: {current_dir}")

    # Crear directorio temporal
    temp_dir = tempfile.mkdtemp()
    print(f"Directorio temporal creado: {temp_dir}")

    # Definir archivo original y copia temporal
    original_file = "contrato_automatizado_rendered.docx"
    original_path = os.path.join(current_dir, original_file)
    temp_file = os.path.join(temp_dir, f"temp_{datetime.now().strftime('%Y%m%d_%H%M%S')}_{original_file}")

    # Verificar existencia del archivo original
    if not os.path.exists(original_path):
        raise FileNotFoundError(f"El archivo original {original_path} no se encuentra.")

    # Copiar archivo a directorio temporal
    shutil.copy2(original_path, temp_file)
    print(f"Archivo original copiado a: {temp_file}")

    # Cargar documento temporal y crear documento destino
    print("Cargando copia temporal del documento...")
    word = Document(temp_file)

    # Trabajo con la portada
    portada_template = DocxTemplate("portada_melipilla_contrato.docx")
    context = {"numero_contrato": "123456789", "involucrados": "MMJ,ASJ,MMA,MVO,PQE"}
    from Jinja_2 import context_for_template2
    portada_template.render(context_for_template2)
    portada_template.save("portada_melipilla_contrato_renderizado.docx")
    doc = Document("portada_melipilla_contrato_renderizado.docx")
    print("Documento temporal cargado.")


    # Titulo
    heading_paragraph = doc.add_heading('', level=0)

    # Add the first part of the text as a run
    run1 = heading_paragraph.add_run("CONTRATACION DEL SUMINISTRO DE INSUMOS Y ACCESORIOS PARA TERAPIA DE PRESIÓN NEGATIVA CON EQUIPOS EN COMODATO PARA EL HOSPITAL SAN JOSÉ DE MELIPILLA ENTRE")
    run1.font.size = Pt(11)

    # Add a line break
    run1.add_break(WD_BREAK.LINE)

    # Add the second part of the text as another run
    run2 = heading_paragraph.add_run("MEDCORP S.A Y HOSPITAL SAN JOSÉ DE MELIPILLA")
    heading_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run2.font.size = Pt(11)


    #--- Introducción del Contrato
    intro = doc.add_paragraph()
    intro.add_run("En Melipilla, a {{ fecha_contrato }}, entre El Hospital San José de Melipilla, RUT N° 61.602.123-0, representado por don Óscar Vargas Duranti, Chileno, RUT Nº 8.578.583-4, de profesión Médico Cirujano, domiciliado para estos efectos en calle O’Higgins Nº 551,  comuna de Melipilla, en adelante “El Hospital” y, por otra parte como proveedor ")
    intro.add_run("“{{ nombre_proveedor }}”").bold=True
    intro.add_run(" en adelante")
    intro.add_run("“el proveedor adjudicado”").bold = True
    intro.add_run(", RUT {{ rut_proveedor }}, representado legalmente por {{ representante_legal }}, cedula nacional de identidad N° {{ rut_representante_legal }}, con domicilio para estos efectos en {{domicilio_representante_legal}}, viene a suscribir el siguiente contrato:")


    doc.add_heading("Primero: Declaraciones.", level = 2)
    parrafo_primero = doc.add_paragraph()
    parrafo_primero.add_run("Se deja constancia que el presente servicio ha sido sometido a licitación pública a través del Sistema de Información de la Dirección de Compras y contratación pública (en adelante, “Portal Mercado Público), según ")
    parrafo_primero.add_run("ID {{ id_licitacion }} ").bold = True
    parrafo_primero.add_run("aprobándose las bases de licitación mediante la Resolución Exenta Nº {{ numero_resolucion_aprobacion }} de fecha {{ fecha_resolucion_aprobacion }} y adjudicándose al proveedor ")
    parrafo_primero.add_run("“{{ nombre_proveedor }}”").bold = True
    parrafo_primero.add_run("en virtud de resolución exenta N°{{ resolucion_exenta }} de fecha {{ fecha_resolucion_exenta }}.")
    parrafo_primero.add_run(r"\n")
    parrafo_primero.add_run("“El Proveedor/Empresa” declara que cuenta con la capacidad técnica y el personal necesario para la realización de los servicios contratados por el Hospital San José de Melipilla.")

    # Segundo
    doc.add_heading("Segundo: Objeto del contrato y valor de la contratación.", level = 2)
    doc.add_paragraph("El presente contrato tiene por objeto la compra de Suministro de Insumos y Accesorios para Terapia de Presión Negativa con Equipos en Comodato para el Hospital San José de Melipilla, a fin de entregar una prestación de salud integral y oportuna a los usuarios del Hospital  y de esta manera dar cumplimiento con el tratamiento de los pacientes.", style = "List Bullet")
    doc.add_paragraph("La adjudicación se realizará a valor unitario por hora, generando un contrato por $350.000.000.- (Impuestos incluidos.",style="List Bullet")

    doc.add_paragraph("Esto en virtud del siguiente cuadro:")
    doc.add_paragraph("[[TABLE_PLACEHOLDER]]")
    # Copiar la primera tabla del documento 'word'
    #


    # Lista de secciones a procesar
    """secciones = [
        "BASES ADMINISTRATIVAS PARA EL SUMINISTRO DE INSUMOS Y ACCESORIOS PARA TERAPIA DE PRESIÓN NEGATIVA CON EQUIPOS EN COMODATO PARA EL HOSPITAL SAN JOSÉ DE MELIPILLA",
        "Condiciones Contractuales, Vigencia de las Condiciones Comerciales, Operatoria de la Licitación y Otras Cláusulas:",
        "BASES TECNICAS PARA EL SUMINISTRO DE INSUMOS Y ACCESORIOS PARA TERAPIA DE PRESIÓN NEGATIVA CON EQUIPOS EN COMODATO PARA EL HOSPITAL SAN JOSÉ DE MELIPILLA"
    ]"""
    secciones = [
        "De los Productos",
        "Sobre los equipos solicitados en comodato para el uso de insumos condición obligatoria",
        "Entrega y Recepción",
        "MOTIVOS DE RECHAZO POR OBSERVACIÓN FÍSICA (ya iniciado en contrato):",
        "Tercero Documentos Integrantes",
        "Cuarto Modificaciones Del Contrato",
        "Quinto Gastos e Impuestos",
        "Sexto Efectos derivados de Incumplimiento del proveedor",
        "Séptimo Garantía de Fiel Cumplimiento de Contrato.",
        "Octavo Cobro de la Garantía de Fiel Cumplimiento de Contrato",
        "Noveno Término anticipado del contrato",
        "Décimo Resciliación de Mutuo Acuerdo",
        "Décimo Primero Procedimiento para Aplicación de Medidas derivadas de incumplimientos",
        "Decimo Segundo Emisión de la Orden de Compra",
        "Décimo Tercero Del Pago",
        "Décimo Cuarto Vigencia del Contrato",
        "Décimo Quinto Administrador del Contrato y/o Referente Técnico.",
        "Décimo Sexto Pacto de Integridad",
        "Décimo Séptimo Comportamiento ético del Adjudicatario.",
        "Décimo Octavo Auditorías",
        "Décimo Noveno Confidencialidad",
        "Vigésimo Propiedad de la información",
        "Vigésimo Primero Saldos insolutos de remuneraciones o cotizaciones de seguridad social.",
        "Vigésimo Segundo Normas Laborales Aplicables",
        "Vigésimo Tercero Cambio de personal del proveedor adjudicado.",
        "Vigésimo Cuarto Cesión y subcontratación.",
        "Vigésimo Quinto Discrepancias",
        "Vigésimo Sexto Constancia",
        "Vigésimo Séptimo  Ejemplares",
        "Vigésimo Octavo La personería"
    ]

    #"Entrega y Recepción",
        #"Documentos integrantes",
        #"Modificación del Contrato",
    # Documentos Integrantes, es Tercero
    # Modificación del Contrato, Cuarto
    # Gastos e Impuestos, Quinto
    # Efectos derivados de Incumplimiento del proveedor, Sexto
    # Garantía de Seriedad de la Oferta, Séptimo



    # Extraer y copiar cada sección
    # Extraer y copiar cada sección
    for seccion in secciones:
        print(f"Procesando sección {seccion}...")
        resultado = extraer_seccion_completa(word, seccion)
        if resultado:
            encabezado, elementos, _ = resultado  # Ignoramos el nivel original de la sección
            # Forzamos el nivel 2 para el encabezado principal de la sección
            copiar_seccion_completa(doc, encabezado, elementos, 2)
            print(f"Sección {seccion} copiada como Nivel 2.")
        else:
            print(f"Sección {seccion} no encontrada.")

    aplicar_formato_global(doc)

    # Guardar documento resultante
    output_file = f"contrato_automatizado_over.docx"
    output_path = os.path.join(current_dir, output_file)
    print(f"Guardando documento nuevo como: {output_path}...")
    doc.save(output_path)
    print("Documento nuevo guardado exitosamente.")
    print("Nota: El archivo original no ha sido modificado. Solo se trabajó con una copia temporal.")
    prototipo_path = os.path.abspath("prototipo_tabla_rellenado.docx")
    final_with_tables = os.path.join(current_dir, "contrato_automatizado_tablas.docx")

    final_with_tables = "contrato_automatizado_tablas.docx"
    copiar_tablas_con_win32(original_path, output_path, prototipo_path, final_with_tables)
    print(f"Documento final con tablas guardado en: {final_with_tables}")

    # Limpiar directorio temporal
    try:
        shutil.rmtree(temp_dir)
        print("Directorio temporal y archivos temporales eliminados.")
    except Exception as e:
        print(f"Error al eliminar el directorio temporal: {e}")

# Now I want

if __name__ == "__main__":
    main()
