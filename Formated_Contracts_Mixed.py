import os
import re
import random
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def configurar_directorio_trabajo():
    """Configura el directorio de trabajo en la subcarpeta 'Files'."""
    cwd = os.getcwd()
    target_dir_name = "Files"
    wd = os.path.join(cwd, target_dir_name)
    pattern = r"Files\\Files"
    if re.search(pattern, wd):
        wd = wd.replace(r"\Files\Files", r"\Files")
    if os.path.isdir(wd):
        os.chdir(wd)
        print(f"Directorio de trabajo cambiado a: {wd}")
    else:
        print(f"Advertencia: El directorio '{wd}' no existe o no es válido.")

def crear_numeracion(doc):
    """Crea un formato de numeración y devuelve su ID."""
    part = doc._part
    if not hasattr(part, 'numbering_part'):
        part._add_numbering_part()
    return random.randint(1000, 9999)

def aplicar_numeracion(parrafo, num_id, nivel=0):
    """Aplica numeración a un párrafo con el ID y nivel especificados."""
    p = parrafo._p
    pPr = p.get_or_add_pPr()
    for child in pPr.iterchildren():
        if child.tag.endswith('numPr'):
            pPr.remove(child)
    numPr = OxmlElement('w:numPr')
    ilvl = OxmlElement('w:ilvl')
    ilvl.set(qn('w:val'), str(nivel))
    numPr.append(ilvl)
    numId = OxmlElement('w:numId')
    numId.set(qn('w:val'), str(num_id))
    numPr.append(numId)
    pPr.append(numPr)
    ind = OxmlElement('w:ind')
    ind.set(qn('w:left'), '720')
    ind.set(qn('w:hanging'), '360')
    pPr.append(ind)
    return parrafo

def centrar_verticalmente_tabla(tabla):
    """Aplica alineación vertical centrada a todas las celdas de una tabla."""
    for fila in tabla.rows:
        for celda in fila.cells:
            tc = celda._tc
            tcPr = tc.get_or_add_tcPr()
            for vAlign in tcPr.findall(qn('w:vAlign')):
                tcPr.remove(vAlign)
            vAlign = OxmlElement('w:vAlign')
            vAlign.set(qn('w:val'), 'center')
            tcPr.append(vAlign)

def agregar_parrafo_con_texto(doc, texto, estilo=None, negrita=False):
    """Agrega un párrafo con texto y aplica estilo o formato si se especifica."""
    p = doc.add_paragraph(texto, style=estilo)
    if negrita and p.runs:
        p.runs[0].bold = True
    return p

def agregar_parrafo_con_runs(doc, partes, estilo=None):
    """Agrega un párrafo con múltiples runs y aplica formatos específicos."""
    p = doc.add_paragraph(style=estilo)
    for texto, formato in partes:
        run = p.add_run(texto)
        if formato:
            for fmt in formato.split(','):
                if fmt == 'bold':
                    run.bold = True
                elif fmt == 'underline':
                    run.underline = True
    return p

def crear_tabla(doc, datos, estilo='Table Grid', centrar=True):
    """Crea una tabla con los datos proporcionados y aplica formato."""
    filas = len(datos)
    columnas = len(datos[0]) if filas > 0 else 0
    tabla = doc.add_table(filas, columnas, style=estilo)
    for i, fila in enumerate(datos):
        for j, texto in enumerate(fila):
            celda = tabla.cell(i, j)
            celda.text = texto
    if centrar:
        centrar_verticalmente_tabla(tabla)
    return tabla

def agregar_contenido_celda(tabla, fila, columna, contenidos):
    """Agrega contenido a una celda con múltiples párrafos o runs formateados."""
    celda = tabla.cell(fila, columna)
    if celda.paragraphs and not celda.paragraphs[0].text:
        p_element = celda.paragraphs[0]._p
        celda._element.remove(p_element)
    for contenido in contenidos:
        if isinstance(contenido, str):
            celda.add_paragraph(contenido)
        else:
            p = celda.add_paragraph()
            for item in contenido:
                if isinstance(item, tuple) and len(item) == 2:
                    texto, formato = item
                    run = p.add_run(texto)
                    if formato:
                        for fmt in formato.split(','):
                            if fmt == 'bold':
                                run.bold = True
                            elif fmt == 'underline':
                                run.underline = True
                else:
                    # If it's not a tuple or has unexpected length, treat it as plain text
                    run = p.add_run(str(item))

def main():
    configurar_directorio_trabajo()
    doc = Document()
    list_style = 'List Number'
    if list_style not in doc.styles:
        doc.styles.add_style(list_style, WD_STYLE_TYPE.PARAGRAPH)

    # Crear IDs para numeración
    num_id_vistos = crear_numeracion(doc)
    num_id_resolucion = crear_numeracion(doc)
    num_id_bases_p1 = crear_numeracion(doc)
    num_id_requisitos = crear_numeracion(doc)
    num_id_consultas = crear_numeracion(doc)
    administrado_contrato_id_lista = crear_numeracion(doc)
    pacto_integridad_id = crear_numeracion(doc)

    # Títulos y secciones principales
    doc.add_heading("RESOLUCIÓN EXENTA Nº1", level=0)
    doc.add_heading("VISTOS", level=2)
    agregar_parrafo_con_texto(doc, "Lo dispuesto en la Ley Nº 19.886 de Bases sobre Contratos Administrativos de Suministro y Prestación de Servicios; el Decreto Supremo Nº 250 /04 modificado por los Decretos Supremos Nº 1763/09, 1383/11 y 1410/14 todos del Ministerio de Hacienda; D. S. 38/2005, Reglamento Orgánico de los Establecimientos de Menor Complejidad y de los Establecimientos de Autogestión en Red; en uso de las atribuciones que me confieren el D.F.L. Nº 1/2.005, en virtud del cual se fija el texto refundido, coordinado y sistematizado del D.L. 2.763/79 y de las leyes 18.933 y 18.469; lo establecido en los Decretos Supremos Nos 140/04, Reglamento Orgánico de los Servicios de Salud; la Resolución Exenta RA 116395/343/2024 de fecha 12/08/2024 del SSMOCC., la cual nombra Director del Hospital San José de Melipilla al suscrito; lo dispuesto por las Resoluciones 10/2017, 7/2019 y 8/2019 ambas de la Contraloría General de la República, y,")

    # Sección CONSIDERANDO
    doc.add_heading("CONSIDERANDO", level=2)
    vistos_items = [
        ("Visto: La Ley N° 19.880, de 2003, que establece normas sobre los actos administrativos...", list_style),
        ("Que, el Hospital de San José de Melipilla perteneciente a la red de salud...", list_style),
        ("Que, dada la naturaleza del Establecimiento...", list_style),
    ]
    for texto, estilo in vistos_items:
        p = agregar_parrafo_con_texto(doc, texto, estilo)
        aplicar_numeracion(p, num_id_vistos)

    vistos_p4 = doc.add_paragraph(style=list_style)
    vistos_p4.add_run("Que, existe la necesidad ")
    run_bold = vistos_p4.add_run("suministro de insumos y accesorios para terapia de presión negativa con equipos en comodato")
    run_bold.bold = True
    vistos_p4.add_run(", a fin de entregar una prestación de salud integral y oportuna...")
    aplicar_numeracion(vistos_p4, num_id_vistos)

    for texto in [
        "Que corresponde asegurar la transparencia en este proceso...",
        "Que, considerando los montos de la contratación...",
        "Que revisado el catálogo de bienes y servicios...",
        "Que, en consecuencia y en mérito de lo expuesto...",
        "Que, en razón de lo expuesto y la normativa vigente;"
    ]:
        p = agregar_parrafo_con_texto(doc, texto, list_style)
        aplicar_numeracion(p, num_id_vistos)

    # Sección RESOLUCIÓN
    doc.add_heading("RESOLUCIÓN", level=2)
    resolucion_p1 = doc.add_paragraph(style=list_style)
    resolucion_p1.add_run("LLÁMASE ").bold = True
    resolucion_p1.add_run("a Licitación Pública Nacional a través del Portal Mercado Público, para la compra de ")
    resolucion_p1.add_run("Suministro de Insumos y Accesorios para Terapia de Presión Negativa con Equipos en Comodato ").bold = True
    resolucion_p1.add_run("para el Hospital San José de Melipilla.")
    aplicar_numeracion(resolucion_p1, num_id_resolucion)

    resolucion_p2 = doc.add_paragraph(style=list_style)
    resolucion_p2.add_run("ACOGIENDOSE ").bold = True
    resolucion_p2.add_run("al Art.º 25 del decreto 250 que aprueba el reglamento de la ley Nº 19.886...")
    aplicar_numeracion(resolucion_p2, num_id_resolucion)

    resolucion_p3 = doc.add_paragraph(style=list_style)
    resolucion_p3.add_run("APRUÉBENSE las bases administrativas, técnicas y anexos N.º 1, 2, 3, 4, 5, 6, 7, 8 y 9 ").bold = True
    resolucion_p3.add_run("desarrollados para efectuar el llamado a licitación, que se transcriben a continuación:")
    aplicar_numeracion(resolucion_p3, num_id_resolucion)

    # Sección BASES ADMINISTRATIVAS
    doc.add_heading("BASES ADMINISTRATIVAS PARA EL SUMINISTRO DE INSUMOS Y ACCESORIOS PARA TERAPIA DE PRESIÓN NEGATIVA CON EQUIPOS EN COMODATO PARA EL HOSPITAL SAN JOSÉ DE MELIPILLA", level=1)
    doc.add_heading("REQUISITOS", level=2)
    agregar_parrafo_con_texto(doc, "En Santiago, a 1 de enero de 2023, se resuelve lo siguiente:")

    bases_p1 = doc.add_paragraph(style=list_style)
    bases_p1.add_run("Antecedentes Básicos de la ENTIDAD LICITANTE").bold = True
    aplicar_numeracion(bases_p1, num_id_bases_p1)

    tabla_p1_datos = [
        ["Razón Social del organismo", "Hospital San José de Melipilla"],
        ["Unidad de Compra", "Unidad de Abastecimiento de Bienes y Servicios"],
        ["R.U.T. del organismo", "61.602.123-0"],
        ["Dirección", "O'Higgins #551"],
        ["Comuna", "Melipilla"],
        ["Región en que se genera la Adquisición", "Región Metropolitana"]
    ]
    tabla_p1 = crear_tabla(doc, tabla_p1_datos)

    bases_p2 = doc.add_paragraph(style=list_style)
    bases_p2.add_run("Antecedentes Administrativos").bold = True
    aplicar_numeracion(bases_p2, num_id_bases_p1)

    tabla_p2_datos = [
        ["Nombre Adquisición", "Suministro de Insumos y Accesorios para Terapia de Presión Negativa con Equipos en Comodato"],
        ["Descripción", "El Hospital requiere generar un convenio por el SUMINISTRO DE INSUMOS Y ACCESORIOS PARA TERAPIA DE PRESIÓN NEGATIVA CON EQUIPOS EN COMODATO PARA EL HOSPITAL SAN JOSÉ DE MELIPILLA, en adelante “EL HOSPITAL”. El convenio tendrá una vigencia de 36 meses."],
        ["Tipo de Convocatoria", "Abierta"],
        ["Moneda o Unidad reajustable", "Pesos Chilenos"],
        ["Presupuesto Referencial", "$350.000.000.- (Impuestos incluidos)"],
        ["Etapas del Proceso de Apertura", "Una Etapa (Etapa de Apertura Técnica y Etapa de Apertura Económica en una misma instancia)."],
        ["Opciones de pago", "Transferencia Electrónica"],
        ["Tipo de Adjudicación", "Adjudicación por la totalidad"]
    ]
    tabla_p2 = crear_tabla(doc, tabla_p2_datos)

    bases_p2_runs = doc.add_paragraph()
    bases_p2_runs.add_run("* Presupuesto referencial:").underline = True
    bases_p2_runs.add_run(" El Hospital se reserva el derecho de aumentar, previo acuerdo entre las partes, hasta un 30% el presupuesto referencial estipulado en las presentes bases de licitación.")

    # Definiciones
    bases_p2_definiciones = doc.add_paragraph()
    bases_p2_definiciones.add_run("Definiciones").bold = True

    definiciones = [
        ("Proponente u oferente:", "El proveedor o prestador que participa en el proceso de licitación mediante la presentación de una propuesta, en la forma y condiciones establecidas en las Bases.", "List Number 3"),
        ("Administrador o coordinador Externo del Contrato", "Persona designada por el oferente adjudicado, quien actuará como contraparte ante el Hospital.", "List Number 3"),
        ("Dias Hábiles:", "Son todos los días de la semana, excepto los sábados, domingos y festivos.", "List Number 3"),
        ("Días Corridos:", "Son los días de la semana que se computan uno a uno en forma correlativa. Salvo que se exprese lo contrario, los plazos de días señalados en las presentes bases de licitación son días corridos. En caso que el plazo expire en días sábados, domingos o festivos se entenderá prorrogados para el día hábil siguiente.", "List Number 3"),
        ("Administrador del Contrato y/o Referente Técnico:", "Es el funcionario designado por el Hospital para supervisar la correcta ejecución del contrato, solicitar órdenes de compra, validar prefacturas, gestionar multas y/o toda otra labor que guarde relación con la ejecución del contrato.", "List Number 3"),
        ("Gestor de Contrato:", "Es el funcionario a cargo de la ejecución del presente proceso de Licitación, desde la publicación de las Bases hasta la generación del contrato en formato documental, como la elaboración de ficha en la plataforma “gestor de contrato” en el portal de mercado público, además de ser el responsable de dar seguimiento y cumplimiento a los procesos y plazos establecidos.", "List Number 3")
    ]
    for titulo, descripcion, estilo in definiciones:
        p = doc.add_paragraph(style=estilo)
        p.add_run(f"{titulo}").bold = True
        p.add_run(f" {descripcion}")

    # Tabla de tipos de licitación
    tabla_licitaciones_datos = [
        ["RANGO (en UTM)", "TIPO LICITACION PUBLICA", "PLAZO PUBLICACION EN DIAS CORRIDOS"],
        ["<100", "L1", "5"],
        ["<=100 y <1000", "LE", "10, rebajable a 5"],
        ["<=1000 y <2000", "LP", "20, rebajable a 10"],
        ["<=2000 y <5000", "LQ", "20, rebajable a 10"],
        ["<=5000", "LR", "30"]
    ]
    crear_tabla(doc, tabla_licitaciones_datos)

    bases_p3 = doc.add_paragraph(style=list_style)
    bases_p3.add_run("Etapas y plazos:").bold = True
    aplicar_numeracion(bases_p3, num_id_bases_p1)

    tabla_plazos_datos = [
        ["VIGENCIA DE LA PUBLICACION 10 DIAS CORRIDOS", "", ""],
        ["Consultas", "Hasta las 15:00 Horas de 4º (cuarto) día corrido de publicada la Licitación.", "Deben ser ingresadas al portal www.mercadopublico.cl"],
        ["Respuestas a Consultas", "Hasta las 17:00 Horas de 7º (séptimo) día corrido de publicada la Licitación.", "Deben ser ingresadas al portal www.mercadopublico.cl"],
        ["Aclaratorias", "Hasta 1 días corrido antes del cierre de recepción de ofertas.", "Deben ser ingresadas al portal www.mercadopublico.cl"],
        ["Recepción de ofertas", "Hasta las 17:00 Horas de 10º (décimo) día corrido de publicada la Licitación.", "Deben ser ingresadas al portal www.mercadopublico.cl"],
        ["Evaluación de las Ofertas", "Máximo 40 días corridos a partir del cierre de la Licitación.", ""],
        ["Plazo Adjudicaciones", "Máximo 20 días corridos a partir de la fecha del acta de evaluación de las ofertas.", ""],
        ["Suscripción de Contrato", "Máximo de 20 días hábiles desde la Adjudicación de la Licitación.", ""],
        ["Consideración", "Los plazos de días establecidos en la cláusula 3, Etapas y Plazos, son de días corridos, excepto el plazo para emitir la orden de compra, el que se considerará en días hábiles, entendiéndose que son inhábiles los sábados, domingos y festivos en Chile, sin considerar los feriados regionales.", ""]
    ]
    tabla_plazos = crear_tabla(doc, tabla_plazos_datos)
    tabla_plazos.cell(0, 0).merge(tabla_plazos.cell(0, 2))
    tabla_plazos.cell(0, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    tabla_plazos.cell(0, 0).paragraphs[0].runs[0].bold = True

    # Sección Consultas, Aclaraciones y Modificaciones
    doc.add_heading("Consultas, Aclaraciones y modificaciones a las bases.", level=2)
    consultas_items = [
        [("Las consultas de los participantes se deberán realizar únicamente a través del portal ", ""),
         (" www.mercadopublico.cl", "bold"),
         (" conforme el cronograma de actividades de esta licitación señalado en el punto 3 precedente. A su vez, las respuestas y aclaraciones estarán disponibles a través del portal de Mercado Público, en los plazos indicados en el cronograma señalado precedentemente, información que se entenderá conocida por todos los interesados desde el momento de su publicación.", "")],
        [("No serán admitidas las consultas formuladas fuera de plazo o por un conducto diferente al señalado.", "bold")],
        [("“EL HOSPITAL” realizará las aclaraciones a las Bases comunicando las respuestas a través del Portal Web de Mercado Público, sitio", ""),
         (" www.mercadopublico.cl", "bold")],
        [("Las aclaraciones, derivadas de este proceso de consultas, formarán parte integrante de las Bases, teniéndose por conocidas y aceptadas por todos los participantes aun cuando el oferente no las hubiere solicitado, por lo que los proponentes no podrán alegar desconocimiento de las mismas.", "")],
        [("“EL HOSPITAL” podrá modificar las presentes bases y sus anexos previa autorización por acto administrativo, durante el periodo de presentación de las ofertas, hasta antes de fecha de cierre de recepción de ofertas. Estas modificaciones, que se lleven a cabo, serán informadas a través del portal", ""),
         (" www.mercadopublico.cl", "bold")],
        [("Estas consultas, aclaratorias y modificaciones formaran parte integra de las bases y estarán vigentes desde la total tramitación del acto administrativo que las apruebe. Junto con aprobar las modificaciones, deberá establecer un nuevo plazo prudencial cuando lo amerite para el cierre o recepción de las propuestas, a fin de que los potenciales oferentes puedan adecuar sus ofertas.", "")],
        [("No se aceptarán consultas realizadas por otros medios, tales como correos electrónicos, fax u otros.", "")]
    ]
    for item in consultas_items:
        p = agregar_parrafo_con_runs(doc, item, estilo='List Bullet')

    # Requisitos Mínimos para Participar
    doc.add_heading("Requisitos Mínimos para Participar.", level=3)
    requisitos_items = [
        "No haber sido condenado por prácticas antisindicales, infracción a los derechos fundamentales del trabajador o por delitos concursales establecidos en el Código Penal dentro de los dos últimos años anteriores a la fecha de presentación de la oferta, de conformidad con lo dispuesto en el artículo 4° de la ley N° 19.886.",
        "No haber sido condenado por el Tribunal de Defensa de la Libre Competencia a la medida dispuesta en la letra d) del artículo 26 del Decreto con Fuerza de Ley N°1, de 2004, del Ministerio de Economía, Fomento y Reconstrucción, que Fija el texto refundido, coordinado y sistematizado del Decreto Ley N° 211, de 1973, que fija normas para la defensa de la libre competencia, hasta por el plazo de cinco años contado desde que la sentencia definitiva quede ejecutoriada.",
        "No ser funcionario directivo de la respectiva entidad compradora; o una persona unida a aquél por los vínculos de parentesco descritos en la letra b) del artículo 54 de la ley N° 18.575; o una sociedad de personas de las que aquél o esta formen parte; o una sociedad comandita por acciones o anónima cerrada en que aquélla o esta sea accionista; o una sociedad anónima abierta en que aquél o esta sean dueños de acciones que representen el 10% o más del capital; o un gerente, administrador, representante o director de cualquiera de las sociedades antedichas.",
        "Tratándose exclusivamente de una persona jurídica, no haber sido condenada conforme a la ley N° 20.393 a la pena de prohibición de celebrar actos y contratos con el Estado, mientras esta pena esté vigente.",
        "A fin de acreditar el cumplimiento de dichos requisitos, los oferentes deberán presentar una “Declaración jurada de requisitos para ofertar”, la cual será generada completamente en línea a través de www.mercadopublico.cl en el módulo de presentación de las ofertas. Sin perjuicio de lo anterior, la entidad licitante podrá verificar la veracidad de la información entregada en la declaración, en cualquier momento, a través de los medios oficiales disponibles."
    ]
    for texto in requisitos_items:
        p = agregar_parrafo_con_texto(doc, texto)
        aplicar_numeracion(p, num_id_requisitos)

    req_p6 = doc.add_paragraph()
    req_p6.add_run("En caso de que los antecedentes administrativos solicitados en esta sección no sean entregados y/o completados en forma correcta y oportuna, se desestimará la propuesta, no será evaluada y será declarada ")
    req_p6.add_run("inadmisible").bold = True
    req_p6.add_run(".")
    aplicar_numeracion(req_p6, num_id_requisitos)

    # Instrucciones para la Presentación de Ofertas
    doc.add_heading("Instrucciones para la Presentación de Ofertas.", level=2)
    tabla_ofertas = doc.add_table(rows=4, cols=2, style='Table Grid')
    tabla_ofertas.cell(0, 0).text = "Presentar Ofertas por Sistema."
    tabla_ofertas.cell(0, 1).text = "Obligatorio."

    agregar_contenido_celda(tabla_ofertas, 1, 0, ["Anexos Administrativos."])
    agregar_contenido_celda(tabla_ofertas, 1, 1, [
        [("Anexo N° 1 Identificación del Oferente.", "bold")],
        [("Anexo N° 2 Declaración Jurada de Habilidad.", "bold")],
        [("Anexo N° 3 Declaración Jurada de Cumplimiento de Obligaciones Laborales y Previsionales.", "bold")],
        [("Declaración jurada online:", "bold"), (" Los oferentes deberán presentar una ", ""), ("Declaración jurada de requisitos para ofertar", "bold"), (", la cual será generada completamente en línea a través de www.mercadopublico.cl en el módulo de presentación de las ofertas.", "")],
        [("Unión Temporal de Proveedores (UTP):", "bold"), (" Solo en el caso de que la oferta sea presentada por una unión temporal de proveedores deberán presentar obligatoriamente la siguiente documentación en su totalidad, en caso contrario, ésta no será sujeta a aclaración y la oferta será declarada ", ""), ("inadmisible", "bold"), (".", "")],
        [("Anexo N°4. Declaración para Uniones Temporales de Proveedores:", "bold"), (" Debe ser presentado por el miembro de la UTP que presente la oferta en el Sistema de Información y quien realiza la declaración a través de la “Declaración jurada de requisitos para ofertar” electrónica presentada junto a la oferta.", "")],
        ["Las ofertas presentadas por una Unión Temporal de Proveedores (UTP) deberán contar con un apoderado, el cual debe corresponder a un integrante de la misma, ya sea persona natural o jurídica. En el caso que el apoderado sea una persona jurídica, ésta deberá actuar a través de su representante legal para ejercer sus facultades."],
        [("En caso de no presentarse debidamente la declaración jurada online constatando la ausencia de conflictos de interés e inhabilidades por condenas, o no presentarse el Anexo N°4, la oferta será declarada ", ""), ("inadmisible", "bold"), (".", "")]
    ])

    agregar_contenido_celda(tabla_ofertas, 2, 0, ["Anexos\nEconómicos."])
    agregar_contenido_celda(tabla_ofertas, 2, 1, [
        [("Anexo N°5: Oferta económica", "bold")],
        ["El anexo referido debe ser ingresado a través del sistema www.mercadopublico.cl , en la sección Anexos Económicos."],
        [("En caso de que no se presente debidamente el Anexo N°5 “Oferta económica”, la oferta será declarada ", ""), ("inadmisible", "bold")]
    ])

    agregar_contenido_celda(tabla_ofertas, 3, 0, ["Anexos Técnicos."])
    agregar_contenido_celda(tabla_ofertas, 3, 1, [
        [("Anexo N°6: Evaluación Técnica", "bold")],
        [("Anexo N°7: Ficha Técnica", "bold")],
        [("Anexo N°8: Plazo de Entrega", "bold")],
        [("Anexo N°9: Servicio Post-venta", "bold")],
        ["Los anexos referidos deben ser ingresados a través del sistema www.mercadopublico.cl. en la sección Anexos Técnicos."],
        [("En el caso que no se presente debidamente los Anexos N°7, N°8 y N°9 la oferta será declarada ", ""), ("inadmisible", "bold")]
    ])
    centrar_verticalmente_tabla(tabla_ofertas)

    # Observaciones
    doc.add_heading("Observaciones", level=3)
    agregar_parrafo_con_runs(doc, [
        ("Los oferentes deberán presentar su oferta a través de su cuenta en el Sistema de Información www.mercadopublico.cl. De existir discordancia entre el oferente o los antecedentes de su oferta y la cuenta a través de la cual la presenta, esta no será evaluada, siendo desestimada del proceso y declarada como", ""),
        (" inadmisible", "bold")
    ])

    agregar_parrafo_con_runs(doc, [
        ("Las únicas ofertas válidas serán las presentadas a través del portal", ""),
        (" www.mercadopublico.cl", "bold"),
        (", en la forma en que se solicita en estas bases. No se aceptarán ofertas que se presenten por un medio distinto al establecido en estas Bases, a menos que se acredite la indisponibilidad técnica del sistema, de conformidad con el artículo 62 del Reglamento de la Ley de Compras. Será responsabilidad de los oferentes adoptar las precauciones necesarias para ingresar oportuna y adecuadamente sus ofertas.", "")
    ])

    agregar_parrafo_con_texto(doc, "Los oferentes deben constatar que el envío de su oferta a través del portal electrónico de compras públicas haya sido realizado con éxito, incluyendo el previo ingreso de todos los formularios y anexos requeridos completados de acuerdo con lo establecido en las presentes bases. Debe verificar que los archivos que se ingresen contengan efectivamente los anexos solicitados.")

    agregar_parrafo_con_texto(doc, "Asimismo, se debe comprobar siempre, luego de que se finalice la última etapa de ingreso de la oferta respectiva, que se produzca el despliegue automático del “Comprobante de Envío de Oferta” que se entrega en dicho Sistema, el cual puede ser impreso por el proponente para su resguardo. En dicho comprobante será posible visualizar los anexos adjuntos, cuyo contenido es de responsabilidad del oferente.")

    agregar_parrafo_con_texto(doc, "El hecho de que el oferente haya obtenido el “Comprobante de envío de ofertas” señalado, únicamente acreditará el envío de ésta a través del Sistema, pero en ningún caso certificará la integridad o la completitud de ésta, lo cual será evaluado por la comisión evaluadora. En caso de que, antes de la fecha de cierre de la licitación, un proponente edite una oferta ya enviada, deberá asegurarse de enviar nuevamente la oferta una vez haya realizado los ajustes que estime, debiendo descargar un nuevo Comprobante.")

    agregar_parrafo_con_runs(doc, [
        ("Si la propuesta económica subida al portal, presenta diferencias entre el valor del anexo económico solicitado y el valor indicado en la línea de la plataforma", ""),
        (" www.mercadopublico.cl", "bold"),
        (", prevalecerá la oferta del anexo económico solicitado en bases. Sin embargo, el Hospital San José de Melipilla, podrá solicitar aclaraciones de las ofertas realizadas a través del portal.", "")
    ])

    # Antecedentes legales para poder ser contratado
    doc.add_heading("Antecedentes legales para poder ser contratado.", level=2)
    tabla_legal = doc.add_table(rows=7, cols=3, style='Table Grid')

    start_natural_row = 0
    end_natural_row = 3
    tabla_legal.cell(start_natural_row, 0).text = "Si el oferente\nes Persona\nNatural"
    tabla_legal.cell(start_natural_row, 2).text = "Acreditar en\nel Registro de\nProveedores"
    tabla_legal.cell(start_natural_row, 1).add_paragraph().add_run("Inscripción (en estado hábil) en el Registro electrónico oficial de contratistas de la Administración, en adelante “").bold = True
    tabla_legal.cell(start_natural_row, 1).paragraphs[0].add_run("Registro de Proveedores").bold = True
    tabla_legal.cell(start_natural_row, 1).paragraphs[0].add_run("”.").bold = True
    tabla_legal.cell(start_natural_row + 1, 1).add_paragraph().add_run("Anexo N°3. Declaración Jurada de Cumplimiento de Obligaciones Laborales y Previsionales.").bold = True
    tabla_legal.cell(start_natural_row + 2, 1).text = "Todos los Anexos deben ser firmados por la persona natural respectiva."
    tabla_legal.cell(start_natural_row + 3, 1).text = "Fotocopia de su cédula de identidad."
    tabla_legal.cell(start_natural_row, 0).merge(tabla_legal.cell(end_natural_row, 0))
    tabla_legal.cell(start_natural_row, 2).merge(tabla_legal.cell(end_natural_row, 2))

    start_nonatural_row = end_natural_row + 1
    end_nonatural_row = start_nonatural_row + 2
    tabla_legal.cell(start_nonatural_row, 0).text = "Si el oferente\nno es\nPersona\nNatural"
    tabla_legal.cell(start_nonatural_row, 2).text = "Acreditar en\nel Registro de\nProveedores"
    tabla_legal.cell(start_nonatural_row, 1).add_paragraph().add_run("Inscripción (en estado hábil) en el Registro de Proveedores.").bold = True
    tabla_legal.cell(start_nonatural_row + 1, 1).text = "Certificado de Vigencia del poder del representante legal, con una antigüedad no superior a 60 días corridos, contados desde la fecha de notificación de la adjudicación, otorgado por el Conservador de Bienes Raíces correspondiente o, en los casos que resulte procedente, cualquier otro antecedente que acredite la vigencia del poder del representante legal del oferente, a la época de presentación de la oferta."
    tabla_legal.cell(start_nonatural_row + 2, 1).text = "Certificado de Vigencia de la Sociedad con una antigüedad no superior a 60 días corridos, contados desde la fecha de notificación de la adjudicación"
    tabla_legal.cell(start_nonatural_row, 0).merge(tabla_legal.cell(end_nonatural_row, 0))
    tabla_legal.cell(start_nonatural_row, 2).merge(tabla_legal.cell(end_nonatural_row, 2))
    centrar_verticalmente_tabla(tabla_legal)

    # Observaciones (continuación)
    doc.add_heading("Observaciones", level=3)
    for texto in [
        "Los antecedentes legales para poder ser contratado, sólo se requerirán respecto del adjudicatario y deberán estar disponibles en el Registro de Proveedores.",
        "Lo señalado en el párrafo precedente no resultará aplicable a la garantía de fiel cumplimiento de contrato, la cual podrá ser entregada físicamente en los términos que indican las presentes bases en aquellos casos que aplique su entrega.",
        "En los casos en que se otorgue de manera electrónica, deberá ajustarse a la ley N° 19.799 sobre documentos electrónicos, firma electrónica y servicios de certificación de dicha firma, y remitirse en la forma señalada en la cláusula 8.2 de estas bases."
    ]:
        agregar_parrafo_con_texto(doc, texto)

    observ_parafo_2 = doc.add_paragraph()
    observ_parafo_2.add_run("Si el respectivo proveedor no entrega la totalidad de los antecedentes requeridos para ser contratado, dentro del plazo fatal de 10 días hábiles contados desde la notificación de la resolución de adjudicación o no suscribe el contrato en los plazos establecidos en estas bases, la entidad licitante podrá readjudicar de conformidad a lo establecido en la")
    observ_parafo_2.add_run(" cláusula 9 letra i")
    observ_parafo_2.add_run(" de las presentes bases. Además, tales incumplimientos darán origen al cobro de la garantía de seriedad de la oferta, si la hubiere.")

    # Inscripción en el Registro de Proveedores
    doc.add_heading("Inscripción en el registro de proveedores", level=2)
    for texto in [
        "En caso de que el proveedor que resulte adjudicado no se encuentre inscrito en el Registro Electrónico Oficial de Contratistas de la Administración (Registro de Proveedores), deberá inscribirse dentro del plazo de 15 días hábiles, contados desde la notificación de la resolución de adjudicación.",
        "Tratándose de los adjudicatarios de una Unión Temporal de Proveedores, cada integrante de ésta deberá inscribirse en el Registro de Proveedores, dentro del plazo de 15 días hábiles, contados desde la notificación de la resolución de adjudicación."
    ]:
        agregar_parrafo_con_texto(doc, texto)

    # Naturaleza y monto de las garantías
    doc.add_heading("Naturaleza y monto de las garatías", level=2)
    doc.add_heading("Evaluación y adjudicación de las ofertas", level=3)

    comis_eval_p1 = doc.add_paragraph()
    comis_eval_p1.add_run("Comisión Evaluadora: ").bold = True
    comis_eval_p1.add_run("La Dirección del Hospital San José de Melipilla designa como integrantes de la Comisión de Evaluación de la propuesta a los siguientes funcionarios: el Subdirector(a) Administrativo, Subdirector(a) Médico de Atención Abierta, Subdirector(a) Médico de Atención Cerrada, Subdirector(a) de Gestión del Cuidado de Enfermería, Subdirector(a) de Gestión y Desarrollo de las Personas, Subdirector(a) de Matronería, Subdirector(a) de Análisis de Información para la Gestión, Subdirector(a) de Apoyo Clínico o sus subrogantes. Para los efectos del quórum para sesionar se requerirá un mínimo de tres miembros. Lo anterior en conformidad con lo dispuesto en el artículo 37 del Decreto Nº 250 que establece el Reglamento de la Ley Nº 19.886.Los miembros de la Comisión Evaluadora no podrán:")

    for texto, estilo in [
        ("La comisión evaluadora verificará el cumplimiento de los requisitos mínimos de participación.", 'List Bullet'),
        ("Se evaluarán los criterios técnicos y económicos según la ponderación definida en las bases.", 'List Bullet')
    ]:
        agregar_parrafo_con_texto(doc, texto, estilo)

    eval_p3 = doc.add_paragraph(style='List Bullet')
    eval_p3.add_run("La adjudicación se realizará al oferente que obtenga el ")
    eval_p3.add_run("mayor puntaje total").bold = True
    eval_p3.add_run(".")

    agregar_parrafo_con_texto(doc, "La misma Comisión estudiará los antecedentes de la Propuesta y elaborará un informe fundado para el Director de este Establecimiento, quien podrá declarar, mediante resolución fundada, admisible aquellas ofertas que cumplan con los requisitos establecidos en las bases de licitación, como también podrá declarar, mediante resolución fundada, inadmisible aquellas ofertas que no cumplan los requisitos establecidos en las bases. En caso de no presentarse oferentes o cuando las ofertas no resulten convenientes para los intereses del Establecimiento, podrá declarar desierta la licitación, fundándose en razones objetivas y no discriminatorias.Esta Comisión Evaluadora podrá invitar a profesionales técnicos para colaborar en el proceso de adjudicación")

    consider_general_p1 = doc.add_paragraph()
    consider_general_p1.add_run("Consideraciones Generales: ").bold = True
    consider_general_p1.add_run("Se exigirá el cumplimiento de los requerimientos establecidos en la cláusula 6, “Instrucciones para Presentación de Ofertas”, de las presentes Bases de Licitación. Aquellas ofertas que no fueran presentadas a través del portal, en los términos solicitados, se declararán como propuestas inadmisibles, por tanto, no serán consideradas en la evaluación. Lo anterior, sin perjuicio de que concurra y se acredite algunas de las causales de excepción establecidas en el artículo 62 del Reglamento de la Ley de Compras.")

    for texto in [
        "La entidad licitante declarará inadmisible cualquiera de las ofertas presentadas que no cumplan los requisitos o condiciones establecidos en las presentes bases, sin perjuicio de la facultad de la entidad licitante de solicitar a los oferentes que salven errores u omisiones formales, de acuerdo con lo establecido en el artículo 40 del Reglamento de la Ley N°19.886 y en las presentes bases.",
        "Los documentos solicitados por la entidad licitante deben estar vigentes a la fecha de cierre de la presentación de las ofertas indicado en la cláusula 3 de las presentes bases y ser presentados como copias simples, legibles y firmadas por el representante legal de la empresa o persona natural. Sin perjuicio de ello, la entidad licitante podrá verificar la veracidad de la información entregada por el proveedor. En el caso en que el proveedor esté inscrito y habilitado por el Registro de Proveedores, serán suficientes los antecedentes que se encuentren en dicho Registro, en la medida que se haya dado cumplimiento a las normas de actualización de documentos que establece el Registro de Proveedores."
    ]:
        agregar_parrafo_con_texto(doc, texto)

    subsan_err_p1 = doc.add_paragraph()
    subsan_err_p1.add_run("Subsanación de errores u omisiones formales: ").bold = True
    subsan_err_p1.add_run("vicios u omisiones no les confieran a esos oferentes una situación de privilegio respecto de los demás competidores, esto es, en tanto no se afecten los principios de estricta sujeción a las bases y de igualdad de los oferentes, y se informe de dicha solicitud al resto de los oferentes, a través del Sistema de Información")
    subsan_err_p1.add_run("www.mercadopublico.cl.").bold = True

    agregar_parrafo_con_texto(doc, "El plazo que tendrán los oferentes, en este caso para dar cumplimiento a lo solicitado por el mandante, no será inferior a las 24 horas, contadas desde la fecha de publicación de la solicitud por parte del Hospital, la que se informará a través del Sistema de información www.mercadopublico.cl. La responsabilidad de revisar oportunamente dicho sistema durante el período de evaluación recae exclusivamente en los respectivos oferentes.")

    inadmisibilidad_p1 = doc.add_paragraph()
    inadmisibilidad_p1.add_run("Inadmisibilidad de las ofertas y declaración de desierta de la licitación: ").bold = True
    inadmisibilidad_p1.add_run("La entidad licitante declarará inadmisible las ofertas presentadas que no cumplan los requisitos mínimos establecidos en los Anexos N°5, N°6, N°7, N°8 y N°9 y/o las condiciones establecidas en las presentes bases de licitación, sin perjuicio de la facultad para solicitar a los oferentes que salven errores u omisiones formales de acuerdo con lo establecido en las presentes bases.")

    agregar_parrafo_con_texto(doc, "La entidad licitante podrá, además, declarar desierta la licitación cuando no se presenten ofertas o cuando éstas no resulten convenientes a sus intereses.Dichas declaraciones deberán materializarse a través de la dictación de una resolución fundada y no darán derecho a indemnización alguna a los oferentes.")

    criterios_eval_p1 = doc.add_paragraph()
    criterios_eval_p1.add_run("Criterios de evaluación y procedimientos de las ofertas: ").bold = True
    criterios_eval_p1.add_run("La evaluación de las ofertas se realizará en una etapa, utilizando criterios técnicos, económicos y administrativos.")

    criterios_eval_p2 = doc.add_paragraph()
    criterios_eval_p2.add_run("La evaluación de las ofertas presentadas para el ")
    criterios_eval_p2.add_run("SUMINISTRO DE INSUMOS Y ACCESORIOS PARA TERAPIA DE PRESIÓN NEGATIVA CON EQUIPOS EN COMODATO PARA EL HOSPITAL SAN JOSÉ DE MELIPILLA, ")
    criterios_eval_p2.add_run(" se regirá por las siguientes ponderaciones y criterios a evaluar:")

    # Tabla de criterios de evaluación
    tabla_criterios_datos = [
        ["CRITERIOS", "", "PONDERACIÓN", "EVALUADO\nSEGÚN ANEXO"],
        ["ECONÓMICO", "OFERTA ECONÓMICA", "60%", "ANEXO N°5"],
        ["TÉCNICOS", "EVALUACIÓN TÉCNICA", "20%", "ANEXO N°6"],
        ["", "PLAZO DE ENTREGA", "10%", "ANEXO N°8"],
        ["", "SERVICIO POST-VENTA", "10%", "ANEXO N°9"]
    ]
    tabla_criterios = crear_tabla(doc, tabla_criterios_datos)
    tabla_criterios.cell(0, 0).merge(tabla_criterios.cell(0, 1))
    tabla_criterios.cell(0, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    tabla_criterios.cell(0, 0).paragraphs[0].runs[0].bold = True
    tabla_criterios.cell(0, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    tabla_criterios.cell(0, 2).paragraphs[0].runs[0].bold = True
    tabla_criterios.cell(0, 3).paragraphs[0].text = ""
    tabla_criterios.cell(0, 3).paragraphs[0].add_run("EVALUADO").bold = True
    tabla_criterios.cell(0, 3).paragraphs[0].add_run("\n")
    tabla_criterios.cell(0, 3).paragraphs[0].add_run("SEGÚN ANEXO").bold = True
    tabla_criterios.cell(0, 3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    tabla_criterios.cell(2, 0).merge(tabla_criterios.cell(4, 0))
    tabla_criterios.cell(2, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    tabla_criterios.cell(2, 0).paragraphs[0].runs[0].bold = True
    for i in range(1, 5):
        tabla_criterios.cell(i, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        tabla_criterios.cell(i, 3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    tabla_criterios.cell(1, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    tabla_criterios.cell(1, 0).paragraphs[0].runs[0].bold = True

    # Cálculo del puntaje
    calculo_puntaje = doc.add_paragraph()
    calculo_puntaje.add_run("Cálculo del puntaje total de evaluación: ").bold = True

    for texto in [
        "El Puntaje de la Evaluación Final estará dado por el siguiente polinomio:",
        "Puntaje Evaluación Técnica + Puntaje Evaluación Económica",
        "Donde el Puntaje Evaluación Técnica = Evaluación Técnica + Plazo de Entrega + Servicio Post-Venta.",
        "Donde Puntaje Evaluación Económica = Precio."
    ]:
        agregar_parrafo_con_texto(doc, texto)

    # Criterios de evaluación
    doc.add_heading("Criterios de evaluación", level=2)
    doc.add_heading("Criterios Económicos", level=3)
    crit_econ = doc.add_paragraph()
    crit_econ.add_run("OFERTA ECONÓMICA 60%").bold = True
    agregar_parrafo_con_texto(doc, "Valor ítem ofertado. Para calcular el puntaje correspondiente al precio se utilizará la siguiente fórmula: A este puntaje se le aplicará la ponderación del 60 %. El oferente deberá declarar en Anexo N°5, los valores ofertados considerando todos los gastos involucrados e impuestos que apliquen.")

    doc.add_heading("Criterios Técnicos", level=3)
    crit_tecn_para = doc.add_paragraph(style="List Bullet")
    crit_tecn_run = crit_tecn_para.add_run("EVALUACIÓN TÉCNICA 20%: ")
    crit_tecn_run.bold = True
    crit_tecn_para.add_run("Se evaluará según información presentada para el Anexo N°7 que deberá ser adjuntada en su oferta en el Portal de Mercado Público, junto con la pauta de evaluación del Anexo N°6. Se evaluará por producto ofertado, donde el puntaje total será el promedio de la evaluación de todos los insumos ofertados.")

    plazo_entrega_para = doc.add_paragraph(style="List Bullet")
    plazo_entrega_run = plazo_entrega_para.add_run("PLAZO DE ENTREGA 10%: ")
    plazo_entrega_run.bold = True
    plazo_entrega_para.add_run("Se evaluará según información presentada en el Anexo N° 8 de la presente base de licitación.")

    serv_post_venta_para = doc.add_paragraph(style="List Bullet")
    serv_post_venta_run = serv_post_venta_para.add_run("Servicio Post-Venta 10%: ")
    serv_post_venta_run.bold = True
    serv_post_venta_para.add_run("Se evaluará según información presentada en el Anexo N° 9 de la presente base de licitación.")

    # Adjudicación y otras secciones
    doc.add_heading("Adjudicación", level=2)
    agregar_parrafo_con_texto(doc, "Se adjudicará al oferente que obtenga el mayor puntaje, en los términos descritos en las presentes bases. La presente licitación se adjudicará a través de una resolución dictada por la autoridad competente, la que será publicada en www.mercadopublico.cl, una vez que se encuentre totalmente tramitada.")

    doc.add_heading("Mecanismo de Resolución de empates.", level=2)
    agregar_parrafo_con_texto(doc, "En el evento de que, una vez culminado el proceso de evaluación de ofertas, hubiese dos o más proponentes que hayan obtenido el mismo puntaje en la evaluación final, quedando más de uno en condición de resultar adjudicado, se optará por aquella oferta que cuente con un mayor puntaje de acuerdo con la secuencia de los criterios que resulten aplicables, de acuerdo al siguiente orden: EVALUACION TECNICA, seguido por PLAZO DE ENTREGA, seguido por SERVICIO POST-VENTA, seguido por CRITERIO ECONOMICO. Finalmente, de mantenerse la igualdad, se adjudicará a aquel oferente que haya ingresado primero su propuesta en el portal Mercado Público considerándose la hora en que aquello se efectúe.")

    doc.add_heading("Resolución de consultas respecto de la Adjudicación.", level=2)
    resolucion_consultas_par = doc.add_paragraph()
    resolucion_consultas_par.add_run("Las consultas sobre la adjudicación deberán realizarse dentro del plazo fatal de 5 días hábiles contados desde la publicación de la resolución en el Sistema de Información ")
    resolucion_consultas_par.add_run("www.mercadopublico.cl").bold = True
    resolucion_consultas_par.add_run("a través del siguiente enlace: ")
    resolucion_consultas_par.add_run("http://ayuda.mercadopublico.cl ").bold = True

    doc.add_heading("Readjudicación", level=2)
    agregar_parrafo_con_texto(doc, "Si el adjudicatario se desistiere de firmar el contrato o de aceptar la orden de compra, o no cumpliese con las demás condiciones y requisitos establecidos en las presentes bases para la suscripción o aceptación de los referidos documentos, la entidad licitante podrá, junto con dejar sin efecto la adjudicación original, adjudicar la licitación al oferente que le seguía en puntaje, o a los que le sigan sucesivamente, dentro del plazo de 60 días corridos contados desde la publicación de la adjudicación original.")

    doc.add_section()
    doc.add_heading("Condiciones Contractuales, Vigencia de las Condiciones Comerciales, Operatoria de la Licitación y Otras Cláusulas:", level=1)
    doc.add_heading("Documentos integrantes", level=2)
    agregar_parrafo_con_texto(doc, "La relación contractual que se genere entre la entidad licitante y el adjudicatario se ceñirá a los siguientes documentos:")
    for texto in [
        "Bases de licitación y sus anexos.",
        "Aclaraciones, respuestas y modificaciones a las Bases, si las hubiere.",
        "Oferta.",
        "Orden de compra"
    ]:
        agregar_parrafo_con_texto(doc, texto, estilo="List Bullet")

    agregar_parrafo_con_texto(doc, "Todos los documentos antes mencionados forman un todo integrado y se complementan recíprocamente, especialmente respecto de las obligaciones que aparezcan en uno u otro de los documentos señalados. Se deja constancia que se considerará el principio de preeminencia de las Bases.")

    doc.add_heading("Validez de la Oferta", level=2)
    for texto in [
        "La oferta tendrá validez de ciento veinte días (120) días corridos, contados desde la fecha de apertura de la propuesta. La oferta cuyo periodo de validez sea menor que el requerido, será rechazada de inmediato.",
        "Si vencido el plazo señalado precedentemente, el Hospital San José de Melipilla no ha realizado la adjudicación, podrá solicitar a los Proponentes la prórroga de sus ofertas y garantías. Los proponentes podrán ratificar sus ofertas o desistir de ellas, formalizando su decisión mediante comunicación escrita dirigida al Hospital. Se devolverá la garantía a aquellos que no accedan a la prórroga."
    ]:
        agregar_parrafo_con_texto(doc, texto)

    doc.add_heading("Suscripción del Contrato", level=2)
    for texto in [
        "Para suscribir el contrato o aceptar la orden de compra contemplada en el artículo 63 del reglamento de la Ley de Compras, el adjudicado deberá estar inscrito en el Registro de Proveedores.",
        "Para formalizar las adquisiciones de bienes y servicios regidas por la ley Nº 19.886, se requerirá la suscripción de un contrato, la que en este caso se verá reflejada por la sola aceptación de la respectiva Orden de Compras.",
        "El respectivo contrato deberá suscribirse dentro de los 20 días hábiles siguientes a la notificación de la resolución de adjudicación totalmente tramitada. Asimismo, cuando corresponda, la orden de compra que formaliza la adquisición deberá ser aceptada por el adjudicatario dentro de ese mismo plazo.",
        "Si por cualquier causa que no sea imputable a la entidad licitante, el contrato no se suscribe dentro de dicho plazo, o no se acepta la orden de compra que formaliza la adquisición dentro de ese mismo término, se entenderá desistimiento de la oferta, pudiendo readjudicar la licitación al oferente que le seguía en puntaje, o a los que le sigan sucesivamente."
    ]:
        agregar_parrafo_con_texto(doc, texto)

    doc.add_heading("Modificación del Contrato", level=2)
    agregar_parrafo_con_texto(doc, "Las partes de común acuerdo podrán modificar el contrato aumentando o disminuyendo los Bienes o servicios licitados, como también se podrán pactar nuevos bienes o servicios que no alteren la naturaleza del contrato. Estas modificaciones podrán ser hasta un 30% el presupuesto disponible estipulado en las presentes bases de licitación. En el caso de aumentar los bienes o servicios contratados, la garantía fiel cumplimiento de contrato también podrá readecuarse en proporción al monto de la modificación que se suscriba según aquellos casos que apliquen. En caso de aumentar o disminuir los bienes o servicios contratados, los valores a considerar, serán aquellos ofertados en el anexo oferta económica.Con todo, las eventuales modificaciones que se pacten no producirán efecto alguno sino desde la total tramitación del acto administrativo que las apruebe.")

    doc.add_heading("Gastos e Impuestos", level=2)
    agregar_parrafo_con_texto(doc, "Todos los gastos e impuestos que se generen o produzcan por causa o con ocasión de este Contrato, tales como los gastos notariales de celebración de contratos y/o cualesquiera otros que se originen en el cumplimiento de obligaciones que, según las Bases, ha contraído el oferente adjudicado, serán de cargo exclusivo de éste.")

    doc.add_heading("Efectos derivados de Incumplimiento del proveedor", level=2)
    agregar_parrafo_con_texto(doc, "En función de la gravedad de la infracción cometida por el adjudicatario, se le aplicarán las siguientes sanciones:")

    clasificacion_par1 = doc.add_paragraph()
    clasificacion_par1.add_run("Amonestación: ").bold = True
    clasificacion_par1.add_run("Corresponde a un registro escrito, que dejará de manifiesto cualquier falta menor cometida por el adjudicado. Se entenderá por falta menor aquella que no ponga en riesgo de forma alguna la prestación del servicio o la vida e integridad psíquica y física de los pacientes, que se vinculen a temas administrativos y técnicos y que no sea constitutiva de multa. La amonestación no estará afecta a sanción pecuniaria.")

    clasificacion_par2 = doc.add_paragraph()
    clasificacion_par2.add_run("Multa: ").bold = True
    clasificacion_par2.add_run("Corresponde a la sanción de cualquier falta, de gravedad leve, moderada o grave en que incurra el adjudicado, cada vez que éste no dé cumplimiento a cualquiera de las obligaciones contempladas en las presentes bases. Se expresará en Unidades Tributarias Mensuales (UTM).El monto de cada multa, dependerá de la gravedad de la infracción cometida, en este sentido las multas se clasifican en:")

    p_leve = doc.add_paragraph()
    p_leve.add_run("Multa Leve: ").bold = True
    p_leve.add_run("Sera considerada LEVE aquella situación originada por una falta de carácter menor, que no origina riesgos a las personas, ni daños a los bienes de la Institución o a su imagen. Su importe será de 3 Unidades Tributarias Mensuales (UTM). Las conductas que pueden estar afectas a multa leve son:")

    for texto in [
        "Entrega de productos con atraso de hasta dos (2) días hábiles, contados desde la fecha de entrega estipulada en el contrato, sin que estos se hubiesen re-pactado en su plazo de entrega.",
        "Conducta o trato irrespetuoso de parte del personal del Oferente adjudicado o su cadena de distribución.",
        "La acumulación de dos amonestaciones.",
        "Incumplimiento del contrato que no origine riesgos a las personas o daño a los bienes del establecimiento o a su imagen."
    ]:
        agregar_parrafo_con_texto(doc, texto, estilo='List Bullet')

    p_moderada = doc.add_paragraph()
    p_moderada.add_run("Multa Moderada: ").bold = True
    p_moderada.add_run("Sera considerada MODERADA, aquella situación originada por una falta que afecte o ponga en riesgo, directa o indirectamente a personas o a la Institución o que limite significativamente la atención y calidad del servicio, pero que es factible de ser corregida. Su importe será de 6 Unidades Tributarias Mensuales (UTM). Las conductas que puedan estar afectas a multa moderada son:")

    for texto in [
        "No aceptar la Orden de Compra dentro de los dos (4) días hábiles siguientes al envío de la orden a través del portal de Mercado Publico.",
        "Entrega de los productos con atraso de entre tres (3) y seis (6) días hábiles inclusive, contados desde la fecha de entrega estipulada en el contrato, sin que estos se hubiesen re-pactado en su plazo de entrega.",
        "Despacho de productos en lugares no autorizados por el Hospital.",
        "La acumulación de dos multas leves trimestres móviles.",
        "cualquier falta que afecte o ponga en riesgo, directa o indirectamente, a personas o a la institución, o que limite significativamente la atención y calidad del servicio, pero que es factible de ser corregida."
    ]:
        agregar_parrafo_con_texto(doc, texto, estilo='List Bullet')

    p_grave = doc.add_paragraph()
    p_grave.add_run("Multa Grave: ").bold = True
    p_grave.add_run("Sera considerada GRAVE, aquella situación originada por una falta que atente, directa o indirectamente con la atención y calidad del servicio. Su importe será de 10 Unidades Tributarias Mensuales (UTM). Las conductas que pueden estar afectas a multa grave son:")

    for texto in [
        "Incumplimiento de la totalidad de lo requerido en la orden de compra.",
        "Entrega de productos con atraso de más de seis (6) días hábiles, contados desde la fecha de entrega estipulada en el contrato, sin que estos se hubiesen re-pactado en su plazo de entrega.",
        "Rechazo total de los productos por no ajustarse a las especificaciones técnicas.",
        "El incumplimiento en el recambio de los productos que presenten problemas de estabilidad, empaque, envases en mal estado, conservación inadecuada, vencida, defectuosa o dañada en un periodo máximo de cuarenta y ocho (48) horas.",
        "La acumulación de dos multas moderadas en 2 trimestres móviles",
        "Si el adjudicatario no inicia sus labores en la fecha acordada",
        "Si se acreditaren acciones y/u omisiones maliciosas y/o negligentes que comprometan la eficiencia y eficacia del servicio o la seguridad y/o bienes dispuestos por el Hospital para la correcta ejecución del convenio.",
        "No cumplir con lo señalado en los requerimientos técnicos.",
        "Vulnerar normas referidas al uso de información reservada o confidencial al que se tenga acceso en razón del servicio especializado contratado.",
        "No cumplir con los horarios establecidos en las bases",
        "No cumplir con los tiempos de respuestas establecidos en bases de licitación según sea el caso.",
        "Incumplimiento de las normas y procedimientos internos vigentes, tanto técnicas como administrativas impartidas por el Hospital, a través del administrador interno del contrato",
        "cualquier falta que atente, directa o indirectamente, contra la integridad física de los pacientes o funcionarios, que implique u obstruye la atención, calidad del servicio y/o según lo establecido en las bases."
    ]:
        agregar_parrafo_con_texto(doc, texto, estilo='List Bullet')

    for texto in [
        "Las referidas multas, en total, no podrán sobrepasar el 20% del valor total neto del contrato. Igualmente, no se le podrán cursar más de 6 multas totalmente tramitadas en un período de 6 meses consecutivos. En ambos casos, superado cada límite, se configurará una causal de término anticipado del contrato.",
        "Las multas deberán ser pagadas en el plazo máximo de 5 días hábiles contados desde la notificación de la resolución que aplica la multa.",
        "Cuando el cálculo del monto de la respectiva multa, convertido a pesos chilenos, resulte un número con decimales, éste se redondeará al número entero más cercano. La fecha de conversión de la UTM será la del día de emisión del respectivo acto administrativo que origina el cobro de la multa",
        "Las multas se aplicarán sin perjuicio del derecho de la entidad licitante de recurrir ante los Tribunales Ordinarios de Justicia ubicados en la ciudad de Melipilla, a fin de hacer efectiva la responsabilidad del contratante incumplidor.",
        "No procederá el cobro de las multas señaladas en este punto, si el incumplimiento se debe a un caso fortuito o fuerza mayor, de acuerdo con los artículos 45 y 1547 del Código Civil o una causa enteramente ajena a la voluntad de las partes, el cual será calificado como tal por la Entidad Licitante, en base al estudio de los antecedentes por los cuales el oferente adjudicado acredite el hecho que le impide cumplir."
    ]:
        agregar_parrafo_con_texto(doc, texto)

    doc.add_heading("Cobro de la Garantía de Fiel Cumplimiento de Contrato", level=3)
    agregar_parrafo_con_texto(doc, "Al Adjudicatario le podrá ser aplicada la medida de cobro de la Garantía por Fiel Cumplimiento del Contrato por la entidad licitante, en los siguientes casos:")
    for texto in [
        "No pago de multas dentro de los plazos establecidos en las presentes bases y/o el respectivo contrato.",
        "Incumplimientos de las exigencias técnicas de los bienes y servicios (en caso de que hayan sido requeridos) adjudicados establecidos en el Contrato.",
        "Cualquiera de las causales señaladas en el N°10.6.3 sobre “Término Anticipado del Contrato”, a excepción del numeral 3) y numeral 16), en todas estas causales señaladas, se procederá al cobro de la garantía de fiel cumplimiento del contrato, si se hubiere exigido dicha caución en las Bases."
    ]:
        agregar_parrafo_con_texto(doc, texto, estilo='List Bullet')

    doc.add_heading("Término anticipado del contrato", level=3)
    agregar_parrafo_con_texto(doc, "El hospital está facultado para declarar administrativamente mediante resolución fundada el término anticipado del contrato, en cualquier momento, sin derecho a indemnización alguna para el adjudicado, si concurre alguna de las causales que se señalan a continuación:")

    termino_contrato_items = [
        "Por incumplimiento grave de las obligaciones contraídas por el proveedor adjudicado, cuando sea imputable a éste. Se entenderá por incumplimiento grave la no ejecución o la ejecución parcial por parte del adjudicatario de las obligaciones contractuales, descritas en las presentes Bases, sin que exista alguna causal que le exima de responsabilidad, y cuando dicho incumplimiento le genere al hospital un perjuicio en el cumplimiento de sus funciones. Alguno de estos motivos puede ser:",
        "Si el adjudicado se encuentra en estado de notoria insolvencia o fuere declarado deudor en un procedimiento concursal de liquidación. En el caso de una UTP, aplica para cualquiera de sus integrantes. En este caso no procederá el término anticipado si se mejoran las cauciones entregadas o las existentes sean suficientes para garantizar el cumplimiento del contrato.",
        "Por exigirlo la necesidad del servicio, el interés público o la seguridad nacional.",
        "Registrar, a la mitad del período de ejecución contractual, con un máximo de seis meses, saldos insolutos de remuneraciones o cotizaciones de seguridad social con sus actuales trabajadores o con trabajadores contratados en los últimos 2 años.",
        "Si se disuelve la sociedad o empresa adjudicada, o en caso de fallecimiento del contratante, si se trata de una persona natural.",
        "Incumplimiento de uno o más de los compromisos asumidos por los adjudicatarios, en virtud del “Pacto de integridad\" contenido en estas bases. Cabe señalar que en el caso que los antecedentes den cuenta de una posible afectación a la libre competencia, el organismo licitante pondrá dichos antecedentes en conocimiento de la Fiscalía Nacional Económica.",
        "Sin perjuicio de lo señalado en el “Pacto de integridad”, si el adjudicatario, sus representantes, o el personal dependiente de aquél, no observaren el más alto estándar ético exigible, durante la ejecución de la licitación, y propiciaren prácticas corruptas, tales como:",
        "No renovación oportuna de la Garantía de Fiel Cumplimiento, según lo establecido en la cláusula 8.2 de las bases de licitación cuando aplique.",
        "La comprobación de la falta de idoneidad, de fidelidad o de completitud de los antecedentes aportados por el proveedor adjudicado, para efecto de ser adjudicado o contratado.",
        "La comprobación de que el adjudicatario, al momento de presentar su oferta contaba con información o antecedentes relacionados con el proceso de diseño de las bases, encontrándose a consecuencia de ello en una posición de privilegio en relación al resto de los oferentes, ya sea que dicha información hubiese sido conocida por el proveedor en razón de un vínculo laboral o profesional entre éste y las entidades compradoras, o bien, como resultado de prácticas contrarias al ordenamiento jurídico.",
        "En caso de ser el adjudicatario de una Unión Temporal de Proveedores (UTP):",
        "En caso de infracción de lo dispuesto en la cláusula sobre “Cesión de contrato y Subcontratación”",
        "En caso de que las multas cursadas, en total, sobrepasen el 20 % del valor total contratado con impuestos incluidos o se apliquen más de 6 multas totalmente tramitadas en un periodo de 6 meses consecutivos.",
        "Por el no pago de las multas aplicadas.",
        "Por la aplicación de dos multas graves en que incurra el adjudicatario en virtud del incumplimiento de las obligaciones reguladas en las bases y del presente contrato.",
        "Si el Hospital San José de Melipilla cesara su funcionamiento en lugar de origen por cambio de ubicación de sus dependencias.",
        "Por la comprobación de la inhabilidad del adjudicatario para contratar con la Administración del Estado en portal de mercado público, durante la ejecución del presente contrato. Solo en el caso que el proveedor desde la notificación de esta situación no regularice su registro en un plazo superior a 15 días hábiles.",
        "Por incumplimiento de obligaciones de confidencialidad establecidas en las respectivas Bases."
    ]
    for texto in termino_contrato_items:
        p = agregar_parrafo_con_texto(doc, texto, estilo='List Number')
        aplicar_numeracion(p, num_id_consultas, nivel=0)

    # Sub-items para ciertos puntos
    subitems_1 = [
        "La aplicación de dos o más Multas Graves en un periodo de seis meses móviles.",
        "Si el proveedor fuese condenado a algún delito que tuviera pena aflictiva o tratándose de una empresa, sus socios, o en el caso de una sociedad anónima, algunos de los miembros del directorio o el gerente de la sociedad.",
        "Si el proveedor delega, cede, aporta o transfiere el presente convenio a cualquier título efectúa asociaciones u otorga concesiones o subconcesiones.",
        "Si la sociedad se disolviere por Quiebra o cesación de pagos del proveedor."
    ]
    for texto in subitems_1:
        p = agregar_parrafo_con_texto(doc, texto, estilo='List Number')
        aplicar_numeracion(p, num_id_consultas, nivel=1)

    subitems_7 = [
        "Dar u ofrecer obsequios, regalías u ofertas especiales al personal del hospital, que pudiere implicar un conflicto de intereses, presente o futuro, entre el respectivo adjudicatario y el servicio hospitalario.",
        "Dar u ofrecer cualquier cosa de valor con el fin de influenciar la actuación de un funcionario público durante la relación contractual objeto de la presente licitación.",
        "Tergiversar hechos, con el fin de influenciar decisiones de la entidad licitante."
    ]
    for texto in subitems_7:
        p = agregar_parrafo_con_texto(doc, texto, estilo='List Number')
        aplicar_numeracion(p, num_id_consultas, nivel=1)

    subitems_11 = [
        "Inhabilidad sobreviniente de uno de los integrantes de la UTP en el Registro de Proveedores, que signifique que la UTP no pueda continuar ejecutando el contrato con los restantes miembros en los mismos términos adjudicados.",
        "De constatarse que los integrantes de la UTP constituyeron dicha figura con el objeto de vulnerar la libre competencia. En este caso, deberán remitirse los antecedentes pertinentes a la Fiscalía Nacional Económica.",
        "Retiro de algún integrante de la UTP que hubiere reunido una o más características objeto de la evaluación de la oferta.",
        "Cuando el número de integrantes de una UTP sea inferior a dos y dicha circunstancia ocurre durante la ejecución del contrato.",
        "Disolución de la UTP."
    ]
    for texto in subitems_11:
        p = agregar_parrafo_con_texto(doc, texto, estilo='List Number')
        aplicar_numeracion(p, num_id_consultas, nivel=1)

    for texto in [
        "De concurrir cualquiera de las causales anteriormente señaladas como término anticipado del contrato, exceptuando las causales número 3 y número 16, se procederá al cobro de la garantía de fiel cumplimiento del contrato, siempre y cuando se hubiere exigido dicha caución en las Bases.",
        "El término anticipado por incumplimientos se aplicará siguiendo el procedimiento establecido en la cláusula “sobre aplicación de Medidas derivadas de incumplimientos.”"
    ]:
        agregar_parrafo_con_texto(doc, texto)

    doc.add_heading("Resciliación de Mutuo Acuerdo", level=2)
    agregar_parrafo_con_texto(doc, "Sin perjuicio de lo anterior, la entidad licitante y el respectivo adjudicatario podrán poner término al contrato en cualquier momento, de común acuerdo, sin constituir una medida por incumplimiento.")

    doc.add_heading("Procedimiento para Aplicación de Medidas derivadas de incumplimientos", level=2)
    for texto in [
        "Detectada una situación que amerite la aplicación de una multa u otra medida derivada de incumplimientos contemplada en las presentes bases, o que constituya una causal de término anticipado, con excepción de la resciliación, el referente técnico o administrador del contrato notificará de ello al oferente adjudicado, informándole sobre la medida a aplicar y sobre los hechos que la fundamentan.",
        "A contar de la notificación singularizada en el párrafo anterior, el proveedor adjudicado tendrá un plazo de 5 días hábiles para efectuar sus descargos por escrito, acompañando todos los antecedentes que lo fundamenten. Vencido el plazo indicado sin que se hayan presentados descargos, la Dirección del Hospital resolverá según la naturaleza de la infracción, notificando al proveedor la resolución del caso por parte del Hospital.",
        "Si el proveedor adjudicado ha presentado sus descargos dentro del plazo establecido para estos efectos, el Hospital tendrá un plazo de 30 días hábiles, contados desde la recepción de los descargos del proveedor, para rechazarlos o acogerlos, total o parcialmente. Al respecto, el rechazo total o parcial de los descargos del respectivo proveedor deberá formalizarse a través de la dictación de una resolución fundada del hospital, en la cual deberá detallarse el contenido y las características de la medida. La indicada resolución será notificada al proveedor adjudicado.",
        "Con todo, el adjudicatario solo será responsable por hechos imputables a su incumplimiento directo y no por indisponibilidades de servicio ocasionadas por fallas ajenas a su gestión y control, lo que deberá, en todo caso, acreditarse debidamente. Sin perjuicio de lo anterior, el adjudicatario deberá adoptar medidas que ofrezcan continuidad operativa a los servicios materia de la respectiva licitación.",
        "Una vez finalizados los trámites administrativos señalados precedentemente y para el evento de que esta conlleve la aplicación de una multa o sanción, el Hospital San José de Melipilla podrá realizar el cobro de la multa o sanción que será debidamente notificado junto con el acto administrativo que lo autoriza. El monto de las multas podría ser rebajado del pago, que el Hospital deba efectuar al proveedor, en el estado de pago más próximo a la notificación del acto administrativo, pudiéndose aplicar tanto en la emisión de la orden de compra, como también en la aplicación del descuento en el pago de facturas. De no ser suficiente este monto o en caso de no existir pagos pendientes, el proveedor deberá pagar directamente al Hospital San José de Melipilla, el monto indicado en el acto administrativo previamente notificado, este pago no podrá ser superior a los 5 días hábiles desde su notificación. Si el proveedor no paga dentro de dicho plazo, se hará efectivo el cobro de la garantía de fiel cumplimiento del contrato, debiendo reponer una nueva boleta de garantía por un monto igual al original, en un plazo no superior a 5 días hábiles en caso que aplique la solicitud de dicha caución.",
        "En el caso de no reponer la boleta de garantía, el hospital podrá proceder a tramitar el termino anticipado del contrato en aquellos casos que aplique con la solicitud de dicha caución.",
        "El valor de la UTM a considerar será el equivalente a su valor en pesos del mes en el cual se aplicó la multa."
    ]:
        agregar_parrafo_con_texto(doc, texto)

    doc.add_heading("Emisión de la Orden de Compra", level=3)
    for texto in [
        "Las órdenes de compra se emitirán previa solicitud del administrador del contrato, quien, en función de la necesidad y demanda del servicio, realizara los pedidos correspondientes.",
        "La orden de compra sólo se emitirá en los casos que el proveedor este en estado hábil para ser contratado por el Estado de Chile y sólo se emitirá el documento a nombre del proveedor adjudicado por el Hospital.",
        "Al inicio del convenio, por registros en la plataforma y tramites del “gestor de contratos” se emitirá una orden de compras por un monto mínimo, la que solo debe ser aceptada por el proveedor, sin tramitar dicho servicio. Todo cambio respecto a este punto, será informado con la respectiva anticipación."
    ]:
        agregar_parrafo_con_texto(doc, texto)

    doc.add_heading("Del Pago", level=2)
    for texto in [
        "El pago se efectuará una vez que el “Hospital” haya recibido oportunamente y a su entera satisfacción dichos bienes o servicios y desde la recepción conforme de la factura u otro instrumento de cobro.",
        "El pago será efectuado dentro de los 30 días corridos siguientes, contados desde la recepción de la factura respectiva, salvo las excepciones indicadas en el artículo 79 bis del Reglamento de la Ley N° 19.886.",
        "El proveedor solo podrá facturar los bienes o servicios efectivamente entregados y recibidos conforme por este organismo comprador, una vez que el administrador del contrato por parte del organismo comprador autorice la facturación en virtud de la recepción conforme de los bienes o servicios. “El Hospital” rechazará todas las facturas que hayan sido emitidas sin contar con la recepción conforme de los bienes o servicios y la autorización expresa de facturar por parte de éste.",
        "Para efectos del pago, el proveedor adjudicado deberá indicar en la factura el número de orden de compra, además, no podrá superar el monto de la orden de compra, de lo contrario, se cancelará la factura por “forma”."
    ]:
        agregar_parrafo_con_texto(doc, texto)

    del_pago_correo = doc.add_paragraph()
    del_pago_correo.add_run("La factura electrónica deberá ser enviada al correo: ")
    del_pago_correo.add_run("facturas.hjsm@redsalud.gov.cl").bold = True
    del_pago_correo.add_run(" con copia al correo ")
    del_pago_correo.add_run("dipresrecepcion@custodium.com").bold = True
    del_pago_correo.add_run("(En formato PDF y XML)")

    for texto in [
        "El valor del convenio se reajustará anualmente de acuerdo con la variación que haya experimentado el Índice de Precios al Consumidor IPC, obtenido del promedio de la sumatoria de los IPC de los doce meses inmediatamente anteriores al mes en que se efectúa su cálculo. Este reajuste es de exclusiva responsabilidad de la empresa adjudicada; si por alguna razón no lo aplicare, no se permitirá su cobro en forma retroactiva. Su precio se pagará conforme a lo establecido.",
        "En ningún caso procederán cobros adicionales por bienes o servicios no convenidos previamente, ni por tiempos en que el proveedor no preste los servicios.",
        "Cabe señalar que, cuando el resultado del monto a facturar resulte un número con decimales, éste se redondeará al número entero siguiente en caso de que la primera cifra decimal sea igual o superior a 5. En caso contrario el monto deberá ser redondeado al número entero anterior."
    ]:
        agregar_parrafo_con_texto(doc, texto)

    doc.add_heading("Vigencia del Contrato", level=2)
    agregar_parrafo_con_texto(doc, "El contrato tendrá una duración de treinta y seis (36) meses contados desde la total tramitación del acto administrativo que aprueba la adjudicación o hasta que se cumpla con el monto estipulado en las presentes bases, lo que suceda primero y sin perjuicio, que por razones de buen servicio las prestaciones materia de la licitación podrían iniciarse desde el momento de la suscripción del mismo, sin que proceda pago alguno en el tiempo intermedio.")

    doc.add_heading("Administrador del Contrato y/o Referente Técnico.", level=2)
    agregar_parrafo_con_texto(doc, "Con el objeto de supervisar y verificar el cumplimiento materia de la presente licitación, El Hospital designará a (la) Enfermera Supervisora(o) del Servicio de Pabellón y al Jefe(a) de Farmacia o su subrogante, para coordinar y fiscalizar la efectiva ejecución del contrato en términos administrativos.")

    administrado_contrato = doc.add_paragraph()
    administrado_contrato.add_run("El adjudicatario").bold = True
    administrado_contrato.add_run("deberá nombrar un coordinador del contrato, cuya identidad deberá ser informada al Hospital.")
    agregar_parrafo_con_texto(doc, "En el desempeño de su cometido, el coordinador del contrato deberá, a lo menos:")

    for texto in [
        "Informar oportunamente al órgano comprador de todo hecho relevante que pueda afectar el cumplimiento del contrato.",
        "Representar al proveedor en la discusión de las materias relacionadas con la ejecución del contrato.",
        "Coordinar las acciones que sean pertinentes para la operación y cumplimiento de este contrato."
    ]:
        p = agregar_parrafo_con_texto(doc, texto)
        aplicar_numeracion(p, administrado_contrato_id_lista)

    agregar_parrafo_con_texto(doc, "La designación del coordinador y todo cambio posterior deberá ser informado por el adjudicatario al responsable de administrar el contrato y/o referente técnico por parte del órgano comprador, a más tardar dentro de las 24 horas siguientes de efectuada la designación o el cambio, por medio del correo electrónico institucional del funcionario.")

    doc.add_heading("Pacto de Integridad", level=2)
    agregar_parrafo_con_texto(doc, "El oferente declara que, por el sólo hecho de participar en la presente licitación, acepta expresamente el presente pacto de integridad, obligándose a cumplir con todas y cada una de las estipulaciones contenidas en el mismo, sin perjuicio de las que se señalen en el resto de las bases de licitación y demás documentos integrantes. Especialmente, el oferente acepta el suministrar toda la información y documentación que sea considerada necesaria y exigida de acuerdo con las presentes bases de licitación, asumiendo expresamente los siguientes compromisos:")

    pacto_items = [
        "El oferente se compromete a respetar los derechos fundamentales de sus trabajadores, entendiéndose por éstos los consagrados en la Constitución Política de la República en su artículo 19, números 1º, 4º, 5º, 6º, 12º, y 16º, en conformidad al artículo 485 del Código del Trabajo. Asimismo, el oferente se compromete a respetar los derechos humanos, lo que significa que debe evitar dar lugar o contribuir a efectos adversos en los derechos humanos mediante sus actividades, bienes o servicios, y subsanar esos efectos cuando se produzcan, de acuerdo con los Principios Rectores de Derechos Humanos y Empresas de Naciones Unidas.",
        "El oferente se obliga a no ofrecer ni conceder, ni intentar ofrecer o conceder, sobornos, regalos, premios, dádivas o pagos, cualquiera fuese su tipo, naturaleza y/o monto, a ningún funcionario público en relación con su oferta, con el proceso de licitación pública, ni con la ejecución de el o los contratos que eventualmente se deriven de la misma, ni tampoco a ofrecerlas o concederlas a terceras personas que pudiesen influir directa o indirectamente en el proceso licitatorio, en su toma de decisiones o en la posterior adjudicación y ejecución del o los contratos que de ello se deriven.",
        "El oferente se obliga a no intentar ni efectuar acuerdos o realizar negociaciones, actos o conductas que tengan por objeto influir o afectar de cualquier forma la libre competencia, cualquiera fuese la conducta o acto específico, y especialmente, aquellos acuerdos, negociaciones, actos o conductas de tipo o naturaleza colusiva, en cualquiera de sus tipos o formas.",
        "El oferente se obliga a revisar y verificar toda la información y documentación, que deba presentar para efectos del presente proceso licitatorio, tomando todas las medidas que sean necesarias para asegurar su veracidad, integridad, legalidad, consistencia, precisión y vigencia.",
        "El oferente se obliga a ajustar su actuar y cumplir con los principios de legalidad, probidad y transparencia en el presente proceso licitatorio.",
        "El oferente manifiesta, garantiza y acepta que conoce y respetará las reglas y condiciones establecidas en las bases de licitación, sus documentos integrantes y él o los contratos que de ellos se derivase.",
        "El oferente reconoce y declara que la oferta presentada en el proceso licitatorio es una propuesta seria, con información fidedigna y en términos técnicos y económicos ajustados a la realidad, que aseguren la posibilidad de cumplir con la misma en las condiciones y oportunidad ofertadas.",
        "El oferente se obliga a tomar todas las medidas que fuesen necesarias para que las obligaciones anteriormente señaladas sean asumidas y cabalmente cumplidas por sus empleados, dependientes, asesores y/o agentes y, en general, todas las personas con que éste o éstos se relacionen directa o indirectamente en virtud o como efecto de la presente licitación, incluidos sus subcontratistas, haciéndose plenamente responsable de las consecuencias de su infracción, sin perjuicio de las responsabilidades individuales que también procediesen y/o fuesen determinadas por los organismos correspondientes."
    ]
    for texto in pacto_items:
        p = agregar_parrafo_con_texto(doc, texto)
        aplicar_numeracion(p, pacto_integridad_id)

    doc.add_heading("Comportamiento ético del Adjudicatario.", level=3)
    agregar_parrafo_con_texto(doc, "El adjudicatario que preste los servicios deberá observar, durante toda la época de ejecución del contrato, el más alto estándar ético exigible a los funcionarios públicos. Tales estándares de probidad deben entenderse equiparados a aquellos exigidos a los funcionarios de la Administración Pública, en conformidad con el Título III de la ley N° 18.575, Orgánica Constitucional de Bases Generales de la Administración del Estado.")

    doc.add_heading("Auditorías", level=2)
    agregar_parrafo_con_texto(doc, "El adjudicatario podrá ser sometido a auditorías externas, contratadas por la entidad licitante a empresas auditoras independientes, con la finalidad de velar por el cumplimiento de las obligaciones contractuales y de las medidas de seguridad comprometidas por el adjudicatario en su oferta. Si el resultado de estas auditorías evidencia incumplimientos contractuales por parte del adjudicatario, el proveedor quedará sujeto a las medidas que corresponda aplicar la entidad licitante, según las presentes bases.")

    doc.add_heading("Confidencialidad", level=2)
    for texto in [
        "El adjudicatario no podrá utilizar para ninguna finalidad ajena a la ejecución del contrato, la documentación, los antecedentes y, en general, cualquier información, que haya conocido o a la que haya accedido, en virtud de cualquier actividad relacionada con el contrato.",
        "El adjudicatario, así como su personal dependiente que se haya vinculado a la ejecución del contrato, en cualquiera de sus etapas, deben guardar confidencialidad sobre los antecedentes relacionados con el proceso licitatorio y el respectivo contrato.",
        "El adjudicatario debe adoptar medidas para el resguardo de la confidencialidad de la información, reservándose el órgano comprador el derecho de ejercer las acciones legales que correspondan, de acuerdo con las normas legales vigentes, en caso de divulgación no autorizada, por cualquier medio, de la totalidad o parte de la información referida.",
        "La divulgación, por cualquier medio, de la totalidad o parte de la información referida en los párrafos anteriores, por parte del proveedor, durante la vigencia del contrato o dentro de los 5 años siguientes después de finalizado éste, podrá dar pie a que la Entidad entable en su contra las acciones judiciales que correspondan. Con todo, tratándose de bases de datos de carácter personal, la obligación de confidencialidad dura indefinidamente, de acuerdo con la Ley N°19.628, sobre Protección de la Vida Privada."
    ]:
        agregar_parrafo_con_texto(doc, texto)

    doc.add_heading("Propiedad de la información", level=2)
    agregar_parrafo_con_texto(doc, "La entidad licitante será la titular de todos los datos de transacciones, bitácoras (logs), parámetros, documentos electrónicos y archivos adjuntos y, en general, de las bases de datos y de toda información contenida en la infraestructura física y tecnológica que le suministre el proveedor contratado y que se genere en virtud de la ejecución de los servicios objeto de la presente licitación. El proveedor no podrá utilizar la información indicada en el párrafo anterior, durante la ejecución del contrato ni con posterioridad al término de su vigencia, sin autorización escrita de la entidad licitante. Por tal motivo, una vez que el proveedor entregue dicha información a la entidad o al finalizar la relación contractual, deberá borrarla de sus registros lógicos y físicos.")

    doc.add_heading("Saldos insolutos de remuneraciones o cotizaciones de seguridad social.", level=2)
    for texto in [
        "Durante la vigencia del respectivo contrato el adjudicatario deberá acreditar que no registra saldos insolutos de obligaciones laborales y sociales con sus actuales trabajadores o con trabajadores contratados en los últimos dos años.",
        "El órgano comprador podrá requerir al adjudicatario, en cualquier momento, los antecedentes que estime necesarios para acreditar el cumplimiento de las obligaciones laborales y sociales antes señaladas.",
        "En caso de que la empresa adjudicada registre saldos insolutos de remuneraciones o cotizaciones de seguridad social con sus actuales trabajadores o con trabajadores contratados en los últimos dos años, los primeros estados de pago de los bienes y servicios de esta licitación deberán ser destinados al pago de dichas obligaciones, debiendo la empresa acreditar que la totalidad de las obligaciones se encuentran liquidadas al cumplirse la mitad del período de ejecución de las prestaciones, con un máximo de seis meses.",
        "La entidad licitante deberá exigir que la empresa adjudicada proceda a dichos pagos y le presente los comprobantes y planillas que demuestren el total cumplimiento de la obligación. El incumplimiento de estas obligaciones por parte de la empresa adjudicataria dará derecho a terminar la relación contractual, pudiendo llamarse a una nueva licitación en la que la empresa referida no podrá participar."
    ]:
        agregar_parrafo_con_texto(doc, texto)

    doc.add_heading("Normas Laborales Aplicables", level=2)
    for texto in [
        "El adjudicatario, en su calidad de empleador, será responsable exclusivo del cumplimiento íntegro y oportuno de las normas del Código del Trabajo y leyes complementarias, leyes sociales, de previsión, de seguros, de enfermedades profesionales, de accidentes del trabajo y demás pertinentes respecto de sus trabajadores y/o integrantes de sus respectivos equipos de trabajo.",
        "En consecuencia, el adjudicatario será responsable, en forma exclusiva, y sin que la enumeración sea taxativa, del pago oportuno de las remuneraciones, honorarios, indemnizaciones, desahucios, gratificaciones, gastos de movilización, beneficios y, en general, de toda suma de dinero que, por cualquier concepto, deba pagarse a sus trabajadores y/o integrantes de sus respectivos equipos de trabajo.",
        "El Hospital se reserva el derecho a exigir al contratista, a simple requerimiento de la contraparte técnica, y sin perjuicio de lo dispuesto en el artículo 4° de la Ley de Compras y el artículo 183-C del Código del Trabajo, un certificado que acredite el monto y estado de cumplimiento de las obligaciones laborales y previsionales emitido por la Inspección del Trabajo respectiva, o bien, por medios idóneos que garanticen la veracidad de dicho monto y estado de cumplimiento, respecto de sus trabajadores. Ello, con el propósito de hacer efectivo por parte del órgano comprador, su derecho a ser informado y el derecho de retención, consagrados en los incisos segundo y tercero del artículo 183-C del Código del Trabajo, en el marco de la responsabilidad subsidiaria derivada de dichas obligaciones laborales y previsionales, a la que alude el artículo 183-D del mismo Código.",
        "Por otra parte, se deja expresa constancia que la suscripción del contrato respectivo no significará en caso alguno que el adjudicatario, sus trabajadores o integrantes de los equipos presentados por éstos, adquieran la calidad de funcionarios públicos, no existiendo vínculo alguno de subordinación o dependencia de ellos con el órgano comprador."
    ]:
        agregar_parrafo_con_texto(doc, texto)

    doc.add_heading("Cambio de personal del proveedor adjudicado.", level=2)
    for texto in [
        "El Hospital San José de Melipilla podrá, por razones de buen servicio, solicitar el cambio de trabajadores, expresando la causa del derecho a cambiar al personal del proveedor, entendiéndose como el derecho a prohibir unilateralmente la continuidad de funciones de un trabajador que implique un potencial riesgo a los pacientes, funcionarios, bienes e imagen de la organización.",
        "El Proveedor adjudicado deberá reemplazar al personal, dentro del plazo que se le indique. La decisión del Hospital San José de Melipilla se comunicará por escrito al Proveedor precisando las causas que motivan la solicitud, con a lo menos 5 días de anticipación a la fecha en que se solicita deje de prestar servicios en sus dependencias, el trabajador que se indique."
    ]:
        agregar_parrafo_con_texto(doc, texto)

    doc.add_heading("Cesión y subcontratación.", level=2)
    for texto in [
        "El proveedor adjudicado no podrá ceder ni transferir en forma alguna, total ni parcialmente, los derechos y obligaciones que nacen del desarrollo de esta licitación, y, en especial, los establecidos en los respectivos contratos que se celebren con los órganos públicos mandantes.",
        "La infracción de esta prohibición será causal inmediata de término del contrato, sin perjuicio de las acciones legales que procedan ante esta situación.",
        "Durante la ejecución del contrato, y previa autorización por escrito del Hospital, el adjudicatario sólo podrá efectuar aquellas subcontrataciones que sean indispensables para la realización de tareas específicas, todo lo cual será calificado por el coordinador del contrato. En todo caso, el adjudicatario seguirá siendo el único responsable de las obligaciones contraídas en virtud del respectivo contrato suscrito con el Hospital.",
        "Así mismo, el subcontratista debe encontrarse hábil en el registro de Proveedores del Estado y tratándose de servicios, acreditar el cumplimiento de obligaciones laborales, conforme lo establece el artículo 4° inciso 2° de la Ley N°19.886.",
        "En todos los casos es el oferente y eventual adjudicatario el único responsable del pleno cumplimiento de lo señalado en estas bases (Art. N° 76, Reglamento de la Ley N° 19.886)."
    ]:
        agregar_parrafo_con_texto(doc, texto)

    doc.add_heading("Discrepancias", level=2)
    for texto in [
        "Si con motivo de la ejecución del contrato se presentaran denuncias, querellas o demandas ante el Ministerio Público o los Tribunales Ordinarios de Justicia; o reclamos ante el Consejo de Defensa del Estado por el cuestionamiento en la prestación otorgada y que corresponda al objeto del contrato celebrado, será el proveedor el único responsable por tales actos, por lo que, sí el Hospital fuese condenado a pagar una multa o indemnización, en razón de los actos precedentemente enunciados o el Hospital tuviera que pagar alguna transacción judicial o extrajudicial que deba celebrarse en razón de las situaciones antes enunciadas, el proveedor deberá reembolsar al Hospital el total del monto resultante de un fallo ejecutoriado o de una transacción judicial o extrajudicial o de un procedimiento de medición de acuerdo a la Ley Nº 19.966.",
        "Asimismo, serán responsables de todos los daños, pérdidas, deterioros o perjuicios de bienes muebles e inmuebles del Hospital, producto del mal uso ocasionado en virtud de la prestación de servicio, debiendo restituir al Hospital los costos en que deba incurrir para reparar los daños producidos por este motivo. Esta obligación se mantendrá aun cuando el presente contrato que al efecto se suscriba se dé por terminado ya sea por expiración del plazo establecido o por decisión del Hospital."
    ]:
        agregar_parrafo_con_texto(doc, texto)

    doc.add_heading("Constancia", level=2)
    agregar_parrafo_con_texto(doc, "Se deja expresa constancia que todas y cada una de las cláusulas contenidas en las presentes Bases, Anexos y aclaratorias, se entienden incorporadas sin necesidad de mención expresa en el correspondiente contrato que se materialice con el adjudicado y éste se hace responsable del cumplimiento de las obligaciones de tales documentos, Bases Administrativas y Contrato que se deriven.")

    # Guardar el documento
    doc_path = 'resolucion_numerada.docx'
    doc.save(doc_path)
    print(f"Documento guardado como: {doc_path}")

if __name__ == "__main__":
    main()

