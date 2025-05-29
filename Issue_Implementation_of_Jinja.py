from Formated_Base_PEP8 import configurar_directorio_trabajo  # Your existing import
import docx
from docx.oxml.ns import qn  # For qualified XML tag names (e.g., w:bookmarkStart)
import os  # For path operations


# --- Function to extract bookmark names and their enclosed text ---
def get_bookmark_text_data(doc_object):
    """
    Extracts bookmark names and their enclosed text from a docx.Document object.

    This function iterates through paragraphs and their underlying XML to find
    bookmark start and end tags, and collects the text runs between them.

    Args:
        doc_object (docx.Document): The python-docx Document object to process.

    Returns:
        dict: A dictionary where keys are bookmark names (str) and
              values are the text (str) enclosed by those bookmarks.
              Example: {'ClienteNombre': 'John Doe', 'FechaContrato': '2023-01-15'}
    """
    bookmark_data = {}  # To store final results: {name: text}

    # Buffer for text being collected for currently "open" (active) bookmarks
    # Key: bookmark_id (str), Value: dict {'name': str, 'text_parts': list[str]}
    active_bookmarks_info = {}

    # doc.paragraphs includes paragraphs from the main document body,
    # including those within table cells.
    for paragraph in doc_object.paragraphs:
        # paragraph._p is the lxml _Element object for the <w:p> tag
        for element in paragraph._p:

            # --- Handle Bookmark Start (<w:bookmarkStart ... />) ---
            if element.tag == qn('w:bookmarkStart'):
                bm_id = element.get(qn('w:id'))  # Get the bookmark ID
                bm_name = element.get(qn('w:name'))  # Get the bookmark name

                # A bookmark starts; prepare to collect its text.
                # We assume bookmarks are well-formed and non-overlapping for simplicity.
                if bm_id not in active_bookmarks_info:
                    active_bookmarks_info[bm_id] = {'name': bm_name, 'text_parts': []}

            # --- Handle Bookmark End (<w:bookmarkEnd ... />) ---
            elif element.tag == qn('w:bookmarkEnd'):
                bm_id = element.get(qn('w:id'))  # Get the ID of the ending bookmark

                if bm_id in active_bookmarks_info:
                    # This bookmark has ended; consolidate its collected text.
                    info = active_bookmarks_info[bm_id]
                    name = info['name']
                    text = "".join(info['text_parts'])  # Join all collected text parts

                    # Store the extracted name and text.
                    # If a bookmark name appears multiple times, this logic prioritizes
                    # the first non-empty text found. Word UI usually enforces unique bookmark names.
                    if name not in bookmark_data or (not bookmark_data.get(name) and text):
                        bookmark_data[name] = text

                    # This bookmark is no longer active for collection.
                    del active_bookmarks_info[bm_id]

            # --- Handle Runs (<w:r> ... </w:r>) where text content resides ---
            elif element.tag == qn('w:r'):
                run_text = ""
                # A run can contain multiple <w:t> (text) elements. Concatenate them.
                for t_element in element.findall(qn('w:t')):
                    if t_element.text:
                        run_text += t_element.text

                if run_text:
                    # If this run contains text, append it to ALL currently active bookmarks.
                    # This handles bookmarks that span across multiple runs or paragraphs.
                    for bm_id_active in active_bookmarks_info:  # Iterate over keys
                        active_bookmarks_info[bm_id_active]['text_parts'].append(run_text)

    # After processing all paragraphs, check for any bookmarks that started but didn't end.
    # This might indicate a malformed document or bookmarks spanning into headers/footers
    # not covered by doc.paragraphs.
    if active_bookmarks_info:
        # print(f"Warning: Unclosed bookmarks detected at end of paragraph processing: {list(active_bookmarks_info.keys())}")
        for bm_id, info in active_bookmarks_info.items():
            name = info['name']
            text = "".join(info['text_parts'])
            # Store if the name hasn't been stored yet, or if it was stored as empty and this has text
            if name not in bookmark_data or (not bookmark_data.get(name) and text):
                # print(f"Warning: Storing text for potentially unclosed bookmark '{name}' (ID: {bm_id})")
                bookmark_data[name] = text

    return bookmark_data


# --- End of get_bookmark_text_data function ---


# --- Main script execution ---
if __name__ == "__main__":
    configurar_directorio_trabajo()  # Your function to set the working directory

    doc_filename_base = "contrato_automatizado_con_marcadores"
    doc_path = doc_filename_base + ".docx"

    # Check if the document exists
    if not os.path.exists(doc_path):
        print(f"Error: El archivo '{doc_path}' no fue encontrado en el directorio de trabajo: {os.getcwd()}")
        # You could create a dummy document here for testing if needed:
        # print("Creando un documento de prueba...")
        # dummy_doc = docx.Document()
        # p = dummy_doc.add_paragraph()
        # # Add bookmarkStart for BM_Prueba
        # bm_start_el = docx.oxml.OxmlElement('w:bookmarkStart')
        # bm_start_el.set(docx.oxml.ns.qn('w:name'), 'BM_Prueba')
        # bm_start_el.set(docx.oxml.ns.qn('w:id'), '0')
        # p._p.append(bm_start_el)
        # # Add a run with text
        # p.add_run("Este es el texto para BM_Prueba.")
        # # Add bookmarkEnd for BM_Prueba
        # bm_end_el = docx.oxml.OxmlElement('w:bookmarkEnd')
        # bm_end_el.set(docx.oxml.ns.qn('w:id'), '0')
        # p._p.append(bm_end_el)
        # dummy_doc.save(doc_path)
        # print(f"Documento de prueba '{doc_path}' creado con un marcador 'BM_Prueba'.")
        exit()

    try:
        doc = docx.Document(doc_path)
    except Exception as e:
        print(f"Error al abrir el documento '{doc_path}': {e}")
        exit()

    # Call the function to extract all bookmark names and their texts
    # This replaces your previous `marcadores = obtener_marcadores(doc)` if that
    # function didn't directly return the text. If it did, you can use it instead.
    all_bookmarks_data = get_bookmark_text_data(doc)

    # Now, 'all_bookmarks_data' is a dictionary.
    # You can iterate through it to access each bookmark's name and text.
    print("\n--- Datos de los Marcadores Extraídos ---")
    if all_bookmarks_data:
        for bookmark_name, bookmark_text in all_bookmarks_data.items():
            # 'bookmark_name' and 'bookmark_text' are the variables for each bookmark's data
            bookmark_text = "{{" + bookmark_text + "}}"
            print(f"Nombre del Marcador: {bookmark_name}")
            print(f"Texto Asociado     : '{bookmark_text}'")  # Using '' to show if text is empty
            print("-" * 30)

            # You can now use these variables as needed. For example, store them,
            # print them, or use them in further processing.
            # If you need to assign them to specific, uniquely named Python variables
            # (e.g., cliente_nombre = "John Doe"), you would typically do that
            # by checking the bookmark_name:
            #
            # if bookmark_name == "NombreCliente":
            #     nombre_cliente_variable = bookmark_text
            #     print(f"Variable 'nombre_cliente_variable' asignada con: '{nombre_cliente_variable}'")
            # elif bookmark_name == "FechaDocumento":
            #     fecha_documento_variable = bookmark_text
            #     print(f"Variable 'fecha_documento_variable' asignada con: '{fecha_documento_variable}'")

        # The dictionary 'all_bookmarks_data' itself stores all the data.
        # You can access specific bookmark texts like this:
        # texto_especifico = all_bookmarks_data.get("NombreDelMarcadorBuscado")
        # if texto_especifico is not None:
        #     print(f"\nTexto para 'NombreDelMarcadorBuscado': '{texto_especifico}'")
        # else:
        #     print("\nMarcador 'NombreDelMarcadorBuscado' no encontrado.")
        doc.save(doc_filename_base + "_testing_v1.docx")

    else:
        print("No se encontraron marcadores en el documento, o no se pudo extraer texto de ellos.")

# For i in range
"""for nombre in all_bookmarks_data:
    print(f"Nombre del Marcador: {nombre}")

valores = all_bookmarks_data.values()
import pandas as pd
valores = all_bookmarks_data.values()
import pandas as pd
df = pd.DataFrame(
    list(all_bookmarks_data.items()),
    columns=['Marcador', 'Texto']
)

df["Texto_a_ingresar"] = "Primero " + df["Texto"]
# Export to xlsx
df.to_excel('marcadores.xlsx', index=False)"""

import os
import docx
from docx.oxml.ns import nsmap
from docx.oxml import OxmlElement

# Define the WordprocessingML namespace for easier XML element checking
W_NS = nsmap['w']


def iter_elements(parent_element):
    """
    Recursively yields XML elements from a parent element in document order.
    Helps traverse complex structures like paragraphs, runs, tables, SDTs, etc.
    """
    if parent_element is None:
        return

    # Iterate directly over the children of the parent element
    for child in parent_element.getchildren():
        # Yield the child element itself FIRST
        yield child

        # Then, if the child is a container type that can hold other elements
        # we care about (like runs, hyperlinks, fields, etc.), recurse into it.
        # This list of tags covers common inline containers and fields.
        if child.tag in (
                OxmlElement('{%s}r' % W_NS).tag,  # Run
                OxmlElement('{%s}hyperlink' % W_NS).tag,  # Hyperlink
                OxmlElement('{%s}fldSimple' % W_NS).tag,  # Simple field
                OxmlElement('{%s}smartTag' % W_NS).tag,  # Deprecated Smart Tag
                # OxmlElement('{%s}sdt' % W_NS).tag,     # SDT handled below
        ):
            yield from iter_elements(child)

        # Handle Structured Document Tags (SDTs) - need to find the content part
        elif child.tag == OxmlElement('{%s}sdt' % W_NS).tag:
            sdtContent = child.find(OxmlElement('{%s}sdtContent' % W_NS).tag)
            if sdtContent is not None:
                yield from iter_elements(sdtContent)  # Recurse into content

        # Handle tables - need to recurse into cells (within rows)
        elif child.tag == OxmlElement('{%s}tbl' % W_NS).tag:
            # Iterate through rows and cells and recurse into cell content
            # Use xpath to find all tr and tc within the table element
            for row in child.xpath('.//w:tr', namespaces={'w': W_NS}):
                for cell in row.xpath('.//w:tc', namespaces={'w': W_NS}):
                    yield from iter_elements(cell)


def get_bookmark_text_data(doc):
    """
    Extracts text content associated with each bookmark in the document.

    Args:
        doc (docx.document.Document): The python-docx document object.

    Returns:
        dict: A dictionary where keys are bookmark names (str) and values
              are the extracted text (str) between the bookmark start and end.
              Note: Whitespace/newlines within the bookmark might not be
              perfectly preserved depending on the document's XML structure
              and how w:t elements handle spacing. This function concatenates
              text from all w:t and w:instrText elements found within an
              active bookmark range.
    """
    bookmark_data = {}
    active_bookmarks = {}  # {bookmark_id: bookmark_name}

    # Iterate through all XML elements in the document body in order
    # doc.element.body is the root of the main document content XML
    for element in iter_elements(doc.element.body):

        # Check if it's a bookmarkStart element
        if element.tag == OxmlElement('{%s}bookmarkStart' % W_NS).tag:
            bkm_name = element.get(OxmlElement('{%s}name' % W_NS).tag)
            bkm_id = element.get(OxmlElement('{%s}id' % W_NS).tag)
            if bkm_name and bkm_id:
                # Activate this bookmark for text collection
                active_bookmarks[bkm_id] = bkm_name
                # Initialize the text entry if it doesn't exist
                bookmark_data[bkm_name] = bookmark_data.get(bkm_name, "")

        # Check if it's a bookmarkEnd element
        elif element.tag == OxmlElement('{%s}bookmarkEnd' % W_NS).tag:
            bkm_id = element.get(OxmlElement('{%s}id' % W_NS).tag)
            # Deactivate this bookmark by removing its ID
            # This assumes typical bookmarks; complex nested cases might need more logic
            if bkm_id in active_bookmarks:
                del active_bookmarks[bkm_id]

        # Check if it's a text element (w:t) or field instruction text (w:instrText)
        # Field instruction text is common within bookmarks used for merge fields
        elif element.tag == OxmlElement('{%s}t' % W_NS).tag or element.tag == OxmlElement('{%s}instrText' % W_NS).tag:
            text = element.text
            if text is not None and active_bookmarks:  # Only add text if any bookmark is currently active
                # Append this text to the accumulated text for ALL currently active bookmarks
                # This handles nested bookmarks by duplicating text, which is a common
                # way text inside nested bookmarks is treated for extraction purposes.
                for bkm_name in active_bookmarks.values():
                    bookmark_data[bkm_name] += text

        # Note: This simple text concatenation approach doesn't perfectly handle
        # complex whitespace, tabs (w:tab), breaks (w:br), or complex field codes.
        # It focuses on extracting raw text between the start and end tags.

    return bookmark_data


# --- Your existing script starts here ---

doc_filename_base = "contrato_automatizado_con_marcadores"
doc_path = doc_filename_base + ".docx"

# Check if the document exists
if not os.path.exists(doc_path):
    print(f"Error: El archivo '{doc_path}' no fue encontrado en el directorio de trabajo: {os.getcwd()}")
    # The commented-out dummy document creation code is a good way to test
    # if you don't have a real document ready. Uncomment if needed.
    # print("Creando un documento de prueba...")
    # try:
    #     dummy_doc = docx.Document()
    #     p = dummy_doc.add_paragraph()
    #     # Add bookmarkStart for BM_Prueba
    #     bm_start_el = docx.oxml.OxmlElement('w:bookmarkStart')
    #     bm_start_el.set(docx.oxml.ns.qn('w:name'), 'BM_Prueba')
    #     bm_start_el.set(docx.oxml.ns.qn('w:id'), '0')
    #     p._p.append(bm_start_el)
    #     # Add a run with text
    #     p.add_run("Este es el texto para ")
    #     # Add bookmarkEnd for BM_Prueba
    #     bm_end_el = docx.oxml.OxmlElement('w:bookmarkEnd')
    #     bm_end_el.set(docx.oxml.ns.qn('w:id'), '0')
    #     p._p.append(bm_end_el)
    #     # Add more text in another run after the bookmark
    #     p.add_run("BM_Prueba.")

    #     # Add a second paragraph with another bookmark
    #     p2 = dummy_doc.add_paragraph()
    #     bm_start_el2 = docx.oxml.OxmlElement('w:bookmarkStart')
    #     bm_start_el2.set(docx.oxml.ns.qn('w:name'), 'OtroMarcador')
    #     bm_start_el2.set(docx.oxml.ns.qn('w:id'), '1')
    #     p2._p.append(bm_start_el2)
    #     p2.add_run("Texto del segundo marcador.")
    #     bm_end_el2 = docx.oxml.OxmlElement('w:bookmarkEnd')
    #     bm_end_el2.set(docx.oxml.ns.qn('w:id'), '1')
    #     p2._p.append(bm_end_el2)

    #     dummy_doc.save(doc_path)
    #     print(f"Documento de prueba '{doc_path}' creado con marcadores 'BM_Prueba' y 'OtroMarcador'.")
    # except Exception as create_e:
    #     print(f"Error creating dummy document: {create_e}")
    exit()  # Exit if the file doesn't exist and we aren't creating a dummy

try:
    doc = docx.Document(doc_path)
except Exception as e:
    print(f"Error al abrir el documento '{doc_path}': {e}")
    exit()

# Call the function to extract all bookmark names and their texts
all_bookmarks_data = get_bookmark_text_data(doc)

# Now, 'all_bookmarks_data' is a dictionary like {'BookmarkName': 'Text found inside'}.
# You can iterate through it to access each bookmark's name and text.
print("\n--- Datos de los Marcadores Extraídos ---")
if all_bookmarks_data:
    # Sort bookmarks by name for consistent output (optional)
    sorted_bookmark_names = sorted(all_bookmarks_data.keys())

    for bookmark_name in sorted_bookmark_names:
        bookmark_text = all_bookmarks_data[bookmark_name]

        # Format the text by wrapping it in double curly braces
        formatted_text = "{{" + bookmark_text + "}}"

        # Print the extracted data
        print(f"Nombre del Marcador: {bookmark_name}")
        print(f"Texto Asociado     : '{formatted_text}'")  # Using '' to show if text is empty
        print("-" * 30)

        # You can now use these variables as needed. For example, store them,
        # print them, or use them in further processing.

        # Example: Assigning specific bookmark texts to unique variables
        # You would typically do this by checking the bookmark_name:
        #
        # if bookmark_name == "NombreCliente":
        #     nombre_cliente_variable = bookmark_text # Use the original text if you don't need {{}} yet
        #     print(f"(Asignado a 'nombre_cliente_variable': '{nombre_cliente_variable}')")
        # elif bookmark_name == "FechaDocumento":
        #     fecha_documento_variable = bookmark_text
        #     print(f"(Asignado a 'fecha_documento_variable': '{fecha_documento_variable}')")
        # elif bookmark_name == "MontoContrato":
        #      monto_contrato_variable = bookmark_text
        #      print(f"(Asignado a 'monto_contrato_variable': '{monto_contrato_variable}')")

    # The dictionary 'all_bookmarks_data' itself stores all the data.
    # You can access specific bookmark texts directly like this:
    # (Using .get() is safer as it returns None if the key doesn't exist, avoiding errors)
    texto_cliente = all_bookmarks_data.get("NombreCliente", "Marcador 'NombreCliente' no encontrado")
    print(f"\nAcceso directo a 'NombreCliente': '{texto_cliente}'")

    texto_fecha = all_bookmarks_data.get("FechaDocumento", "Marcador 'FechaDocumento' no encontrado")
    print(f"Acceso directo a 'FechaDocumento': '{texto_fecha}'")

    # You could also create a new dictionary with the formatted text if needed
    formatted_bookmarks_data = {name: "{{" + text + "}}" for name, text in all_bookmarks_data.items()}
    # print("\nFormatted Dictionary:", formatted_bookmarks_data)

    # The original document object 'doc' is not modified by get_bookmark_text_data,
    # so saving it here will save the original document as is.
    # If you were *replacing* bookmark text, you would save *after* replacement.
    # doc.save(doc_filename_base + "_testing_v1.docx") # Saving the original doc
    # If you needed to save a *new* doc with the formatted text *replacing* bookmarks,
    # that would be a separate step involving finding bookmarks and replacing runs.


else:
    print("No se encontraron marcadores en el documento, o no se pudo extraer texto de ellos.")

print("\nProceso de extracción completado.")