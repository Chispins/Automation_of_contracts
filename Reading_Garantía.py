import re

from Formated_Base_PEP8 import configurar_directorio_trabajo
import os

configurar_directorio_trabajo()

"""lista_elements = os.listdir()
regex = r"(GARANT).{2}"

coincidencia = re.match(regex, lista_elements[0])
if coincidencia:
    print("Coincidencia encontrada:", coincidencia.group())

for i in lista_elements:
    coincidencia = re.search(regex, i)
    if coincidencia:
        print("Coincidencia encontrada:", coincidencia.group())
        print("Nombre del archivo:", i)
        file_name = str(i)
        break"""

#regex_pdf = r"(\w+\.pdf)"
from PIL import Image
# Now I want to open that pdf and read it with OCR
"""
import pytesseract
# Open the pdf
pdf = Image.open(file_name)
# Read the pdf with OCR
text = pytesseract.image_to_string(pdf)
print(text)"""



from Formated_Base_PEP8 import configurar_directorio_trabajo
configurar_directorio_trabajo()
# --- Save the document ---



#file_name = "CERTIFICADO GARANTIA HOSPITAL MELIPILLA.png"
file_name = "Anotación 2025-05-12 101902.jpg"
image = Image.open(file_name)


google_api_key = "AIzaSyAgXYnDNJ6gQmMVUk88G6wGPdXz-qG2Gbw"
from openai import OpenAI
OpenAI.api_key = google_api_key


model = "gemini-2.5-flash-preview-04-17"
# read the image




import google.generativeai as genai
image = Image.open(file_name)

# Set your Google API key
genai.configure(api_key=google_api_key)

# Initialize the model (use the specific model name you're targeting)
model = genai.GenerativeModel('models/gemini-2.5-flash-preview-04-17')  # Adjust if the exact model name differs

prompt = (
    """
Markdown

Extrae los siguientes datos del documento proporcionado y devuelve la información estrictamente en formato JSON, sin texto adicional antes ni después.

Los datos a extraer son:

Certificado de fianza Web
Fecha (día, mes y año)
Nombre de afianzado
R.U.T. N° del afianzado
Domicilio del afianzado
Nombre del mandante
R.U.T. N° del mandante
Domicilio del mandante
Obligación caucionada
Monto
Glosa
El formato de salida debe ser un array JSON que contenga exactamente dos elementos:

Un objeto JSON donde las claves sean los nombres exactos de los datos solicitados (tal como se listan arriba) y los valores correspondan a una lista y sean los datos extraídos junto con su nivel de confianza.
Asegúrate de que la respuesta sea solo el array JSON, comenzando con [ y terminando con ]. No incluyas prefijos (como "Confianza:"), explicaciones o cualquier otro texto. El objetivo es que la salida sea directamente parseable como un array donde el primer elemento sea el diccionario de datos.

Ejemplo de estructura esperada (los valores y confianzas serán los extraídos del documento):
[
{
"Certificado de fianza Web": {"valor extraído": Confianza},
"Fecha": "valor extraído": Confianza,
"Nombre de afianzado": "valor extraído:Confianza": Confianza,
"R.U.T. N° del afianzado": "valor extraído":Confianza,
"Domicilio del afianzado": "valor extraído":Confianza,
"Nombre del mandante": "valor extraído":Confianza,
"R.U.T. N° del mandante": "valor extraído":Confianza,
"Domicilio del mandante": "valor extraído":Confianza,
"Obligación caucionada": "valor extraído":Confianza,
"Monto": "valor extraído":Confianza,
"Glosa": "valor extraído":Confianza
},
    """
)

google_flash = model.generate_content([prompt, image])
google_flash_text = google_flash.text
import json


import ast
import json

input_string = 'n\n[\n{\n"Certificado de fianza Web": "Certificado de Fianza Fiel cumplimiento de contrato",\n"Fecha": "06/02/2025",\n"Nombre de afianzado": "CENTRO RADIOLÓGICO FLEMING S.A.",\n"R.U.T. N° del afianzado": "88.316.700-7",\n"Domicilio del afianzado": "",\n"Nombre del mandante": "HOSPITAL DE MELIPILLA",\n"R.U.T. N° del mandante": "61.602.123-0",\n"Domicilio del mandante": "",\n"Obligación caucionada": "\\"PARA GARANTIZAR EL FIEL CUMPLIMIENTO DEL CONTRATO DENOMINADO: COMPRA DE EXÁMENES DE RESONANCIA PARA PACIENTES DEL HOSPITAL SAN JOSÉ DE MELIPILLA ID 1057480-70-LR24 Y/O DE LAS OBLIGACIONES LABORALES Y SOCIALES DEL ADJUDICATARIO\\".",\n"Monto": "$ 30.000.000",\n"Glosa": "PAGADERO A PRIMER REQUERIMIENTO."\n},\n[0.95, 1.0, 1.0, 1.0, 0.9, 1.0, 1.0, 0.9, 1.0, 1.0, 0.95]\n]\n```'

# Clean the input string
# Remove the 'n\n' and the trailing '```'
cleaned_string = input_string.replace('n\n', '', 1).strip().rstrip('```')

# The cleaned string is almost a Python literal representation of a list
# containing a dictionary and a list. We can use ast.literal_eval to safely parse it.
# json.loads expects a JSON string, and while the dictionary part is JSON-like,
# the overall structure `[{...}, [...]]` with the list of numbers is not strictly
# a single JSON object or array in the way json.loads would expect without
# potential errors due to the mix and the surrounding Python list syntax.
# ast.literal_eval is suitable for evaluating string literals of Python structures.

try:
    # Safely evaluate the string as a Python literal
    parsed_list = ast.literal_eval(cleaned_string)

    # Check if the parsed result is a list with two elements
    if isinstance(parsed_list, list) and len(parsed_list) == 2:
        # The first element is the dictionary
        dict_data = parsed_list[0]
        # The second element is the list of numbers
        list_numbers = parsed_list[1]

        # Now you have your two lists (one is a dictionary, the other a list of floats)
        print("Dictionary List:")
        print(dict_data)
        print("\nNumbers List:")
        print(list_numbers)
    else:
        print("The parsed string did not result in a list with two elements.")

except (ValueError, SyntaxError) as e:
    print(f"Error parsing the string: {e}")


import json, re

# 1. Extraer el bloque JSON entre corchetes
m = re.search(r'(\[.*\])', google_flash_text, re.S)
json_text = m.group(1) if m else google_flash_text.strip()

# 2. Parsear y mostrar
datos = json.loads(json_text)
for entrada in datos:
    for campo, (valor) in entrada.items():
        print(f"{campo}: {valor}")
datos = datos[0]

for entrada in datos:
    for campo, (valor) in entrada.items():
        print(f"{campo}: {valor}")

from docx import Document

# Supongamos que `datos` es la lista resultante de json.loads()
# y que el primer elemento es el diccionario de interés:
dict_data = datos[0]

document = Document()
table = document.add_table(rows=1, cols=2)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Campo'
hdr_cells[1].text = 'Valor'

for campo, (valor, _) in dict_data.items():
    row_cells = table.add_row().cells
    row_cells[0].text = campo
    row_cells[1].text = valor or ''

document.save('resultado.docx')



datos_1 = datos[0]
for element in datos_1:
    print(element, datos_1[element][0])

import docx
doc = docx.Document()
# Now make the table with the list
# --- Create Document ---
datos_1["Certificado de fianza Web"]


datos[1]

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
# --- Create Document ---
document = Document()

# --- Add the first block of information using a table ---
# Use a 4-column table to align the pairs side-by-side
# Column 0: Label 1
# Column 1: Value 1
# Column 2: Label 2
# Column 3: Value 2
# The default table created by python-docx usually has no visible borders
table1 = document.add_table(rows=0, cols=4)
table1.autofit = False # Prevent auto-fitting
# Set approximate column widths for better alignment (adjust as needed)
# These widths help control the horizontal spacing and alignment
table1.columns[0].width = Inches(1.5) # Label 1 width
table1.columns[1].width = Inches(2.0) # Value 1 width
table1.columns[2].width = Inches(1.5) # Label 2 width
table1.columns[3].width = Inches(2.5) # Value 2 width

# Add rows and content to the first table
row = table1.add_row().cells
row[0].text = 'Certificado de Fianza:'
row[1].text = "FJEH124U2U33"
row[2].text = 'Fecha de Emisión:'
row[3].text = '27 de junio de 2024.'



row = table1.add_row().cells
row[0].text = 'Tomador:'
# Merge the remaining cells for a single value spanning columns 1-3
# Note: merging can sometimes be tricky with default cell objects, but this is the standard way.
# If you encounter issues with merging text, ensure text is added *after* merging.
merged_cell = row[1].merge(row[2]).merge(row[3])
merged_cell.text = 'DISEÑO Y MANTENCIÓN DE JARDINES PABLO NAVARRETE EIRL'

row = table1.add_row().cells
row[0].text = 'R.U.T. N°:'
merged_cell = row[1].merge(row[2]).merge(row[3])
merged_cell.text = '76.594.422-8'

row = table1.add_row().cells
row[0].text = 'Dirección del Tomador :'
merged_cell = row[1].merge(row[2]).merge(row[3])
merged_cell.text = 'Condominio almendros norte, casa 47, comuna de Calera de Tango.'

row = table1.add_row().cells
row[0].text = 'Beneficiario:'
row[1].text = 'Hospital San José de Melipilla'
row[2].text = 'R.U.T. N°'
row[3].text = '61.602.123-0'


# --- Add COBERTURA heading ---
# Add a paragraph below the first table
p = document.add_paragraph()
run = p.add_run('COBERTURA')
run.bold = True
p.alignment = WD_ALIGN_PARAGRAPH.CENTER # Center the heading

# --- Add the COBERTURA details using a second table ---
# Use a 4-column table again for the side-by-side layout
# This table will also typically have no visible borders by default
table2 = document.add_table(rows=0, cols=4)
table2.autofit = False
# Use the same column widths as the first table for consistency
table2.columns[0].width = Inches(1.5)
table2.columns[1].width = Inches(2.0)
table2.columns[2].width = Inches(1.5)
table2.columns[3].width = Inches(2.5)


# Add rows and content to the second table
row = table2.add_row().cells
row[0].text = 'Fecha de inicio:'
row[1].text = '21-06-2024'
row[2].text = 'Fecha de termino:'
row[3].text = '23-10-2025'

row = table2.add_row().cells
row[0].text = 'Valor Asegurado:'
merged_cell = row[1].merge(row[2]).merge(row[3])
merged_cell.text = 'UF 56,50'

row = table2.add_row().cells
row[0].text = 'Finalidad:'
merged_cell = row[1].merge(row[2]).merge(row[3])
merged_cell.text = 'Fiel cumplimiento.'

document.save('insurance_policy_snippet.docx')