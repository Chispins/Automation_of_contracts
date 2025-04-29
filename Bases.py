import os
import docx
from docx.enum.section import WD_SECTION_START
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import re

# Construye la ruta normalizada al directorio "Files"
def configurar_directorio_trabajo():
    """Configura el directorio de trabajo en la subcarpeta 'Files'."""
    # Construye la ruta normalizada al directorio "Files"
    cwd = os.getcwd()
    target_dir_name = "Files"
    wd = os.path.join(cwd, target_dir_name)

    # Define el patrón específico a buscar (duplicación de \Files)
    # Se asume separador de Windows (\). Se usan dobles barras invertidas en el patrón regex.
    pattern = r"Files\\Files"

    # Busca el patrón en la ruta generada
    if re.search(pattern, wd):
        # Si se encuentra, reemplaza la primera ocurrencia de la duplicación
        wd = wd.replace(r"\Files\Files", r"\Files")

    # Cambia al directorio destino, verificando primero si es un directorio válido
    if os.path.isdir(wd):
        os.chdir(wd)
        print(f"Directorio de trabajo cambiado a: {wd}") # Opcional: confirmar cambio
    else:
        print(f"Advertencia: El directorio '{wd}' no existe o no es válido. No se cambió el directorio de trabajo.")
        # Aquí podrías decidir crear el directorio si no existe, o manejar el error.
        # Ejemplo para crearlo:
        # try:
        #     os.makedirs(wd, exist_ok=True) # Crea el directorio si no existe
        #     os.chdir(wd)
        #     print(f"Directorio '{wd}' creado y establecido como directorio de trabajo.")
        # except OSError as e:
        #     print(f"Error al crear o acceder al directorio '{wd}': {e}")

# Llamar a la función para configurar el directorio de trabajo
configurar_directorio_trabajo()


def crear_numeracion(doc):
    """Crea un formato de numeración y devuelve su ID"""
    # Asegurarse de que exista la parte de numeración
    part = doc._part
    if not hasattr(part, 'numbering_part'):
        part._add_numbering_part()

    # Crear un ID único para esta numeración
    import random
    num_id = random.randint(1000, 9999)  # ID aleatorio para evitar conflictos
    return num_id
# Función para aplicar numeración a un párrafo
def aplicar_numeracion(parrafo, num_id, nivel=0):
    """Aplica numeración a un párrafo"""
    p = parrafo._p
    pPr = p.get_or_add_pPr()

    # Eliminar numeración previa si existe
    for child in pPr.iterchildren():
        if child.tag.endswith('numPr'):
            pPr.remove(child)

    # Crear nuevo numPr
    numPr = OxmlElement('w:numPr')

    # Establecer nivel
    ilvl = OxmlElement('w:ilvl')
    ilvl.set(qn('w:val'), str(nivel))
    numPr.append(ilvl)

    # Establecer ID de numeración
    numId = OxmlElement('w:numId')
    numId.set(qn('w:val'), str(num_id))
    numPr.append(numId)

    # Agregar al párrafo
    pPr.append(numPr)

    # Configurar sangría
    ind = OxmlElement('w:ind')
    ind.set(qn('w:left'), '720')
    ind.set(qn('w:hanging'), '360')
    pPr.append(ind)

    return parrafo

# Configuración del documento

base_partida = "Base_en_Blanco.docx"
doc = docx.Document()

# Crear un único estilo para listas numeradas (o usar el predefinido)
list_style = 'List Number'
if list_style not in doc.styles:
    doc.styles.add_style(list_style, WD_STYLE_TYPE.PARAGRAPH)

# Crear IDs para numeración
num_id_vistos = crear_numeracion(doc)
num_id_resolucion = crear_numeracion(doc)
num_id_bases_p1 = crear_numeracion(doc)

# Centrar para tablas
def centrar_verticalmente_tabla(tabla):
    """Aplica alineación vertical centrada a todas las celdas de una tabla"""
    for fila in tabla.rows:
        for celda in fila.cells:
            # Acceder al elemento XML de la celda
            tc = celda._tc

            # Buscar o crear el elemento tcPr (propiedades de celda)
            tcPr = tc.get_or_add_tcPr()

            # Eliminar cualquier configuración previa de alineación vertical
            for vAlign in tcPr.findall(qn('w:vAlign')):
                tcPr.remove(vAlign)

            # Crear nuevo elemento de alineación vertical
            vAlign = OxmlElement('w:vAlign')
            vAlign.set(qn('w:val'), 'center')  # Valor 'center' para centrado vertical

            # Agregar el elemento a las propiedades de la celda
            tcPr.append(vAlign)


# Titulos Nivel 0
doc.add_heading("RESOLUCIÓN EXENTA Nº1", level=0)

# SECCIÓN VISTOS
doc.add_heading("VISTOS", level=2)
doc.add_paragraph("Lo dispuesto en la Ley Nº 19.886 de Bases sobre Contratos Administrativos de Suministro y Prestación de Servicios; el Decreto Supremo Nº 250 /04 modificado por los Decretos Supremos Nº 1763/09, 1383/11 y 1410/14 todos del Ministerio de Hacienda; D. S. 38/2005, Reglamento Orgánico de los Establecimientos de Menor Complejidad y de los Establecimientos de Autogestión en Red; en uso de las atribuciones que me confieren el D.F.L. Nº 1/2.005, en virtud del cual se fija el texto refundido, coordinado y sistematizado del D.L. 2.763/79 y de las leyes 18.933 y 18.469; lo establecido en los Decretos Supremos Nos 140/04, Reglamento Orgánico de los Servicios de Salud; la Resolución Exenta RA 116395/343/2024 de fecha 12/08/2024 del SSMOCC., la cual nombra Director del Hospital San José de Melipilla al suscrito; lo dispuesto por las Resoluciones 10/2017, 7/2019 y 8/2019 ambas de la Contraloría General de la República, y,")

# SECCIÓN CONSIDERANDO
doc.add_heading("CONSIDERANDO", level=2)

# Lista numerada para CONSIDERANDO (usando el mismo ID que antes usábamos para VISTOS)
p1 = doc.add_paragraph("Visto: La Ley N° 19.880, de 2003, que establece normas sobre los actos administrativos...",
                    style=list_style)
aplicar_numeracion(p1, num_id_vistos)

p2 = doc.add_paragraph("Que, el Hospital de San José de Melipilla perteneciente a la red de salud...",
                    style=list_style)
aplicar_numeracion(p2, num_id_vistos)

p3 = doc.add_paragraph("Que, dada la naturaleza del Establecimiento...", style=list_style)
aplicar_numeracion(p3, num_id_vistos)

# Párrafo con formato en negrita
vistos_p4 = doc.add_paragraph(style=list_style)
vistos_p4.add_run("Que, existe la necesidad ")
run_bold = vistos_p4.add_run("suministro de insumos y accesorios para terapia de presión negativa con equipos en comodato")
run_bold.bold = True
vistos_p4.add_run(", a fin de entregar una prestación de salud integral y oportuna...")
aplicar_numeracion(vistos_p4, num_id_vistos)

p5 = doc.add_paragraph("Que corresponde asegurar la transparencia en este proceso...", style=list_style)
aplicar_numeracion(p5, num_id_vistos)

p6 = doc.add_paragraph("Que, considerando los montos de la contratación...", style=list_style)
aplicar_numeracion(p6, num_id_vistos)

p7 = doc.add_paragraph("Que revisado el catálogo de bienes y servicios...", style=list_style)
aplicar_numeracion(p7, num_id_vistos)

p8 = doc.add_paragraph("Que, en consecuencia y en mérito de lo expuesto...", style=list_style)
aplicar_numeracion(p8, num_id_vistos)

p9 = doc.add_paragraph("Que, en razón de lo expuesto y la normativa vigente;", style=list_style)
aplicar_numeracion(p9, num_id_vistos)

# SECCIÓN RESOLUCIÓN
doc.add_heading("RESOLUCIÓN", level=2)

# Lista numerada para RESOLUCIÓN (nuevo ID = nueva secuencia)
resolucion_p1 = doc.add_paragraph(style=list_style)
resolucion_p1.add_run("LLÁMASE ").bold = True
resolucion_p1.add_run("a Licitación Pública Nacional a través del Portal Mercado Público, para la compra de ")
resolucion_p1.add_run("Suministro de Insumos y Accesorios para Terapia de Presión Negativa con Equipos en Comodato ").bold = True
resolucion_p1.add_run("para el Hospital San José de Melipilla.")
aplicar_numeracion(resolucion_p1, num_id_resolucion)

resolucion_p2 = doc.add_paragraph(style=list_style)
resolucion_p2.add_run("ACOGIENDOSE ").bold = True
resolucion_p2.add_run("al Art.º 25 del decreto 250 que aprueba el reglamento de la ley Nº 19.886...")
aplicar_numeracion(resolucion_p2, num_id_resolucion)

resolucion_p3 = doc.add_paragraph(style=list_style)
resolucion_p3.add_run("APRUÉBENSE las bases administrativas, técnicas y anexos N.º 1, 2, 3, 4, 5, 6, 7, 8 y 9 ").bold = True
resolucion_p3.add_run("desarrollados para efectuar el llamado a licitación, que se transcriben a continuación:")
aplicar_numeracion(resolucion_p3, num_id_resolucion)

# Nueva Sección
doc.add_heading("BASES ADMINISTRATIVAS PARA EL SUMINISTRO DE INSUMOS Y ACCESORIOS PARA TERAPIA DE PRESIÓN NEGATIVA CON EQUIPOS EN COMODATO PARA EL HOSPITAL SAN JOSÉ DE MELIPILLA",level=1)
doc.add_heading("REQUISITOS", level=2)
doc.add_paragraph("En Santiago, a 1 de enero de 2023, se resuelve lo siguiente:")

# Lista numerada para BASES (nuevo ID = nueva secuencia)
bases_p1 = doc.add_paragraph(style=list_style)
bases_p1.add_run("Antecedentes Básicos de la ENTIDAD LICITANTE").bold = True
aplicar_numeracion(bases_p1, num_id_bases_p1)

# Tabla
tabla_p1 = doc.add_table(6, 2, style="Table Grid")
tabla_p1.cell(0, 0).text = "Razón Social del organismo"
tabla_p1.cell(0, 1).text = "Hospital San José de Melipilla"
tabla_p1.cell(1, 0).text = "Unidad de Compra"
tabla_p1.cell(1, 1).text = "Unidad de Abastecimiento de Bienes y Servicios"
tabla_p1.cell(2, 0).text = "R.U.T. del organismo"
tabla_p1.cell(2, 1).text = "61.602.123-0"
tabla_p1.cell(3, 0).text = "Dirección"
tabla_p1.cell(3, 1).text = "O'Higgins #551"
tabla_p1.cell(4, 0).text = "Comuna"
tabla_p1.cell(4, 1).text = "Melipilla"
tabla_p1.cell(5, 0).text = "Región en que se genera la Adquisición"
tabla_p1.cell(5, 1).text = "Región Metropolitana"
centrar_verticalmente_tabla(tabla_p1)

bases_p2 = doc.add_paragraph(style=list_style)
bases_p2.add_run("Antecedentes Administrativos").bold = True
aplicar_numeracion(bases_p2, num_id_bases_p1)

tabla_p2 = doc.add_table(8, 2, style="Table Grid")
tabla_p2.cell(0, 0).text = "Nombre Adquisición"
tabla_p2.cell(0, 1).text = "Suministro de Insumos y Accesorios para Terapia de Presión Negativa con Equipos en Comodato"
tabla_p2.cell(1, 0).text = "Descripción"
tabla_p2.cell(1, 1).text = "El Hospital requiere generar un convenio por el SUMINISTRO DE INSUMOS Y ACCESORIOS PARA TERAPIA DE PRESIÓN NEGATIVA CON EQUIPOS EN COMODATO PARA EL HOSPITAL SAN JOSÉ DE MELIPILLA, en adelante “EL HOSPITAL”. El convenio tendrá una vigencia de 36 meses."
tabla_p2.cell(2, 0).text = "Tipo de Convocatoria"
tabla_p2.cell(2, 1).text = "Abierta"
tabla_p2.cell(3, 0).text = "Moneda o Unidad reajustable"
tabla_p2.cell(3, 1).text = "Pesos Chilenos"
tabla_p2.cell(4, 0).text = "Presupuesto Referencial"
tabla_p2.cell(4, 1).text = "$350.000.000.- (Impuestos incluidos)"
tabla_p2.cell(5, 0).text = "Etapas del Proceso de Apertura"
tabla_p2.cell(5, 1).text = "Una Etapa (Etapa de Apertura Técnica y Etapa de Apertura Económica en una misma instancia)."
tabla_p2.cell(6, 0).text = "Opciones de pago"
tabla_p2.cell(6, 1).text = "Transferencia Electrónica"
tabla_p2.cell(7, 0).text = "Tipo de Adjudicación"
tabla_p2.cell(7, 1).text = "Adjudicación por la totalidad"
centrar_verticalmente_tabla(tabla_p2)

bases_p2_runs = doc.add_paragraph()
bases_p2_runs.add_run("* Presupuesto referencial:").underline = True
bases_p2_runs.add_run(" " + "El Hospital se reserva el derecho de aumentar, previo acuerdo entre las partes, hasta un 30% el presupuesto referencial estipulado en las presentes bases de licitación.")

# -- Definiciones --
bases_p2_definiciones = doc.add_paragraph()
bases_p2_definiciones.add_run("Definiciones").bold = True

bases_p2_definiciones_def_a = doc.add_paragraph(style = "List Number 3")
bases_p2_definiciones_def_a.add_run("Proponente u oferente:").bold = True
bases_p2_definiciones_def_a.add_run(" " + "El proveedor o prestador que participa en el proceso de licitación mediante la presentación de una propuesta, en la forma y condiciones establecidas en las Bases.")

bases_p2_definiciones_def_b = doc.add_paragraph(style = "List Number 3")
bases_p2_definiciones_def_b.add_run("Administrador o coordinador Externo del Contrato").bold = True
bases_p2_definiciones_def_b.add_run(" " + "Persona designada por el oferente adjudicado, quien actuará como contraparte ante el Hospital.")

bases_p2_definiciones_def_c = doc.add_paragraph(style = "List Number 3")
bases_p2_definiciones_def_c.add_run("Dias Hábiles:").bold = True
bases_p2_definiciones_def_c.add_run(" " + "Son todos los días de la semana, excepto los sábados, domingos y festivos.")

bases_p2_definiciones_def_d = doc.add_paragraph(style = "List Number 3")
bases_p2_definiciones_def_d.add_run("Días Corridos:").bold = True
bases_p2_definiciones_def_d.add_run(" " + "Son los días de la semana que se computan uno a uno en forma correlativa. Salvo que se exprese lo contrario, los plazos de días señalados en las presentes bases de licitación son días corridos. En caso que el plazo expire en días sábados, domingos o festivos se entenderá prorrogados para el día hábil siguiente.")

bases_p2_definiciones_def_e = doc.add_paragraph(style = "List Number 3")
bases_p2_definiciones_def_e.add_run("Administrador del Contrato y/o Referente Técnico: ").bold = True
bases_p2_definiciones_def_e.add_run(" " + " Es el funcionario designado por el Hospital para supervisar la correcta ejecución del contrato, solicitar órdenes de compra, validar prefacturas, gestionar multas y/o toda otra labor que guarde relación con la ejecución del contrato.")

bases_p2_definiciones_def_f = doc.add_paragraph(style = "List Number 3")
bases_p2_definiciones_def_f.add_run("Gestor de Contrato:").bold = True
bases_p2_definiciones_def_f.add_run(" " + "Es el funcionario a cargo de la ejecución del presente proceso de Licitación, desde la publicación de las Bases hasta la generación del contrato en formato documental, como la elaboración de ficha en la plataforma “gestor de contrato” en el portal de mercado público, además de ser el responsable de dar seguimiento y cumplimiento a los procesos y plazos establecidos. ")


# Tabla de definiciones de tipos de licitación
tabla_p2_licitaciones = doc.add_table(6, 3, style="Table Grid")

# Encabezados de la tabla
tabla_p2_licitaciones.cell(0, 0).text = "RANGO (en UTM)"
tabla_p2_licitaciones.cell(0, 1).text = "TIPO LICITACION PUBLICA"
tabla_p2_licitaciones.cell(0, 2).text = "PLAZO PUBLICACION EN DIAS CORRIDOS"

# Filas de datos
tabla_p2_licitaciones.cell(1, 0).text = "<100"
tabla_p2_licitaciones.cell(1, 1).text = "L1"
tabla_p2_licitaciones.cell(1, 2).text = "5"

tabla_p2_licitaciones.cell(2, 0).text = "<=100 y <1000"
tabla_p2_licitaciones.cell(2, 1).text = "LE"
tabla_p2_licitaciones.cell(2, 2).text = "10, rebajable a 5"

tabla_p2_licitaciones.cell(3, 0).text = "<=1000 y <2000"
tabla_p2_licitaciones.cell(3, 1).text = "LP"
tabla_p2_licitaciones.cell(3, 2).text = "20, rebajable a 10"

tabla_p2_licitaciones.cell(4, 0).text = "<=2000 y <5000"
tabla_p2_licitaciones.cell(4, 1).text = "LQ"
tabla_p2_licitaciones.cell(4, 2).text = "20, rebajable a 10"

tabla_p2_licitaciones.cell(5, 0).text = "<=5000"
tabla_p2_licitaciones.cell(5, 1).text = "LR"
tabla_p2_licitaciones.cell(5, 2).text = "30"
centrar_verticalmente_tabla(tabla_p2_licitaciones)

bases_p3 = doc.add_paragraph(style=list_style)
bases_p3.add_run("Etapas y plazos:").bold = True
aplicar_numeracion(bases_p3, num_id_bases_p1)

# Crear tabla de etapas y plazos con una fila adicional para el título
tabla_p3_plazos = doc.add_table(9, 3, style="Table Grid")

# Combinar las celdas de la primera fila para el título
tabla_p3_plazos.cell(0, 0).merge(tabla_p3_plazos.cell(0, 2))

# Agregar el título en la celda combinada
titulo_celda = tabla_p3_plazos.cell(0, 0)
titulo_celda.text = "VIGENCIA DE LA PUBLICACION 10 DIAS CORRIDOS"
# Hacer que el texto del título esté en negrita
for paragraph in titulo_celda.paragraphs:
    for run in paragraph.runs:
        run.bold = True
    # Centrar el título
    paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

# Fila 1 (ahora es la segunda fila en la tabla)
tabla_p3_plazos.cell(1, 0).text = "Consultas"
tabla_p3_plazos.cell(1, 1).text = "Hasta las 15:00 Horas de 4º (cuarto) día corrido de publicada la Licitación."
tabla_p3_plazos.cell(1, 2).text = "Deben ser ingresadas al portal www.mercadopublico.cl"

# Fila 2
tabla_p3_plazos.cell(2, 0).text = "Respuestas a Consultas"
tabla_p3_plazos.cell(2, 1).text = "Hasta las 17:00 Horas de 7º (séptimo) día corrido de publicada la Licitación."
tabla_p3_plazos.cell(2, 2).text = "Deben ser ingresadas al portal www.mercadopublico.cl"

# Fila 3
tabla_p3_plazos.cell(3, 0).text = "Aclaratorias"
tabla_p3_plazos.cell(3, 1).text = "Hasta 1 días corrido antes del cierre de recepción de ofertas."
tabla_p3_plazos.cell(3, 2).text = "Deben ser ingresadas al portal www.mercadopublico.cl"

# Fila 4
tabla_p3_plazos.cell(4, 0).text = "Recepción de ofertas"
tabla_p3_plazos.cell(4, 1).text = "Hasta las 17:00 Horas de 10º (décimo) día corrido de publicada la Licitación."
tabla_p3_plazos.cell(4, 2).text = "Deben ser ingresadas al portal www.mercadopublico.cl"

# Fila 5
tabla_p3_plazos.cell(5, 0).text = "Evaluación de las Ofertas"
tabla_p3_plazos.cell(5, 1).text = "Máximo 40 días corridos a partir del cierre de la Licitación."
tabla_p3_plazos.cell(5, 2).text = ""

# Fila 6
tabla_p3_plazos.cell(6, 0).text = "Plazo Adjudicaciones"
tabla_p3_plazos.cell(6, 1).text = "Máximo 20 días corridos a partir de la fecha del acta de evaluación de las ofertas."
tabla_p3_plazos.cell(6, 2).text = ""

# Fila 7
tabla_p3_plazos.cell(7, 0).text = "Suscripción de Contrato"
tabla_p3_plazos.cell(7, 1).text = "Máximo de 20 días hábiles desde la Adjudicación de la Licitación."
tabla_p3_plazos.cell(7, 2).text = ""

# Fila 8
tabla_p3_plazos.cell(8, 0).text = "Consideración"
tabla_p3_plazos.cell(8, 1).text = "Los plazos de días establecidos en la cláusula 3, Etapas y Plazos, son de días corridos, excepto el plazo para emitir la orden de compra, el que se considerará en días hábiles, entendiéndose que son inhábiles los sábados, domingos y festivos en Chile, sin considerar los feriados regionales."
tabla_p3_plazos.cell(8, 2).text = ""

centrar_verticalmente_tabla(tabla_p3_plazos)

doc.add_heading("Consultas, Aclaraciones y modificaciones a las bases.", level=2)

# Párrafo 1
p_consultas1 = doc.add_paragraph()
p_consultas1.add_run("Las consultas de los participantes se deberán realizar únicamente a través del portal")
p_consultas1.add_run(" www.mercadopublico.cl").bold = True
p_consultas1.add_run(" conforme el cronograma de actividades de esta licitación señalado en el punto 3 precedente. A su vez, las respuestas y aclaraciones estarán disponibles a través del portal de Mercado Público, en los plazos indicados en el cronograma señalado precedentemente, información que se entenderá conocida por todos los interesados desde el momento de su publicación.")
p_consultas1.style = 'List Bullet' # Apply the bullet style

# Párrafo 2 (now a list item)
p_consultas2 = doc.add_paragraph()
p_consultas2.add_run("No serán admitidas las consultas formuladas fuera de plazo o por un conducto diferente al señalado.").bold = True
p_consultas2.style = 'List Bullet' # Apply the bullet style

# Párrafo 3 (now a list item)
p_consultas3 = doc.add_paragraph()
p_consultas3.add_run("“EL HOSPITAL” realizará las aclaraciones a las Bases comunicando las respuestas a través del Portal Web de Mercado Público, sitio")
p_consultas3.add_run(" www.mercadopublico.cl").bold = True
p_consultas3.style = 'List Bullet' # Apply the bullet style

# Párrafo 4 (now a list item)
p_consultas4 = doc.add_paragraph()
p_consultas4.add_run("Las aclaraciones, derivadas de este proceso de consultas, formarán parte integrante de las Bases, teniéndose por conocidas y aceptadas por todos los participantes aun cuando el oferente no las hubiere solicitado, por lo que los proponentes no podrán alegar desconocimiento de las mismas.")
p_consultas4.style = 'List Bullet' # Apply the bullet style

# Párrafo 5 (now a list item)
p_consultas5 = doc.add_paragraph()
p_consultas5.add_run("“EL HOSPITAL” podrá modificar las presentes bases y sus anexos previa autorización por acto administrativo, durante el periodo de presentación de las ofertas, hasta antes de fecha de cierre de recepción de ofertas. Estas modificaciones, que se llewen a cabo, serán informadas a través del portal")
p_consultas5.add_run(" www.mercadopublico.cl").bold = True
p_consultas5.style = 'List Bullet' # Apply the bullet style

# Párrafo 6 (now a list item)
p_consultas6 = doc.add_paragraph()
p_consultas6.add_run("Estas consultas, aclaratorias y modificaciones formaran parte integra de las bases y estarán vigentes desde la total tramitación del acto administrativo que las apruebe. Junto con aprobar las modificaciones, deberá establecer un nuevo plazo prudencial cuando lo amerite para el cierre o recepción de las propuestas, a fin de que los potenciales oferentes puedan adecuar sus ofertas.")
p_consultas6.style = 'List Bullet' # Apply the bullet style

# Párrafo 7 (now a list item)
p_consultas7 = doc.add_paragraph()
p_consultas7.add_run("No se aceptarán consultas realizadas por otros medios, tales como correos electrónicos, fax u otros.")
p_consultas7.style = 'List Bullet' # Apply the bullet style

# Requisitos Participación
# Crear IDs para numeración
num_id_vistos = crear_numeracion(doc)
num_id_resolucion = crear_numeracion(doc)
num_id_bases_p1 = crear_numeracion(doc)
# Crear un nuevo ID para la sección de consultas
num_id_consultas = crear_numeracion(doc)
# Crear un nuevo ID para la sección de requisitos mínimos
num_id_requisitos = crear_numeracion(doc)

# ... (código anterior para VISTOS, CONSIDERANDO, RESOLUCIÓN, BASES ADMINISTRATIVAS, CONSULTAS) ...

# Requisitos Participación
doc.add_heading("Requisitos Mínimos para Participar.", level=3)

# Párrafos numerados para Requisitos Mínimos
req_p1 = doc.add_paragraph(
    "No haber sido condenado por prácticas antisindicales, infracción a los derechos fundamentales del trabajador o por delitos concursales establecidos en el Código Penal dentro de los dos últimos años anteriores a la fecha de presentación de la oferta, de conformidad con lo dispuesto en el artículo 4° de la ley N° 19.886."
)
aplicar_numeracion(req_p1, num_id_requisitos) # Aplicar numeración

req_p2 = doc.add_paragraph(
    "No haber sido condenado por el Tribunal de Defensa de la Libre Competencia a la medida dispuesta en la letra d) del artículo 26 del Decreto con Fuerza de Ley N°1, de 2004, del Ministerio de Economía, Fomento y Reconstrucción, que Fija el texto refundido, coordinado y sistematizado del Decreto Ley N° 211, de 1973, que fija normas para la defensa de la libre competencia, hasta por el plazo de cinco años contado desde que la sentencia definitiva quede ejecutoriada."
)
aplicar_numeracion(req_p2, num_id_requisitos) # Aplicar numeración

req_p3 = doc.add_paragraph(
    "No ser funcionario directivo de la respectiva entidad compradora; o una persona unida a aquél por los vínculos de parentesco descritos en la letra b) del artículo 54 de la ley N° 18.575; o una sociedad de personas de las que aquél o esta formen parte; o una sociedad comandita por acciones o anónima cerrada en que aquélla o esta sea accionista; o una sociedad anónima abierta en que aquél o esta sean dueños de acciones que representen el 10% o más del capital; o un gerente, administrador, representante o director de cualquiera de las sociedades antedichas."
)
aplicar_numeracion(req_p3, num_id_requisitos) # Aplicar numeración

req_p4 = doc.add_paragraph(
    "Tratándose exclusivamente de una persona jurídica, no haber sido condenada conforme a la ley N° 20.393 a la pena de prohibición de celebrar actos y contratos con el Estado, mientras esta pena esté vigente."
)
aplicar_numeracion(req_p4, num_id_requisitos) # Aplicar numeración

req_p5 = doc.add_paragraph("A fin de acreditar el cumplimiento de dichos requisitos, los oferentes deberán presentar una “Declaración jurada de requisitos para ofertar”, la cual será generada completamente en línea a través de www.mercadopublico.cl en el módulo de presentación de las ofertas. Sin perjuicio de lo anterior, la entidad licitante podrá verificar la veracidad de la información entregada en la declaración, en cualquier momento, a través de los medios oficiales disponibles.")
aplicar_numeracion(req_p5, num_id_requisitos) # Aplicar numeración

req_p6 = doc.add_paragraph()
req_p6.add_run("En caso de que los antecedentes administrativos solicitados en esta sección no sean entregados y/o completados en forma correcta y oportuna, se desestimará la propuesta, no será evaluada y será declarada ")
req_p6.add_run("inadmisible").bold = True
req_p6.add_run(".")
aplicar_numeracion(req_p6, num_id_requisitos) # Aplicar numeración






from docx import Document

# Assume 'doc' is an existing Document object, like:
# doc = Document()

import docx
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from Bases import configurar_directorio_trabajo, centrar_verticalmente_tabla # Assuming these are defined in Bases.py

# Configure working directory (if needed)
configurar_directorio_trabajo()

# Assume 'doc' is an existing Document object or create a new one
# doc = Document() # Uncomment if you need a new document

# Add the heading before the table
doc.add_heading("Instrucciones para la Presentación de Ofertas.", level=2)

# Create a table with exactly 4 rows and 2 columns
table = doc.add_table(rows=4, cols=2)
table.style = 'Table Grid' # Optional: Add grid lines

# --- Row 1: Header ---
cell_r1_c1 = table.cell(0, 0)
cell_r1_c2 = table.cell(0, 1)
cell_r1_c1.text = "Presentar Ofertas por Sistema."
cell_r1_c2.text = "Obligatorio."
# Optional: Center header text if needed
# cell_r1_c1.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
# cell_r1_c2.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

# --- Row 2: Anexos Administrativos ---
# --- Row 2: Anexos Administrativos ---
cell_r2_c1 = table.cell(1, 0)
cell_r2_c2 = table.cell(1, 1)
cell_r2_c1.text = "Anexos Administrativos."

# Add content to the right cell (Row 2, Col 2)
# Clear the default paragraph if it exists and is empty
if cell_r2_c2.paragraphs and not cell_r2_c2.paragraphs[0].text:
    p_element = cell_r2_c2.paragraphs[0]._p
    cell_r2_c2._element.remove(p_element)


# Item 1: Anexo N° 1
p1 = cell_r2_c2.add_paragraph()
p1.add_run("Anexo N° 1 Identificación del Oferente.").bold = True

# Item 2: Anexo N° 2
p2 = cell_r2_c2.add_paragraph()
p2.add_run("Anexo N° 2 Declaración Jurada de Habilidad.").bold = True

# Item 3: Anexo N° 3 (CORREGIDO)
p3 = cell_r2_c2.add_paragraph() # Crear el párrafo primero
p3.add_run("Anexo N° 3 Declaración Jurada de Cumplimiento de Obligaciones Laborales y Previsionales.").bold = True # Añadir run y poner en negrita

# Item 4: Declaración jurada online
p4 = cell_r2_c2.add_paragraph()
p4.add_run("Declaración jurada online:").bold = True
p4.add_run(" Los oferentes deberán presentar una ") # Espacio añadido al inicio
p4.add_run("Declaración jurada de requisitos para ofertar").bold = True
p4.add_run(", la cual será generada completamente en línea a través de www.mercadopublico.cl en el módulo de presentación de las ofertas.") # Coma y espacio añadidos

# Item 5: Unión Temporal de Proveedores (UTP)
p5 = cell_r2_c2.add_paragraph()
p5.add_run("Unión Temporal de Proveedores (UTP):").bold = True
p5.add_run(" Solo en el caso de que la oferta sea presentada por una unión temporal de proveedores deberán presentar obligatoriamente la siguiente documentación en su totalidad, en caso contrario, ésta no será sujeta a aclaración y la oferta será declarada ")
p5.add_run("inadmisible").bold = True
p5.add_run(".") # Punto añadido

# Item 6: Anexo N°4 for UTP
p6 = cell_r2_c2.add_paragraph()
p6.add_run("Anexo N°4. Declaración para Uniones Temporales de Proveedores:").bold = True
p6.add_run(" Debe ser presentado por el miembro de la UTP que presente la oferta en el Sistema de Información y quien realiza la declaración a través de la “Declaración jurada de requisitos para ofertar” electrónica presentada junto a la oferta.") # Espacio añadido al inicio

# Item 7: UTP Offers and Apoderado
p7 = cell_r2_c2.add_paragraph()
p7.add_run("Las ofertas presentadas por una Unión Temporal de Proveedores (UTP) deberán contar con un apoderado, el cual debe corresponder a un integrante de la misma, ya sea persona natural o jurídica. En el caso que el apoderado sea una persona jurídica, ésta deberá actuar a través de su representante legal para ejercer sus facultades.")

# Item 8: Inadmissibility condition
p8 = cell_r2_c2.add_paragraph()
p8.add_run("En caso de no presentarse debidamente la declaración jurada online constatando la ausencia de conflictos de interés e inhabilidades por condenas, o no presentarse el Anexo N°4, la oferta será declarada ")
p8.add_run("inadmisible").bold = True
p8.add_run(".") # Punto añadido

# --- Row 3: Anexos Económicos ---
cell_r3_c1 = table.cell(2, 0)
cell_r3_c2 = table.cell(2, 1)
cell_r3_c1.text = "Anexos\nEconómicos." # Using \n for line break in the cell text

# Add content to the right cell (Row 3, Col 2)
# Clear the default paragraph if it exists and is empty
if cell_r3_c2.paragraphs and not cell_r3_c2.paragraphs[0].text:
    p_element = cell_r3_c2.paragraphs[0]._p
    cell_r3_c2._element.remove(p_element)

# Item 1: Anexo N°5
ep1 = cell_r3_c2.add_paragraph()
ep1.add_run("Anexo N°5: Oferta económica").bold = True

# Item 2: Entering via system
ep2 = cell_r3_c2.add_paragraph()
ep2.add_run("El anexo referido debe ser ingresado a través del sistema www.mercadopublico.cl , en la sección Anexos Económicos.")

# Item 3: Inadmissibility condition
ep3 = cell_r3_c2.add_paragraph()
ep3.add_run("En caso de que no se presente debidamente el Anexo N°5 “Oferta económica”, la oferta será declarada ")
ep3.add_run("inadmisible").bold = True

# --- Row 4: Anexos Técnicos ---
cell_r4_c1 = table.cell(3, 0)
cell_r4_c2 = table.cell(3, 1)
cell_r4_c1.text = "Anexos Técnicos." # Using \n for line break

# Add content to the right cell (Row 4, Col 2)
# Clear the default paragraph if it exists and is empty
if cell_r4_c2.paragraphs and not cell_r4_c2.paragraphs[0].text:
    p_element = cell_r4_c2.paragraphs[0]._p
    cell_r4_c2._element.remove(p_element)

# Item 1: Anexo N°6
tp1 = cell_r4_c2.add_paragraph()
tp1.add_run("Anexo N°6: Evaluación Técnica").bold = True

# Item 2: Anexo N°7
tp2 = cell_r4_c2.add_paragraph()
tp2.add_run("Anexo N°7: Ficha Técnica").bold = True

# Item 3: Anexo N°8
tp3 = cell_r4_c2.add_paragraph()
tp3.add_run("Anexo N°8: Plazo de Entrega").bold = True

# Item 4: Anexo N°9
tp4 = cell_r4_c2.add_paragraph()
tp4.add_run("Anexo N°9: Servicio Post-venta").bold = True

# Item 5: Entering via system
tp5 = cell_r4_c2.add_paragraph()
tp5.add_run("Los anexos referidos deben ser ingresados a través del sistema www.mercadopublico.cl. en la sección Anexos Técnicos.")

# Item 6: Inadmissibility condition
tp6 = cell_r4_c2.add_paragraph()
tp6.add_run("En el caso que no se presente debidamente los Anexos N°7, N°8 y N°9 la oferta será declarada ")
tp6.add_run("inadmisible").bold = True

# Apply vertical centering to the entire table (optional, requires your function)
centrar_verticalmente_tabla(table)

# Save the document
doc.add_heading("Observaciones", level = 3)
parrafos_observaciones = doc.add_paragraph()
parrafos_observaciones.add_run("Los oferentes deberán presentar su oferta a través de su cuenta en el Sistema de Información www.mercadopublico.cl. De existir discordancia entre el oferente o los antecedentes de su oferta y la cuenta a través de la cual la presenta, esta no será evaluada, siendo desestimada del proceso y declarada como")
parrafos_observaciones.add_run(" " + "inadmisible").bold = True


p1 = doc.add_paragraph()
p1.add_run("Las únicas ofertas válidas serán las presentadas a través del portal")
p1.add_run(" www.mercadopublico.cl").bold = True
p1.add_run(", en la forma en que se solicita en estas bases. No se aceptarán ofertas que se presenten por un medio distinto al establecido en estas Bases, a menos que se acredite la indisponibilidad técnica del sistema, de conformidad con el artículo 62 del Reglamento de la Ley de Compras. Será responsabilidad de los oferentes adoptar las precauciones necesarias para ingresar oportuna y adecuadamente sus ofertas.")

# Paragraph 2
p2 = doc.add_paragraph()
p2.add_run("Los oferentes deben constatar que el envío de su oferta a través del portal electrónico de compras públicas haya sido realizado con éxito, incluyendo el previo ingreso de todos los formularios y anexos requeridos completados de acuerdo con lo establecido en las presentes bases. Debe verificar que los archivos que se ingresen contengan efectivamente los anexos solicitados.")

# Paragraph 3
p3 = doc.add_paragraph()
p3.add_run("Asimismo, se debe comprobar siempre, luego de que se finalice la última etapa de ingreso de la oferta respectiva, que se produzca el despliegue automático del “Comprobante de Envío de Oferta” que se entrega en dicho Sistema, el cual puede ser impreso por el proponente para su resguardo. En dicho comprobante será posible visualizar los anexos adjuntos, cuyo contenido es de responsabilidad del oferente.")

# Paragraph 4
p4 = doc.add_paragraph()
p4.add_run("El hecho de que el oferente haya obtenido el “Comprobante de envío de ofertas” señalado, únicamente acreditará el envío de ésta a través del Sistema, pero en ningún caso certificará la integridad o la completitud de ésta, lo cual será evaluado por la comisión evaluadora. En caso de que, antes de la fecha de cierre de la licitación, un proponente edite una oferta ya enviada, deberá asegurarse de enviar nuevamente la oferta una vez haya realizado los ajustes que estime, debiendo descargar un nuevo Comprobante.")

# Paragraph 5
p5 = doc.add_paragraph()
p5.add_run("Si la propuesta económica subida al portal, presenta diferencias entre el valor del anexo económico solicitado y el valor indicado en la línea de la plataforma")
p5.add_run(" www.mercadopublico.cl").bold = True
p5.add_run(", prevalecerá la oferta del anexo económico solicitado en bases. Sin embargo, el Hospital San José de Melipilla, podrá solicitar aclaraciones de las ofertas realizadas a través del portal.")


doc.add_heading("Antecedentes legales para poder ser contratado.", level=2)


table = doc.add_table(rows=7, cols=3)
table.style = 'Table Grid' # Apply grid lines style

# Optional: Adjust column widths for better layout
# table.columns[0].width = Inches(1.5)
# table.columns[1].width = Inches(4.0)
# table.columns[2].width = Inches(1.5)

# --- Section: Si el oferente es Persona Natural (Rows 0-3) ---
start_natural_row = 0
end_natural_row = 3 # Rows 0, 1, 2, 3 (4 rows total)

# Cell (0, 0): "Si el oferente es Persona Natural" - Merged
cell_0_0 = table.cell(start_natural_row, 0)
# Use paragraphs and runs for better control over potential future formatting or line breaks
p_0_0 = cell_0_0.paragraphs[0]
p_0_0.add_run("Si el oferente\nes Persona\nNatural")

# Cell (0, 2): "Acreditar en el Registro de Proveedores" - Merged
cell_0_2 = table.cell(start_natural_row, 2)
p_0_2 = cell_0_2.paragraphs[0]
p_0_2.add_run("Acreditar en\nel Registro de\nProveedores")

# Populate middle column (Column 1) for Persona Natural section
# Cell (0, 1): Requirement 1
cell_0_1 = table.cell(start_natural_row, 1)
p_0_1 = cell_0_1.paragraphs[0]
# Text is bold in the image for "Inscripción..." and "Registro..."
p_0_1.add_run("Inscripción (en estado hábil) en el Registro electrónico oficial de contratistas de la Administración, en adelante “").bold = True
p_0_1.add_run("Registro de Proveedores").bold = True
p_0_1.add_run("”.").bold = True # The closing quote and period also appear bold


# Cell (1, 1): Requirement 2 - Anexo N°3
cell_1_1 = table.cell(start_natural_row + 1, 1)
p_1_1 = cell_1_1.paragraphs[0]
p_1_1.add_run("Anexo N°3. Declaración Jurada de Cumplimiento de Obligaciones Laborales y Previsionales.").bold = True # Whole line appears bold

# Cell (2, 1): Requirement 3 - Todos los Anexos...
cell_2_1 = table.cell(start_natural_row + 2, 1)
cell_2_1.text = "Todos los Anexos deben ser firmados por la persona natural respectiva."

# Cell (3, 1): Requirement 4 - Fotocopia...
cell_3_1 = table.cell(start_natural_row + 3, 1)
cell_3_1.text = "Fotocopia de su cédula de identidad."

# Perform merges for the Persona Natural section
cell_0_0.merge(table.cell(end_natural_row, 0)) # Merge column 0 (rows 0 to 3)
cell_0_2.merge(table.cell(end_natural_row, 2)) # Merge column 2 (rows 0 to 3)


# --- Section: Si el oferente no es Persona Natural (Rows 4-6) ---
start_nonatural_row = end_natural_row + 1 # Starts at row 4
end_nonatural_row = start_nonatural_row + 2 # Rows 4, 5, 6 (3 rows total)

# Cell (4, 0): "Si el oferente no es Persona Natural" - Merged
cell_4_0 = table.cell(start_nonatural_row, 0)
p_4_0 = cell_4_0.paragraphs[0]
p_4_0.add_run("Si el oferente\nno es\nPersona\nNatural")

# Cell (4, 2): "Acreditar en el Registro de Proveedores" - Merged
cell_4_2 = table.cell(start_nonatural_row, 2)
p_4_2 = cell_4_2.paragraphs[0]
p_4_2.add_run("Acreditar en\nel Registro de\nProveedores")

# Populate middle column (Column 1) for Persona no Natural section
# Cell (4, 1): Requirement 1 - Inscripción...
cell_4_1 = table.cell(start_nonatural_row, 1)
p_4_1 = cell_4_1.paragraphs[0]
p_4_1.add_run("Inscripción (en estado hábil) en el Registro de Proveedores.").bold = True # Whole line appears bold

# Cell (5, 1): Requirement 2 - Certificado de Vigencia...
cell_5_1 = table.cell(start_nonatural_row + 1, 1)
cell_5_1.text = "Certificado de Vigencia del poder del representante legal, con una antigüedad no superior a 60 días corridos, contados desde la fecha de notificación de la adjudicación, otorgado por el Conservador de Bienes Raíces correspondiente o, en los casos que resulte procedente, cualquier otro antecedente que acredite la vigencia del poder del representante legal del oferente, a la época de presentación de la oferta."

# Cell (6, 1): Requirement 3 - Certificado de Vigencia de la Sociedad...
cell_6_1 = table.cell(start_nonatural_row + 2, 1)
cell_6_1.text = "Certificado de Vigencia de la Sociedad con una antigüedad no superior a 60 días corridos, contados desde la fecha de notificación de la" # Text cut off in image

# Perform merges for the Persona no Natural section
cell_4_0.merge(table.cell(end_nonatural_row, 0)) # Merge column 0 (rows 4 to 6)
cell_4_2.merge(table.cell(end_nonatural_row, 2)) # Merge column 2 (rows 4 to 6)


# Observaciones
doc.add_heading("Observaciones", level = 3)
doc.add_paragraph("Los antecedentes legales para poder ser contratado, sólo se requerirán respecto del adjudicatario y deberán estar disponibles en el Registro de Proveedores.")
doc.add_paragraph("Lo señalado en el párrafo precedente no resultará aplicable a la garantía de fiel cumplimiento de contrato, la cual podrá ser entregada físicamente en los términos que indican las presentes bases en aquellos casos que aplique su entrega.")
doc.add_paragraph("En los casos en que se otorgue de manera electrónica, deberá ajustarse a la ley N° 19.799 sobre documentos electrónicos, firma electrónica y servicios de certificación de dicha firma, y remitirse en la forma señalada en la cláusula 8.2 de estas bases.")
observ_parafo_2 = doc.add_paragraph()
observ_parafo_2.add_run("Si el respectivo proveedor no entrega la totalidad de los antecedentes requeridos para ser contratado, dentro del plazo fatal de 10 días hábiles administrativos contados desde la notificación de la resolución de adjudicación o no suscribe el contrato en los plazos establecidos en estas bases, la entidad licitante podrá readjudicar de conformidad a lo establecido en la")
observ_parafo_2.add_run(" " + "cláusula 9 letra i")
observ_parafo_2.add_run(" de las presentes bases. Además, tales incumplimientos darán origen al cobro de la garantía de seriedad de la oferta, si la hubiere.")

doc.add_heading("Inscripción en el registro de proveedores", level =2 )
doc.add_paragraph("En caso de que el proveedor que resulte adjudicado no se encuentre inscrito en el Registro Electrónico Oficial de Contratistas de la Administración (Registro de Proveedores), deberá inscribirse dentro del plazo de 15 días hábiles, contados desde la notificación de la resolución de adjudicación.")
doc.add_paragraph("Tratándose de los adjudicatarios de una Unión Temporal de Proveedores, cada integrante de ésta deberá inscribirse en el Registro de Proveedores, dentro del plazo de 15 días hábiles, contados desde la notificación de la resolución de adjudicación. ")

doc.add_heading("Naturaleza y monto de las garatías", level = 2)
doc.add_heading("Evaluación y adjudicación de las ofertas", level = 3)

comis_eval_p1 = doc.add_paragraph()
comis_eval_p1.add_run("Comisión Evaluadora: ").bold = True
comis_eval_p1.add_run("La Dirección del Hospital San José de Melipilla designa como integrantes de la Comisión de Evaluación de la propuesta a los siguientes funcionarios: el Subdirector(a) Administrativo, Subdirector(a) Médico de Atención Abierta, Subdirector(a) Médico de Atención Cerrada, Subdirector(a) de Gestión del Cuidado de Enfermería, Subdirector(a) de Gestión y Desarrollo de las Personas, Subdirector(a) de Matronería, Subdirector(a) de Análisis de Información para la Gestión, Subdirector(a) de Apoyo Clínico o sus subrogantes. Para los efectos del quórum para sesionar se requerirá un mínimo de tres miembros. Lo anterior en conformidad con lo dispuesto en el artículo 37 del Decreto Nº 250 que establece el Reglamento de la Ley Nº 19.886.Los miembros de la Comisión Evaluadora no podrán:")


# Párrafo 1 de la lista con viñetas

# Elemento 2 de la lista con viñetas
eval_p1 = doc.add_paragraph("La comisión evaluadora verificará el cumplimiento de los requisitos mínimos de participación.", style='List Bullet')

# Elemento 3 de la lista con viñetas
eval_p2 = doc.add_paragraph("Se evaluarán los criterios técnicos y económicos según la ponderación definida en las bases.", style='List Bullet')

# Elemento 4 de la lista con viñetas y negrita
eval_p3 = doc.add_paragraph(style='List Bullet')
eval_p3.add_run("La adjudicación se realizará al oferente que obtenga el ")
eval_p3.add_run("mayor puntaje total").bold = True
eval_p3.add_run(".")

doc.add_paragraph("La misma Comisión estudiará los antecedentes de la Propuesta y elaborará un informe fundado para el Director de este Establecimiento, quien podrá declarar, mediante resolución fundada, admisible aquellas ofertas que cumplan con los requisitos establecidos en las bases de licitación, como también podrá declarar, mediante resolución fundada, inadmisible aquellas ofertas que no cumplan los requisitos establecidos en las bases. En caso de no presentarse oferentes o cuando las ofertas no resulten convenientes para los intereses del Establecimiento, podrá declarar desierta la licitación, fundándose en razones objetivas y no discriminatorias.Esta Comisión Evaluadora podrá invitar a profesionales técnicos para colaborar en el proceso de adjudicación")


consider_general_p1 = doc.add_paragraph()
consider_general_p1.add_run("Consideraciones Generales: ").bold = True
consider_general_p1.add_run("Se exigirá el cumplimiento de los requerimientos establecidos en la cláusula 6, “Instrucciones para Presentación de Ofertas”, de las presentes Bases de Licitación. Aquellas ofertas que no fueran presentadas a través del portal, en los términos solicitados, se declararán como propuestas inadmisibles, por tanto, no serán consideradas en la evaluación. Lo anterior, sin perjuicio de que concurra y se acredite algunas de las causales de excepción establecidas en el artículo 62 del Reglamento de la Ley de Compras.")

doc.add_paragraph("La entidad licitante declarará inadmisible cualquiera de las ofertas presentadas que no cumplan los requisitos o condiciones establecidos en las presentes bases, sin perjuicio de la facultad de la entidad licitante de solicitar a los oferentes que salven errores u omisiones formales, de acuerdo con lo establecido en el artículo 40 del Reglamento de la Ley N°19.886 y en las presentes bases.")

doc.add_paragraph("Los documentos solicitados por la entidad licitante deben estar vigentes a la fecha de cierre de la presentación de las ofertas indicado en la cláusula 3 de las presentes bases y ser presentados como copias simples, legibles y firmadas por el representante legal de la empresa o persona natural. Sin perjuicio de ello, la entidad licitante podrá verificar la veracidad de la información entregada por el proveedor. En el caso en que el proveedor esté inscrito y habilitado por el Registro de Proveedores, serán suficientes los antecedentes que se encuentren en dicho Registro, en la medida que se haya dado cumplimiento a las normas de actualización de documentos que establece el Registro de Proveedores.")

subsan_err_p1 = doc.add_paragraph()
subsan_err_p1.add_run("Subsanación de errores u omisiones formales: ").bold = True
subsan_err_p1.add_run("vicios u omisiones no les confieran a esos oferentes una situación de privilegio respecto de los demás competidores, esto es, en tanto no se afecten los principios de estricta sujeción a las bases y de igualdad de los oferentes, y se informe de dicha solicitud al resto de los oferentes, a través del Sistema de Información")
subsan_err_p1.add_run("www.mercadopublico.cl.").bold = True

doc.add_paragraph("El plazo que tendrán los oferentes, en este caso para dar cumplimiento a lo solicitado por el mandante, no será inferior a las 24 horas, contadas desde la fecha de publicación de la solicitud por parte del Hospital, la que se informará a través del Sistema de información www.mercadopublico.cl. La responsabilidad de revisar oportunamente dicho sistema durante el período de evaluación recae exclusivamente en los respectivos oferentes.")

inadmisibilidad_p1 = doc.add_paragraph()
inadmisibilidad_p1.add_run("Inadmisibilidad de las ofertas y declaración de desierta de la licitación: ").bold = True
inadmisibilidad_p1.add_run("La entidad licitante declarará inadmisible las ofertas presentadas que no cumplan los requisitos mínimos establecidos en los Anexos N°5, N°6, N°7, N°8 y N°9 y/o las condiciones establecidas en las presentes bases de licitación, sin perjuicio de la facultad para solicitar a los oferentes que salven errores u omisiones formales de acuerdo con lo establecido en las presentes bases.")
doc.add_paragraph("La entidad licitante podrá, además, declarar desierta la licitación cuando no se presenten ofertas o cuando éstas no resulten convenientes a sus intereses.Dichas declaraciones deberán materializarse a través de la dictación de una resolución fundada y no darán derecho a indemnización alguna a los oferentes.")

criterios_eval_p1 = doc.add_paragraph()
criterios_eval_p1.add_run("Criterios de evaluación y procedimientos de las ofertas: ").bold = True
criterios_eval_p1.add_run("La evaluación de las ofertas se realizará en una etapa, utilizando criterios técnicos, económicos y administrativos.")

criterios_eval_p2 = doc.add_paragraph()
criterios_eval_p2.add_run("La evaluación de las ofertas presentadas para el ")
criterios_eval_p2.add_run("SUMINISTRO DE INSUMOS Y ACCESORIOS PARA TERAPIA DE PRESIÓN NEGATIVA CON EQUIPOS EN COMODATO PARA EL HOSPITAL SAN JOSÉ DE MELIPILLA, ")
criterios_eval_p2.add_run(" se regirá por las siguientes ponderaciones y criterios a evaluar:")


# Tabla
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT

# Assume 'doc' is an existing Document object, like:
# doc = Document()

# --- Create Table ---
# Needs 5 rows (1 header + 4 data) and 4 columns
num_rows = 5
num_cols = 4
table_criterios = doc.add_table(rows=num_rows, cols=num_cols)
table_criterios.style = 'Table Grid' # Apply borders

# --- Populate Header Row (Row 0) ---
# Merging the first two cells for "CRITERIOS"
header_cell_0 = table_criterios.cell(0, 0)
header_cell_1 = table_criterios.cell(0, 1)
merged_header_cell = header_cell_0.merge(header_cell_1)
merged_header_cell.text = "CRITERIOS"
merged_header_cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
merged_header_cell.paragraphs[0].runs[0].bold = True
merged_header_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

# PONDERACIÓN header
cell_0_2 = table_criterios.cell(0, 2)
cell_0_2.text = "PONDERACIÓN"
cell_0_2.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
cell_0_2.paragraphs[0].runs[0].bold = True
cell_0_2.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

# EVALUADO SEGÚN ANEXO header
cell_0_3 = table_criterios.cell(0, 3)
# Add paragraph and runs for line break
p_0_3 = cell_0_3.paragraphs[0] # Get the existing paragraph
p_0_3.text = "" # Clear existing text if any
p_0_3.add_run("EVALUADO").bold = True
p_0_3.add_run("\n") # Add the line break
p_0_3.add_run("SEGÚN ANEXO").bold = True
p_0_3.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
cell_0_3.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


# --- Populate Data Rows ---

# Row 1: ECONÓMICO / OFERTA ECONÓMICA
cell_1_0 = table_criterios.cell(1, 0)
cell_1_0.text = "ECONÓMICO"
cell_1_0.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
cell_1_0.paragraphs[0].runs[0].bold = True
cell_1_0.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

cell_1_1 = table_criterios.cell(1, 1)
cell_1_1.text = "OFERTA ECONÓMICA"
cell_1_1.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
cell_1_1.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

cell_1_2 = table_criterios.cell(1, 2)
cell_1_2.text = "60%"
cell_1_2.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
cell_1_2.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

cell_1_3 = table_criterios.cell(1, 3)
cell_1_3.text = "ANEXO N°5"
cell_1_3.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
cell_1_3.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

# Row 2: TÉCNICOS (start) / EVALUACIÓN TÉCNICA
cell_2_0 = table_criterios.cell(2, 0) # This is the start of the merge
cell_2_0.text = "TÉCNICOS"
cell_2_0.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
cell_2_0.paragraphs[0].runs[0].bold = True
cell_2_0.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

cell_2_1 = table_criterios.cell(2, 1)
cell_2_1.text = "EVALUACIÓN TÉCNICA"
cell_2_1.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
cell_2_1.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

cell_2_2 = table_criterios.cell(2, 2)
cell_2_2.text = "20%"
cell_2_2.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
cell_2_2.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

cell_2_3 = table_criterios.cell(2, 3)
cell_2_3.text = "ANEXO N°6"
cell_2_3.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
cell_2_3.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

# Row 3: PLAZO DE ENTREGA
# Column 0 is part of the merge, leave empty or it will be overwritten by merge
cell_3_1 = table_criterios.cell(3, 1)
cell_3_1.text = "PLAZO DE ENTREGA"
cell_3_1.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
cell_3_1.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

cell_3_2 = table_criterios.cell(3, 2)
cell_3_2.text = "10%"
cell_3_2.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
cell_3_2.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

cell_3_3 = table_criterios.cell(3, 3)
cell_3_3.text = "ANEXO N°8"
cell_3_3.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
cell_3_3.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

# Row 4: SERVICIO POST-VENTA
# Column 0 is part of the merge, leave empty or it will be overwritten by merge
cell_4_1 = table_criterios.cell(4, 1)
cell_4_1.text = "SERVICIO POST-VENTA"
cell_4_1.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
cell_4_1.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

cell_4_2 = table_criterios.cell(4, 2)
cell_4_2.text = "10%"
cell_4_2.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
cell_4_2.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

cell_4_3 = table_criterios.cell(4, 3)
cell_4_3.text = "ANEXO N°9"
cell_4_3.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
cell_4_3.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

# --- Perform Vertical Merge for "TÉCNICOS" ---
cell_4_0 = table_criterios.cell(4, 0) # Last cell in the merge range
cell_2_0.merge(cell_4_0) # Merge cell(2,0) down to cell(4,0)

# --- Final step: Save the document ---
# doc.save("criterios_table_basic.docx")

calculo_puntaje = doc.add_paragraph()
calculo_puntaje.add_run("Cálculo del puntaje total de evaluación: ").bold = True

calculo_puntaje_2 = doc.add_paragraph()
calculo_puntaje_2.add_run("El Puntaje de la Evaluación Final estará dado por el siguiente polinomio:")

calculo_puntaje_3 = doc.add_paragraph()
calculo_puntaje_3.add_run("Puntaje Evaluación Técnica + Puntaje Evaluación Económica")

calculo_puntaje_4 = doc.add_paragraph()
calculo_puntaje_4.add_run("Donde el Puntaje Evaluación Técnica = Evaluación Técnica + Plazo de Entrega + Servicio Post-Venta.")

calculo_puntaje_5 = doc.add_paragraph()
calculo_puntaje_5.add_run("Donde Puntaje Evaluación Económica = Precio.")

# Ahora paso a la siguiente fase
doc.add_heading("Criterios de evaluación", level = 2)

doc.add_heading("Criterios Económicos", level = 3)

crit_econ = doc.add_paragraph()
crit_econ.add_run("OFERTA ECONÓMICA 60%").bold = True

doc.add_paragraph("Valor ítem ofertado. Para calcular el puntaje correspondiente al precio se utilizará la siguiente fórmula: A este puntaje se le aplicará la ponderación del 60 %. El oferente deberá declarar en Anexo N°5, los valores ofertados considerando todos los gastos involucrados e impuestos que apliquen.")


doc.add_heading("Criterios Técnicos", level = 3)

# Crear el párrafo primero
crit_tecn_para = doc.add_paragraph(style="List Bullet")
# Añadir un Run con el texto en negrita
crit_tecn_run = crit_tecn_para.add_run("EVALUACIÓN TÉCNICA 20%: ")
crit_tecn_run.bold = True
# Añadir el resto del texto al mismo párrafo (sin negrita)
crit_tecn_para.add_run("Se evaluará según información presentada para el Anexo N°7 que deberá ser adjuntada en su oferta en el Portal de Mercado Público, junto con la pauta de evaluación del Anexo N°6. Se evaluará por producto ofertado, donde el puntaje total será el promedio de la evaluación de todos los insumos ofertados.")

# Modificar las siguientes líneas para usar el párrafo correcto si es necesario,
# o crear nuevos párrafos como antes si son elementos de lista separados.
# Si son elementos separados de la lista:
plazo_entrega_para = doc.add_paragraph(style="List Bullet")
plazo_entrega_run = plazo_entrega_para.add_run("PLAZO DE ENTREGA 10%: ")
plazo_entrega_run.bold = True
plazo_entrega_para.add_run("Se evaluará según información presentada en el Anexo N° 8 de la presente base de licitación.")

serv_post_venta_para = doc.add_paragraph(style="List Bullet")
serv_post_venta_run = serv_post_venta_para.add_run("Servicio Post-Venta 10%: ")
serv_post_venta_run.bold = True
serv_post_venta_para.add_run("Se evaluará según información presentada en el Anexo N° 9 de la presente base de licitación.")

doc.add_heading("Adjudicación", level = 2)
doc.add_paragraph("Se adjudicará al oferente que obtenga el mayor puntaje, en los términos descritos en las presentes bases. La presente licitación se adjudicará a través de una resolución dictada por la autoridad competente, la que será publicada en www.mercadopublico.cl, una vez que se encuentre totalmente tramitada.")

doc.add_heading("Mecanismo de Resolución de empates.", level = 2)
doc.add_paragraph("En el evento de que, una vez culminado el proceso de evaluación de ofertas, hubiese dos o más proponentes que hayan obtenido el mismo puntaje en la evaluación final, quedando más de uno en condición de resultar adjudicado, se optará por aquella oferta que cuente con un mayor puntaje de acuerdo con la secuencia de los criterios que resulten aplicables, de acuerdo al siguiente orden: EVALUACION TECNICA, seguido por PLAZO DE ENTREGA, seguido por SERVICIO POST-VENTA, seguido por CRITERIO ECONOMICO. Finalmente, de mantenerse la igualdad, se adjudicará a aquel oferente que haya ingresado primero su propuesta en el portal Mercado Público considerándose la hora en que aquello se efectúe.")

doc.add_heading("Resolución de consultas respecto de la Adjudicación.", level = 2)
resolucion_consultas_par = doc.add_paragraph()
resolucion_consultas_par.add_run("Las consultas sobre la adjudicación deberán realizarse dentro del plazo fatal de 5 días hábiles contados desde la publicación de la resolución en el Sistema de Información ")
resolucion_consultas_par.add_run("www.mercadopublico.cl").bold = True
resolucion_consultas_par.add_run("a través del siguiente enlace: ")
resolucion_consultas_par.add_run("http://ayuda.mercadopublico.cl ").bold = True

doc.add_heading("Readjudicación", level = 2)
doc.add_paragraph("Si el adjudicatario se desistiere de firmar el contrato o de aceptar la orden de compra, o no cumpliese con las demás condiciones y requisitos establecidos en las presentes bases para la suscripción o aceptación de los referidos documentos, la entidad licitante podrá, junto con dejar sin efecto la adjudicación original, adjudicar la licitación al oferente que le seguía en puntaje, o a los que le sigan sucesivamente, dentro del plazo de 60 días corridos contados desde la publicación de la adjudicación original.")

doc.add_section()
doc.add_heading("Condiciones Contractuales, Vigencia de las Condiciones Comerciales, Operatoria de la Licitación y Otras Cláusulas:", level = 1)

doc.add_heading("Documentos integrantes")
documentos_integrantes_p1 = doc.add_paragraph()
doc.add_paragraph("La relación contractual que se genere entre la entidad licitante y el adjudicatario se ceñirá a los siguientes documentos:")

doc.add_paragraph("Bases de licitación y sus anexos.", style = "List Bullet")
doc.add_paragraph("Aclaraciones, respuestas y modificaciones a las Bases, si las hubiere.", style = "List Bullet")
doc.add_paragraph("Oferta.", style = "List Bullet")
doc.add_paragraph("Orden de compra", style = "List Bullet")

doc.add_paragraph("Todos los documentos antes mencionados forman un todo integrado y se complementan recíprocamente, especialmente respecto de las obligaciones que aparezcan en uno u otro de los documentos señalados. Se deja constancia que se considerará el principio de preeminencia de las Bases.")

doc.add_heading("Validez de la Oferta")
doc.add_paragraph("La oferta tendrá validez de ciento veinte días (120) días corridos, contados desde la fecha de apertura de la propuesta. La oferta cuyo periodo de validez sea menor que el requerido, será rechazada de inmediato.")
doc.add_paragraph("Si vencido el plazo señalado precedentemente, el Hospital San José de Melipilla no ha realizado la adjudicación, podrá solicitar a los Proponentes la prórroga de sus ofertas y garantías. Los proponentes podrán ratificar sus ofertas o desistir de ellas, formalizando su decisión mediante comunicación escrita dirigida al Hospital. Se devolverá la garantía a aquellos que no accedan a la prórroga.")

doc.add_heading("Suscripción del Contrato")
doc.add_paragraph("Para suscribir el contrato o aceptar la orden de compra contemplada en el artículo 63 del reglamento de la Ley de Compras, el adjudicado deberá estar inscrito en el Registro de Proveedores.")
doc.add_paragraph("Para formalizar las adquisiciones de bienes y servicios regidas por la ley Nº 19.886, se requerirá la suscripción de un contrato, la que en este caso se verá reflejada por la sola aceptación de la respectiva Orden de Compras.")
doc.add_paragraph("El respectivo contrato deberá suscribirse dentro de los 20 días hábiles siguientes a la notificación de la resolución de adjudicación totalmente tramitada. Asimismo, cuando corresponda, la orden de compra que formaliza la adquisición deberá ser aceptada por el adjudicatario dentro de ese mismo plazo. ")
doc.add_paragraph("Si por cualquier causa que no sea imputable a la entidad licitante, el contrato no se suscribe dentro de dicho plazo, o no se acepta la orden de compra que formaliza la adquisición dentro de ese mismo término, se entenderá desistimiento de la oferta, pudiendo readjudicar la licitación al oferente que le seguía en puntaje, o a los que le sigan sucesivamente. ")
doc.add_paragraph("Para suscribir el contrato o aceptar la orden de compra contemplada en el artículo 63 del reglamento de la Ley de Compras, el adjudicado deberá estar inscrito en el Registro de Proveedores.")

doc.add_heading("Modificación del Contrato")
doc.add_paragraph("Las partes de común acuerdo podrán modificar el contrato aumentando o disminuyendo los Bienes o servicios licitados, como también se podrán pactar nuevos bienes o servicios que no alteren la naturaleza del contrato. Estas modificaciones podrán ser hasta un 30% el presupuesto disponible estipulado en las presentes bases de licitación. En el caso de aumentar los bienes o servicios contratados, la garantía fiel cumplimiento de contrato también podrá readecuarse en proporción al monto de la modificación que se suscriba según aquellos casos que apliquen. En caso de aumentar o disminuir los bienes o servicios contratados, los valores a considerar, serán aquellos ofertados en el anexo oferta económica.Con todo, las eventuales modificaciones que se pacten no producirán efecto alguno sino desde la total tramitación del acto administrativo que las apruebe.")

doc.add_heading("Gastos e Impuestos")
doc.add_paragraph("Todos los gastos e impuestos que se generen o produzcan por causa o con ocasión de este Contrato, tales como los gastos notariales de celebración de contratos y/o cualesquiera otros que se originen en el cumplimiento de obligaciones que, según las Bases, ha contraído el oferente adjudicado, serán de cargo exclusivo de éste.")

doc.add_heading("Efectps derivados de Incumplimiento del proveedor")
doc.add_paragraph("En función de la gravedad de la infracción cometida por el adjudicatario, se le aplicarán las siguientes sanciones:")

clasificacion_par1 = doc.add_paragraph()
clasificacion_par1.add_run("Amonestación: ").bold = True
clasificacion_par1.add_run("Corresponde a un registro escrito, que dejará de manifiesto cualquier falta menor cometida por el adjudicado. Se entenderá por falta menor aquella que no ponga en riesgo de forma alguna la prestación del servicio o la vida e integridad psíquica y física de los pacientes, que se vinculen a temas administrativos y técnicos y que no sea constitutiva de multa. La amonestación no estará afecta a sanción pecuniaria.")

# Crear el párrafo
clasificacion_par2 = doc.add_paragraph()
# Añadir el primer run ("Multa: ") y ponerlo en negrita
clasificacion_par2.add_run("Multa: ").bold = True
# Añadir el segundo run con el resto del texto
clasificacion_par2.add_run("Corresponde a la sanción de cualquier falta, de gravedad leve, moderada o grave en que incurra el adjudicado, cada vez que éste no dé cumplimiento a cualquiera de las obligaciones contempladas en las presentes bases. Se expresará en Unidades Tributarias Mensuales (UTM).El monto de cada multa, dependerá de la gravedad de la infracción cometida, en este sentido las multas se clasifican en:")


# Multa Leve Section
p_leve = doc.add_paragraph()
# Optional: Indent this paragraph slightly if desired
# p_leve.paragraph_format.left_indent = Inches(0.25)
p_leve.add_run("Multa Leve: ").bold = True
p_leve.add_run("Sera considerada LEVE aquella situación originada por una falta de carácter menor, que no origina riesgos a las personas, ni daños a los bienes de la Institución o a su imagen. Su importe será de 3 Unidades Tributarias Mensuales (UTM). Las conductas que pueden estar afectas a multa leve son:")

# Bullet points for Multa Leve
doc.add_paragraph("Entrega de productos con atraso de hasta dos (2) días hábiles, contados desde la fecha de entrega estipulada en el contrato, sin que estos se hubiesen re-pactado en su plazo de entrega.", style='List Bullet')
doc.add_paragraph("Conducta o trato irrespetuoso de parte del personal del Oferente adjudicado o su cadena de distribución.", style='List Bullet')
doc.add_paragraph("La acumulación de dos amonestaciones.", style='List Bullet')
doc.add_paragraph("Incumplimiento del contrato que no origine riesgos a las personas o daño a los bienes del establecimiento o a su imagen.", style='List Bullet')
# Optional: Add space after the list
last_leve_bullet = doc.paragraphs[-1] # Get the last added paragraph

# Multa Moderada Section
p_moderada = doc.add_paragraph()
# Optional: Indent this paragraph slightly if desired
# p_moderada.paragraph_format.left_indent = Inches(0.25)
p_moderada.add_run("Multa Moderada: ").bold = True
p_moderada.add_run("Sera considerada MODERADA, aquella situación originada por una falta que afecte o ponga en riesgo, directa o indirectamente a personas o a la Institución o que limite significativamente la atención y calidad del servicio, pero que es factible de ser corregida. Su importe será de 6 Unidades Tributarias Mensuales (UTM). Las conductas que puedan estar afectas a multa moderada son:")

# Bullet points for Multa Moderada
doc.add_paragraph("No aceptar la Orden de Compra dentro de los dos (4) días hábiles siguientes al envío de la orden a través del portal de Mercado Publico.", style='List Bullet')
doc.add_paragraph("Entrega de los productos con atraso de entre tres (3) y seis (6) días hábiles inclusive, contados desde la fecha de entrega estipulada en el contrato, sin que estos se hubiesen re-pactado en su plazo de entrega.", style='List Bullet')
doc.add_paragraph("Despacho de productos en lugares no autorizados por el Hospital.", style='List Bullet')
doc.add_paragraph("La acumulación de dos multas leves trimestres móviles.", style='List Bullet')
doc.add_paragraph("cualquier falta que afecte o ponga en riesgo, directa o indirectamente, a personas o a la institución, o que limite significativamente la atención y calidad del servicio, pero que es factible de ser corregida.", style='List Bullet')
# Optional: Add space after the list
last_moderada_bullet = doc.paragraphs[-1]

# Multa Grave Section
p_grave = doc.add_paragraph()
# Optional: Indent this paragraph slightly if desired
# p_grave.paragraph_format.left_indent = Inches(0.25)
p_grave.add_run("Multa Grave: ").bold = True
p_grave.add_run("Sera considerada GRAVE, aquella situación originada por una falta que atente, directa o indirectamente con la atención y calidad del servicio. Su importe será de 10 Unidades Tributarias Mensuales (UTM). Las conductas que pueden estar afectas a multa grave son:")# Optional small space before list

# Bullet points for Multa Grave
doc.add_paragraph("Incumplimiento de la totalidad de lo requerido en la orden de compra.", style='List Bullet')
doc.add_paragraph("Entrega de productos con atraso de más de seis (6) días hábiles, contados desde la fecha de entrega estipulada en el contrato, sin que estos se hubiesen re-pactado en su plazo de entrega.", style='List Bullet')
doc.add_paragraph("Rechazo total de los productos por no ajustarse a las especificaciones técnicas.", style='List Bullet')
doc.add_paragraph("El incumplimiento en el recambio de los productos que presenten problemas de estabilidad, empaque, envases en mal estado, conservación inadecuada, vencida, defectuosa o dañada en un periodo máximo de cuarenta y ocho (48) horas.", style='List Bullet')
doc.add_paragraph("La acumulación de dos multas moderadas en 2 trimestres móviles", style='List Bullet')
doc.add_paragraph("Si el adjudicatario no inicia sus labores en la fecha acordada", style='List Bullet')
doc.add_paragraph("Si se acreditaren acciones y/u omisiones maliciosas y/o negligentes que comprometan la eficiencia y eficacia del servicio o la seguridad y/o bienes dispuestos por el Hospital para la correcta ejecución del convenio.", style='List Bullet')
doc.add_paragraph("No cumplir con lo señalado en los requerimientos técnicos.", style='List Bullet')
doc.add_paragraph("Vulnerar normas referidas al uso de información reservada o confidencial al que se tenga acceso en razón del servicio especializado contratado.", style='List Bullet')
doc.add_paragraph("No cumplir con los horarios establecidos en las bases", style='List Bullet')
doc.add_paragraph("No cumplir con los tiempos de respuestas establecidos en bases de licitación según sea el caso.", style='List Bullet')
doc.add_paragraph("Incumplimiento de las normas y procedimientos internos vigentes, tanto técnicas como administrativas impartidas por el Hospital, a través del administrador interno del contrato", style='List Bullet')
doc.add_paragraph("cualquier falta que atente, directa o indirectamente, contra la integridad física de los pacientes o funcionarios, que implique u obstruye la atención, calidad del servicio y/o según lo establecido en las bases.", style='List Bullet')
# Optional: Add space after the list

# --- Concluding paragraphs for the Multas section ---
doc.add_paragraph("Las referidas multas, en total, no podrán sobrepasar el 20% del valor total neto del contrato. Igualmente, no se le podrán cursar más de 6 multas totalmente tramitadas en un período de 6 meses consecutivos. En ambos casos, superado cada límite, se configurará una causal de término anticipado del contrato.")
doc.add_paragraph("Las multas deberán ser pagadas en el plazo máximo de 5 días hábiles contados desde la notificación de la resolución que aplica la multa.")
doc.add_paragraph("Cuando el cálculo del monto de la respectiva multa, convertido a pesos chilenos, resulte un número con decimales, éste se redondeará al número entero más cercano. La fecha de conversión de la UTM será la del día de emisión del respectivo acto administrativo que origina el cobro de la multa")
doc.add_paragraph("Las multas se aplicarán sin perjuicio del derecho de la entidad licitante de recurrir ante los Tribunales Ordinarios de Justicia ubicados en la ciudad de Melipilla, a fin de hacer efectiva la responsabilidad del contratante incumplidor.")
doc.add_paragraph("No procederá el cobro de las multas señaladas en este punto, si el incumplimiento se debe a un caso fortuito o fuerza mayor, de acuerdo con los artículos 45 y 1547 del Código Civil o una causa enteramente ajena a la voluntad de las partes, el cual será calificado como tal por la Entidad Licitante, en base al estudio de los antecedentes por los cuales el oferente adjudicado acredite el hecho que le impide cumplir.")


doc.add_heading("Cobro de la Garantía de Fiel Cumplimiento de Contrato", level = 3)
doc.add_paragraph("Al Adjudicatario le podrá ser aplicada la medida de cobro de la Garantía por Fiel Cumplimiento del Contrato por la entidad licitante, en los siguientes casos:")

doc.add_paragraph("Al Adjudicatario le podrá ser aplicada la medida de cobro de la Garantía por Fiel Cumplimiento del Contrato por la entidad licitante, en los siguientes casos:")
doc.add_paragraph("No pago de multas dentro de los plazos establecidos en las presentes bases y/o el respectivo contrato.", style='List Bullet')
doc.add_paragraph("Incumplimientos de las exigencias técnicas de los bienes y servicios (en caso de que hayan sido requeridos) adjudicados establecidos en el Contrato.", style = "List Bullet")
doc.add_paragraph("Cualquiera de las causales señaladas en el N°10.6.3 sobre “Término Anticipado del Contrato”, a excepción del numeral 3) y numeral 16), en todas estas causales señaladas, se procederá al cobro de la garantía de fiel cumplimiento del contrato, si se hubiere exigido dicha caución en las Bases.", style = "List Bullet")

doc.add_heading("Término anticipado del contrato", level = 3)

termino_contrato_p1 = doc.add_paragraph(
    "El hospital está facultado para declarar administrativamente mediante resolución fundada el término anticipado "
    "del contrato, en cualquier momento, sin derecho a indemnización alguna para el adjudicado, si concurre alguna "
    "de las causales que se señalan a continuación:"
)

# --- Start of the numbered list ---

# Item 1 (with sub-items A, B, C, D)
termino_contrato_p2 = doc.add_paragraph(
    "Por incumplimiento grave de las obligaciones contraídas por el proveedor adjudicado, cuando sea imputable a éste. "
    "Se entenderá por incumplimiento grave la no ejecución o la ejecución parcial por parte del adjudicatario de las "
    "obligaciones contractuales, descritas en las presentes Bases, sin que exista alguna causal que le exima de "
    "responsabilidad, y cuando dicho incumplimiento le genere al hospital un perjuicio en el cumplimiento de sus funciones. "
    "Alguno de estos motivos puede ser:",
    style='List Number'
)

# Sub-items for 1 (Applying 'List Number' style handles nesting to A, B, C...)
termino_contrato_p3 = doc.add_paragraph(
    "La aplicación de dos o más Multas Graves en un periodo de seis meses móviles.",
    style='List Number'
)
termino_contrato_p4 = doc.add_paragraph(
    "Si el proveedor fuese condenado a algún delito que tuviera pena aflictiva o tratándose de una empresa, sus socios, "
    "o en el caso de una sociedad anónima, algunos de los miembros del directorio o el gerente de la sociedad.",
    style='List Number'
)
termino_contrato_p5 = doc.add_paragraph(
    "Si el proveedor delega, cede, aporta o transfiere el presente convenio a cualquier título efectúa asociaciones u otorga "
    "concesiones o subconcesiones.",
    style='List Number'
)
termino_contrato_p6 = doc.add_paragraph(
    "Si la sociedad se disolviere por Quiebra o cesación de pagos del proveedor.",
    style='List Number'
)

# Item 2
termino_contrato_p7 = doc.add_paragraph(
    "Si el adjudicado se encuentra en estado de notoria insolvencia o fuere declarado deudor en un procedimiento concursal "
    "de liquidación. En el caso de una UTP, aplica para cualquiera de sus integrantes. En este caso no procederá el "
    "término anticipado si se mejoran las cauciones entregadas o las existentes sean suficientes para garantizar el "
    "cumplimiento del contrato.",
    style='List Number'
)

# Item 3
termino_contrato_p8 = doc.add_paragraph(
    "Por exigirlo la necesidad del servicio, el interés público o la seguridad nacional.",
    style='List Number'
)

# Item 4
termino_contrato_p9 = doc.add_paragraph(
    "Registrar, a la mitad del período de ejecución contractual, con un máximo de seis meses, saldos insolutos de "
    "remuneraciones o cotizaciones de seguridad social con sus actuales trabajadores o con trabajadores contratados "
    "en los últimos 2 años.",
    style='List Number'
)

# Item 5
termino_contrato_p10 = doc.add_paragraph(
    "Si se disuelve la sociedad o empresa adjudicada, o en caso de fallecimiento del contratante, si se trata de una "
    "persona natural.",
    style='List Number'
)

# Item 6
termino_contrato_p11 = doc.add_paragraph(
    "Incumplimiento de uno o más de los compromisos asumidos por los adjudicatarios, en virtud del “Pacto de integridad\" "
    "contenido en estas bases. Cabe señalar que en el caso que los antecedentes den cuenta de una posible afectación a la "
    "libre competencia, el organismo licitante pondrá dichos antecedentes en conocimiento de la Fiscalía Nacional Económica.",
    style='List Number'
)

# Item 7 (with sub-items a, b, c)
termino_contrato_p12 = doc.add_paragraph(
    "Sin perjuicio de lo señalado en el “Pacto de integridad”, si el adjudicatario, sus representantes, o el personal "
    "dependiente de aquél, no observaren el más alto estándar ético exigible, durante la ejecución de la licitación, y "
    "propiciaren prácticas corruptas, tales como:",
    style='List Number'
)

# Sub-items for 7
termino_contrato_p13 = doc.add_paragraph(
    "Dar u ofrecer obsequios, regalías u ofertas especiales al personal del hospital, que pudiere implicar un conflicto "
    "de intereses, presente o futuro, entre el respectivo adjudicatario y el servicio hospitalario.",
    style='List Number'
)
termino_contrato_p14 = doc.add_paragraph(
    "Dar u ofrecer cualquier cosa de valor con el fin de influenciar la actuación de un funcionario público durante la "
    "relación contractual objeto de la presente licitación.",
    style='List Number'
)
termino_contrato_p15 = doc.add_paragraph(
    "Tergiversar hechos, con el fin de influenciar decisiones de la entidad licitante.",
    style='List Number'
)

# Item 8
termino_contrato_p16 = doc.add_paragraph(
    "No renovación oportuna de la Garantía de Fiel Cumplimiento, según lo establecido en la cláusula 8.2 de las bases "
    "de licitación cuando aplique.",
    style='List Number'
)

# Item 9
termino_contrato_p17 = doc.add_paragraph(
    "La comprobación de la falta de idoneidad, de fidelidad o de completitud de los antecedentes aportados por el "
    "proveedor adjudicado, para efecto de ser adjudicado o contratado.",
    style='List Number'
)

# Item 10
termino_contrato_p18 = doc.add_paragraph(
    "La comprobación de que el adjudicatario, al momento de presentar su oferta contaba con información o antecedentes "
    "relacionados con el proceso de diseño de las bases, encontrándose a consecuencia de ello en una posición de privilegio "
    "en relación al resto de los oferentes, ya sea que dicha información hubiese sido conocida por el proveedor en razón "
    "de un vínculo laboral o profesional entre éste y las entidades compradoras, o bien, como resultado de prácticas "
    "contrarias al ordenamiento jurídico.",
    style='List Number'
)

# Item 11 (with sub-items a, b, c, d, e)
termino_contrato_p19 = doc.add_paragraph(
    "En caso de ser el adjudicatario de una Unión Temporal de Proveedores (UTP):",
    style='List Number'
)

# Sub-items for 11
termino_contrato_p20 = doc.add_paragraph(
    "Inhabilidad sobreviniente de uno de los integrantes de la UTP en el Registro de Proveedores, que signifique que la "
    "UTP no pueda continuar ejecutando el contrato con los restantes miembros en los mismos términos adjudicados.",
    style='List Number'
)
termino_contrato_p21 = doc.add_paragraph(
    "De constatarse que los integrantes de la UTP constituyeron dicha figura con el objeto de vulnerar la libre competencia. "
    "En este caso, deberán remitirse los antecedentes pertinentes a la Fiscalía Nacional Económica.",
    style='List Number'
)
termino_contrato_p22 = doc.add_paragraph(
    "Retiro de algún integrante de la UTP que hubiere reunido una o más características objeto de la evaluación de la oferta.",
    style='List Number'
)
termino_contrato_p23 = doc.add_paragraph(
    "Cuando el número de integrantes de una UTP sea inferior a dos y dicha circunstancia ocurre durante la ejecución del contrato.",
    style='List Number'
)
termino_contrato_p24 = doc.add_paragraph(
    "Disolución de la UTP.",
    style='List Number'
)

# Item 12
termino_contrato_p25 = doc.add_paragraph(
    "En caso de infracción de lo dispuesto en la cláusula sobre “Cesión de contrato y Subcontratación”",
    style='List Number'
)

# Item 13
termino_contrato_p26 = doc.add_paragraph(
    "En caso de que las multas cursadas, en total, sobrepasen el 20 % del valor total contratado con impuestos incluidos o "
    "se apliquen más de 6 multas totalmente tramitadas en un periodo de 6 meses consecutivos.",
    style='List Number'
)

# Item 14
termino_contrato_p27 = doc.add_paragraph(
    "Por el no pago de las multas aplicadas.",
    style='List Number'
)

# Item 15
termino_contrato_p28 = doc.add_paragraph(
    "Por la aplicación de dos multas graves en que incurra el adjudicatario en virtud del incumplimiento de las obligaciones "
    "reguladas en las bases y del presente contrato.",
    style='List Number'
)

# Item 16
termino_contrato_p29 = doc.add_paragraph(
    "Si el Hospital San José de Melipilla cesara su funcionamiento en lugar de origen por cambio de ubicación de sus dependencias.",
    style='List Number'
)

# Item 17
termino_contrato_p30 = doc.add_paragraph(
    "Por la comprobación de la inhabilidad del adjudicatario para contratar con la Administración del Estado en portal de "
    "mercado público, durante la ejecución del presente contrato. Solo en el caso que el proveedor desde la notificación "
    "de esta situación no regularice su registro en un plazo superior a 15 días hábiles.",
    style='List Number'
)

# Item 18
termino_contrato_p31 = doc.add_paragraph(
    "Por incumplimiento de obligaciones de confidencialidad establecidas en las respectivas Bases.",
    style='List Number'
)

# --- End of the numbered list ---

# Add concluding paragraphs (these are not part of the list)
termino_contrato_p32 = doc.add_paragraph(
    "De concurrir cualquiera de las causales anteriormente señaladas como término anticipado del contrato, exceptuando las "
    "causales número 3 y número 16, se procederá al cobro de la garantía de fiel cumplimiento del contrato, siempre y "
    "cuando se hubiere exigido dicha caución en las Bases."
)
termino_contrato_p33 = doc.add_paragraph(
    "El término anticipado por incumplimientos se aplicará siguiendo el procedimiento establecido en la cláusula "
    "“sobre aplicación de Medidas derivadas de incumplimientos.”"
)

doc.add_heading("Resciliación de Mutuo Acuerdo")
doc.add_paragraph("Sin perjuicio de lo anterior, la entidad licitante y el respectivo adjudicatario podrán poner término al contrato en cualquier momento, de común acuerdo, sin constituir una medida por incumplimiento.")

doc.add_heading("Procedimiento para Aplicación de Medidas derivadas de incumplimientos")
doc.add_paragraph("Detectada una situación que amerite la aplicación de una multa u otra medida derivada de incumplimientos contemplada en las presentes bases, o que constituya una causal de término anticipado, con excepción de la resciliación, el referente técnico o administrador del contrato notificará de ello al oferente adjudicado, informándole sobre la medida a aplicar y sobre los hechos que la fundamentan.")
doc.add_paragraph("A contar de la notificación singularizada en el párrafo anterior, el proveedor adjudicado tendrá un plazo de 5 días hábiles para efectuar sus descargos por escrito, acompañando todos los antecedentes que lo fundamenten. Vencido el plazo indicado sin que se hayan presentado descargos, la Dirección del Hospital resolverá según la naturaleza de la infracción, notificando al proveedor la resolución del caso por parte del Hospital.")
doc.add_paragraph("Si el proveedor adjudicado ha presentado sus descargos dentro del plazo establecido para estos efectos, el Hospital tendrá un plazo de 30 días hábiles, contados desde la recepción de los descargos del proveedor, para rechazarlos o acogerlos, total o parcialmente. Al respecto, el rechazo total o parcial de los descargos del respectivo proveedor deberá formalizarse a través de la dictación de una resolución fundada del hospital, en la cual deberá detallarse el contenido y las características de la medida. La indicada resolución será notificada al proveedor adjudicado.")
doc.add_paragraph("Con todo, el adjudicatario solo será responsable por hechos imputables a su incumplimiento directo y no por indisponibilidades de servicio ocasionadas por fallas ajenas a su gestión y control, lo que deberá, en todo caso, acreditarse debidamente. Sin perjuicio de lo anterior, el adjudicatario deberá adoptar medidas que ofrezcan continuidad operativa a los servicios materia de la respectiva licitación.")
doc.add_paragraph("Una vez finalizados los trámites administrativos señalados precedentemente y para el evento de que esta conlleve la aplicación de una multa o sanción, el Hospital San José de Melipilla podrá realizar el cobro de la multa o sanción que será debidamente notificado junto con el acto administrativo que lo autoriza. El monto de las multas podría ser rebajado del pago, que el Hospital deba efectuar al proveedor, en el estado de pago más próximo a la notificación del acto administrativo, pudiéndose aplicar tanto en la emisión de la orden de compra, como también en la aplicación del descuento en el pago de facturas. De no ser suficiente este monto o en caso de no existir pagos pendientes, el proveedor deberá pagar directamente al Hospital San José de Melipilla, el monto indicado en el acto administrativo previamente notificado, este pago no podrá ser superior a los 5 días hábiles desde su notificación. Si el proveedor no paga dentro de dicho plazo, se hará efectivo el cobro de la garantía de fiel cumplimiento del contrato, debiendo reponer una nueva boleta de garantía por un monto igual al original, en un plazo no superior a 5 días hábiles en caso que aplique la solicitud de dicha caución.")
doc.add_paragraph("En el caso de no reponer la boleta de garantía, el hospital podrá proceder a tramitar el termino anticipado del contrato en aquellos casos que aplique con la solicitud de dicha caución.")
doc.add_paragraph("El valor de la UTM a considerar será el equivalente a su valor en pesos del mes en el cual se aplicó la multa.")

doc.add_heading("Emisión de la Orden de Compra", level = 3)
doc.add_paragraph("Las órdenes de compra se emitirán previa solicitud del administrador del contrato, quien, en función de la necesidad y demanda del servicio, realizara los pedidos correspondientes.")
doc.add_paragraph("La orden de compra sólo se emitirá en los casos que el proveedor este en estado hábil para ser contratado por el Estado de Chile y sólo se emitirá el documento a nombre del proveedor adjudicado por el Hospital.")
doc.add_paragraph("Al inicio del convenio, por registros en la plataforma y tramites del “gestor de contratos” se emitirá una orden de compras por un monto mínimo, la que solo debe ser aceptada por el proveedor, sin tramitar dicho servicio. Todo cambio respecto a este punto, será informado con la respectiva anticipación.")

doc.add_heading("Del Pago")
doc.add_paragraph("El pago se efectuará una vez que el “Hospital” haya recibido oportunamente y a su entera satisfacción dichos bienes o servicios y desde la recepción conforme de la factura u otro instrumento de cobro.")
doc.add_paragraph("El pago será efectuado dentro de los 30 días corridos siguientes, contados desde la recepción de la factura respectiva, salvo las excepciones indicadas en el artículo 79 bis del Reglamento de la Ley N° 19.886.")
doc.add_paragraph("El proveedor solo podrá facturar los bienes o servicios efectivamente entregados y recibidos conforme por este organismo comprador, una vez que el administrador del contrato por parte del organismo comprador autorice la facturación en virtud de la recepción conforme de los bienes o servicios. “El Hospital” rechazará todas las facturas que hayan sido emitidas sin contar con la recepción conforme de los bienes o servicios y la autorización expresa de facturar por parte de éste.")
doc.add_paragraph("Para efectos del pago, el proveedor adjudicado deberá indicar en la factura el número de orden de compra, además, no podrá superar el monto de la orden de compra, de lo contrario, se cancelará la factura por “forma”.")
del_pago_correo = doc.add_paragraph()
del_pago_correo.add_run("La factura electrónica deberá ser enviada al correo: ")
del_pago_correo.add_run("facturas.hjsm@redsalud.gov.cl").bold = True
del_pago_correo.add_run(" con copia al correo ")
del_pago_correo.add_run("dipresrecepcion@custodium.com").bold = True
del_pago_correo.add_run("(En formato PDF y XML)")

doc.add_paragraph("El valor del convenio se reajustará anualmente de acuerdo con la variación que haya experimentado el Índice de Precios al Consumidor IPC, obtenido del promedio de la sumatoria de los IPC de los doce meses inmediatamente anteriores al mes en que se efectúa su cálculo. Este reajuste es de exclusiva responsabilidad de la empresa adjudicada; si por alguna razón no lo aplicare, no se permitirá su cobro en forma retroactiva. Su precio se pagará conforme a lo establecido.")
doc.add_paragraph("El valor del convenio se reajustará anualmente de acuerdo con la variación que haya experimentado el Índice de Precios al Consumidor IPC, obtenido del promedio de la sumatoria de los IPC de los doce meses inmediatamente anteriores al mes en que se efectúa su cálculo. Este reajuste es de exclusiva responsabilidad de la empresa adjudicada; si por alguna razón no lo aplicare, no se permitirá su cobro en forma retroactiva. Su precio se pagará conforme a lo establecido.")
doc.add_paragraph("En ningún caso procederán cobros adicionales por bienes o servicios no convenidos previamente, ni por tiempos en que el proveedor no preste los servicios. ")
doc.add_paragraph("Cabe señalar que, cuando el resultado del monto a facturar resulte un número con decimales, éste se redondeará al número entero siguiente en caso de que la primera cifra decimal sea igual o superior a 5. En caso contrario el monto deberá ser redondeado al número entero anterior.")

doc.add_heading("Vigencia del Contrato")
doc.add_paragraph("El contrato tendrá una duración de treinta y seis (36) meses contados desde la total tramitación del acto administrativo que aprueba la adjudicación o hasta que se cumpla con el monto estipulado en las presentes bases, lo que suceda primero y sin perjuicio, que por razones de buen servicio las prestaciones materia de la licitación podrían iniciarse desde el momento de la suscripción del mismo, sin que proceda pago alguno en el tiempo intermedio.")

doc.add_heading("Administrador del Contrato y/o Referente Técnico.")
doc.add_paragraph("Con el objeto de supervisar y verificar el cumplimiento materia de la presente licitación, El Hospital designará a (la) Enfermera Supervisora(o) del Servicio de Pabellón y al Jefe(a) de Farmacia o su subrogante, para coordinar y fiscalizar la efectiva ejecución del contrato en términos administrativos.")

administrado_contrato = doc.add_paragraph()
administrado_contrato.add_run("El adjudicatario").bold = True
administrado_contrato.add_run("deberá nombrar un coordinador del contrato, cuya identidad deberá ser informada al Hospital.")
doc.add_paragraph("En el desempeño de su cometido, el coordinador del contrato deberá, a lo menos:")


# doc.add_

"""1. Informar oportunamente al órgano comprador de todo hecho relevante que pueda afectar el cumplimiento del contrato.
2. Representar al proveedor en la discusión de las materias relacionadas con la ejecución del contrato.
3. Coordinar las acciones que sean pertinentes para la operación y cumplimiento de este contrato.
La designación del coordinador y todo cambio posterior deberá ser informado por el adjudicatario al responsable de administrar el contrato y/o referente técnico por parte del órgano comprador, a más tardar dentro de las 24 horas siguientes de efectuada la designación o el cambio, por medio del correo electrónico institucional del funcionario.
"""
# Numeracion
# Numeracion para la lista del administrador de contrato
administrado_contrato_id_lista = crear_numeracion(doc)

# Añadir los elementos de la lista numerada usando el ID correcto y sin estilo de párrafo
administrado_contrato_p1 = doc.add_paragraph("Informar oportunamente al órgano comprador de todo hecho relevante que pueda afectar el cumplimiento del contrato.")
aplicar_numeracion(administrado_contrato_p1, administrado_contrato_id_lista)

administrado_contrato_p2 = doc.add_paragraph("Representar al proveedor en la discusión de las materias relacionadas con la ejecución del contrato.")
aplicar_numeracion(administrado_contrato_p2, administrado_contrato_id_lista)

administrado_contrato_p3 = doc.add_paragraph("Coordinar las acciones que sean pertinentes para la operación y cumplimiento de este contrato.")
aplicar_numeracion(administrado_contrato_p3, administrado_contrato_id_lista)

# Añadir el párrafo después de la lista como un párrafo normal
doc.add_paragraph("La designación del coordinador y todo cambio posterior deberá ser informado por el adjudicatario al responsable de administrar el contrato y/o referente técnico por parte del órgano comprador, a más tardar dentro de las 24 horas siguientes de efectuada la designación o el cambio, por medio del correo electrónico institucional del funcionario.")

# Guardar el documento (ya existente al final de tu script)
# doc_path = 'resolucion_numerada.docx'
# doc.save(doc_path)
# print(f"Documento guardado como: {doc_path}")
doc.add_heading("Pacto de Integridad", level = 2)
doc.add_paragraph("El oferente declara que, por el sólo hecho de participar en la presente licitación, acepta expresamente el presente pacto de integridad, obligándose a cumplir con todas y cada una de las estipulaciones contenidas en el mismo, sin perjuicio de las que se señalen en el resto de las bases de licitación y demás documentos integrantes. Especialmente, el oferente acepta el suministrar toda la información y documentación que sea considerada necesaria y exigida de acuerdo con las presentes bases de licitación, asumiendo expresamente los siguientes compromisos:")

pacto_integridad_id = crear_numeracion(doc)
pacto_integridad_p1 = doc.add_paragraph("El oferente se compromete a respetar los derechos fundamentales de sus trabajadores, entendiéndose por éstos los consagrados en la Constitución Política de la República en su artículo 19, números 1º, 4º, 5º, 6º, 12º, y 16º, en conformidad al artículo 485 del Código del Trabajo. Asimismo, el oferente se compromete a respetar los derechos humanos, lo que significa que debe evitar dar lugar o contribuir a efectos adversos en los derechos humanos mediante sus actividades, bienes o servicios, y subsanar esos efectos cuando se produzcan, de acuerdo con los Principios Rectores de Derechos Humanos y Empresas de Naciones Unidas.")
aplicar_numeracion(pacto_integridad_p1, pacto_integridad_id)

pacto_integridad_p2 = doc.add_paragraph("El oferente se obliga a no ofrecer ni conceder, ni intentar ofrecer o conceder, sobornos, regalos, premios, dádivas o pagos, cualquiera fuese su tipo, naturaleza y/o monto, a ningún funcionario público en relación con su oferta, con el proceso de licitación pública, ni con la ejecución de el o los contratos que eventualmente se deriven de la misma, ni tampoco a ofrecerlas o concederlas a terceras personas que pudiesen influir directa o indirectamente en el proceso licitatorio, en su toma de decisiones o en la posterior adjudicación y ejecución del o los contratos que de ello se deriven.")
aplicar_numeracion(pacto_integridad_p2, pacto_integridad_id)

# Numeracion para la lista del Pacto de Integridad (ya creada antes)
# pacto_integridad_id = crear_numeracion(doc) # No es necesario volver a crearla

# Añadir los elementos de la lista numerada (continuación)

# Item iii
pacto_integridad_p3 = doc.add_paragraph("El oferente se obliga a no intentar ni efectuar acuerdos o realizar negociaciones, actos o conductas que tengan por objeto influir o afectar de cualquier forma la libre competencia, cualquiera fuese la conducta o acto específico, y especialmente, aquellos acuerdos, negociaciones, actos o conductas de tipo o naturaleza colusiva, en cualquiera de sus tipos o formas.")
aplicar_numeracion(pacto_integridad_p3, pacto_integridad_id)

# Item iv
pacto_integridad_p4 = doc.add_paragraph("El oferente se obliga a revisar y verificar toda la información y documentación, que deba presentar para efectos del presente proceso licitatorio, tomando todas las medidas que sean necesarias para asegurar su veracidad, integridad, legalidad, consistencia, precisión y vigencia.")
aplicar_numeracion(pacto_integridad_p4, pacto_integridad_id)

# Item v
pacto_integridad_p5 = doc.add_paragraph("El oferente se obliga a ajustar su actuar y cumplir con los principios de legalidad, probidad y transparencia en el presente proceso licitatorio.")
aplicar_numeracion(pacto_integridad_p5, pacto_integridad_id)

# Item vi
pacto_integridad_p6 = doc.add_paragraph("El oferente manifiesta, garantiza y acepta que conoce y respetará las reglas y condiciones establecidas en las bases de licitación, sus documentos integrantes y él o los contratos que de ellos se derivase.")
aplicar_numeracion(pacto_integridad_p6, pacto_integridad_id)

# Item vii
pacto_integridad_p7 = doc.add_paragraph("El oferente reconoce y declara que la oferta presentada en el proceso licitatorio es una propuesta seria, con información fidedigna y en términos técnicos y económicos ajustados a la realidad, que aseguren la posibilidad de cumplir con la misma en las condiciones y oportunidad ofertadas.")
aplicar_numeracion(pacto_integridad_p7, pacto_integridad_id)

# Item viii
pacto_integridad_p8 = doc.add_paragraph("El oferente se obliga a tomar todas las medidas que fuesen necesarias para que las obligaciones anteriormente señaladas sean asumidas y cabalmente cumplidas por sus empleados, dependientes, asesores y/o agentes y, en general, todas las personas con que éste o éstos se relacionen directa o indirectamente en virtud o como efecto de la presente licitación, incluidos sus subcontratistas, haciéndose plenamente responsable de las consecuencias de su infracción, sin perjuicio de las responsabilidades individuales que también procediesen y/o fuesen determinadas por los organismos correspondientes.")
aplicar_numeracion(pacto_integridad_p8, pacto_integridad_id)

# El resto del código continúa aquí...

doc.add_heading("Comportamiento ético del Adjudicatario.", level = 3)
doc.add_paragraph("El adjudicatario que preste los servicios deberá observar, durante toda la época de ejecución del contrato, el más alto estándar ético exigible a los funcionarios públicos. Tales estándares de probidad deben entenderse equiparados a aquellos exigidos a los funcionarios de la Administración Pública, en conformidad con el Título III de la ley N° 18.575, Orgánica Constitucional de Bases Generales de la Administración del Estado.")

doc.add_heading("Auditorías")
doc.add_paragraph("El adjudicatario podrá ser sometido a auditorías externas, contratadas por la entidad licitante a empresas auditoras independientes, con la finalidad de velar por el cumplimiento de las obligaciones contractuales y de las medidas de seguridad comprometidas por el adjudicatario en su oferta. Si el resultado de estas auditorías evidencia incumplimientos contractuales por parte del adjudicatario, el proveedor quedará sujeto a las medidas que corresponda aplicar la entidad licitante, según las presentes bases.")

doc.add_heading("Confidencialidad")
doc.add_paragraph("El adjudicatario no podrá utilizar para ninguna finalidad ajena a la ejecución del contrato, la documentación, los antecedentes y, en general, cualquier información, que haya conocido o a la que haya accedido, en virtud de cualquier actividad relacionada con el contrato.")
doc.add_paragraph("El adjudicatario, así como su personal dependiente que se haya vinculado a la ejecución del contrato, en cualquiera de sus etapas, deben guardar confidencialidad sobre los antecedentes relacionados con el proceso licitatorio y el respectivo contrato.")
doc.add_paragraph("El adjudicatario debe adoptar medidas para el resguardo de la confidencialidad de la información, reservándose el órgano comprador el derecho de ejercer las acciones legales que correspondan, de acuerdo con las normas legales vigentes, en caso de divulgación no autorizada, por cualquier medio, de la totalidad o parte de la información referida.")
doc.add_paragraph("La divulgación, por cualquier medio, de la totalidad o parte de la información referida en los párrafos anteriores, por parte del proveedor, durante la vigencia del contrato o dentro de los 5 años siguientes después de finalizado éste, podrá dar pie a que la Entidad entable en su contra las acciones judiciales que correspondan. Con todo, tratándose de bases de datos de carácter personal, la obligación de confidencialidad dura indefinidamente, de acuerdo con la Ley N°19.628, sobre Protección de la Vida Privada.")

doc.add_heading("Propiedad de la información")
doc.add_paragraph("La entidad licitante será la titular de todos los datos de transacciones, bitácoras (logs), parámetros, documentos electrónicos y archivos adjuntos y, en general, de las bases de datos y de toda información contenida en la infraestructura física y tecnológica que le suministre el proveedor contratado y que se genere en virtud de la ejecución de los servicios objeto de la presente licitación. El proveedor no podrá utilizar la información indicada en el párrafo anterior, durante la ejecución del contrato ni con posterioridad al término de su vigencia, sin autorización escrita de la entidad licitante. Por tal motivo, una vez que el proveedor entregue dicha información a la entidad o al finalizar la relación contractual, deberá borrarla de sus registros lógicos y físicos.")

doc.add_heading("Saldos insolutos de remuneraciones o cotizaciones de seguridad social.")
doc.add_paragraph("Durante la vigencia del respectivo contrato el adjudicatario deberá acreditar que no registra saldos insolutos de obligaciones laborales y sociales con sus actuales trabajadores o con trabajadores contratados en los últimos dos años.")
doc.add_paragraph("El órgano comprador podrá requerir al adjudicatario, en cualquier momento, los antecedentes que estime necesarios para acreditar el cumplimiento de las obligaciones laborales y sociales antes señaladas.")
doc.add_paragraph("En caso de que la empresa adjudicada registre saldos insolutos de remuneraciones o cotizaciones de seguridad social con sus actuales trabajadores o con trabajadores contratados en los últimos dos años, los primeros estados de pago de los bienes y servicios de esta licitación deberán ser destinados al pago de dichas obligaciones, debiendo la empresa acreditar que la totalidad de las obligaciones se encuentran liquidadas al cumplirse la mitad del período de ejecución de las prestaciones, con un máximo de seis meses.")
doc.add_paragraph("La entidad licitante deberá exigir que la empresa adjudicada proceda a dichos pagos y le presente los comprobantes y planillas que demuestren el total cumplimiento de la obligación. El incumplimiento de estas obligaciones por parte de la empresa adjudicataria dará derecho a terminar la relación contractual, pudiendo llamarse a una nueva licitación en la que la empresa referida no podrá participar.")

doc.add_heading("Normas Laborales Aplicables")
doc.add_paragraph("El adjudicatario, en su calidad de empleador, será responsable exclusivo del cumplimiento íntegro y oportuno de las normas del Código del Trabajo y leyes complementarias, leyes sociales, de previsión, de seguros, de enfermedades profesionales, de accidentes del trabajo y demás pertinentes respecto de sus trabajadores y/o integrantes de sus respectivos equipos de trabajo.")
doc.add_paragraph("En consecuencia, el adjudicatario será responsable, en forma exclusiva, y sin que la enumeración sea taxativa, del pago oportuno de las remuneraciones, honorarios, indemnizaciones, desahucios, gratificaciones, gastos de movilización, beneficios y, en general, de toda suma de dinero que, por cualquier concepto, deba pagarse a sus trabajadores y/o integrantes de sus respectivos equipos de trabajo.")
doc.add_paragraph("El Hospital se reserva el derecho a exigir al contratista, a simple requerimiento de la contraparte técnica, y sin perjuicio de lo dispuesto en el artículo 4° de la Ley de Compras y el artículo 183-C del Código del Trabajo, un certificado que acredite el monto y estado de cumplimiento de las obligaciones laborales y previsionales emitido por la Inspección del Trabajo respectiva, o bien, por medios idóneos que garanticen la veracidad de dicho monto y estado de cumplimiento, respecto de sus trabajadores. Ello, con el propósito de hacer efectivo por parte del órgano comprador, su derecho a ser informado y el derecho de retención, consagrados en los incisos segundo y tercero del artículo 183-C del Código del Trabajo, en el marco de la responsabilidad subsidiaria derivada de dichas obligaciones laborales y previsionales, a la que alude el artículo 183-D del mismo Código.")
doc.add_paragraph("Por otra parte, se deja expresa constancia que la suscripción del contrato respectivo no significará en caso alguno que el adjudicatario, sus trabajadores o integrantes de los equipos presentados por éstos, adquieran la calidad de funcionarios públicos, no existiendo vínculo alguno de subordinación o dependencia de ellos con el órgano comprador.")

doc.add_heading("Cambio de personal del proveedor adjudicado.")
doc.add_paragraph("El Hospital San José de Melipilla podrá, por razones de buen servicio, solicitar el cambio de trabajadores, expresando la causa del derecho a cambiar al personal del proveedor, entendiéndose como el derecho a prohibir unilateralmente la continuidad de funciones de un trabajador que implique un potencial riesgo a los pacientes, funcionarios, bienes e imagen de la organización.")
doc.add_paragraph("El Proveedor adjudicado deberá reemplazar al personal, dentro del plazo que se le indique. La decisión del Hospital San José de Melipilla se comunicará por escrito al Proveedor precisando las causas que motivan la solicitud, con a lo menos 5 días de anticipación a la fecha en que se solicita deje de prestar servicios en sus dependencias, el trabajador que se indique.")

doc.add_heading("Cesión y subcontratación.")
doc.add_paragraph("El proveedor adjudicado no podrá ceder ni transferir en forma alguna, total ni parcialmente, los derechos y obligaciones que nacen del desarrollo de esta licitación, y, en especial, los establecidos en los respectivos contratos que se celebren con los órganos públicos mandantes. ")
doc.add_paragraph("La infracción de esta prohibición será causal inmediata de término del contrato, sin perjuicio de las acciones legales que procedan ante esta situación. ")
doc.add_paragraph("Durante la ejecución del contrato, y previa autorización por escrito del Hospital, el adjudicatario sólo podrá efectuar aquellas subcontrataciones que sean indispensables para la realización de tareas específicas, todo lo cual será calificado por el coordinador del contrato. En todo caso, el adjudicatario seguirá siendo el único responsable de las obligaciones contraídas en virtud del respectivo contrato suscrito con el Hospital. ")
doc.add_paragraph("Así mismo, el subcontratista debe encontrarse hábil en el registro de Proveedores del Estado y tratándose de servicios, acreditar el cumplimiento de obligaciones laborales, conforme lo establece el artículo 4° inciso 2° de la Ley N°19.886. ")
doc.add_paragraph("En todos los casos es el oferente y eventual adjudicatario el único responsable del pleno cumplimiento de lo señalado en estas bases (Art. N° 76, Reglamento de la Ley N° 19.886).")

doc.add_heading("Discrepancias")
doc.add_paragraph("Si con motivo de la ejecución del contrato se presentaran denuncias, querellas o demandas ante el Ministerio Público o los Tribunales Ordinarios de Justicia; o reclamos ante el Consejo de Defensa del Estado por el cuestionamiento en la prestación otorgada y que corresponda al objeto del contrato celebrado, será el proveedor el único responsable por tales actos, por lo que, sí el Hospital fuese condenado a pagar una multa o indemnización, en razón de los actos precedentemente enunciados o el Hospital tuviera que pagar alguna transacción judicial o extrajudicial que deba celebrarse en razón de las situaciones antes enunciadas, el proveedor deberá reembolsar al Hospital el total del monto resultante de un fallo ejecutoriado o de una transacción judicial o extrajudicial o de un procedimiento de medición de acuerdo a la Ley Nº 19.966.")
doc.add_paragraph("Asimismo, serán responsables de todos los daños, pérdidas, deterioros o perjuicios de bienes muebles e inmuebles del Hospital, producto del mal uso ocasionado en virtud de la prestación de servicio, debiendo restituir al Hospital los costos en que deba incurrir para reparar los daños producidos por este motivo. Esta obligación se mantendrá aun cuando el presente contrato que al efecto se suscriba se dé por terminado ya sea por expiración del plazo establecido o por decisión del Hospital.")

doc.add_heading("Constancia")
doc.add_paragraph("Se deja expresa constancia que todas y cada una de las cláusulas contenidas en las presentes Bases, Anexos y aclaratorias, se entienden incorporadas sin necesidad de mención expresa en el correspondiente contrato que se materialice con el adjudicado y éste se hace responsable del cumplimiento de las obligaciones de tales documentos, Bases Administrativas y Contrato que se deriven.")


doc_path = 'resolucion_numerada.docx' # Or your desired output file name
doc.save(doc_path)
print(f"Documento guardado como: {doc_path}")