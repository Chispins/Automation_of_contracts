from Bases import configurar_directorio_trabajo # Línea añadida

configurar_directorio_trabajo() # Línea añadida

# Inside your Reading_file.py (or wherever the original functions are)
from docx import Document
from docx.shared import Inches

# Crear un nuevo documento
document = Document()

# Título Principal del Documento (Opcional, nivel 0)
# document.add_heading('RESOLUCIÓN EXENTA Nº [NÚMERO]', level=0)

# --- Sección Principal 1 ---
document.add_heading('PARTE A: BASES ADMINISTRATIVAS', level=1) # Título 1

# --- Subsección 1.1 ---
document.add_heading('7. EVALUACIÓN Y ADJUDICACIÓN DE LAS OFERTAS', level=2) # Título 2
document.add_paragraph(
    'La evaluación de las ofertas se realizará en una etapa, utilizando criterios técnicos, económicos y administrativos.'
)

# --- Subsección 1.2 ---
document.add_heading('8. CONDICIONES CONTRACTUALES', level=2) # Título 2

# --- Sub-Subsección 1.2.1 (usando Título 3) ---
document.add_heading('8.1 Documentos Integrantes del Contrato', level=3) # Título 3
document.add_paragraph(
    'La relación contractual se ceñirá a los siguientes documentos:'
)
# Ahora usamos viñetas para la lista bajo 8.1
document.add_paragraph('Bases de licitación y sus anexos.', style='List Bullet')
document.add_paragraph('Aclaraciones, respuestas y modificaciones.', style='List Bullet')
document.add_paragraph('Oferta.', style='List Bullet')
document.add_paragraph('Orden de compra.', style='List Bullet')


# --- Sub-Subsección 1.2.2 (usando Título 3 y luego viñetas) ---
document.add_heading('8.6 Efectos derivados de Incumplimientos del Proveedor', level=3) # Título 3
document.add_paragraph(
    'Se podrán aplicar diversas medidas ante incumplimientos:'
)

# Aquí usamos viñetas directamente bajo el Título 3 para listar los tipos de efectos
p = document.add_paragraph('Multas', style='List Bullet') # Viñeta Nivel 1

# Viñetas anidadas (Nivel 2) para detallar las multas
# Nota: Los nombres exactos de los estilos de viñeta pueden variar ligeramente ('List Bullet 2', 'List Bullet 3', etc.)
# Si 'List Bullet 2' no funciona, prueba inspeccionando los estilos en un documento Word de ejemplo.
document.add_paragraph('Clasificación de las sanciones (Amonestación, Multa)', style='List Bullet 2')
document.add_paragraph('Tipos de Multa (Leve, Moderada, Grave)', style='List Bullet 2')
document.add_paragraph('Límites y Pago de Multas', style='List Bullet 2')

p = document.add_paragraph('Cobro de la Garantía de Fiel Cumplimiento de Contrato', style='List Bullet') # Viñeta Nivel 1
document.add_paragraph('Se ejecutará por causales específicas como no pago de multas, etc.', style='List Bullet 2')

p = document.add_paragraph('Término Anticipado del Contrato', style='List Bullet') # Viñeta Nivel 1
document.add_paragraph('Incumplimiento grave', style='List Bullet 2')
document.add_paragraph('Insolvencia o procedimiento concursal', style='List Bullet 2')
document.add_paragraph('Necesidad del servicio', style='List Bullet 2')
document.add_paragraph('Saldos insolutos de remuneraciones', style='List Bullet 2')
document.add_paragraph('(Otras causales listadas...)', style='List Bullet 2')

p = document.add_paragraph('Resciliación o término de mutuo acuerdo', style='List Bullet') # Viñeta Nivel 1

# --- Sección Principal 2 ---
document.add_heading('PARTE B: BASES TÉCNICAS', level=1) # Título 1

# --- Subsección 2.1 ---
document.add_heading('10. DISPOSICIONES ESPECÍFICAS DE LA LICITACIÓN', level=2) # Título 2

# --- Sub-Subsección 2.1.1 (usando Título 3 y luego viñetas) ---
document.add_heading('10.2 De los Productos', level=3) # Título 3
document.add_paragraph(
    'La licitación se enfoca en los siguientes productos (cantidades referenciales):'
)
# Usamos viñetas para listar algunos productos
document.add_paragraph('Ítem 1: Recolector 300 ml', style='List Bullet')
document.add_paragraph('Ítem 4: Kit apósito espuma negra LARGE', style='List Bullet')
document.add_paragraph('Ítem 13: Kit apósito abdominal', style='List Bullet')
document.add_paragraph('(etc...)', style='List Bullet')
document.add_paragraph(
    'Nota: Adjuntar ficha técnica en español es obligatorio.'
    , style='Normal' # Asegurarse que este párrafo no tenga estilo de viñeta
)


# Guardar el documento
file_name = 'ejemplo_estructura_licitacion_1.docx'
document.save(file_name)

import docx
from Contrato import extraer_seccion_completa, copiar_seccion_completa # Asegúrate que ambas estén importadas

# Carga el documento original
doc_original = docx.Document(file_name)

# Texto del encabezado a buscar (ajusta si es necesario)
# Nota: En tu código original usas "a)\tEVALUACIÓN...", pero en Reading_file.py es "7. EVALUACIÓN..."
# Asegúrate de usar el texto correcto que está en el DOCX.
texto_encabezado = "7. EVALUACIÓN Y ADJUDICACIÓN DE LAS OFERTAS"

# Extrae la sección del documento original
resultados_Evaluacion = extraer_seccion_completa(doc_original, texto_encabezado)

if resultados_Evaluacion:
    print(f"Sección '{texto_encabezado}' extraída correctamente.")

    # Crea un NUEVO documento para guardar solo la sección
    doc_nuevo = docx.Document()

    # Extrae las partes del resultado
    encabezado, elementos, nivel = resultados_Evaluacion

    # Copia la sección extraída al NUEVO documento
    copiar_seccion_completa(doc_nuevo, encabezado, elementos, nivel)

    # Guarda el NUEVO documento que contiene solo la sección
    archivo_salida_seccion = "seccion_evaluacion_extraida.docx"
    doc_nuevo.save(archivo_salida_seccion)
    print(f"La sección extraída ha sido guardada en '{archivo_salida_seccion}'")

else:
    print(f"No se encontró la sección '{texto_encabezado}'. No se generó archivo de salida.")

# Ya no necesitas guardar el 'doc_original' si solo querías la sección
# doc_original.save("ejemplo_estructura_licitacion_1_paste.docx") # Comentado o eliminado


# Abrir documentos
archivo_original = "ejemplo_estructura_licitacion_1.docx"
original = docx.Document(archivo_original)

title = "PARTE A: BASES ADMINISTRATIVAS"
extraccion_compl = extraer_seccion_completa(original, title)
encabezado, elementos, nivel = extraccion_compl

import os
from docx import Document

file_path = "vacio.docx"  # Or use an absolute path, e.g., "C:\\Users\\Usuario\\PycharmProjects\\Automation_of_contracts\\Files\\vacio.docx"

if os.path.exists(file_path):
    vacio = Document(file_path)
else:
    print(f"Error: File '{file_path}' not found. Please create it or check the path.")
    # Optionally, create a new document if needed
    vacio = Document()  # This creates a new empty document


copiar_seccion_completa(vacio, encabezado, elementos, nivel)


filename = "BASE N°140 VAC.docx"
titulo_interes = "Antecedentes  y Plazos"

doc = docx.Document(filename)
# Extraer la sección completa
resultado = extraer_seccion_completa(doc, titulo_interes)
encabezado, elementos, nivel = resultado
# Crear un nuevo documento
nuevo_doc = docx.Document()
# Copiar la sección extraída al nuevo documento
copiar_seccion_completa(nuevo_doc, encabezado, elementos, nivel)
# Guardar el nuevo documento
nuevo_doc.save("seccion_extraida.docx")










import win32com.client as win32
import pythoncom
import os
from Bases import configurar_directorio_trabajo  # Assuming this is needed

pythoncom.CoInitialize()  # Initialize COM at the start

# Initialize Word application
word_app = win32.Dispatch("Word.Application")
word_app.Visible = False  # Set to True for debugging

configurar_directorio_trabajo()  # Keep this as is


# Function to create numbering (using Word's ListTemplates)
def crear_numeracion(doc):
    """Creates a numbering format and returns its index"""
    list_template = doc.ListTemplates.Add(OutlineNumbers=True)  # Create a simple outline numbering
    list_template.Name = "CustomNumbering"
    return list_template.Index  # Return the index for reference


# Function to apply numbering to a paragraph
def aplicar_numeracion(paragraph, num_index, level=0):
    """Applies numbering to a paragraph"""
    try:
        # Apply numbering using the list template
        paragraph.Range.ListFormat.ApplyListTemplate(
            ListTemplate=doc.ListTemplates(num_index),
            Level=level,
            ContinuePreviousList=False  # Start a new list
        )

        # Set indentation (approximation for left indent and hanging)
        paragraph.Range.ParagraphFormat.LeftIndent = 720  # In points (equivalent to 0.5 inches)
        paragraph.Range.ParagraphFormat.FirstLineIndent = -360  # Hanging indent
    except Exception as e:
        print(f"Error applying numbering: {e}")
    return paragraph


def extraer_seccion_completa(doc, titulo_seccion):
    """
    Extracts a full section including all paragraphs and tables until the next heading of the same or higher level.
    """
    seccion_heading = None
    elementos_seccion = []  # List of tuples (type, element, index)
    nivel_seccion = None
    indice_inicio = -1
    indice_fin = None

    # Search for the heading
    for i, paragraph in enumerate(doc.Paragraphs):
        if paragraph.Range.Text.strip() == titulo_seccion and "Heading" in paragraph.Style.Name:  # Check for heading style
            seccion_heading = paragraph
            indice_inicio = i
            # Extract level from style (e.g., Heading 1 -> level 1)
            if "Heading" in paragraph.Style.Name:
                try:
                    nivel_seccion = int(paragraph.Style.Name.replace("Heading ", ""))
                except ValueError:
                    nivel_seccion = 1  # Default to level 1
            break

    if indice_inicio == -1:
        print(f"No se encontró la sección '{titulo_seccion}'")
        return None

    # Find the end of the section
    for i in range(indice_inicio + 1, doc.Paragraphs.Count):
        next_paragraph = doc.Paragraphs(i + 1)  # Note: COM is 1-based
        if "Heading" in next_paragraph.Style.Name:
            try:
                nivel_actual = int(next_paragraph.Style.Name.replace("Heading ", ""))
                if nivel_actual <= nivel_seccion:
                    indice_fin = i
                    break
            except ValueError:
                pass

    if indice_fin is None:
        indice_fin = doc.Paragraphs.Count

    # Add paragraphs (except headings) between start and end
    for i in range(indice_inicio + 1, indice_fin):
        paragraph = doc.Paragraphs(i + 1)  # 1-based index
        if "Heading" not in paragraph.Style.Name:
            elementos_seccion.append(('parrafo', paragraph, i))

    # Add tables (approximate by checking their range)
    for i in range(doc.Tables.Count):
        table = doc.Tables(i + 1)  # 1-based
        table_range_start = table.Range.Start
        if indice_inicio < table_range_start < indice_fin:
            elementos_seccion.append(('tabla', table, table_range_start))

    # Sort elements by their position (using index or start position)
    elementos_seccion.sort(key=lambda x: x[2])

    # Strip indices for return
    elementos_seccion = [(tipo, elem) for tipo, elem in elementos_seccion]

    if seccion_heading is not None:
        return seccion_heading, elementos_seccion, nivel_seccion
    else:
        return None


def extraer_seccion_completa(doc, titulo_seccion):
    """
    Extracts a full section including all paragraphs and tables until the next heading of the same or higher level.
    Improved with fallback checks for heading detection.
    """
    seccion_heading = None
    elementos_seccion = []  # List of tuples (type, element, index)
    nivel_seccion = None
    indice_inicio = -1
    indice_fin = None

    print(f"Searching for section: '{titulo_seccion}' in document...")  # Debugging log

    for i, paragraph in enumerate(doc.Paragraphs):
        try:
            # Primary check: Use Style.Name if available
            style_name = paragraph.Style.Name
            if paragraph.Range.Text.strip() == titulo_seccion and "Heading" in style_name:
                seccion_heading = paragraph
                indice_inicio = i
                # Extract level from style
                if "Heading" in style_name:
                    try:
                        nivel_seccion = int(style_name.replace("Heading ", ""))
                    except ValueError:
                        nivel_seccion = 1  # Default to level 1
                print(f"Found section '{titulo_seccion}' at index {i} with level {nivel_seccion}")
                break
        except AttributeError:
            # Fallback 1: Check if it's a built-in heading style via Description
            try:
                if paragraph.Style.BuiltIn and "Heading" in paragraph.Style.Description:
                    if paragraph.Range.Text.strip() == titulo_seccion:
                        seccion_heading = paragraph
                        indice_inicio = i
                        nivel_seccion = 1  # Approximate level
                        print(f"Fallback: Found section '{titulo_seccion}' at index {i} (approximated level {nivel_seccion})")
                        break
            except AttributeError:
                # Fallback 2: Check formatting (e.g., bold or larger font) as a last resort
                try:
                    range_font = paragraph.Range.Font
                    is_bold = getattr(range_font, 'Bold', 0)  # Check if bold
                    font_size = getattr(range_font, 'Size', 12)  # Default to 12 if not available
                    if paragraph.Range.Text.strip() == titulo_seccion and (is_bold == -1 or font_size > 14):  # Assuming headings are bold or >14pt
                        seccion_heading = paragraph
                        indice_inicio = i
                        nivel_seccion = 1  # Default level for fallback
                        print(f"Advanced Fallback: Found section '{titulo_seccion}' at index {i} based on formatting")
                        break
                except AttributeError:
                    print(f"Warning: Paragraph {i} has no accessible Style or Font properties. Skipping.")
                    continue  # Skip to the next paragraph

    if indice_inicio == -1:
        print(f"No se encontró la sección '{titulo_seccion}' after checking all paragraphs and fallbacks.")
        return None

    # The rest of the function remains similar...
    for i in range(indice_inicio + 1, doc.Paragraphs.Count):
        next_paragraph = doc.Paragraphs(i + 1)  # 1-based index
        try:
            next_style_name = next_paragraph.Style.Name
            if "Heading" in next_style_name:
                try:
                    nivel_actual = int(next_style_name.replace("Heading ", ""))
                    if nivel_actual <= nivel_seccion:
                        indice_fin = i
                        break
                except ValueError:
                    pass
        except AttributeError:
            # Use fallbacks here as well for consistency
            try:
                if next_paragraph.Style.BuiltIn and "Heading" in next_paragraph.Style.Description:
                    nivel_actual = 1  # Approximate
                    if nivel_actual <= nivel_seccion:
                        indice_fin = i
                        break
            except AttributeError:
                continue

    if indice_fin is None:
        indice_fin = doc.Paragraphs.Count

    for i in range(indice_inicio + 1, indice_fin):
        paragraph = doc.Paragraphs(i + 1)
        elementos_seccion.append(('parrafo', paragraph, i))

    for i in range(doc.Tables.Count):
        table = doc.Tables(i + 1)
        table_range_start = table.Range.Start
        if indice_inicio < table_range_start < indice_fin:
            elementos_seccion.append(('tabla', table, table_range_start))

    elementos_seccion.sort(key=lambda x: x[2])
    elementos_seccion = [(tipo, elem) for tipo, elem in elementos_seccion]

    return seccion_heading, elementos_seccion, nivel_seccion


def copiar_tablas_con_win32(source_path, intermediate_path, output_path):
    """
    Uses win32com to copy tables from the source document to the intermediate one, replacing placeholders.
    """
    try:
        source_doc = word_app.Documents.Open(os.path.abspath(source_path))
        intermediate_doc = word_app.Documents.Open(os.path.abspath(intermediate_path))

        table_index = 0
        for i in range(intermediate_doc.Paragraphs.Count):
            paragraph = intermediate_doc.Paragraphs(i + 1)
            if "[[TABLE_PLACEHOLDER]]" in paragraph.Range.Text:
                if table_index < source_doc.Tables.Count:
                    source_table = source_doc.Tables(table_index + 1)
                    source_table.Range.Copy()
                    paragraph.Range.Paste()  # Replace placeholder
                    table_index += 1

        intermediate_doc.SaveAs(os.path.abspath(output_path))
        intermediate_doc.Close()
        source_doc.Close()
    except Exception as e:
        print(f"Error in table copy: {e}")
    finally:
        if 'source_doc' in locals():
            source_doc.Close()
        if 'intermediate_doc' in locals():
            intermediate_doc.Close()


# Main code
output_numered_cor = "resolucion_numerada.docx"
source_doc = word_app.Documents.Open(os.path.abspath(output_numered_cor))  # Load source
dest_doc = word_app.Documents.Add()  # Create new document

# Extract and copy sections
resultado_vistos = extraer_seccion_completa(source_doc, "VISTOS")
if resultado_vistos:
    encabezado_vistos, parrafos_vistos, nivel_vistos = resultado_vistos
    copiar_seccion_completa(dest_doc, encabezado_vistos, parrafos_vistos, nivel_vistos)

resultado_considerando = extraer_seccion_completa(source_doc, "CONSIDERANDO")
if resultado_considerando:
    encabezado_cons, parrafos_cons, nivel_cons = resultado_considerando
    copiar_seccion_completa(dest_doc, encabezado_cons, parrafos_cons, nivel_cons)

resultado_resolucion = extraer_seccion_completa(source_doc, "RESOLUCIÓN")
if resultado_resolucion:
    encabezado_res, parrafos_res, nivel_res = resultado_resolucion
    copiar_seccion_completa(dest_doc, encabezado_res, parrafos_res, nivel_res)

resultado_requisitos = extraer_seccion_completa(source_doc, "REQUISITOS")
if resultado_requisitos:
    encabezado_req, parrafos_req, nivel_req = resultado_requisitos
    copiar_seccion_completa(dest_doc, encabezado_req, parrafos_req, nivel_req)

# Save intermediate document
intermediate_file = "onlydocx.docx"
dest_doc.SaveAs(os.path.abspath(intermediate_file))

# Copy tables
final_output = "seccion_completa_copiada.docx"
copiar_tablas_con_win32(output_numered_cor, intermediate_file, final_output)

dest_doc.Close()  # Close the destination document
source_doc.Close()  # Close the source document
word_app.Quit()  # Quit Word application
pythoncom.CoUninitialize()  # Uninitialize COM at the end

print(f"Documento final guardado como: {final_output}")


