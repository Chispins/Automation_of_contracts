import os
import sys
import time
from watchdog.observers import Observer
from watchdog.events import FileSystemEventHandler
import shutil
import re
import shutil
import importlib.util
from docxtpl import DocxTemplate
import pandas as pd

# Definimos el directorio a monitorear
directorio = r"\\10.5.130.24\Abastecimiento\Compartido Abastecimiento\Otros\Licitaciones_Testing"
import docx
doc = docx.Document()
cwd = os.getcwd()

class Handler(FileSystemEventHandler):
    def on_modified(self, event):
        # Verificamos si el evento es un archivo y no un directorio
        if event.is_directory:
            return

        file_name = os.path.basename(event.src_path)
        # Filtramos archivos temporales y ocultos
        if not file_name.startswith("~$") and not file_name.endswith(".tmp"):
            if file_name == "Libro1.xlsx":
                try:
                    # Leemos el archivo Excel
                    data = pd.read_excel(event.src_path, sheet_name = "Datos_Base")
                    data_contrato = pd.read_excel(event.src_path, sheet_name = "Datos_Contrato_P1")
                    # Verificamos si existe la fila 3, columna "si"
                    # Nota: pandas usa índices base-0, así que fila 3 es índice 2
                    # Asumiendo que buscas en la columna 1 (índice 0)
                    directorio_carpeta_monitoreada = os.path.dirname(event.src_path)

                    if data.iloc[1,3] == "si":
                        if "base_automatizada.docx" not in os.listdir(directorio_carpeta_monitoreada):
                            print("Base automatizada inicada.")
                            generar_base()
                            render_base()
                        else:
                            print("Base automatizada ya existe.")
                    elif (data_contrato.iloc[1,3] == "si"):
                        if "contrato_automatizado" not in os.listdir(directorio_carpeta_monitoreada):
                            print("Contrato automatizado iniciado.")
                            generar_contrato()
                            render_contrato()
                        else:
                            print("Contrato automatizado ya existe.")

                except Exception as e:
                    print(f"Error al procesar {file_name}: {e}")

    def on_created(self, event):
        directorio_carpeta_creada = event.src_path
        print(f"Carpeta monitoreada es {directorio_carpeta_creada}")

        if event.is_directory:
            try:
                # Creamos documentos separados para cada caso
                doc_base = docx.Document()
                doc_contrato = docx.Document()
                name_file_origin = "Libro1.xlsx"
                file_origin = os.path.join(r"\\10.5.130.24\Abastecimiento\Compartido Abastecimiento\Otros\NO_MODIFICAR",
                                           name_file_origin)
                destination = os.path.join(directorio_carpeta_creada, name_file_origin)

                # Verificar si el archivo origen existe
                if not os.path.exists(file_origin):
                    print(f"Error: Archivo origen {file_origin} no existe.")
                    # Intentar buscar en otras ubicaciones
                    alternative_origin = os.path.join(cwd, name_file_origin)
                    if os.path.exists(alternative_origin):
                        file_origin = alternative_origin
                        print(f"Se encontró archivo alternativo en: {file_origin}")
                    else:
                        print("No se encontró archivo Libro1.xlsx en ninguna ubicación.")
                        return

                if os.path.exists(destination):
                    print(f"El archivo {name_file_origin} ya existe en {destination}.")
                else:
                    shutil.copy2(file_origin, destination)
                    print(f"Archivo copiado a {destination}")


                if "portada_melipilla_base" not in os.listdir(directorio_carpeta_creada):
                    print("Creando portada melipilla base...")
                    ruta_completa = os.path.join(directorio_carpeta_creada, "contrato_base.docx")
                    doc_base.save(ruta_completa)
                    print(f"Archivo guardado en: {ruta_completa}")
                    crear_portada_melipilla_base()

                if "portada_melipilla_contrato" not in os.listdir(directorio_carpeta_creada):
                    print("Creando portada melipilla contrato...")
                    ruta_completa = os.path.join(directorio_carpeta_creada, "contrato_portada.docx")
                    doc_contrato.save(ruta_completa)
                    print(f"Archivo guardado en: {ruta_completa}")
                    crear_portada_melipilla_contrato()

            except Exception as e:
                print(f"Error al procesar directorio: {e}")


def generar_base():
    try:
        dir_original = os.getcwd()

        # Buscar el módulo Formated_Base_PEP8.py
        posibles_rutas = [
            os.path.join(cwd, "Formated_Base_PEP8.py"),
            os.path.join(os.path.dirname(cwd), "Formated_Base_PEP8.py"),
            os.path.join(directorio, "Formated_Base_PEP8.py"),
            r"\\10.5.130.24\Abastecimiento\Compartido Abastecimiento\Otros\NO_MODIFICAR\Formated_Base_PEP8.py"
        ]

        ruta_formated_bases = None
        for ruta in posibles_rutas:
            print(f"Buscando módulo en: {ruta}")
            if os.path.exists(ruta):
                ruta_formated_bases = ruta
                break

        if not ruta_formated_bases:
            raise FileNotFoundError("No se pudo encontrar Formated_Base_PEP8.py en ninguna ubicación conocida")

        print(f"Módulo encontrado en: {ruta_formated_bases}")

        # Buscar el archivo Libro1.xlsx en las subcarpetas del directorio
        ruta_excel = None

        # Primero intentar en la carpeta principal
        posibles_rutas_excel = [
            os.path.join(directorio, "Libro1.xlsx"),
            os.path.join(dir_original, "Libro1.xlsx"),
        ]

        # Luego buscar en todas las subcarpetas (sólo un nivel)
        if os.path.exists(directorio):
            for item in os.listdir(directorio):
                item_path = os.path.join(directorio, item)
                if os.path.isdir(item_path):
                    excel_path = os.path.join(item_path, "Libro1.xlsx")
                    posibles_rutas_excel.append(excel_path)

        for ruta in posibles_rutas_excel:
            print(f"Buscando Excel en: {ruta}")
            if os.path.exists(ruta):
                ruta_excel = ruta
                break

        if not ruta_excel:
            raise FileNotFoundError(f"No se pudo encontrar Libro1.xlsx en ninguna ubicación conocida")

        print(f"Excel encontrado en: {ruta_excel}")

        modulo_formated_bases = cargar_modulo_desde_archivo(ruta_formated_bases, "formated_bases")

        print("Generando base automatizada...")
        modulo_formated_bases.main_bases(archivo="base", wd=os.path.dirname(ruta_excel), monitoring=True,
                                         excel_path=ruta_excel)

        print("Base automatizada generada correctamente.")
    except Exception as e:
        print(f"Error al generar la base automatizada: {e}")
        import traceback
        traceback.print_exc()

def generar_contrato():
    """Genera el contrato automatizado usando el módulo Formated_Contracts_PEP8"""
    try:
        # Guardar el directorio actual
        dir_original = os.getcwd()

        # Buscar el archivo en diferentes ubicaciones posibles
        posibles_rutas = [
            os.path.join(cwd, "Formated_Contracts_PEP8.py"),
            os.path.join(os.path.dirname(cwd), "Formated_Contracts_PEP8.py"),
            os.path.join(directorio, "Formated_Contracts_PEP8.py"),
            r"\\10.5.130.24\Abastecimiento\Compartido Abastecimiento\Otros\NO_MODIFICAR\Formated_Contracts_PEP8.py"
        ]

        ruta_formated_contracts = None
        for ruta in posibles_rutas:
            print(f"Buscando módulo en: {ruta}")
            if os.path.exists(ruta):
                ruta_formated_contracts = ruta
                break

        if not ruta_formated_contracts:
            raise FileNotFoundError("No se pudo encontrar Formated_Contracts_PEP8.py en ninguna ubicación conocida")

        print(f"Módulo encontrado en: {ruta_formated_contracts}")

        # Cargar el módulo sin ejecutar su código principal
        modulo_formated_contracts = cargar_modulo_desde_archivo(ruta_formated_contracts, "formated_contracts")

        print("Generando contrato automatizado...")
        # Llamar a la función principal del módulo con el directorio actual
        modulo_formated_contracts.main(wd=dir_original, monitoring=True)

        print("Contrato automatizado generado correctamente.")
    except Exception as e:
        print(f"Error al generar el contrato automatizado: {e}")

def render_base():
    print("render")

def render_contrato():
    print("render contrato")

def crear_portada_melipilla_base():
    print("Creando portada melipilla base...")

def crear_portada_melipilla_contrato():
    print("Creando portada melipilla contrato...")

def iniciar_monitoreo(path):
    # Nos aseguramos de que el directorio exista
    if not os.path.exists(path):
        print(f"El directorio {path} no existe.")
        return

    # Cambiamos al directorio especificado
    os.chdir(path)

    observer = Observer()
    observer.schedule(Handler(), path=".", recursive=True)
    observer.start()

    print(f"Monitoreando cambios en {path}...")

    try:
        while True:
            time.sleep(1)
    except KeyboardInterrupt:
        print("Monitoreo detenido por el usuario")
        observer.stop()

    observer.join()


# Sugerencias de ChatGPT:

def configurar_directorio_trabajo(reset=False, directorio_base=None):
    """
    Configura el directorio de trabajo en la subcarpeta 'Files'.

    Args:
        reset: Si es True, restaura el directorio original después de realizar operaciones
        directorio_base: Directorio base alternativo para usar en lugar del CWD

    Returns:
        El directorio original si reset=True, para poder restaurarlo después
    """
    directorio_original = os.getcwd()

    if directorio_base:
        cwd = directorio_base
    else:
        cwd = directorio_original

    target_dir_name = "Files"
    wd = os.path.join(cwd, target_dir_name)

    # Evitar rutas duplicadas
    pattern = r"Files\\Files"
    if re.search(pattern, wd):
        wd = wd.replace(r"\Files\Files", r"\Files")

    if os.path.isdir(wd):
        os.chdir(wd)
        print(f"Directorio de trabajo cambiado a: {wd}")
    else:
        print(f"Advertencia: El directorio '{wd}' no existe o no es válido.")

    if reset:
        return directorio_original
    return None

def cargar_modulo_desde_archivo(ruta_archivo, nombre_modulo):
    """
    Carga un módulo Python desde una ruta de archivo sin ejecutar su código principal.

    Args:
        ruta_archivo: Ruta al archivo Python
        nombre_modulo: Nombre para el módulo importado

    Returns:
        El módulo cargado
    """
    spec = importlib.util.spec_from_file_location(nombre_modulo, ruta_archivo)
    modulo = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(modulo)
    return modulo


def procesar_documento_completo():
    """Función principal que integra la funcionalidad de los distintos módulos"""

    # Guardar el directorio original
    dir_original = os.getcwd()

    try:
        # 1. Configurar directorio de trabajo inicial
        configurar_directorio_trabajo()

        # 2. Importar los módulos necesarios (sin ejecutar el código principal)
        # Asegúrate de que las rutas sean correctas según tu estructura de archivos
        modulo_formated_contracts = cargar_modulo_desde_archivo("../Formated_Contracts_PEP8.py", "formated_contracts")
        modulo_formated_bases = cargar_modulo_desde_archivo("../Formated_Base_PEP8.py", "formated_bases")

        # 3. Ejecutar la funcionalidad de Formatted_Contracts_PEP8
        # Accede a las funciones del módulo pero sin ejecutar su __main__
        print("Procesando contrato...")
        modulo_formated_contracts.main()

        # 4. Ejecutar la funcionalidad de Jinja (aprovechando el código existente)
        print("Aplicando plantillas Jinja...")
        # Aquí podría ir el código de procesamiento de Jinja adaptado para usar rutas absolutas
        excel_name = os.path.abspath("Libro1.xlsx")
        template_name_1 = os.path.abspath("base_automatizada.docx")
        template_name_3 = os.path.abspath("contrato_automatizado_tablas.docx")

        # Lógica de procesamiento de Jinja...
        # (Incluir aquí el código del procesamiento de Jinja adaptado para usar rutas absolutas)

        print("Proceso completo finalizado con éxito!")

    finally:
        # Restaurar el directorio original al finalizar
        os.chdir(dir_original)
        print(f"Directorio restaurado a: {dir_original}")


def verificar_archivos_necesarios():
    """Verifica que los archivos necesarios existan"""
    archivos = {
        "Libro1.xlsx": r"\\10.5.130.24\Abastecimiento\Compartido Abastecimiento\Otros\NO_MODIFICAR\Libro1.xlsx"
    }

    for nombre, ruta in archivos.items():
        if not os.path.exists(ruta):
            print(f"Advertencia: El archivo {nombre} no existe en {ruta}")
            # Crear una copia desde el directorio local si existe
            local_path = os.path.join(cwd, nombre)
            if os.path.exists(local_path):
                os.makedirs(os.path.dirname(ruta), exist_ok=True)
                shutil.copy2(local_path, ruta)
                print(f"Se ha copiado {nombre} desde {local_path} a {ruta}")
            else:
                print(f"Error: No se encontró {nombre} en ninguna ubicación.")

# Ejecutamos la función de monitoreo
if __name__ == "__main__":
    verificar_archivos_necesarios()  # Añadir esta línea
    iniciar_monitoreo(directorio)