import docx
import os
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from win32com import client
import pythoncom
from Bases import configurar_directorio_trabajo # Línea añadida

configurar_directorio_trabajo() # Línea añadida

# Función para crear numeración
def crear_numeracion(doc):
    """Crea un formato de numeración y devuelve su ID"""
    part = doc._part
    if not hasattr(part, 'numbering_part'):
        part._add_numbering_part()
    import random
    num_id = random.randint(1000, 9999)
    return num_id

# Función para aplicar numeración a un párrafo
def aplicar_numeracion(parrafo, num_id, nivel=0):
    """Aplica numeración a un párrafo"""
    p = parrafo._p
    pPr = p.get_or_add_pPr()

    # Eliminar numeración previa si existe
    for child in pPr.iterchildren():
        if child.tag.endswith('numPr'):
            pPr.remove(child)

    # Crear nuevo numPr
    numPr = OxmlElement('w:numPr')

    # Establecer nivel
    ilvl = OxmlElement('w:ilvl')
    ilvl.set(qn('w:val'), str(nivel))
    numPr.append(ilvl)

    # Establecer ID de numeración
    numId = OxmlElement('w:numId')
    numId.set(qn('w:val'), str(num_id))
    numPr.append(numId)

    # Agregar al párrafo
    pPr.append(numPr)

    # Configurar sangría
    ind = OxmlElement('w:ind')
    ind.set(qn('w:left'), '720')
    ind.set(qn('w:hanging'), '360')
    pPr.append(ind)

    return parrafo

def extraer_seccion_completa(doc_cargado, titulo_seccion):
    """
    Extrae una sección completa incluyendo todos sus párrafos y tablas hasta el siguiente encabezado
    del mismo nivel o superior.
    """
    seccion_heading = None
    elementos_seccion = []  # Lista de tuplas (tipo, elemento, indice)
    nivel_seccion = None
    indice_inicio = -1
    indice_fin = None

    # Buscar el encabezado
    for i, element in enumerate(doc_cargado.paragraphs):
        if element.style.name.startswith('Heading') and element.text.strip() == titulo_seccion:
            seccion_heading = element
            indice_inicio = i
            try:
                nivel_seccion = int(element.style.name.split()[-1])
            except (ValueError, IndexError):
                nivel_seccion = 2  # Valor por defecto
            break

    if indice_inicio == -1:
        print(f"No se encontró la sección '{titulo_seccion}'")
        return None

    # Encontrar el índice del fin de la sección
    for i in range(indice_inicio + 1, len(doc_cargado.paragraphs)):
        parrafo = doc_cargado.paragraphs[i]
        if parrafo.style.name.startswith('Heading'):
            try:
                nivel_actual = int(parrafo.style.name.split()[-1])
                if nivel_actual <= nivel_seccion:
                    indice_fin = i
                    break
            except (ValueError, IndexError):
                pass

    if indice_fin is None:
        indice_fin = len(doc_cargado.paragraphs)

    # Agregar párrafos (excepto encabezados) entre inicio y fin
    for i in range(indice_inicio + 1, indice_fin):
        parrafo = doc_cargado.paragraphs[i]
        if not parrafo.style.name.startswith('Heading'):
            elementos_seccion.append(('parrafo', parrafo, i))

    # Crear un mapa de la posición de cada párrafo en el documento XML
    parrafos_posicion_xml = {}
    for i, p in enumerate(doc_cargado.paragraphs):
        parrafos_posicion_xml[p._p] = i

    # Extraer la posición de cada tabla en el XML
    tablas_posicion = []
    for i, tabla in enumerate(doc_cargado.tables):
        tabla_p = tabla._tbl
        elemento_anterior = tabla_p.getprevious()
        pos_anterior = -1
        while elemento_anterior is not None:
            if elemento_anterior in parrafos_posicion_xml:
                pos_anterior = parrafos_posicion_xml[elemento_anterior]
                break
            elemento_anterior = elemento_anterior.getprevious()
        if indice_inicio <= pos_anterior < indice_fin:
            tablas_posicion.append((i, tabla, pos_anterior + 0.5))

    # Agregar las tablas que están dentro de la sección
    for i, tabla, pos in tablas_posicion:
        elementos_seccion.append(('tabla', tabla, pos))

    # Ordenar elementos por su posición en el documento
    elementos_seccion.sort(key=lambda x: x[2])

    # Eliminar los índices para el retorno
    elementos_seccion = [(tipo, elem) for tipo, elem, _ in elementos_seccion]

    if seccion_heading is not None:
        return seccion_heading, elementos_seccion, nivel_seccion
    else:
        return None

def copiar_seccion_completa(doc_destino, seccion_heading, elementos_seccion, nivel_seccion):
    """
    Copia una sección completa al documento destino manteniendo una única numeración para todos los párrafos de "List Number".
    No copia tablas en esta etapa, solo texto.
    """
    # Copiar encabezado
    doc_destino.add_heading(seccion_heading.text, level=nivel_seccion)

    # Crear un único ID de numeración para toda la sección
    seccion_num_id = None

    # Copiar cada elemento de la sección
    for tipo, elemento in elementos_seccion:
        if tipo == 'parrafo':
            parrafo = elemento
            nuevo_parrafo = doc_destino.add_paragraph(style=parrafo.style.name)

            # Copiar cada run con su formato
            for run in parrafo.runs:
                nuevo_run = nuevo_parrafo.add_run(run.text)
                nuevo_run.bold = run.bold
                nuevo_run.italic = run.italic
                nuevo_run.underline = run.underline
                if run.font.color.rgb:
                    nuevo_run.font.color.rgb = run.font.color.rgb
                if run.font.name:
                    nuevo_run.font.name = run.font.name
                if run.font.size:
                    nuevo_run.font.size = run.font.size

            # Aplicar numeración solo a párrafos de tipo lista
            if parrafo.style.name == "List Number":
                if seccion_num_id is None:
                    seccion_num_id = crear_numeracion(doc_destino)
                aplicar_numeracion(nuevo_parrafo, seccion_num_id)

        elif tipo == 'tabla':
            # No copiar tablas en esta etapa, solo agregar un marcador de posición
            doc_destino.add_paragraph(f"[[TABLE_PLACEHOLDER]]")

def copiar_tablas_con_win32(source_path, intermediate_path, output_path):
    """
    Usa win32com para copiar las tablas del documento original al documento intermedio,
    reemplazando los marcadores de posición.
    """
    pythoncom.CoInitialize()
    try:
        word_app = client.Dispatch("Word.Application")
        word_app.Visible = False

        # Abrir documentos
        source_doc = word_app.Documents.Open(os.path.abspath(source_path))
        intermediate_doc = word_app.Documents.Open(os.path.abspath(intermediate_path))

        # Contar tablas en el documento original
        source_tables = source_doc.Tables
        table_index = 0

        # Buscar marcadores de posición en el documento intermedio
        for para in intermediate_doc.Paragraphs:
            if "[[TABLE_PLACEHOLDER]]" in para.Range.Text:
                if table_index < source_tables.Count:
                    # Copiar la tabla correspondiente del documento original
                    source_table = source_tables.Item(table_index + 1)
                    source_table.Range.Copy()

                    # Pegar la tabla en el lugar del marcador
                    para.Range.Paste()
                    table_index += 1

        # Guardar el documento final
        intermediate_doc.SaveAs(os.path.abspath(output_path))
        intermediate_doc.Close()
        source_doc.Close()
        word_app.Quit()

    finally:
        pythoncom.CoUninitialize()

# Configuración del documento
output_numered_cor = "resolucion_numerada.docx"
word = docx.Document(output_numered_cor)
doc = docx.Document()

secciones = [
    "BASES ADMINISTRATIVAS PARA EL SUMINISTRO DE INSUMOS Y ACCESORIOS PARA TERAPIA DE PRESIÓN NEGATIVA CON EQUIPOS EN COMODATO PARA EL HOSPITAL SAN JOSÉ DE MELIPILLA",
    "Condiciones Contractuales, Vigencia de las Condiciones Comerciales, Operatoria de la Licitación y Otras Cláusulas:",
    "BASES TECNICAS PARA EL SUMINISTRO DE INSUMOS Y ACCESORIOS PARA TERAPIA DE PRESIÓN NEGATIVA CON EQUIPOS EN COMODATO PARA EL HOSPITAL SAN JOSÉ DE MELIPILLA"
]
# Extraer y copiar sección VISTOS completa
resultado_vistos = extraer_seccion_completa(word, secciones[0])
if resultado_vistos:
    encabezado_vistos, parrafos_vistos, nivel_vistos = resultado_vistos
    copiar_seccion_completa(doc, encabezado_vistos, parrafos_vistos, nivel_vistos)

# Extraer y copiar sección CONSIDERANDO completa
resultado_considerando = extraer_seccion_completa(word, secciones[1])
if resultado_considerando:
    encabezado_cons, parrafos_cons, nivel_cons = resultado_considerando
    copiar_seccion_completa(doc, encabezado_cons, parrafos_cons, nivel_cons)

# Extraer y copiar sección RESOLUCIÓN completa
resultado_resolucion = extraer_seccion_completa(word, secciones[2])
if resultado_resolucion:
    encabezado_res, parrafos_res, nivel_res = resultado_resolucion
    copiar_seccion_completa(doc, encabezado_res, parrafos_res, nivel_res)

# Extraer y copiar sección REQUISITOS completa


# Guardar documento intermedio con solo texto
intermediate_file = "onlydocx.docx"
doc.save(intermediate_file)

# Copiar tablas usando win32com
final_output = "seccion_completa_copiada.docx"
copiar_tablas_con_win32(output_numered_cor, intermediate_file, final_output)

print(f"Documento final guardado como: {final_output}")
