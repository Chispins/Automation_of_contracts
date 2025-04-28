import os
import docx
from docx.oxml.ns import qn
from docx.oxml.shared import OxmlElement
from Bases import configurar_directorio_trabajo

configurar_directorio_trabajo()

def obtener_marcadores(documento):
    """Obtiene todos los marcadores de un documento Word."""
    marcadores = {}
    # Buscar todos los elementos de inicio de marcador
    for elemento in documento._element.xpath('//w:bookmarkStart'):
        nombre_marcador = elemento.get(qn('w:name'))
        id_marcador = elemento.get(qn('w:id'))
        if nombre_marcador:
            marcadores[nombre_marcador] = {
                'id': id_marcador,
                'elemento': elemento
            }
    return marcadores


def crear_marcador(documento, nombre_marcador, texto, parrafo=None):
    """Crea un nuevo marcador en el documento con el texto especificado."""
    # Si no se proporciona un párrafo, crear uno nuevo
    if parrafo is None:
        parrafo = documento.add_paragraph()

    run = parrafo.add_run()
    tag = run._r

    # Crear elemento de inicio del marcador
    start = OxmlElement('w:bookmarkStart')
    start.set(qn('w:id'), str(len(documento._element.xpath('//w:bookmarkStart')) + 1))
    start.set(qn('w:name'), nombre_marcador)
    tag.append(start)

    # Agregar texto
    text_run = parrafo.add_run(texto)

    # Crear elemento de fin del marcador
    end = OxmlElement('w:bookmarkEnd')
    end.set(qn('w:id'), start.get(qn('w:id')))
    tag.append(end)

    return True


def modificar_texto_marcador(documento, nombre_marcador, nuevo_texto):
    """Modifica el texto contenido en un marcador."""
    marcadores = obtener_marcadores(documento)

    if nombre_marcador not in marcadores:
        print(f"No se encontró el marcador '{nombre_marcador}'")
        print("Intentando crear un nuevo marcador...")
        return crear_marcador(documento, nombre_marcador, nuevo_texto)

    # Obtener el elemento de inicio
    elemento_inicio = marcadores[nombre_marcador]['elemento']
    id_marcador = marcadores[nombre_marcador]['id']

    # Buscar el elemento final del marcador
    elemento_fin = None
    for elem in documento._element.xpath(f'//w:bookmarkEnd[@w:id="{id_marcador}"]'):
        elemento_fin = elem
        break

    if elemento_fin is None:
        print(f"Marcador '{nombre_marcador}' incompleto (falta el final)")
        return False

    # Buscar el párrafo que contiene el marcador
    for paragraph in documento.paragraphs:
        if elemento_inicio in paragraph._p.xpath('.//*'):
            # Crear un nuevo run con el texto deseado
            paragraph.text = ""  # Limpiar el párrafo
            run = paragraph.add_run()
            run._r.append(elemento_inicio)  # Añadir el inicio del marcador
            run = paragraph.add_run(nuevo_texto)  # Añadir el texto
            run = paragraph.add_run()
            run._r.append(elemento_fin)  # Añadir el fin del marcador
            return True

    print(f"No se pudo localizar el párrafo del marcador '{nombre_marcador}'")
    return False


# Cargar el documento
doc = docx.Document("resolucion_numerada.docx")

# Listar marcadores disponibles
marcadores = obtener_marcadores(doc)
print("Marcadores disponibles en el documento:")
for nombre in marcadores:
    print(f"- {nombre} (ID: {marcadores[nombre]['id']})")

# Si no hay marcadores o si se menciona un error de referencia
if "origen de la referencia" in doc.paragraphs[0].text if doc.paragraphs else False:
    print("\nSe detectó un error de referencia. Creando un nuevo marcador...")
    crear_marcador(doc, "Referencia", "funcionó")
    print("Marcador 'Referencia' creado correctamente")
elif marcadores:
    nombre_marcador = list(marcadores.keys())[0]  # Usar el primer marcador disponible
    print(f"\nModificando el marcador: '{nombre_marcador}'")
    if modificar_texto_marcador(doc, nombre_marcador, "funcionó"):
        print(f"Se modificó el marcador '{nombre_marcador}' correctamente")
else:
    print("No se encontraron marcadores. Creando uno nuevo...")
    crear_marcador(doc, "NuevoMarcador", "funcionó")
    print("Marcador 'NuevoMarcador' creado correctamente")

# Guardar el documento con cambios
doc.save("resolucion_modificada.docx")
print("Documento guardado como 'resolucion_modificada.docx'")
