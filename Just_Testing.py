import os
import re
import random
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches

def configurar_directorio_trabajo():
    """Configura el directorio de trabajo en la subcarpeta 'Files'."""
    cwd = os.getcwd()
    target_dir_name = "Files"
    wd = os.path.join(cwd, target_dir_name)
    # Avoid creating nested Files/Files directory if already in Files
    pattern = r"[/\\]" + re.escape(target_dir_name) + r"[/\\]" + re.escape(target_dir_name)
    if re.search(pattern, wd):
        wd = wd.replace(os.path.join(target_dir_name, target_dir_name), target_dir_name)
    if os.path.isdir(wd):
        os.chdir(wd)
        print(f"Directorio de trabajo cambiado a: {wd}")
    else:
        try:
            os.makedirs(wd)
            os.chdir(wd)
            print(f"Directorio '{wd}' creado y establecido como directorio de trabajo.")
        except OSError as e:
            print(f"Error al crear o cambiar al directorio '{wd}': {e}")
            print(f"Directorio de trabajo actual: {cwd}") # Stay in current directory if creation fails
            pass # Allow execution to continue even if directory change fails

def crear_numeracion(doc):
    """Crea un formato de numeración y devuelve su ID."""
    # This is a simplification. python-docx numbering is complex.
    # For simple numbered lists, using built-in styles or adding a custom style is better.
    # Let's rely on built-in 'List Number' style and try to manage levels via Oxml if needed.
    # For robust custom lists, more complex xml manipulation is required.
    # Returning a random ID might not be strictly necessary if relying on styles.
    # Let's return a unique ID for each distinct list sequence needed.
    # A simple counter based on the random ID idea might suffice for demonstration.
    return random.randint(1000, 9999)

def aplicar_numeracion(parrafo, num_id, nivel=0):
    """Aplica numeración a un párrafo con el ID y nivel especificados using Oxml."""
    p = parrafo._p
    pPr = p.get_or_add_pPr()

    # Remove existing numbering properties if any
    for child in pPr.iterchildren(qn('w:numPr')):
        pPr.remove(child)

    numPr = OxmlElement('w:numPr')
    ilvl = OxmlElement('w:ilvl')
    ilvl.set(qn('w:val'), str(nivel))
    numPr.append(ilvl)
    numId = OxmlElement('w:numId')
    numId.set(qn('w:val'), str(num_id))
    numPr.append(numId)
    pPr.append(numPr)

    # Add indentation appropriate for the level (simple fixed values)
    # These values might need tuning depending on default list styles
    indent_val = 720 + (nivel * 360) # Base indent 720 (0.5 inch), hanging indent 360 (0.25 inch) per level
    first_line_val = 360 # Hanging indent

    # Remove existing indentation properties if any
    for child in pPr.iterchildren(qn('w:ind')):
        pPr.remove(child)

    ind = OxmlElement('w:ind')
    ind.set(qn('w:left'), str(indent_val))
    ind.set(qn('w:hanging'), str(first_line_val))
    pPr.append(ind)

    return parrafo


def centrar_verticalmente_tabla(tabla):
    """Aplica alineación vertical centrada a todas las celdas de una tabla."""
    for fila in tabla.rows:
        for celda in fila.cells:
            tc = celda._tc
            tcPr = tc.get_or_add_tcPr()
            # Remove existing vAlign if any
            for vAlign in tcPr.iterchildren(qn('w:vAlign')):
                 tcPr.remove(vAlign)
            vAlign = OxmlElement('w:vAlign')
            vAlign.set(qn('w:val'), 'center')
            tcPr.append(vAlign)

def agregar_parrafo_con_texto(doc, texto, style=None, negrita=False, centrado=False):
    """Agrega un párrafo con texto y aplica style or format if specified."""
    p = doc.add_paragraph(texto, style=style)
    if negrita and p.runs:
        for run in p.runs: # Apply bold to all runs in case of multiple runs created by style
             run.bold = True
        if not p.runs: # If paragraph was empty but text added later
             p.add_run(texto).bold = True
    if centrado:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p

def agregar_parrafo_con_runs(doc, partes, style=None, centrado=False):
    """Agrega un párrafo with multiple runs and applies specific formats."""
    p = doc.add_paragraph(style=style)
    if centrado:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for texto, formato in partes:
        run = p.add_run(texto)
        if formato:
            fmts = formato.split(',')
            if 'bold' in fmts:
                run.bold = True
            if 'underline' in fmts:
                run.underline = True
            # Add other formats as needed (e.g., italic, strikethrough)
    return p

def crear_tabla(doc, datos, style='Table Grid', centrar_vertical=True):
    """Crea una tabla with the provided data and applies formatting."""
    filas = len(datos)
    columnas = len(datos[0]) if filas > 0 else 0
    if filas == 0 or columnas == 0:
         return doc.add_table(1, 1, style=style) # Add a minimal table to avoid errors

    tabla = doc.add_table(filas, columnas, style=style)
    for i, fila in enumerate(datos):
        for j, texto in enumerate(fila):
            celda = tabla.cell(i, j)
            # Use agregar_contenido_celda logic for potentially multi-line text
            # This is a simplified version just setting text directly
            celda.text = str(texto)

    if centrar_vertical:
        centrar_verticalmente_tabla(tabla)

    return tabla

def agregar_contenido_celda(tabla, fila, columna, contenidos):
    """Agrega contenido to a cell with multiple paragraphs or formatted runs."""
    celda = tabla.cell(fila, columna)
    # Clear existing content in the first paragraph
    if celda.paragraphs and celda.paragraphs[0].text == '':
         p_to_use = celda.paragraphs[0]
         # Clear any existing runs in this paragraph
         for run in p_to_use.runs:
             p_to_use._element.remove(run._element)
    else:
        # Add a new paragraph if the first one is not empty or doesn't exist
        p_to_use = celda.add_paragraph()

    # Add content to the chosen paragraph (p_to_use)
    for i, contenido in enumerate(contenidos):
        if isinstance(contenido, str):
            # Add as a simple run in the *current* paragraph
            if i > 0: p_to_use.add_run('\n') # Add newline between string items
            p_to_use.add_run(contenido)
        elif isinstance(contenido, list):
            # Assumes list items are (text, format) tuples for runs
            # Add as multiple runs in the *current* paragraph
             if i > 0: p_to_use.add_run('\n') # Add newline between list items
             for item in contenido:
                 if isinstance(item, tuple) and len(item) == 2:
                     texto, formato = item
                     run = p_to_use.add_run(texto)
                     if formato:
                         fmts = formato.split(',')
                         if 'bold' in fmts: run.bold = True
                         if 'underline' in fmts: run.underline = True
                         # Add other formats as needed
                 else:
                     # If it's not a tuple or has unexpected length, treat it as plain text run
                     p_to_use.add_run(str(item))
        # You could add logic here to create *new* paragraphs within the cell
        # For simplicity, current implementation puts everything into one paragraph with newlines.
        # To add separate paragraphs, you'd need:
        # new_p = celda.add_paragraph()
        # process content into new_p


def add_page_number(document):
    """Adds page numbers to the footer of the document."""
    section = document.sections[-1]
    footer = section.footer
    footer_paragraph = footer.paragraphs[0]
    footer_paragraph.text = "Page "
    run = footer_paragraph.add_run()
    # Add page number field
    run.add_field('PAGE')
    footer_paragraph.add_run(' of ')
    # Add total pages field
    run = footer_paragraph.add_run()
    run.add_field('NUMPAGES')
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT # Align right


def main():
    configurar_directorio_trabajo()
    doc = Document()

    # Add a custom numbering style based on List Number template if needed
    # This is complex, sticking to OXml manipulation per paragraph is simpler for mixed lists
    # Or just use built-in styles like 'List Number' and 'List Bullet'

    # Create unique IDs for different list sequences
    num_id_considerando = crear_numeracion(doc)
    num_id_resolucion = crear_numeracion(doc)
    num_id_bases_admin_main = crear_numeracion(doc)
    num_id_requisitos = crear_numeracion(doc)
    num_id_administrador_contrato_list = crear_numeracion(doc)
    num_id_termino_anticipado = crear_numeracion(doc)
    num_id_anexo4 = crear_numeracion(doc)
    num_id_bases_tecnicas_generalidades = crear_numeracion(doc)
    num_id_final_resolution = crear_numeracion(doc)

    # Header (Skipped - complex to replicate logo/layout)
    # Add top text "Servicio de Salud Metropolitano Occidente", etc. if desired

    # Main Title and Location
    agregar_parrafo_con_texto(doc, "RESOLUCIÓN EXENTA Nº", centrado=True)
    agregar_parrafo_con_texto(doc, "MELIPILLA", centrado=True)

    # VISTOS Section
    doc.add_heading("VISTOS", level=1)
    agregar_parrafo_con_texto(doc, "Lo dispuesto en la Ley Nº 19.886 de Bases sobre Contratos Administrativos de Suministro y Prestación de Servicios; el Decreto Supremo Nº 250 /04 modificado por los Decretos Supremos Nº 1763/09, 1383/11 y 1410/14 todos del Ministerio de Hacienda; D. S. 38/2005, Reglamento Orgánico de los Establecimientos de Menor Complejidad y de los Establecimientos de Autogestión en Red; en uso de las atribuciones que me confieren el D.F.L. Nº 1/2.005, en virtud del cual se fija el texto refundido, coordinado y sistematizado del D.L. 2.763/79 y de las leyes 18.933 y 18.469; lo establecido en los Decretos Supremos Nos 140/04, Reglamento Orgánico de los Servicios de Salud; la Resolución Exenta RA 116395/343/2024 de fecha 12/08/2024 del SSMOCC., la cual nombra Director del Hospital San José de Melipilla al suscrito; lo dispuesto por las Resoluciones 10/2017, 7/2019 y 8/2019 ambas de la Contraloría General de la República, y,")

    # CONSIDERANDO Section
    doc.add_heading("CONSIDERANDO", level=1)
    considerando_items = [
        "Que dada la alta complejidad que caracteriza al Hospital San José de Melipilla, obliga a efectuar mejoras constantes y permanentes a fin de brindar a toda nuestra comunidad el desarrollo de diversas funciones con alta calidad que el sistema público puede brindar.",
        "Que, el Hospital de San José de Melipilla perteneciente a la red de salud del Servicio de Salud Metropolitano Occidente, tiene como misión otorgar una atención integral, oportuna y resolutiva a las personas y familias de la provincia de Melipilla y sus alrededores, con un equipo de salud competente, comprometido y solidario, entregando un servicio de calidad y seguridad, en coordinación con la red asistencial;",
        "Que, dada la naturaleza del Establecimiento, la atención de los beneficiarios requiere una oportuna e inmediata resolución, que no puede en caso alguno diferirse en el tiempo, lo que nos compromete a disponer en forma constante, continua y permanente de los servicios necesarios para responder adecuadamente a la demanda asistencial y administrativa a su población beneficiaria.",
        ("Que, existe la necesidad de un ", "suministro de insumos y accesorios para terapia de presión negativa con equipos en comodato", " a fin de entregar una prestación de salud integral y oportuna a los usuarios del Hospital de San José de Melipilla, y de esta manera dar cumplimiento con el tratamiento de los pacientes."), # Tuple for bold
        "Que corresponde asegurar la transparencia en este proceso y conocer las condiciones de oferta imperantes en el mercado bajo la modalidad de la licitación pública en el sistema de compras y contratación públicas establecido en la Ley Nº 19.886 y su Reglamento.",
        "Que, considerando los montos de la contratación y en virtud de lo establecido en las resoluciones N°7/2019 y 16/2020 de la Contraloría General de la República, la presenta contratación no está sometida al trámite de toma de razón.",
        "Que revisado el catálogo de bienes y servicios ofrecidos en el sistema de información Mercado Público, se ha verificado la ausencia de contratos marcos vigentes para el servicio antes mencionado.",
        "Que, en consecuencia y en mérito de lo expuesto, para esta contratación se requiere llamar a licitación pública, debiendo esta regularse por la Bases Administrativas, Técnicas, Formularios y Anexos que se aprueban a través del presente acto administrativo.",
        "Que, en razón de lo expuesto y la normativa vigente;",
    ]

    for i, item in enumerate(considerando_items):
        p = doc.add_paragraph()
        if isinstance(item, str):
             p.add_run(item)
        elif isinstance(item, tuple):
             p.add_run(item[0])
             p.add_run(item[1]).bold = True
             if len(item) > 2:
                  p.add_run(item[2])

        aplicar_numeracion(p, num_id_considerando, nivel=0)


    # RESOLUCIÓN Section
    doc.add_heading("RESOLUCIÓN", level=1)
    resolucion_items = [
        (("LLÁMASE", "bold"), " a Licitación Pública Nacional a través del Portal Mercado Público, para la compra de ", ("Suministro de Insumos y Accesorios para Terapia de Presión Negativa con Equipos en Comodato", "bold"), " para el Hospital San José de Melipilla."),
        (("ACOGIENDOSE", "bold"), " al Art.º 25 del decreto 250 que aprueba el reglamento de la ley Nº 19.886 de las Bases sobre Contratos Administrativos de Suministros y Prestación de Servicios, se reduce el tiempo de publicación de las bases en el portal de Mercado Público de 20 a 10 días, ya que se trata de la contratación de bienes o servicios de simple y objetiva especificación, y que conlleva un esfuerzo menor en la preparación de ofertas."),
        (("APRUÉBENSE", "bold"), " las bases administrativas, técnicas y anexos N.º 1, 2, 3, 4, 5, 6, 7, 8 y 9 desarrollados para efectuar el llamado a licitación, que se transcriben a continuación:"),
    ]

    for item in resolucion_items:
         p = doc.add_paragraph()
         if isinstance(item, tuple):
             for part in item:
                 if isinstance(part, tuple):
                      p.add_run(part[0]).bold = True
                 else:
                      p.add_run(part)
         else:
             p.add_run(item) # Should not happen based on items list

         aplicar_numeracion(p, num_id_resolucion, nivel=0)

    # BASES ADMINISTRATIVAS Section
    # Add a section break if you want this to potentially start on a new page
    # doc.add_section(WD_SECTION_START.NEW_PAGE)
    doc.add_heading("BASES ADMINISTRATIVAS PARA EL SUMINISTRO DE INSUMOS Y ACCESORIOS PARA TERAPIA DE PRESIÓN NEGATIVA CON EQUIPOS EN COMODATO PARA EL HOSPITAL SAN JOSÉ DE MELIPILLA", level=0)

    doc.add_heading("Antecedentes y Plazos", level=1)

    # Antecedentes Básicos Table
    agregar_parrafo_con_texto(doc, "Antecedentes Básicos de la ENTIDAD LICITANTE", negrita=True)
    tabla_p1_datos = [
        ["Razón Social del organismo", "Hospital San José de Melipilla"],
        ["Unidad de Compra", "Unidad de Abastecimiento de Bienes y Servicios"],
        ["R.U.T. del organismo", "61.602.123-0"],
        ["Dirección", "O'Higgins #551"],
        ["Comuna", "Melipilla"],
        ["Región en que se genera la Adquisición", "Región Metropolitana"]
    ]
    crear_tabla(doc, tabla_p1_datos)

    # Antecedentes Administrativos Table
    agregar_parrafo_con_texto(doc, "Antecedentes Administrativos", negrita=True)
    tabla_p2_datos = [
        ["Nombre Adquisición", "SUMINISTRO DE INSUMOS Y ACCESORIOS PARA TERAPIA DE PRESIÓN NEGATIVA CON EQUIPOS EN COMODATO PARA EL HOSPITAL SAN JOSÉ DE MELIPILLA."],
        ["Descripción", "El Hospital requiere generar un convenio por el SUMINISTRO DE INSUMOS Y ACCESORIOS PARA TERAPIA DE PRESIÓN NEGATIVA CON EQUIPOS EN COMODATO PARA EL HOSPITAL SAN JOSÉ DE MELIPILLA, en adelante “EL HOSPITAL”. El convenio tendrá una vigencia de 36 meses."],
        ["Tipo de Convocatoria", "Abierta."],
        ["Moneda o Unidad reajustable", "Peso Chileno."],
        ["Presupuesto referencial", "$350.000.000.- (Impuestos incluidos)"],
        ["Etapas del Proceso de Apertura", "Una Etapa (Etapa de Apertura Técnica y Etapa de Apertura Económica en una misma instancia)."],
        ["Opciones de pago", "Transferencia electrónica"],
        ["Tipo de Adjudicación", "Por la Totalidad"]
    ]
    crear_tabla(doc, tabla_p2_datos)

    agregar_parrafo_con_runs(doc, [
        ("* ", ""), ("Presupuesto referencial:", "underline"), (" El Hospital se reserva el derecho de aumentar, previo acuerdo entre las partes, hasta un 30% el presupuesto referencial estipulado en las presentes bases de licitación.", "")
    ])

    # Definiciones
    agregar_parrafo_con_texto(doc, "Definiciones:", negrita=True)
    definiciones = [
        ("a) Proponente u oferente:", " El proveedor o prestador que participa en el proceso de licitación mediante la presentación de una propuesta, en la forma y condiciones establecidas en las Bases."),
        ("b) Administrador o coordinador Externo del Contrato:", " Persona designada por el oferente adjudicado, quien actuará como contraparte ante el Hospital."),
        ("c) Días Hábiles:", " Son todos los días de la semana, excepto los sábados, domingos y festivos."),
        ("d) Días Corridos:", " Son los días de la semana que se computan uno a uno en forma correlativa. Salvo que se exprese lo contrario, los plazos de días señalados en las presentes bases de licitación son días corridos. En caso que el plazo expire en días sábados, domingos o festivos se entenderá prorrogados para el día hábil siguiente."),
        ("e) Administrador del Contrato y/o Referente Técnico:", " Es el funcionario designado por el Hospital para supervisar la correcta ejecución del contrato, solicitar órdenes de compra, validar prefacturas, gestionar multas y/o toda otra labor que guarde relación con la ejecución del contrato."),
        ("f) Gestor de Contrato:", " Es el funcionario a cargo de la ejecución del presente proceso de Licitación, desde la publicación de las Bases hasta la generación del contrato en formato documental, como la elaboración de ficha en la plataforma “gestor de contrato” en el portal de mercado público, además de ser el responsable de dar seguimiento y cumplimiento a los procesos y plazos establecidos.")
    ]
    for label, desc in definiciones:
        p = doc.add_paragraph()
        p.add_run(label).bold = True
        p.add_run(desc)


    # RANGO Table
    tabla_licitaciones_datos = [
        ["RANGO (en UTM)", "TIPO LICITACION PUBLICA", "PLAZO PUBLICACION EN DIAS CORRIDOS"],
        ["<100", "L1", "5"],
        ["<=100 y <1000", "LE", "10, rebajable a 5"],
        ["<=1000 y <2000", "LP", "20, rebajable a 10"],
        ["<=2000 y <5000", "LQ", "20, rebajable a 10"],
        ["<=5000", "LR", "30"]
    ]
    crear_tabla(doc, tabla_licitaciones_datos)

    # Etapas y Plazos Table
    agregar_parrafo_con_texto(doc, "Etapas y Plazos.", negrita=True)
    tabla_plazos_datos = [
        ["VIGENCIA DE LA PUBLICACION 10 DIAS CORRIDOS", "", ""],
        ["Consultas", "Hasta las 15:00 Horas de 4º (cuarto) día corrido de publicada la Licitación.", "Deben ser ingresadas al portal www.mercadopublico.cl"],
        ["Respuestas a Consultas", "Hasta las 17:00 Horas de 7º (séptimo) día corrido de publicada la Licitación.", "Deben ser ingresadas al portal www.mercadopublico.cl"],
        ["Aclaratorias", "Hasta 1 días corrido antes del cierre de recepción de ofertas.", "Deben ser ingresadas al portal www.mercadopublico.cl"],
        ["Recepción de ofertas", "Hasta las 17:00 Horas de 10º (décimo) día corrido de publicada la Licitación.", "Deben ser ingresadas al portal www.mercadopublico.cl"],
        ["Evaluación de las Ofertas", "Máximo 40 días corridos a partir del cierre de la Licitación.", ""],
        ["Plazo Adjudicaciones", "Máximo 20 días corridos a partir de la fecha del acta de evaluación de las ofertas.", ""],
        ["Suscripción de Contrato", "Máximo de 20 días hábiles desde la Adjudicación de la Licitación.", ""],
        ["Consideración", "Los plazos de días establecidos en la cláusula 3, Etapas y Plazos, son de días corridos, excepto el plazo para emitir la orden de compra, el que se considerará en días hábiles, entendiéndose que son inhábiles los sábados, domingos y festivos en Chile, sin considerar los feriados regionales.", ""]
    ]
    tabla_plazos = crear_tabla(doc, tabla_plazos_datos)
    tabla_plazos.cell(0, 0).merge(tabla_plazos.cell(0, 2))
    tabla_plazos.cell(0, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    tabla_plazos.cell(0, 0).paragraphs[0].runs[0].bold = True

    # Consultas, Aclaraciones y Modificaciones
    doc.add_heading("Consultas, Aclaraciones y modificaciones a las bases.", level=2)
    consultas_items = [
        [("Las consultas de los participantes se deberán realizar únicamente a través del portal ", ""),
         (" www.mercadopublico.cl", "bold"),
         (" conforme el cronograma de actividades de esta licitación señalado en el punto 3 precedente. A su vez, las respuestas y aclaraciones estarán disponibles a través del portal de Mercado Público, en los plazos indicados en el cronograma señalado precedentemente, información que se entenderá conocida por todos los interesados desde el momento de su publicación.", "")],
        [("No serán admitidas las consultas formuladas fuera de plazo o por un conducto diferente al señalado.", "bold")],
        [("“EL HOSPITAL” realizará las aclaraciones a las Bases comunicando las respuestas a través del Portal Web de Mercado Público, sitio ", ""),
         ("http://www.mercadopublico.cl.", "bold")],
        [("Las aclaraciones, derivadas de este proceso de consultas, formarán parte integrante de las Bases, teniéndose por conocidas y aceptadas por todos los participantes aun cuando el oferente no las hubiere solicitado, por lo que los proponentes no podrán alegar desconocimiento de las mismas.", "")],
        [("“EL HOSPITAL” podrá modificar las presentes bases y sus anexos previa autorización por acto administrativo, durante el periodo de presentación de las ofertas, hasta antes de fecha de cierre de recepción de ofertas. Estas modificaciones, que se lleven a cabo, serán informadas a través del portal ", ""),
         ("www.mercadopublico.cl.", "bold")],
        [("Estas consultas, aclaratorias y modificaciones formaran parte integra de las bases y estarán vigentes desde la total tramitación del acto administrativo que las apruebe. Junto con aprobar las modificaciones, deberá establecer un nuevo plazo prudencial cuando lo amerite para el cierre o recepción de las propuestas, a fin de que los potenciales oferentes puedan adecuar sus ofertas.", "")],
    ]
    for item in consultas_items:
        agregar_parrafo_con_runs(doc, item) # No bullet style explicitly requested by OCR layout


    # Requisitos Mínimos para Participar
    doc.add_heading("Requisitos Mínimos para Participar.", level=2)
    requisitos_items = [
        "No haber sido condenado por prácticas antisindicales, infracción a los derechos fundamentales del trabajador o por delitos concursales establecidos en el Código Penal dentro de los dos últimos años anteriores a la fecha de presentación de la oferta, de conformidad con lo dispuesto en el artículo 4º de la ley N° 19.886.",
        "No haber sido condenado por el Tribunal de Defensa de la Libre Competencia a la medida dispuesta en la letra d) del artículo 26 del Decreto con Fuerza de Ley N°1, de 2004, del Ministerio de Economía, Fomento y Reconstrucción, que Fija el texto refundido, coordinado y sistematizado del Decreto Ley N° 211, de 1973, que fija normas para la defensa de la libre competencia, hasta por el plazo de cinco años contado desde que la sentencia definitiva quede ejecutoriada.",
        "No ser funcionario directivo de la respectiva entidad compradora; o una persona unida a aquél por los vínculos de parentesco descritos en la letra b) del artículo 54 de la ley N° 18.575; o una sociedad de personas de las que aquél o esta formen parte; o una sociedad comandita por acciones o anónima cerrada en que aquélla o esta sea accionista; o una sociedad anónima abierta en que aquél o esta sean dueños de acciones que representen el 10% o más del capital; o un gerente, administrador, representante o director de cualquiera de las sociedades antedichas.",
        "Tratándose exclusivamente de una persona jurídica, no haber sido condenada conforme a la ley N° 20.393 a la pena de prohibición de celebrar actos y contratos con el Estado, mientras esta pena esté vigente.",
    ]
    for i, texto in enumerate(requisitos_items):
        p = doc.add_paragraph()
        p.add_run(texto)
        aplicar_numeracion(p, num_id_requisitos, nivel=0) # Numbered 1, 2, 3, 4

    p_req_online = doc.add_paragraph()
    p_req_online.add_run("A fin de acreditar el cumplimiento de dichos requisitos, los oferentes deberán presentar una “Declaración jurada de requisitos para ofertar”, la cual será generada completamente en línea a través de ")
    p_req_online.add_run("www.mercadopublico.cl").bold = True
    p_req_online.add_run(" en el módulo de presentación de las ofertas. Sin perjuicio de lo anterior, la entidad licitante podrá verificar la veracidad de la información entregada en la declaración, en cualquier momento, a través de los medios oficiales disponibles.")
    # This paragraph is not numbered in the OCR list structure

    p_req_inadmisible = doc.add_paragraph()
    p_req_inadmisible.add_run("En caso de que los antecedentes administrativos solicitados en esta sección no sean entregados y/o completados en forma correcta y oportuna, se desestimará la propuesta, no será evaluada y será declarada ")
    p_req_inadmisible.add_run("inadmisible").bold = True
    p_req_inadmisible.add_run(".")
    # This paragraph is not numbered in the OCR list structure


    # Instrucciones para la Presentación de Ofertas
    doc.add_heading("Instrucciones para la Presentación de Ofertas.", level=2)
    tabla_ofertas_data = [
        ["Presentar Ofertas\npor Sistema.", "Obligatorio."],
        ["Anexos Administrativos.", [
            [("Anexo N° 1 Identificación del Oferente.", "bold")],
            [("Anexo N° 2 Declaración Jurada de Habilidad.", "bold")],
            [("Anexo N° 3 Declaración Jurada de Cumplimiento de Obligaciones Laborales y Previsionales.", "bold")],
            [("Declaración jurada online:", "bold"), (" Los oferentes deberán presentar una ", ""), ("“Declaración jurada de requisitos para ofertar”", "bold"), (", la cual será generada completamente en línea a través de ", ""), ("www.mercadopublico.cl", "bold"), (" en el módulo de presentación de las ofertas.", "")],
            [("Unión Temporal de Proveedores (UTP):", "bold"), (" Solo en el caso de que la oferta sea presentada por una unión temporal de proveedores deberán presentar obligatoriamente la siguiente documentación en su totalidad, en caso contrario, ésta no será sujeta a aclaración y la oferta será declarada ", ""), ("inadmisible", "bold"), (".", "")],
            [("Anexo N°4. Declaración para Uniones Temporales de Proveedores:", "bold"), (" Declaración para Uniones Temporales de Proveedores: Debe ser presentado por el miembro de la UTP que presente la oferta en el Sistema de Información y quien realiza la declaración a través de la “Declaración jurada de requisitos para ofertar” electrónica presentada junto a la oferta.", "")],
            ["Las ofertas presentadas por una Unión Temporal de Proveedores (UTP) deberán contar con un apoderado, el cual debe corresponder a un integrante de la misma, ya sea persona natural o jurídica. En el caso que el apoderado sea una persona jurídica, ésta deberá actuar a través de su representante legal para ejercer sus facultades."],
            [("En caso de no presentarse debidamente la declaración jurada online constatando la ausencia de conflictos de interés e inhabilidades por condenas, o no presentarse el Anexo N°4, la oferta será declarada ", ""), ("inadmisible", "bold"), (".", "")]
        ]],
        ["Anexos Económicos.", [
            [("Anexo N°5: Oferta económica", "bold")],
            ["El anexo referido debe ser ingresado a través del sistema www.mercadopublico.cl en la sección Anexos Económicos."],
            [("En caso de que no se presente debidamente el Anexo N°5 “Oferta económica”, la oferta será declarada ", ""), ("inadmisible", "bold"), (".", "")]
        ]],
        ["Anexos Técnicos.", [
            [("Anexo N°6: Evaluación Técnica", "bold")],
            [("Anexo N°7: Ficha Técnica", "bold")],
            [("Anexo N°8: Plazo de Entrega", "bold")],
            [("Anexo N°9: Servicio Post-venta", "bold")],
            ["Los anexos referidos deben ser ingresados a través del sistema www.mercadopublico.cl, en la sección Anexos Técnicos."],
            [("En el caso que no se presente debidamente los Anexos N°7, N°8 y N°9 la oferta será declarada ", ""), ("inadmisible", "bold"), (".", "")]
        ]]
    ]

    tabla_instrucciones = doc.add_table(rows=len(tabla_ofertas_data), cols=2, style='Table Grid')
    for r_idx, row_data in enumerate(tabla_ofertas_data):
        for c_idx, cell_content in enumerate(row_data):
            if isinstance(cell_content, str):
                tabla_instrucciones.cell(r_idx, c_idx).text = cell_content
            elif isinstance(cell_content, list):
                agregar_contenido_celda(tabla_instrucciones, r_idx, c_idx, cell_content)

    centrar_verticalmente_tabla(tabla_instrucciones)


    # Observaciones (Continuation)
    doc.add_heading("Observaciones", level=3)
    observaciones_cont_items = [
        [("Los oferentes deberán presentar su oferta a través de su cuenta en el Sistema de Información ", ""), ("www.mercadopublico.cl", "bold"), (". De existir discordancia entre el oferente o los antecedentes de su oferta y la cuenta a través de la cual la presenta, esta no será evaluada, siendo desestimada del proceso y declarada como ", ""), ("inadmisible", "bold"), (".", "")],
        [("Las únicas ofertas válidas serán las presentadas a través del portal ", ""), ("www.mercadopublico.cl", "bold"), (", en la forma en que se solicita en estas bases. No se aceptarán ofertas que se presenten por un medio distinto al establecido en estas Bases, a menos que se acredite la indisponibilidad técnica del sistema, de conformidad con el artículo 62 del Reglamento de la Ley de Compras. Será responsabilidad de los oferentes adoptar las precauciones necesarias para ingresar oportuna y adecuadamente sus ofertas.", "")],
        ["Los oferentes deben constatar que el envío de su oferta a través del portal electrónico de compras públicas haya sido realizado con éxito, incluyendo el previo ingreso de todos los formularios y anexos requeridos completados de acuerdo con lo establecido en las presentes bases. Debe verificar que los archivos que se ingresen contengan efectivamente los anexos solicitados."],
        ["Asimismo, se debe comprobar siempre, luego de que se finalice la última etapa de ingreso de la oferta respectiva, que se produzca el despliegue automático del “Comprobante de Envío de Oferta\" que se entrega en dicho Sistema, el cual puede ser impreso por el proponente para su resguardo. En dicho comprobante será posible visualizar los anexos adjuntos, cuyo contenido es de responsabilidad del oferente."],
        ["El hecho de que el oferente haya obtenido el “Comprobante de envío de ofertas\" señalado, únicamente acreditará el envío de ésta a través del Sistema, pero en ningún caso certificará la integridad o la completitud de ésta, lo cual será evaluado por la comisión evaluadora. En caso de que, antes de la fecha de cierre de la licitación, un proponente edite una oferta ya enviada, deberá asegurarse de enviar nuevamente la oferta una vez haya realizado los ajustes que estime, debiendo descargar un nuevo Comprobante."],
        [("Si la propuesta económica subida al portal, presenta diferencias entre el valor del anexo económico solicitado y el valor indicado en la línea de la plataforma ", ""), ("www.mercadopublico.cl", "bold"), (", prevalecerá la oferta del anexo económico solicitado en bases. Sin embargo, el Hospital San José de Melipilla, podrá solicitar aclaraciones de las ofertas realizadas a través del portal.", "")]
    ]
    for item in observaciones_cont_items:
         agregar_parrafo_con_runs(doc, item)

    # Antecedentes legales para poder ser contratado
    doc.add_heading("Antecedentes legales para poder ser contratado.", level=2)
    tabla_legal_data = [
        ["Si el oferente\nes Persona\nNatural", [
            [("Inscripción (en estado hábil) en el Registro electrónico oficial de contratistas de la Administración, en adelante “", ""), ("Registro de Proveedores", "bold"), ("”.", "")],
            [("Anexo N°3. Declaración Jurada de Cumplimiento de Obligaciones Laborales y Previsionales.", "bold")],
            ["Todos los Anexos deben ser firmados por la persona natural respectiva."],
            ["Fotocopia de su cédula de identidad."]
        ], "Acreditar en\nel Registro de\nProveedores"],
        ["Si el oferente\nno es\nPersona\nNatural", [
            [("Inscripción (en estado hábil) en el Registro de Proveedores.", "bold")],
            ["Certificado de Vigencia del poder del representante legal, con una antigüedad no superior a 60 días corridos, contados desde la fecha de notificación de la adjudicación, otorgado por el Conservador de Bienes Raíces correspondiente o, en los casos que resulte procedente, cualquier otro antecedente que acredite la vigencia del poder del representante legal del oferente, a la época de presentación de la oferta."],
            ["Certificado de Vigencia de la Sociedad con una antigüedad no superior a 60 días corridos, contados desde la fecha de notificación de la adjudicación, o el antecedente que acredite la existencia jurídica del oferente."],
             [("Anexo N°3. Declaración Jurada de Cumplimiento de Obligaciones Laborales y Previsionales.", "bold")],
            ["Todos los Anexos deben ser firmados por el representante legal de la persona jurídica."]
        ], "Acreditar en\nel Registro de\nProveedores"]
    ]
    tabla_legal = doc.add_table(rows=len(tabla_legal_data), cols=3, style='Table Grid')
    for r_idx, row_data in enumerate(tabla_legal_data):
         tabla_legal.cell(r_idx, 0).text = row_data[0]
         agregar_contenido_celda(tabla_legal, r_idx, 1, row_data[1])
         tabla_legal.cell(r_idx, 2).text = row_data[2]

    centrar_verticalmente_tabla(tabla_legal)


    # Observaciones (Continuation)
    doc.add_heading("Observaciones.", level=3)
    observaciones_legal_items = [
        "Los antecedentes legales para poder ser contratado, sólo se requerirán respecto del adjudicatario y deberán estar disponibles en el Registro de Proveedores.",
        "Lo señalado en el párrafo precedente no resultará aplicable a la garantía de fiel cumplimiento de contrato, la cual podrá ser entregada físicamente en los términos que indican las presentes bases en aquellos casos que aplique su entrega.",
        "En los casos en que se otorgue de manera electrónica, deberá ajustarse a la ley N° 19.799 sobre documentos electrónicos, firma electrónica y servicios de certificación de dicha firma, y remitirse en la forma señalada en la cláusula 8.2 de estas bases.",
        [("Si el respectivo proveedor no entrega la totalidad de los antecedentes requeridos para ser contratado, dentro del plazo fatal de 10 días hábiles administrativos contados desde la notificación de la resolución de adjudicación o no suscribe el contrato en los plazos establecidos en estas bases, la entidad licitante podrá readjudicar de conformidad a lo establecido en la ", ""), ("cláusula 9 letra i", "bold"), (" de las presentes bases. Además, tales incumplimientos darán origen al cobro de la garantía de seriedad de la oferta, si la hubiere.", "")]
    ]
    for item in observaciones_legal_items:
        if isinstance(item, str):
            agregar_parrafo_con_texto(doc, item)
        elif isinstance(item, list):
            agregar_parrafo_con_runs(doc, item)

    # Inscripción en el Registro de Proveedores
    doc.add_heading("Inscripción en el Registro de Proveedores.", level=2)
    inscripcion_rp_items = [
        "En caso de que el proveedor que resulte adjudicado no se encuentre inscrito en el Registro Electrónico Oficial de Contratistas de la Administración (Registro de Proveedores), deberá inscribirse dentro del plazo de 15 días hábiles, contados desde la notificación de la resolución de adjudicación.",
        "Tratándose de los adjudicatarios de una Unión Temporal de Proveedores, cada integrante de ésta deberá inscribirse en el Registro de Proveedores, dentro del plazo de 15 días hábiles, contados desde la notificación de la resolución de adjudicación."
    ]
    for item in inscripcion_rp_items:
         agregar_parrafo_con_texto(doc, item)

    # Naturaleza y Monto de las Garantías
    doc.add_heading("Naturaleza y Monto de las Garantías.", level=2)
    doc.add_heading("Garantía de Seriedad de la Oferta.", level=3)
    garantia_seriedad_items = [
        [("El oferente deberá presentar junto a su oferta una o más garantías, equivalentes en total, al monto que indique la entidad licitante, la que corresponde al monto de $200.000. Si el oferente presenta más de una propuesta, cada una de ellas deberá estar debidamente caucionada, en los términos indicados en la presente cláusula, mediante instrumentos separados.", "")],
        ["La(s) garantía(s) debe(n) ser entregada(s) en Oficina de Partes del Hospital San José de Melipilla, ubicada en calle O'Higgins Nº 551 comuna de Melipilla, Región Metropolitana, dentro del plazo para presentación de ofertas, si fueran en soporte de papel, en el horario hábil de atención de 8:00 a 15:00 horas. De igual manera deberán publicar en su oferta copia de la garantía con el timbre de recepción de oficina de partes del Hospital San José de Melipilla."],
        ["Si la(s) garantía(s) fuera(n) es obtenida de manera electrónica (garantía emitida por las instituciones de Garantía recíproca (IGR), Internacionalmente conocidas como SGR), se debe(n) presentar en el portal www.mercadopublico.cl, hasta la hora del cierre de la licitación."],
        [("Se aceptará cualquier tipo de instrumento de garantía que asegure su cobro de manera rápida y efectiva, pagadera a la vista y con el carácter de irrevocable, y siempre que cumpla con los requisitos dispuestos por el artículo 31 del reglamento de la ley N° 19.886 El instrumento deberá incluir la glosa que se indica “", ""), ("PARA GARANTIZAR LA SERIEDAD DE LA OFERTA EN LA LICITACIÓN PÚBLICA ID PARA LA ADQUISICIÓN DE SUMINISTRO DE INSUMOS Y ACCESORIOS PARA TERAPIA DE PRESIÓN NEGATIVA CON EQUIPOS EN COMODATO PARA EL HOSPITAL SAN JOSÉ DE MELIPILLA", "bold"), ("” que señala que se otorga para garantizar la seriedad de la oferta, singularizando el respectivo proceso de compra. En caso de que el instrumento no permita la inclusión de la glosa señalada, el oferente deberá dar cumplimiento a la incorporación de ésta en forma manuscrita en el mismo instrumento, o bien, mediante un documento anexo a la garantía. Como ejemplos de garantías se pueden mencionar los siguientes instrumentos: Boleta de Garantía, Certificado de Fianza a la Vista, Vale Vista o Póliza de Seguro, entre otros.", "")],
        ["La(s) garantía(s) deberá(n) tener como vigencia mínima 120 días corridos desde el cierre y apertura de la oferta."],
        ["Toda oferta que no acompañe la garantía de seriedad, en la forma y términos expresados, será rechazada por el Hospital San José de Melipilla."],
        ["Será responsabilidad del oferente mantener vigente la(s) garantía(s), debiendo reemplazarla si por razones sobrevinientes a su presentación, deja de cubrir la vigencia mínima exigida en esta cláusula, como por ejemplo ampliación de fecha de cierre de la licitación o del proceso de evaluación."],
        ["Como beneficiario del instrumento debe figurar la razón social y RUT de la entidad licitante, indicadas en la presente licitación, numeral N°1."],
        ["Si el instrumento que se presenta expresa su monto en unidades de fomento (UF), se considerará para determinar su equivalente en pesos chilenos (CLP), el valor de la UF a la fecha en que se realice la apertura de la oferta, considerando las variaciones en el mercado monto que debe ser detallado en peso en el mismo documento de garantía."],
        ["Esta(s) garantía(s) se otorgará(n) para caucionar la seriedad de la oferta, pudiendo ser ejecutada unilateralmente por vía administrativa por la entidad licitante, siempre que los incumplimientos sean imputables al proveedor, en los siguientes casos:"],
        ["1. Por no suscripción del contrato definitivo o se rechace la orden de compra por parte del proveedor adjudicado, si corresponde;"],
        ["2. Por la no entrega de los antecedentes requeridos para la elaboración del contrato, de acuerdo con las presentes bases, si corresponde;"],
        ["3. Por el desistimiento de la oferta dentro de su plazo de validez establecido en las presentes bases;"],
        ["4. Por la presentación de una oferta no fidedigna, manifiestamente errónea o conducente a error, y que así se justifique mediante resolución fundada del órgano comprador;"],
        ["5. Por la no inscripción en el Registro de Proveedores dentro de los plazos establecidos en las presentes bases;"],
        ["6. Por la no presentación oportuna de la garantía de fiel cumplimiento del contrato, en el caso del proveedor adjudicado."]
    ]
    for item in garantia_seriedad_items:
        if isinstance(item, str):
             agregar_parrafo_con_texto(doc, item)
        elif isinstance(item, list):
             agregar_parrafo_con_runs(doc, item) # Handle potential runs


    doc.add_heading("8.1.1 Forma y oportunidad de restitución de la seriedad de la oferta", level=4)
    garantia_restitucion_items = [
        "En el caso del oferente adjudicado, la garantía de seriedad de la oferta estará disponible una vez tramitada completamente la firma del contrato, veinte (20) días hábiles después de adjudicada la Licitación y contra entrega de la garantía de fiel cumplimiento del contrato.",
        [("En caso del oferente no adjudicado, la garantía de seriedad de la oferta estará disponible previa solicitud vía correo electrónico a: ", ""), ("garantias.hsjm@hospitaldemelipilla.cl", "bold"), (", con copia a : ", ""), ("manuel.lara@hospitaldemelipilla.cl", "bold"), (" para su retiro en el departamento de tesorería del Hospital San José de Melipilla, en el siguiente horario: de lunes a viernes desde las 09:00 a 13:00 horas.", "")],
        ["Para el retiro de la garantía deberá presentarse poder simple timbrado por la persona natural o jurídica, fotocopia de la cédula de identidad de la persona que retira y el Rut la persona natural o jurídica."]
    ]
    for item in garantia_restitucion_items:
        if isinstance(item, str):
             agregar_parrafo_con_texto(doc, item)
        elif isinstance(item, list):
             agregar_parrafo_con_runs(doc, item)

    doc.add_heading("Garantía de Fiel Cumplimiento de Contrato.", level=3)
    garantia_fiel_cumplimiento_items = [
        "Para garantizar el fiel y oportuno cumplimiento del contrato, el adjudicado debe presentar una o más garantías de la misma naturaleza, equivalentes en total, al porcentaje del 5% del valor total del contrato adjudicado.",
        ["La(s) garantía(s) debe(n) ser entregada(s) en la dirección de la entidad licitante indicada: Oficina de Partes del Hospital San José de Melipilla, ubicado en calle O'Higgins Nº 551 comuna de Melipilla, Región Metropolitana, dentro de los 10 días hábiles contados desde la notificación de la adjudicación en horario de 8:00 a 14:00 horas."],
        [("Si la(s) garantía(s) fuera(n) en soporte electrónico (garantía emitida por las instituciones de Garantía recíproca (IGR), Internacionalmente conocidas como SGR), se deberá enviar al correo electrónico ", ""), ("garantias.hsjm@hospitaldemelipilla.cl", "bold"), (", si no se presenta esta garantía en tiempo y forma, el Hospital San José de Melipilla podrá hacer efectiva la garantía de seriedad de la oferta y dejar sin efecto administrativamente la adjudicación, sin perjuicio de otros derechos.", "")],
        [("Se aceptará cualquier tipo de instrumento de garantía que asegure su cobro de manera rápida y efectiva, pagadera a la vista y con el carácter de irrevocable, y siempre que cumpla con los requisitos dispuestos por el artículo 68 del reglamento de la ley N°19.886. El instrumento deberá incluir la glosa: \"", ""), ("Para garantizar el fiel cumplimiento del contrato denominado: ADQUISICIÓN DE SUMINISTRO DE INSUMOS Y ACCESORIOS PARA TERAPIA DE PRESIÓN NEGATIVA CON EQUIPOS EN COMODATO PARA EL HOSPITAL SAN JOSÉ DE MELIPILLA ID ", "bold"), ("______", "bold"), (" y/o de las obligaciones laborales y sociales del adjudicatario\". En caso de que el instrumento no permita la inclusión de la glosa señalada, el oferente deberá dar cumplimiento a la incorporación de ésta en forma manuscrita en el mismo instrumento, o bien, mediante un documento anexo a la garantía. Como ejemplos de garantías se pueden mencionar los siguientes instrumentos: Boleta de Garantía, Certificado de Fianza a la Vista, Vale Vista o Póliza de Seguro, entre otros.", "")],
        ["El adjudicatario podrá constituir la garantía, tal como dispone el artículo 68, inciso tercero, del decreto N°250, de 2004, del Ministerio de Hacienda."],
        ["La(s) garantía(s) deberá(n) tener una vigencia mínima de 120 días corridos posteriores al término de la vigencia del contrato."],
        ["Como beneficiario del instrumento debe figurar la razón social y RUT de la entidad licitante, datos indicados en la cláusula 1 de las bases."],
        ["En caso de cobro de esta garantía, derivado del incumplimiento de las obligaciones contractuales del adjudicatario indicadas en las presentes bases, éste deberá reponer la garantía por igual monto y por el mismo plazo de vigencia que la que reemplaza en un plazo de 15 días hábiles, contados desde la notificación de cobro."],
        ["Será responsabilidad del adjudicatario mantener vigente la garantía de fiel cumplimiento, al menos hasta 120 días corridos después de culminado el contrato. Mientras se encuentre vigente el contrato, las renovaciones de esta garantía serán de exclusiva responsabilidad del proveedor."],
        [("La restitución de esta garantía será realizada una vez que se haya cumplido su fecha de vencimiento, en los términos indicados en la presente base, y su retiro será obligación y responsabilidad exclusiva del contratado previa solicitud por correo electrónico a: ", ""), ("garantias.hsjm@hospitaldemelipilla.cl", "bold"), (" con copia a ", ""), ("manuel.lara@hospitaldemelipilla.cl", "bold"), (" retiro podrá efectuarse en el siguiente horario: de lunes a viernes de 09:00 horas hasta las 16:00 horas. Lo anterior previa confirmación por parte del Establecimiento. Para el retiro de la garantía deberá presentar un poder simple timbrado por la empresa, fotocopia de la cédula de identidad de la persona que retira y el Rut de la empresa, siempre que no existan observaciones pendientes.", "")],
        ["Cabe señalar que toda clase de garantías o cauciones que se constituyan en el contexto de esta cláusula, se enmarcan de acuerdo a lo dispuesto por el artículo 11 de la Ley N°19.886, a partir de lo cual se asegurará el fiel y oportuno cumplimiento del contrato, el pago de las obligaciones laborales y sociales con los trabajadores de los contratantes, y permanecerán vigentes hasta 120 días corridos después de culminado el contrato. Asimismo, con cargo a estas mismas cauciones podrán hacerse efectivas las multas y demás sanciones que afecten a los contratistas adjudicados."]
    ]
    for item in garantia_fiel_cumplimiento_items:
        if isinstance(item, str):
             agregar_parrafo_con_texto(doc, item)
        elif isinstance(item, list):
             agregar_parrafo_con_runs(doc, item)


    # 1. Evaluación y adjudicación de las ofertas.
    doc.add_heading("1. Evaluación y adjudicación de las ofertas.", level=2)
    doc.add_heading("a) Comisión Evaluadora.", level=3)
    comision_eval_items = [
        "La Dirección del Hospital San José de Melipilla designa como integrantes de la Comisión de Evaluación de la propuesta a los siguientes funcionarios: el Subdirector(a) Administrativo, Subdirector(a) Médico de Atención Abierta, Subdirector(a) Médico de Atención Cerrada, Subdirector(a) de Gestión del Cuidado de Enfermería, Subdirector(a) de Gestión y Desarrollo de las Personas, Subdirector(a) de Matronería, Subdirector(a) de Análisis de Información para la Gestión, Subdirector(a) de Apoyo Clínico o sus subrogantes. Para los efectos del quórum para sesionar se requerirá un mínimo de tres miembros. Lo anterior en conformidad con lo dispuesto en el artículo 37 del Decreto Nº 250 que establece el Reglamento de la Ley Nº 19.886. Los miembros de la Comisión Evaluadora no podrán:",
        ["• Tener contactos con los oferentes, salvo en cuanto proceda alguno de mecanismos regulados por los artículos 27, 39 y 40 del reglamento de la ley N° 19.886."],
        ["• Aceptar solicitudes de reunión, de parte de terceros, sobre asuntos vinculados directa o indirectamente con esta licitación, mientras integren la Comisión Evaluadora."],
        ["• Aceptar ningún donativo de parte de terceros. Entiéndase como terceros, entre otros, a las empresas que prestan servicios de asesoría, o bien, sociedades consultoras, asociaciones, gremios o corporaciones."],
        ["La misma Comisión estudiará los antecedentes de la Propuesta y elaborará un informe fundado para el Director de este Establecimiento, quien podrá declarar, mediante resolución fundada, admisible aquellas ofertas que cumplan con los requisitos establecidos en las bases de licitación, como también podrá declarar, mediante resolución fundada, inadmisible aquellas ofertas que no cumplan los requisitos establecidos en las bases. En caso de no presentarse oferentes o cuando las ofertas no resulten convenientes para los intereses del Establecimiento, podrá declarar desierta la licitación, fundándose en razones objetivas y no discriminatorias. Esta Comisión Evaluadora podrá invitar a profesionales técnicos para colaborar en el proceso de adjudicación."]
    ]
    for item in comision_eval_items:
         if isinstance(item, str):
              agregar_parrafo_con_texto(doc, item)
         elif isinstance(item, list):
              # Add bulleted items
              if len(item) > 0 and isinstance(item[0], str) and item[0].startswith("• "):
                   agregar_parrafo_con_runs(doc, [(item[0][2:], "")], style='List Bullet') # Remove bullet and add as bullet style
              else:
                   agregar_parrafo_con_runs(doc, item) # Add other runs

    doc.add_heading("b) Consideraciones Generales.", level=3)
    consideraciones_generales_items = [
        [("Se exigirá el cumplimiento de los requerimientos establecidos en la cláusula 6, “Instrucciones para Presentación de Ofertas", "bold"), ("\", de las presentes Bases de Licitación. Aquellas ofertas que no fueran presentadas a través del portal, en los términos solicitados, se declararán como propuestas inadmisibles, por tanto, no serán consideradas en la evaluación. Lo anterior, sin perjuicio de que concurra y se acredite algunas de las causales de excepción establecidas en el artículo 62 del Reglamento de la Ley de Compras.", "")],
        ["La entidad licitante declarará inadmisible cualquiera de las ofertas presentadas que no cumplan los requisitos o condiciones establecidos en las presentes bases, sin perjuicio de la facultad de la entidad licitante de solicitar a los oferentes que salven errores u omisiones formales, de acuerdo con lo establecido en el artículo 40 del Reglamento de la Ley N°19.886 y en las presentes bases."],
        ["Los documentos solicitados por la entidad licitante deben estar vigentes a la fecha de cierre de la presentación de las ofertas indicado en la cláusula 3 de las presentes bases y ser presentados como copias simples, legibles y firmadas por el representante legal de la empresa o persona natural. Sin perjuicio de ello, la entidad licitante podrá verificar la veracidad de la información entregada por el proveedor. En el caso en que el proveedor esté inscrito y habilitado por el Registro de Proveedores, serán suficientes los antecedentes que se encuentren en dicho Registro, en la medida que se haya dado cumplimiento a las normas de actualización de documentos que establece el Registro de Proveedores."]
    ]
    for item in consideraciones_generales_items:
         agregar_parrafo_con_runs(doc, item)

    doc.add_heading("c) Subsanación de errores u omisiones formales.", level=3)
    subsanacion_items = [
        [("Una vez realizada la apertura electrónica de las ofertas, la entidad licitante podrá solicitar a los oferentes que salven errores u omisiones formales, siempre y cuando las rectificaciones de dichos vicios u omisiones no les confieran a esos oferentes una situación de privilegio respecto de los demás competidores, esto es, en tanto no se afecten los principios de estricta sujeción a las bases y de igualdad de los oferentes, y se informe de dicha solicitud al resto de los oferentes, a través del Sistema de Información ", ""), ("www.mercadopublico.cl.", "bold")],
        ["El plazo que tendrán los oferentes, en este caso para dar cumplimiento a lo solicitado por el mandante, no será inferior a las 24 horas, contadas desde la fecha de publicación de la solicitud por parte del Hospital, la que se informará a través del Sistema de información www.mercadopublico.cl. La responsabilidad de revisar oportunamente dicho sistema durante el período de evaluación recae exclusivamente en los respectivos oferentes."]
    ]
    for item in subsanacion_items:
         agregar_parrafo_con_runs(doc, item)

    doc.add_heading("d) Inadmisibilidad de las ofertas y declaración de desierta de la licitación.", level=3)
    inadmisibilidad_declaracion_items = [
        [("La entidad licitante declarará inadmisible las ofertas presentadas que no cumplan los requisitos mínimos establecidos en los Anexos N°5, N°6, N°7, N°8 y N°9 y/o las condiciones establecidas en las presentes bases de licitación, sin perjuicio de la facultad para solicitar a los oferentes que salven errores u omisiones formales de acuerdo con lo establecido en las presentes bases.", "")],
        ["La entidad licitante podrá, además, declarar desierta la licitación cuando no se presenten ofertas o cuando éstas no resulten convenientes a sus intereses. Dichas declaraciones deberán materializarse a través de la dictación de una resolución fundada y no darán derecho a indemnización alguna a los oferentes."]
    ]
    for item in inadmisibilidad_declaracion_items:
         agregar_parrafo_con_runs(doc, item)

    doc.add_heading("e) Criterios de Evaluación y Procedimiento de Evaluación de las ofertas.", level=3)
    criterios_proc_items = [
        "La evaluación de las ofertas se realizará en una etapa, utilizando criterios técnicos, económicos y administrativos.",
        [("La evaluación de las ofertas presentadas para el ", ""), ("SUMINISTRO DE INSUMOS Y ACCESORIOS PARA TERAPIA DE PRESIÓN NEGATIVA CON EQUIPOS EN COMODATO PARA EL HOSPITAL SAN JOSÉ DE MELIPILLA", "bold"), (", se regirá por las siguientes ponderaciones y criterios a evaluar:", "")],
    ]
    for item in criterios_proc_items:
         if isinstance(item, str):
              agregar_parrafo_con_texto(doc, item)
         elif isinstance(item, list):
              agregar_parrafo_con_runs(doc, item)

    # Criterios Table
    tabla_criterios_data = [
        ["CRITERIOS", "PONDERACIÓN", "EVALUADO\nSEGÚN ANEXO"],
        ["ECONÓMICO", "60%", "ANEXO N°5"],
        ["EVALUACIÓN TECNICA", "20%", "ANEXO N°6"],
        ["TÉCNICOS", "PLAZO DE ENTREGA", "10%", "ANEXO N°8"],
        ["", "SERVICIO POST-VENTA", "10%", "ANEXO N°9"]
    ]
    tabla_criterios = doc.add_table(rows=len(tabla_criterios_data), cols=3, style='Table Grid')
    for r_idx, row_data in enumerate(tabla_criterios_data):
         for c_idx, cell_content in enumerate(row_data):
             tabla_criterios.cell(r_idx, c_idx).text = cell_content
             if r_idx == 0 or (r_idx >= 2 and c_idx == 0): # Bold header and TECNICOS cell
                  tabla_criterios.cell(r_idx, c_idx).paragraphs[0].runs[0].bold = True
             if r_idx == 0: # Center header cells
                 tabla_criterios.cell(r_idx, c_idx).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
             if r_idx > 0 and c_idx > 0: # Center data cells (except first column)
                  tabla_criterios.cell(r_idx, c_idx).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    tabla_criterios.cell(2, 0).merge(tabla_criterios.cell(4, 0)) # Merge TECNICOS cell
    centrar_verticalmente_tabla(tabla_criterios)


    doc.add_heading("Cálculo del Puntaje de Evaluación:", level=3)
    calculo_puntaje_items = [
        "El Puntaje de la Evaluación Final estará dado por el siguiente polinomio:",
        "Puntaje Evaluación Final= Puntaje Evaluación Técnica + Puntaje Evaluación Económica",
        "Donde el Puntaje Evaluación Técnica = Evaluación Técnica + Plazo de Entrega + Servicio Post-Venta.",
        "Donde Puntaje Evaluación Económica = Precio." # Assuming Price is evaluated based on Anexo N°5
    ]
    for item in calculo_puntaje_items:
         agregar_parrafo_con_texto(doc, item)

    doc.add_heading("CRITERIOS DE EVALUACIÓN:", level=3)
    doc.add_heading("CRITERIOS ECONÓMICOS:", level=4)

    crit_econ_par = doc.add_paragraph()
    crit_econ_par.add_run("f) OFERTA ECONÓMICA 60%").bold = True
    crit_econ_par.add_run("\n")
    crit_econ_par.add_run("Valor ítem ofertado. Para calcular el puntaje correspondiente al precio se utilizará la siguiente fórmula: ")
    # Formula is an image or complex layout in OCR, just add text description
    crit_econ_par.add_run("Puntaje = (Min Precio entre Oferentes / Precio Oferente i) * 60. ")
    crit_econ_par.add_run("A este puntaje se le aplicará la ponderación del 60 %. El oferente deberá declarar en Anexo N°5, los valores ofertados considerando todos los gastos involucrados e impuestos que apliquen.")

    doc.add_heading("CRITERIOS DE TÉCNICOS:", level=4)
    crit_tecnicos_items = [
        [("g) EVALUACION TECNICA 20%", "bold"), (": Se evaluará según información presentada para el Anexo N°7 que deberá ser adjuntada en su oferta en el Portal de Mercado Público, junto con la pauta de evaluación del Anexo N°6. Se evaluará por producto ofertado, donde el puntaje total será el promedio de la evaluación de todos los insumos ofertados.", "")],
        [("h) PLAZO DE ENTREGA 10%", "bold"), (": Se evaluará según información presentada en el Anexo N° 8 de la presente base de licitación.", "")],
        [("i) SERVICIO POST-VENTA 10%", "bold"), (": Se evaluará según información presentada en el Anexo N° 9 de la presente base de licitación.", "")],
    ]
    for item in crit_tecnicos_items:
         p = doc.add_paragraph()
         agregar_parrafo_con_runs(p, item)


    # Adjudicación and following sections
    doc.add_heading("Adjudicación.", level=2)
    agregar_parrafo_con_texto(doc, "Se adjudicará al oferente que obtenga el mayor puntaje, en los términos descritos en las presentes bases. La presente licitación se adjudicará a través de una resolución dictada por la autoridad competente, la que será publicada en www.mercadopublico.cl, una vez que se encuentre totalmente tramitada.")

    doc.add_heading("Mecanismo de Resolución de empates.", level=2)
    agregar_parrafo_con_texto(doc, "En el evento de que, una vez culminado el proceso de evaluación de ofertas, hubiese dos o más proponentes que hayan obtenido el mismo puntaje en la evaluación final, quedando más de uno en condición de resultar adjudicado, se optará por aquella oferta que cuente con un mayor puntaje de acuerdo con la secuencia de los criterios que resulten aplicables, de acuerdo al siguiente orden: EVALUACION TECNICA, seguido por PLAZO DE ENTREGA, seguido por SERVICIO POST-VENTA, seguido por CRITERIO ECONOMICO. Finalmente, de mantenerse la igualdad, se adjudicará a aquel oferente que haya ingresado primero su propuesta en el portal Mercado Público considerándose la hora en que aquello se efectúe.")

    doc.add_heading("Resolución de consultas respecto de la Adjudicación.", level=2)
    resolucion_consultas_adj_par = doc.add_paragraph()
    resolucion_consultas_adj_par.add_run("Las consultas sobre la adjudicación deberán realizarse dentro del plazo fatal de 5 días hábiles contados desde la publicación de la resolución en el Sistema de Información ")
    resolucion_consultas_adj_par.add_run("www.mercadopublico.cl,").bold = True
    resolucion_consultas_adj_par.add_run("a través del siguiente enlace: ")
    resolucion_consultas_adj_par.add_run("http://ayuda.mercadopublico.cl").bold = True


    doc.add_heading("Readjudicación.", level=2)
    agregar_parrafo_con_texto(doc, "Si el adjudicatario se desistiere de firmar el contrato o de aceptar la orden de compra, o no cumpliese con las demás condiciones y requisitos establecidos en las presentes bases para la suscripción o aceptación de los referidos documentos, la entidad licitante podrá, junto con dejar sin efecto la adjudicación original, adjudicar la licitación al oferente que le seguía en puntaje, o a los que le sigan sucesivamente, dentro del plazo de 60 días corridos contados desde la publicación de la adjudicación original.")

    # Add section break before the next main section
    doc.add_section(WD_SECTION_START.NEW_PAGE)

    doc.add_heading("2. Condiciones Contractuales, Vigencia de las Condiciones Comerciales, Operatoria de la Licitación y Otras Cláusulas:", level=1)

    doc.add_heading("10.1 Documentos integrantes.", level=2)
    agregar_parrafo_con_texto(doc, "La relación contractual que se genere entre la entidad licitante y el adjudicatario se ceñirá a los siguientes documentos:")
    documentos_integrantes_list = [
        "i) Bases de licitación y sus anexos.",
        "ii) Aclaraciones, respuestas y modificaciones a las Bases, si las hubiere.",
        "iii) Oferta.",
        "iv) Orden de compra.",
    ]
    for item in documentos_integrantes_list:
        agregar_parrafo_con_texto(doc, item, style='List Bullet')

    agregar_parrafo_con_texto(doc, "Todos los documentos antes mencionados forman un todo integrado y se complementan recíprocamente, especialmente respecto de las obligaciones que aparezcan en uno u otro de los documentos señalados. Se deja constancia que se considerará el principio de preeminencia de las Bases.")

    doc.add_heading("10.2 Validez de la oferta.", level=2)
    validez_oferta_items = [
        "La oferta tendrá validez de ciento veinte días (120) días corridos, contados desde la fecha de apertura de la propuesta. La oferta cuyo periodo de validez sea menor que el requerido, será rechazada de inmediato.",
        "Si vencido el plazo señalado precedentemente, el Hospital San José de Melipilla no ha realizado la adjudicación, podrá solicitar a los Proponentes la prórroga de sus ofertas y garantías. Los proponentes podrán ratificar sus ofertas o desistir de ellas, formalizando su decisión mediante comunicación escrita dirigida al Hospital. Se devolverá la garantía a aquellos que no accedan a la prórroga."
    ]
    for item in validez_oferta_items:
         agregar_parrafo_con_texto(doc, item)

    doc.add_heading("10.3 Suscripción del Contrato.", level=2)
    suscripcion_contrato_items = [
        "Para suscribir el contrato o aceptar la orden de compra contemplada en el artículo 63 del reglamento de la Ley de Compras, el adjudicado deberá estar inscrito en el Registro de Proveedores.",
        "Para formalizar las adquisiciones de bienes y servicios regidas por la ley Nº 19.886, se requerirá la suscripción de un contrato, la que en este caso se verá reflejada por la sola aceptación de la respectiva Orden de Compras.",
        "El respectivo contrato deberá suscribirse dentro de los 20 días hábiles siguientes a la notificación de la resolución de adjudicación totalmente tramitada. Asimismo, cuando corresponda, la orden de compra que formaliza la adquisición deberá ser aceptada por el adjudicatario dentro de ese mismo plazo.",
        "Si por cualquier causa que no sea imputable a la entidad licitante, el contrato no se suscribe dentro de dicho plazo, o no se acepta la orden de compra que formaliza la adquisición dentro de ese mismo término, se entenderá desistimiento de la oferta, pudiendo readjudicar la licitación al oferente que le seguía en puntaje, o a los que le sigan sucesivamente.",
        "Para suscribir el contrato o aceptar la orden de compra contemplada en el artículo 63 del reglamento de la Ley de Compras, el adjudicado deberá estar inscrito en el Registro de Proveedores." # Repeated in OCR? Yes.
    ]
    for item in suscripcion_contrato_items:
         agregar_parrafo_con_texto(doc, item)

    doc.add_heading("10.4 Modificación del contrato", level=2)
    modificacion_contrato_items = [
         [("Las partes de común acuerdo podrán modificar el contrato aumentando o disminuyendo los Bienes o servicios licitados, como también se podrán pactar nuevos bienes o servicios que no alteren la naturaleza del contrato. Estas modificaciones podrán ser hasta un 30% el presupuesto disponible estipulado en las presentes bases de licitación. En el caso de aumentar los bienes o servicios contratados, la garantía fiel cumplimiento de contrato también podrá readecuarse en proporción al monto de la modificación que se suscriba según aquellos casos que apliquen. En caso de aumentar o disminuir los bienes o servicios contratados, los valores a considerar, serán aquellos ofertados en el anexo ", ""), ("oferta económica", "bold"), (".Con todo, las eventuales modificaciones que se pacten no producirán efecto alguno sino desde la total tramitación del acto administrativo que las apruebe.", "")],
    ]
    for item in modificacion_contrato_items:
        agregar_parrafo_con_runs(doc, item)

    doc.add_heading("10.5 Gastos e Impuestos", level=2)
    agregar_parrafo_con_texto(doc, "Todos los gastos e impuestos que se generen o produzcan por causa o con ocasión de este Contrato, tales como los gastos notariales de celebración de contratos y/o cualesquiera otros que se originen en el cumplimiento de obligaciones que, según las Bases, ha contraído el oferente adjudicado, serán de cargo exclusivo de éste.")

    doc.add_heading("10.6 Efectos derivados de Incumplimientos del Proveedor", level=2)
    doc.add_heading("10.6.1 Multas", level=3)
    doc.add_heading("1. Clasificación de las sanciones y reglas de aplicación:", level=4)
    agregar_parrafo_con_texto(doc, "En función de la gravedad de la infracción cometida por el adjudicatario, se le aplicarán las siguientes sanciones:")

    # a) Amonestación
    amonestacion_par = doc.add_paragraph()
    amonestacion_par.add_run("a) Amonestación:").bold = True
    amonestacion_par.add_run(" Corresponde a un registro escrito, que dejará de manifiesto cualquier falta menor cometida por el adjudicado. Se entenderá por falta menor aquella que no ponga en riesgo de forma alguna la prestación del servicio o la vida e integridad psíquica y física de los pacientes, que se vinculen a temas administrativos y técnicos y que no sea constitutiva de multa. La amonestación no estará afecta a sanción pecuniaria.")

    # b) Multa
    multa_par_intro = doc.add_paragraph()
    multa_par_intro.add_run("b) Multa:").bold = True
    multa_par_intro.add_run(" Corresponde a la sanción de cualquier falta, de gravedad leve, moderada o grave en que incurra el adjudicado, cada vez que éste no dé cumplimiento a cualquiera de las obligaciones contempladas en las presentes bases. Se expresará en Unidades Tributarias Mensuales (UTM). El monto de cada multa, dependerá de la gravedad de la infracción cometida, en este sentido las multas se clasifican en:")

    # b1) Multa Leve
    multa_leve_par = doc.add_paragraph()
    multa_leve_par.add_run("b1) Multa Leve:").bold = True
    multa_leve_par.add_run(" Sera considerada LEVE aquella situación originada por una falta de carácter menor, que no origina riesgos a las personas, ni daños a los bienes de la Institución o a su imagen. Su importe será de 3 Unidades Tributarias Mensuales (UTM). Las conductas que pueden estar afectas a multa leve son:")
    multa_leve_list = [
        "Entrega de productos con atraso de hasta dos (2) días hábiles, contados desde la fecha de entrega estipulada en el contrato, sin que estos se hubiesen re-pactado en su plazo de entrega.",
        "Conducta o trato irrespetuoso de parte del personal del Oferente adjudicado o su cadena de distribución.",
        "La acumulación de dos amonestaciones.",
        "Incumplimiento del contrato que no origine riesgos a las personas o daño a los bienes del establecimiento o a su imagen.",
    ]
    for item in multa_leve_list:
        agregar_parrafo_con_texto(doc, item, style='List Bullet')

    # b2) Multa Moderada
    multa_moderada_par = doc.add_paragraph()
    multa_moderada_par.add_run("b2) Multa Moderada:").bold = True
    multa_moderada_par.add_run(" Sera considerada MODERADA, aquella situación originada por una falta que afecte o ponga en riesgo, directa o indirectamente a personas o a la Institución o que limite significativamente la atención y calidad del servicio, pero que es factible de ser corregida. Su importe será de 6 Unidades Tributarias Mensuales (UTM). Las conductas que puedan estar afectas a multa moderada son:")
    multa_moderada_list = [
        "No aceptar la Orden de Compra dentro de los dos (4) días hábiles siguientes al envío de la orden a través del portal de Mercado Publico.",
        "Entrega de los productos con atraso de entre tres (3) y seis (6) días hábiles inclusive, contados desde la fecha de entrega estipulada en el contrato, sin que estos se hubiesen re-pactado en su plazo de entrega.",
        "Despacho de productos en lugares no autorizados por el Hospital.",
        "La acumulación de dos multas leves trimestres móviles.",
        "cualquier falta que afecte o ponga en riesgo, directa o indirectamente, a personas o a la institución, o que limite significativamente la atención y calidad del servicio, pero que es factible de ser corregida.",
    ]
    for item in multa_moderada_list:
        agregar_parrafo_con_texto(doc, item, style='List Bullet')


    # b3) Multa Grave
    multa_grave_par = doc.add_paragraph()
    multa_grave_par.add_run("b3) Multa Grave:").bold = True
    multa_grave_par.add_run(" Sera considerada GRAVE, aquella situación originada por una falta que atente, directa o indirectamente con la atención y calidad del servicio. Su importe será de 10 Unidades Tributarias Mensuales (UTM). Las conductas que pueden estar afectas a multa grave son:")
    multa_grave_list = [
        "Incumplimiento de la totalidad de lo requerido en la orden de compra.",
        "Entrega de productos con atraso de más de seis (6) días hábiles, contados desde la fecha de entrega estipulada en el contrato, sin que estos se hubiesen re-pactado en su plazo de entrega.",
        "Rechazo total de los productos por no ajustarse a las especificaciones técnicas.",
        "El incumplimiento en el recambio de los productos que presenten problemas de estabilidad, empaque, envases en mal estado, conservación inadecuada, vencida, defectuosa o dañada en un periodo máximo de cuarenta y ocho (48) horas.",
        "La acumulación de dos multas moderadas en 2 trimestres móviles",
        "Si el adjudicatario no inicia sus labores en la fecha acordada",
        "Si se acreditaren acciones y/u omisiones maliciosas y/o negligentes que comprometan la eficiencia y eficacia del servicio o la seguridad y/o bienes dispuestos por el Hospital para la correcta ejecución del convenio.",
        "No cumplir con lo señalado en los requerimientos técnicos.",
        "Vulnerar normas referidas al uso de información reservada o confidencial al que se tenga acceso en razón del servicio especializado contratado.",
        "No cumplir con los horarios establecidos en las bases",
        "No cumplir con los tiempos de respuestas establecidos en bases de licitación según sea el caso.",
        "Incumplimiento de las normas y procedimientos internos vigentes, tanto técnicas como administrativas impartidas por el Hospital, a través del administrador interno del contrato",
        "cualquier falta que atente, directa o indirectamente, contra la integridad física de los pacientes o funcionarios, que implique u obstruye la atención, calidad del servicio y/o según lo establecido en las bases.",
    ]
    for item in multa_grave_list:
        agregar_parrafo_con_texto(doc, item, style='List Bullet')

    multa_summary_items = [
        "Las referidas multas, en total, no podrán sobrepasar el 20% del valor total neto del contrato. Igualmente, no se le podrán cursar más de 6 multas totalmente tramitadas en un período de 6 meses consecutivos. En ambos casos, superado cada límite, se configurará una causal de término anticipado del contrato.",
        "Las multas deberán ser pagadas en el plazo máximo de 5 días hábiles contados desde la notificación de la resolución que aplica la multa.",
        "Cuando el cálculo del monto de la respectiva multa, convertido a pesos chilenos, resulte un número con decimales, éste se redondeará al número entero más cercano. La fecha de conversión de la UTM será la del día de emisión del respectivo acto administrativo que origina el cobro de la multa",
        "Las multas se aplicarán sin perjuicio del derecho de la entidad licitante de recurrir ante los Tribunales Ordinarios de Justicia ubicados en la ciudad de Melipilla, a fin de hacer efectiva la responsabilidad del contratante incumplidor.",
        "No procederá el cobro de las multas señaladas en este punto, si el incumplimiento se debe a un caso fortuito o fuerza mayor, de acuerdo con los artículos 45 y 1547 del Código Civil o una causa enteramente ajena a la voluntad de las partes, el cual será calificado como tal por la Entidad Licitante, en base al estudio de los antecedentes por los cuales el oferente adjudicado acredite el hecho que le impide cumplir."
    ]
    for item in multa_summary_items:
         agregar_parrafo_con_texto(doc, item)


    doc.add_heading("10.6.2 Cobro de la Garantía de Fiel Cumplimiento de Contrato", level=3)
    cobro_garantia_items = [
        "Al Adjudicatario le podrá ser aplicada la medida de cobro de la Garantía por Fiel Cumplimiento del Contrato por la entidad licitante, en los siguientes casos:",
        ["i. No pago de multas dentro de los plazos establecidos en las presentes bases y/o el respectivo contrato."],
        ["ii. Incumplimientos de las exigencias técnicas de los bienes y servicios (en caso de que hayan sido requeridos) adjudicados establecidos en el Contrato."],
        ["iii. Cualquiera de las causales señaladas en el N°10.6.3 sobre “Término Anticipado del Contrato”, a excepción del numeral 3) y numeral 16), en todas estas causales señaladas, se procederá al cobro de la garantía de fiel cumplimiento del contrato, si se hubiere exigido dicha caución en las Bases."]
    ]
    for item in cobro_garantia_items:
         if isinstance(item, str):
              agregar_parrafo_con_texto(doc, item)
         elif isinstance(item, list):
             # Use bullet for these i, ii, iii items as per OCR indentation
             agregar_parrafo_con_runs(doc, item, style='List Bullet')

    doc.add_heading("10.6.3 Término Anticipado del Contrato", level=3)
    agregar_parrafo_con_texto(doc, "El hospital está facultado para declarar administrativamente mediante resolución fundada el término anticipado del contrato, en cualquier momento, sin derecho a indemnización alguna para el adjudicado, si concurre alguna de las causales que se señalan a continuación:")

    # Termino Anticipado List (Complex - needs different levels/styles)
    termino_anticipado_items_data = [
         ("1) Por incumplimiento grave de las obligaciones contraídas por el proveedor adjudicado, cuando sea imputable a éste. Se entenderá por incumplimiento grave la no ejecución o la ejecución parcial por parte del adjudicatario de las obligaciones contractuales, descritas en las presentes Bases, sin que exista alguna causal que le exima de responsabilidad, y cuando dicho incumplimiento le genere al hospital un perjuicio en el cumplimiento de sus funciones. Alguno de estos motivos puede ser:", 0),
         (("A) La aplicación de dos o más Multas Graves en un periodo de seis meses móviles.",), 1),
         (("B) Si el proveedor fuese condenado a algún delito que tuviera pena aflictiva o tratándose de una empresa, sus socios, o en el caso de una sociedad anónima, algunos de los miembros del directorio o el gerente de la sociedad.",), 1),
         (("C) Si el proveedor delega, cede, aporta o transfiere el presente convenio a cualquier título efectúa asociaciones u otorga concesiones o subconcesiones.",), 1),
         (("D) Si la sociedad se disolviere por Quiebra o cesación de pagos del proveedor.",), 1),
         ("2) Si el adjudicado se encuentra en estado de notoria insolvencia o fuere declarado deudor en un procedimiento concursal de liquidación. En el caso de una UTP, aplica para cualquiera de sus integrantes. En este caso no procederá el término anticipado si se mejoran las cauciones entregadas o las existentes sean suficientes para garantizar el cumplimiento del contrato.", 0),
         ("3) Por exigirlo la necesidad del servicio, el interés público o la seguridad nacional.", 0),
         ("4) Registrar, a la mitad del período de ejecución contractual, con un máximo de seis meses, saldos insolutos de remuneraciones o cotizaciones de seguridad social con sus actuales trabajadores o con trabajadores contratados en los últimos 2 años.", 0),
         ("5) Si se disuelve la sociedad o empresa adjudicada, o en caso de fallecimiento del contratante, si se trata de una persona natural.", 0),
         ('6) Incumplimiento de uno o más de los compromisos asumidos por los adjudicatarios, en virtud del "Pacto de integridad" contenido en estas bases. Cabe señalar que en el caso que los antecedentes den cuenta de una posible afectación a la libre competencia, el organismo licitante pondrá dichos antecedentes en conocimiento de la Fiscalía Nacional Económica.', 0),
         ('7) Sin perjuicio de lo señalado en el “Pacto de integridad", si el adjudicatario, sus representantes, o el personal dependiente de aquél, no observaren el más alto estándar ético exigible, durante la ejecución de la licitación, y propiciaren prácticas corruptas, tales como:', 0),
         (("a.- Dar u ofrecer obsequios, regalías u ofertas especiales al personal del hospital, que pudiere implicar un conflicto de intereses, presente o futuro, entre el respectivo adjudicatario y el servicio hospitalario.",), 1),
         (("b.- Dar u ofrecer cualquier cosa de valor con el fin de influenciar la actuación de un funcionario público durante la relación contractual objeto de la presente licitación.",), 1),
         (("c.- Tergiversar hechos, con el fin de influenciar decisiones de la entidad licitante.",), 1),
         ("8) No renovación oportuna de la Garantía de Fiel Cumplimiento, según lo establecido en la cláusula 8.2 de las bases de licitación cuando aplique.", 0),
         ("9) La comprobación de la falta de idoneidad, de fidelidad o de completitud de los antecedentes aportados por el proveedor adjudicado, para efecto de ser adjudicado o contratado.", 0),
         ("10) La comprobación de que el adjudicatario, al momento de presentar su oferta contaba con información o antecedentes relacionados con el proceso de diseño de las bases, encontrándose a consecuencia de ello en una posición de privilegio en relación al resto de los oferentes, ya sea que dicha información hubiese sido conocida por el proveedor en razón de un vínculo laboral o profesional entre éste y las entidades compradoras, o bien, como resultado de prácticas contrarias al ordenamiento jurídico.", 0),
         ("11) En caso de ser el adjudicatario de una Unión Temporal de Proveedores (UTP):", 0),
         (("a. Inhabilidad sobreviniente de uno de los integrantes de la UTP en el Registro de Proveedores, que signifique que la UTP no pueda continuar ejecutando el contrato con los restantes miembros en los mismos términos adjudicados.",), 1),
         (("b. De constatarse que los integrantes de la UTP constituyeron dicha figura con el objeto de vulnerar la libre competencia. En este caso, deberán remitirse los antecedentes pertinentes a la Fiscalía Nacional Económica.",), 1),
         (("c. Retiro de algún integrante de la UTP que hubiere reunido una o más características objeto de la evaluación de la oferta.",), 1),
         (("d. Cuando el número de integrantes de una UTP sea inferior a dos y dicha circunstancia ocurre durante la ejecución del contrato.",), 1),
         (("e. Disolución de la UTP.",), 1),
         ('12) En el caso de infracción de lo dispuesto en la cláusula sobre “Cesión de contrato y Subcontratación"', 0),
         ("13) En caso de que las multas cursadas, en total, sobrepasen el 20 % del valor total contratado con impuestos incluidos o se apliquen más de 6 multas totalmente tramitadas en un periodo de 6 meses consecutivos.", 0),
         ("14) Por el no pago de las multas aplicadas.", 0),
         ("15) Por la aplicación de dos multas graves en que incurra el adjudicatario en virtud del incumplimiento de las obligaciones reguladas en las bases y del presente contrato.", 0),
         ("16) Si el Hospital San José de Melipilla cesara su funcionamiento en lugar de origen por cambio de ubicación de sus dependencias.", 0),
         ("17) Por la comprobación de la inhabilidad del adjudicatario para contratar con la Administración del Estado en portal de mercado público, durante la ejecución del presente contrato. Solo en el caso que el proveedor desde la notificación de esta situación no regularice su registro en un plazo superior a 15 días hábiles.", 0),
         ("18) Por incumplimiento de obligaciones de confidencialidad establecidas en las respectivas Bases.", 0),
    ]
    for item_data, level in termino_anticipado_items_data:
        p = doc.add_paragraph()
        if isinstance(item_data, str):
             p.add_run(item_data)
        elif isinstance(item_data, tuple): # Handle items that are tuples for runs
             for part in item_data:
                 if isinstance(part, tuple): # If a part is a (text, format) tuple
                      text, fmt = part
                      run = p.add_run(text)
                      if 'bold' in fmt: run.bold = True
                      if 'underline' in fmt: run.underline = True
                 else: # If a part is just text
                      p.add_run(str(part))
        aplicar_numeracion(p, num_id_termino_anticipado, nivel=level) # Apply numbering


    termino_anticipado_closing_items = [
        ["De concurrir cualquiera de las causales anteriormente señaladas como término anticipado del contrato, exceptuando las causales número 3 y número 16, se procederá al cobro de la garantía de fiel cumplimiento del contrato, siempre y cuando se hubiere exigido dicha caución en las Bases."],
        ["El término anticipado por incumplimientos se aplicará siguiendo el procedimiento establecido en la cláusula “sobre aplicación de Medidas derivadas de incumplimientos.”"]
    ]
    for item in termino_anticipado_closing_items:
         agregar_parrafo_con_runs(doc, item)

    doc.add_heading("10.6.4 Resciliación o término de mutuo acuerdo", level=3)
    agregar_parrafo_con_texto(doc, "Sin perjuicio de lo anterior, la entidad licitante y el respectivo adjudicatario podrán poner término al contrato en cualquier momento, de común acuerdo, sin constituir una medida por incumplimiento.")

    doc.add_heading("10.7 Procedimiento para Aplicación de Medidas derivadas de incumplimientos", level=2)
    procedimiento_incumplimiento_items = [
        "Detectada una situación que amerite la aplicación de una multa u otra medida derivada de incumplimientos contemplada en las presentes bases, o que constituya una causal de término anticipado, con excepción de la resciliación, el referente técnico o administrador del contrato notificará de ello al oferente adjudicado, informándole sobre la medida a aplicar y sobre los hechos que la fundamentan.",
        "A contar de la notificación singularizada en el párrafo anterior, el proveedor adjudicado tendrá un plazo de 5 días hábiles para efectuar sus descargos por escrito, acompañando todos los antecedentes que lo fundamenten. Vencido el plazo indicado sin que se hayan presentados descargos, la Dirección del Hospital resolverá según la naturaleza de la infracción, notificando al proveedor la resolución del caso por parte del Hospital.",
        "Si el proveedor adjudicado ha presentado sus descargos dentro del plazo establecido para estos efectos, el Hospital tendrá un plazo de 30 días hábiles, contados desde la recepción de los descargos del proveedor, para rechazarlos o acogerlos, total o parcialmente. Al respecto, el rechazo total o parcial de los descargos del respectivo proveedor deberá formalizarse a través de la dictación de una resolución fundada del hospital, en la cual deberá detallarse el contenido y las características de la medida. La indicada resolución será notificada al proveedor adjudicado.",
        "Con todo, el adjudicatario solo será responsable por hechos imputables a su incumplimiento directo y no por indisponibilidades de servicio ocasionadas por fallas ajenas a su gestión y control, lo que deberá, en todo caso, acreditarse debidamente. Sin perjuicio de lo anterior, el adjudicatario deberá adoptar medidas que ofrezcan continuidad operativa a los servicios materia de la respectiva licitación.",
        "Una vez finalizados los trámites administrativos señalados precedentemente y para el evento de que esta conlleve la aplicación de una multa o sanción, el Hospital San José de Melipilla podrá realizar el cobro de la multa o sanción que será debidamente notificado junto con el acto administrativo que lo autoriza. El monto de las multas podría ser rebajado del pago, que el Hospital deba efectuar al proveedor, en el estado de pago más próximo a la notificación del acto administrativo, pudiéndose aplicar tanto en la emisión de la orden de compra, como también en la aplicación del descuento en el pago de facturas. De no ser suficiente este monto o en caso de no existir pagos pendientes, el proveedor deberá pagar directamente al Hospital San José de Melipilla, el monto indicado en el acto administrativo previamente notificado, este pago no podrá ser superior a los 5 días hábiles desde su notificación. Si el proveedor no paga dentro de dicho plazo, se hará efectivo el cobro de la garantía de fiel cumplimiento del contrato, debiendo reponer una nueva boleta de garantía por un monto igual al original, en un plazo no superior a 5 días hábiles en caso que aplique la solicitud de dicha caución.",
        "En el caso de no reponer la boleta de garantía, el hospital podrá proceder a tramitar el termino anticipado del contrato en aquellos casos que aplique con la solicitud de dicha caución.",
        "El valor de la UTM a considerar será el equivalente a su valor en pesos del mes en el cual se aplicó la multa."
    ]
    for item in procedimiento_incumplimiento_items:
         agregar_parrafo_con_texto(doc, item)

    doc.add_heading("10.8 Emisión de la Orden de Compras", level=2)
    emision_oc_items = [
        "Las órdenes de compra se emitirán previa solicitud del administrador del contrato, quien, en función de la necesidad y demanda del servicio, realizara los pedidos correspondientes.",
        "La orden de compra sólo se emitirá en los casos que el proveedor este en estado hábil para ser contratado por el Estado de Chile y sólo se emitirá el documento a nombre del proveedor adjudicado por el Hospital.",
        "Al inicio del convenio, por registros en la plataforma y tramites del “gestor de contratos” se emitirá una orden de compras por un monto mínimo, la que solo debe ser aceptada por el proveedor, sin tramitar dicho servicio. Todo cambio respecto a este punto, será informado con la respectiva anticipación."
    ]
    for item in emision_oc_items:
         agregar_parrafo_con_texto(doc, item)

    doc.add_heading("10.9 Del Pago", level=2)
    del_pago_items = [
        "El pago se efectuará una vez que el “Hospital\" haya recibido oportunamente y a su entera satisfacción dichos bienes o servicios y desde la recepción conforme de la factura u otro instrumento de cobro.",
        "El pago será efectuado dentro de los 30 días corridos siguientes, contados desde la recepción de la factura respectiva, salvo las excepciones indicadas en el artículo 79 bis del Reglamento de la Ley N° 19.886.",
        "El proveedor solo podrá facturar los bienes o servicios efectivamente entregados y recibidos conforme por este organismo comprador, una vez que el administrador del contrato por parte del organismo comprador autorice la facturación en virtud de la recepción conforme de los bienes o servicios. “El Hospital” rechazará todas las facturas que hayan sido emitidas sin contar con la recepción conforme de los bienes o servicios y la autorización expresa de facturar por parte de éste.",
        "Para efectos del pago, el proveedor adjudicado deberá indicar en la factura el número de orden de compra, además, no podrá superar el monto de la orden de compra, de lo contrario, se cancelará la factura por \"forma\".",
        [("La factura electrónica deberá ser enviada al correo: ", ""), ("facturas.hjsm@redsalud.gov.cl", "bold"), (" con copia al correo ", ""), ("dipresrecepcion@custodium.com", "bold"), (" (En formato PDF y XML)", "")],
        "El valor del convenio se reajustará anualmente de acuerdo con la variación que haya experimentado el Índice de Precios al Consumidor IPC, obtenido del promedio de la sumatoria de los IPC de los doce meses inmediatamente anteriores al mes en que se efectúa su cálculo. Este reajuste es de exclusiva responsabilidad de la empresa adjudicada; si por alguna razón no lo aplicare, no se permitirá su cobro en forma retroactiva. Su precio se pagará conforme a lo establecido.",
        "En ningún caso procederán cobros adicionales por bienes o servicios no convenidos previamente, ni por tiempos en que el proveedor no preste los servicios.",
        "Cabe señalar que, cuando el resultado del monto a facturar resulte un número con decimales, éste se redondeará al número entero siguiente en caso de que la primera cifra decimal sea igual o superior a 5. En caso contrario el monto deberá ser redondeado al número entero anterior."
    ]
    for item in del_pago_items:
        if isinstance(item, str):
             agregar_parrafo_con_texto(doc, item)
        elif isinstance(item, list):
             agregar_parrafo_con_runs(doc, item)


    doc.add_heading("10.10 Vigencia del Contrato.", level=2)
    agregar_parrafo_con_texto(doc, "El contrato tendrá una duración de treinta y seis (36) meses contados desde la total tramitación del acto administrativo que aprueba la adjudicación o hasta que se cumpla con el monto estipulado en las presentes bases, lo que suceda primero y sin perjuicio, que por razones de buen servicio las prestaciones materia de la licitación podrían iniciarse desde el momento de la suscripción del mismo, sin que proceda pago alguno en el tiempo intermedio.")

    doc.add_heading("10.11 Administrador del Contrato y/o Referente Técnico.", level=2)
    administrador_rt_items = [
        "Con el objeto de supervisar y verificar el cumplimiento materia de la presente licitación, El Hospital designará a (la) Enfermera Supervisora(o) del Servicio de Pabellón y al Jefe(a) de Farmacia o su subrogante, para coordinar y fiscalizar la efectiva ejecución del contrato en términos administrativos.",
        [("El adjudicatario", "bold"), (" deberá nombrar un coordinador del contrato, cuya identidad deberá ser informada al Hospital.", "")],
        "En el desempeño de su cometido, el coordinador del contrato deberá, a lo menos:",
        ["1. Informar oportunamente al órgano comprador de todo hecho relevante que pueda afectar el cumplimiento del contrato."],
        ["2. Representar al proveedor en la discusión de las materias relacionadas con la ejecución del contrato."],
        ["3. Coordinar las acciones que sean pertinentes para la operación y cumplimiento de este contrato."],
        ["La designación del coordinador y todo cambio posterior deberá ser informado por el adjudicatario al responsable de administrar el contrato y/o referente técnico por parte del órgano comprador, a más tardar dentro de las 24 horas siguientes de efectuada la designación o el cambio, por medio del correo electrónico institucional del funcionario."]
    ]
    for i, item in enumerate(administrador_rt_items):
         if isinstance(item, str):
              agregar_parrafo_con_texto(doc, item)
         elif isinstance(item, list):
             if i >= 3 and i <= 5: # Apply numbering to items 1, 2, 3
                 p = agregar_parrafo_con_runs(doc, item)
                 aplicar_numeracion(p, num_id_administrador_contrato_list, nivel=0)
             else:
                 agregar_parrafo_con_runs(doc, item)


    doc.add_heading("10.12 Pacto de Integridad.", level=2)
    pacto_integridad_items = [
         "El oferente declara que, por el sólo hecho de participar en la presente licitación, acepta expresamente el presente pacto de integridad, obligándose a cumplir con todas y cada una de las estipulaciones contenidas en el mismo, sin perjuicio de las que se señalen en el resto de las bases de licitación y demás documentos integrantes. Especialmente, el oferente acepta el suministrar toda la información y documentación que sea considerada necesaria y exigida de acuerdo con las presentes bases de licitación, asumiendo expresamente los siguientes compromisos:",
         ["i. El oferente se compromete a respetar los derechos fundamentales de sus trabajadores, entendiéndose por éstos los consagrados en la Constitución Política de la República en su artículo 19, números 1º, 4º, 5º, 6º, 12º, y 16º, en conformidad al artículo 485 del Código del Trabajo. Asimismo, el oferente se compromete a respetar los derechos humanos, lo que significa que debe evitar dar lugar o contribuir a efectos adversos en los derechos humanos mediante sus actividades, bienes o servicios, y subsanar esos efectos cuando se produzcan, de acuerdo con los Principios Rectores de Derechos Humanos y Empresas de Naciones Unidas."],
         ["ii. El oferente se obliga a no ofrecer ni conceder, ni intentar ofrecer o conceder, sobornos, regalos, premios, dádivas o pagos, cualquiera fuese su tipo, naturaleza y/o monto, a ningún funcionario público en relación con su oferta, con el proceso de licitación pública, ni con la ejecución de el o los contratos que eventualmente se deriven de la misma, ni tampoco a ofrecerlas o concederlas a terceras personas que pudiesen influir directa o indirectamente en el proceso licitatorio, en su toma de decisiones o en la posterior adjudicación y ejecución del o los contratos que de ello se deriven."],
         ["iii. El oferente se obliga a no intentar ni efectuar acuerdos o realizar negociaciones, actos o conductas que tengan por objeto influir o afectar de cualquier forma la libre competencia, cualquiera fuese la conducta o acto específico, y especialmente, aquellos acuerdos, negociaciones, actos o conductas de tipo o naturaleza colusiva, en cualquiera de sus tipos o formas."],
         ["iv. El oferente se obliga a revisar y verificar toda la información y documentación, que deba presentar para efectos del presente proceso licitatorio, tomando todas las medidas que sean necesarias para asegurar su veracidad, integridad, legalidad, consistencia, precisión y vigencia."],
         ["V. El oferente se obliga a ajustar su actuar y cumplir con los principios de legalidad, probidad y transparencia en el presente proceso licitatorio."],
         ["vi. El oferente manifiesta, garantiza y acepta que conoce y respetará las reglas y condiciones establecidas en las bases de licitación, sus documentos integrantes y él o los contratos que de ellos se derivase."],
         ["vii. El oferente reconoce y declara que la oferta presentada en el proceso licitatorio es una propuesta seria, con información fidedigna y en términos técnicos y económicos ajustados a la realidad, que aseguren la posibilidad de cumplir con la misma en las condiciones y oportunidad ofertadas."],
         ["viii. El oferente se obliga a tomar todas las medidas que fuesen necesarias para que las obligaciones anteriormente señaladas sean asumidas y cabalmente cumplidas por sus empleados, dependientes, asesores y/o agentes y, en general, todas las personas con que éste o éstos se relacionen directa o indirectamente en virtud o como efecto de la presente licitación, incluidos sus subcontratistas, haciéndose plenamente responsable de las consecuencias de su infracción, sin perjuicio de las responsabilidades individuales que también procediesen y/o fuesen determinadas por los organismos correspondientes."]
    ]
    for item in pacto_integridad_items:
        if isinstance(item, str):
             agregar_parrafo_con_texto(doc, item)
        elif isinstance(item, list):
            # Attempt to manually format the Roman numeral/lettered start
            p = doc.add_paragraph()
            if item[0] and isinstance(item[0], str):
                 first_part = item[0]
                 match = re.match(r'([ivxVIX\d]+[).]\s+)(.*)', first_part) # Regex for roman, numbers, followed by ) or . and space
                 if match:
                     label = match.group(1)
                     rest_text = match.group(2)
                     p.add_run(label).bold = True # Make the label bold
                     p.add_run(rest_text)
                 else:
                     p.add_run(first_part) # Add as regular text if no match

            for remaining_part in item[1:]: # Add any subsequent parts as runs
                 if isinstance(remaining_part, tuple):
                      text, fmt = remaining_part
                      run = p.add_run(text)
                      if 'bold' in fmt: run.bold = True
                 else:
                      p.add_run(str(remaining_part))


    doc.add_heading("10.13 Comportamiento ético del Adjudicatario.", level=3)
    agregar_parrafo_con_texto(doc, "El adjudicatario que preste los servicios deberá observar, durante toda la época de ejecución del contrato, el más alto estándar ético exigible a los funcionarios públicos. Tales estándares de probidad deben entenderse equiparados a aquellos exigidos a los funcionarios de la Administración Pública, en conformidad con el Título III de la ley N° 18.575, Orgánica Constitucional de Bases Generales de la Administración del Estado.")

    doc.add_heading("10.14 Auditorías.", level=2)
    agregar_parrafo_con_texto(doc, "El adjudicatario podrá ser sometido a auditorías externas, contratadas por la entidad licitante a empresas auditoras independientes, con la finalidad de velar por el cumplimiento de las obligaciones contractuales y de las medidas de seguridad comprometidas por el adjudicatario en su oferta. Si el resultado de estas auditorías evidencia incumplimientos contractuales por parte del adjudicatario, el proveedor quedará sujeto a las medidas que corresponda aplicar la entidad licitante, según las presentes bases.")

    doc.add_heading("10.15 Confidencialidad.", level=2)
    confidencialidad_items = [
        "El adjudicatario no podrá utilizar para ninguna finalidad ajena a la ejecución del contrato, la documentación, los antecedentes y, en general, cualquier información, que haya conocido o a la que haya accedido, en virtud de cualquier actividad relacionada con el contrato.",
        "El adjudicatario, así como su personal dependiente que se haya vinculado a la ejecución del contrato, en cualquiera de sus etapas, deben guardar confidencialidad sobre los antecedentes relacionados con el proceso licitatorio y el respectivo contrato.",
        "El adjudicatario debe adoptar medidas para el resguardo de la confidencialidad de la información, reservándose el órgano comprador el derecho de ejercer las acciones legales que correspondan, de acuerdo con las normas legales vigentes, en caso de divulgación no autorizada, por cualquier medio, de la totalidad o parte de la información referida.",
        "La divulgación, por cualquier medio, de la totalidad o parte de la información referida en los párrafos anteriores, por parte del proveedor, durante la vigencia del contrato o dentro de los 5 años siguientes después de finalizado éste, podrá dar pie a que la Entidad entable en su contra las acciones judiciales que correspondan. Con todo, tratándose de bases de datos de carácter personal, la obligación de confidencialidad dura indefinidamente, de acuerdo con la Ley N°19.628, sobre Protección de la Vida Privada."
    ]
    for item in confidencialidad_items:
         agregar_parrafo_con_texto(doc, item)


    doc.add_heading("10.16 Propiedad de la Información.", level=2)
    agregar_parrafo_con_texto(doc, "La entidad licitante será la titular de todos los datos de transacciones, bitácoras (logs), parámetros, documentos electrónicos y archivos adjuntos y, en general, de las bases de datos y de toda información contenida en la infraestructura física y tecnológica que le suministre el proveedor contratado y que se genere en virtud de la ejecución de los servicios objeto de la presente licitación. El proveedor no podrá utilizar la información indicada en el párrafo anterior, durante la ejecución del contrato ni con posterioridad al término de su vigencia, sin autorización escrita de la entidad licitante. Por tal motivo, una vez que el proveedor entregue dicha información a la entidad o al finalizar la relación contractual, deberá borrarla de sus registros lógicos y físicos.")

    doc.add_heading("10.17 Saldos insolutos de remuneraciones o cotizaciones de seguridad social.", level=2)
    saldos_insolutos_items = [
        "Durante la vigencia del respectivo contrato el adjudicatario deberá acreditar que no registra saldos insolutos de obligaciones laborales y sociales con sus actuales trabajadores o con trabajadores contratados en los últimos dos años.",
        "El órgano comprador podrá requerir al adjudicatario, en cualquier momento, los antecedentes que estime necesarios para acreditar el cumplimiento de las obligaciones laborales y sociales antes señaladas.",
        "En caso de que la empresa adjudicada registre saldos insolutos de remuneraciones o cotizaciones de seguridad social con sus actuales trabajadores o con trabajadores contratados en los últimos dos años, los primeros estados de pago de los bienes y servicios de esta licitación deberán ser destinados al pago de dichas obligaciones, debiendo la empresa acreditar que la totalidad de las obligaciones se encuentran liquidadas al cumplirse la mitad del período de ejecución de las prestaciones, con un máximo de seis meses.",
        "La entidad licitante deberá exigir que la empresa adjudicada proceda a dichos pagos y le presente los comprobantes y planillas que demuestren el total cumplimiento de la obligación. El incumplimiento de estas obligaciones por parte de la empresa adjudicataria dará derecho a terminar la relación contractual, pudiendo llamarse a una nueva licitación en la que la empresa referida no podrá participar."
    ]
    for item in saldos_insolutos_items:
         agregar_parrafo_con_texto(doc, item)


    doc.add_heading("10.18 Normas laborales aplicables.", level=2)
    normas_laborales_items = [
        "El adjudicatario, en su calidad de empleador, será responsable exclusivo del cumplimiento íntegro y oportuno de las normas del Código del Trabajo y leyes complementarias, leyes sociales, de previsión, de seguros, de enfermedades profesionales, de accidentes del trabajo y demás pertinentes respecto de sus trabajadores y/o integrantes de sus respectivos equipos de trabajo.",
        "En consecuencia, el adjudicatario será responsable, en forma exclusiva, y sin que la enumeración sea taxativa, del pago oportuno de las remuneraciones, honorarios, indemnizaciones, desahucios, gratificaciones, gastos de movilización, beneficios y, en general, de toda suma de dinero que, por cualquier concepto, deba pagarse a sus trabajadores y/o integrantes de sus respectivos equipos de trabajo.",
        "El Hospital se reserva el derecho a exigir al contratista, a simple requerimiento de la contraparte técnica, y sin perjuicio de lo dispuesto en el artículo 4º de la Ley de Compras y el artículo 183-C del Código del Trabajo, un certificado que acredite el monto y estado de cumplimiento de las obligaciones laborales y previsionales emitido por la Inspección del Trabajo respectiva, o bien, por medios idóneos que garanticen la veracidad de dicho monto y estado de cumplimiento, respecto de sus trabajadores. Ello, con el propósito de hacer efectivo por parte del órgano comprador, su derecho a ser informado y el derecho de retención, consagrados en los incisos segundo y tercero del artículo 183-C del Código del Trabajo, en el marco de la responsabilidad subsidiaria derivada de dichas obligaciones laborales y previsionales, a la que alude el artículo 183-D del mismo Código.",
        "Por otra parte, se deja expresa constancia que la suscripción del contrato respectivo no significará en caso alguno que el adjudicatario, sus trabajadores o integrantes de los equipos presentados por éstos, adquieran la calidad de funcionarios públicos, no existiendo vínculo alguno de subordinación o dependencia de ellos con el órgano comprador."
    ]
    for item in normas_laborales_items:
         agregar_parrafo_con_texto(doc, item)

    doc.add_heading("10.19 Cambio de personal del proveedor adjudicado.", level=2)
    cambio_personal_items = [
        "El Hospital San José de Melipilla podrá, por razones de buen servicio, solicitar el cambio de trabajadores, expresando la causa del derecho a cambiar al personal del proveedor, entendiéndose como el derecho a prohibir unilateralmente la continuidad de funciones de un trabajador que implique un potencial riesgo a los pacientes, funcionarios, bienes e imagen de la organización.",
        "El Proveedor adjudicado deberá reemplazar al personal, dentro del plazo que se le indique. La decisión del Hospital San José de Melipilla se comunicará por escrito al Proveedor precisando las causas que motivan la solicitud, con a lo menos 5 días de anticipación a la fecha en que se solicita deje de prestar servicios en sus dependencias, el trabajador que se indique."
    ]
    for item in cambio_personal_items:
         agregar_parrafo_con_texto(doc, item)

    doc.add_heading("10.20 Cesión y subcontratación.", level=2)
    cesion_subcontratacion_items = [
        "El proveedor adjudicado no podrá ceder ni transferir en forma alguna, total ni parcialmente, los derechos y obligaciones que nacen del desarrollo de esta licitación, y, en especial, los establecidos en los respectivos contratos que se celebren con los órganos públicos mandantes.",
        "La infracción de esta prohibición será causal inmediata de término del contrato, sin perjuicio de las acciones legales que procedan ante esta situación.",
        "Durante la ejecución del contrato, y previa autorización por escrito del Hospital, el adjudicatario sólo podrá efectuar aquellas subcontrataciones que sean indispensables para la realización de tareas específicas, todo lo cual será calificado por el coordinador del contrato. En todo caso, el adjudicatario seguirá siendo el único responsable de las obligaciones contraídas en virtud del respectivo contrato suscrito con el Hospital.",
        "Así mismo, el subcontratista debe encontrarse hábil en el registro de Proveedores del Estado y tratándose de servicios, acreditar el cumplimiento de obligaciones laborales, conforme lo establece el artículo 4º inciso 2º de la Ley N°19.886.",
        "En todos los casos es el oferente y eventual adjudicatario el único responsable del pleno cumplimiento de lo señalado en estas bases (Art. N° 76, Reglamento de la Ley N° 19.886)."
    ]
    for item in cesion_subcontratacion_items:
         agregar_parrafo_con_texto(doc, item)


    doc.add_heading("10.21 Discrepancias.", level=2)
    discrepancias_items = [
        "Si con motivo de la ejecución del contrato se presentaran denuncias, querellas o demandas ante el Ministerio Público o los Tribunales Ordinarios de Justicia; o reclamos ante el Consejo de Defensa del Estado por el cuestionamiento en la prestación otorgada y que corresponda al objeto del contrato celebrado, será el proveedor el único responsable por tales actos, por lo que, sí el Hospital fuese condenado a pagar una multa o indemnización, en razón de los actos precedentemente enunciados o el Hospital tuviera que pagar alguna transacción judicial o extrajudicial que deba celebrarse en razón de las situaciones antes enunciadas, el proveedor deberá reembolsar al Hospital el total del monto resultante de un fallo ejecutoriado o de una transacción judicial o extrajudicial o de un procedimiento de medición de acuerdo a la Ley Nº 19.966.",
        "Asimismo, serán responsables de todos los daños, pérdidas, deterioros o perjuicios de bienes muebles e inmuebles del Hospital, producto del mal uso ocasionado en virtud de la prestación de servicio, debiendo restituir al Hospital los costos en que deba incurrir para reparar los daños producidos por este motivo. Esta obligación se mantendrá aun cuando el presente contrato que al efecto se suscriba se dé por terminado ya sea por expiración del plazo establecido o por decisión del Hospital."
    ]
    for item in discrepancias_items:
         agregar_parrafo_con_texto(doc, item)

    doc.add_heading("10.22. Constancia.", level=2)
    agregar_parrafo_con_texto(doc, "Se deja expresa constancia que todas y cada una de las cláusulas contenidas en las presentes Bases, Anexos y aclaratorias, se entienden incorporadas sin necesidad de mención expresa en el correspondiente contrato que se materialice con el adjudicado y éste se hace responsable del cumplimiento de las obligaciones de tales documentos, Bases Administrativas y Contrato que se deriven.")

    # BASES TÉCNICAS Section
    doc.add_heading("BASES TECNICAS PARA EL SUMINISTRO DE INSUMOS Y ACCESORIOS PARA TERAPIA DE PRESIÓN NEGATIVA CON EQUIPOS EN COMODATO PARA EL HOSPITAL SAN JOSÉ DE MELIPILLA", level=0)

    doc.add_heading("I. REQUISITOS PARA ADJUDICARSE DISPOSICIONES GENERALES", level=1)
    requisitos_generales_intro = doc.add_paragraph()
    requisitos_generales_intro.add_run("Para adjudicarse el presente proceso de licitación, los oferentes participantes deberán cumplir con lo siguiente:")
    requisitos_generales_list = [
        [("• Haber llenado y presentado los Anexos Administrativos N°1, N°2, N°3, N°4, el presentar estos anexos habilita al proveedor a participar en la presente licitación.", "")],
        [("• Deben presentar el Económico N°5 y los Anexos Técnicos N°7, N°8 y N°9 con toda la información requerida, debidamente firmados por el representante legal de la empresa o la persona natural, según corresponda.", "")],
        [("• Deberán entregar toda la información necesaria para poder evaluar a la empresa en cada uno de los ítems de los Criterios de Evaluación.", "")],
        [("• Deberán dar respuesta a los requisitos generados por foro inverso en los plazos y/o periodos establecidos en las presentes Bases de Licitación.", "")],
        [("• Presentar ficha técnica y certificados de los productos ofertados.", "")],
        [("• Entregar muestras de los productos solicitados y comodato ofertado.", "")],
        [("• Entregar garantías de la oferta.", "")],
    ]
    for item in requisitos_generales_list:
         agregar_parrafo_con_runs(doc, item, style='List Bullet')

    nota_inadmisible_tec = doc.add_paragraph()
    nota_inadmisible_tec.add_run("Nota:").bold = True
    nota_inadmisible_tec.add_run(" Los oferentes que no cumplan con estos requisitos no serán evaluados, declarándose ")
    nota_inadmisible_tec.add_run("inadmisible").bold = True
    nota_inadmisible_tec.add_run(" su oferta.")


    doc.add_heading("II. DISPOSICIONES DE LA LICITACION", level=1)
    agregar_parrafo_con_texto(doc, "Determinar las directrices y características técnicas necesarias para el suministro de insumos y accesorios para terapia de presión negativa con equipos en comodato.")

    doc.add_heading("1. GENERALIDADES:", level=2)
    generalidades_list = [
        [("a) La Licitación será adjudicada por la ", "bold"), ("totalidad", "bold"), (" y se podrá aumentar o disminuir hasta un 30% por cada línea adjudicada sin superar el monto total presupuestado para la Licitación.", "")],
        ["b) La propuesta deberá contemplar todos los costos de trasporte para el despacho de los productos. El Hospital no cancelará ningún costo asociado a esta temática."],
        ["c) El proveedor deberá mantener la calidad de los productos ofertados para cada solicitud de compra bajo esta Licitación, situación que será constantemente evaluada por el administrador del contrato."],
        ["d) El proveedor deberá permitir la revisión de los productos entregados al Hospital por el personal que se disponga por parte del Establecimiento, para así dar a una correcta recepción conforme."], # Corrected "dar a una" from "dar"
        ["e) En los casos que los productos sean despachados por empresas de transporte estos deberán permitir la revisión de los productos, en caso contrario los productos serán rechazados."],
        ["f) Toda entrega deberá adjuntar un documento que acredite la compra (Guía de despacho, Factura u Orden de Compra)."],
        ["g) La adquisición de estos productos será de forma parcializada durante un periodo máximo de 36 meses, o hasta la duración del monto estipulado en base."],
        ["h) Deberán entregar toda la información necesaria para poder evaluar a la empresa en cada uno de los ítems de los Criterios de Evaluación."],
        ["i) El administrador técnico del contrato será la Enfermera Supervisora de Pabellón y el encargado en aspectos administrativos será el Jefe de Farmacia o quien lo subrogue."]
    ]
    for item in generalidades_list:
        p = doc.add_paragraph()
        if isinstance(item, list):
             agregar_parrafo_con_runs(p, item) # Add runs to the paragraph
        else:
             p.add_run(item) # Add plain text

        # Simulate lettered list a), b), c)
        # This is tricky with python-docx styles programmatically.
        # A simple manual approach for demonstration:
        label_match = re.match(r'([a-z])\)\s', p.text)
        if label_match:
            label = label_match.group(0) # Includes a)
            rest_of_text = p.text[len(label):]
            p.text = '' # Clear original text
            p.add_run(label).bold = True
            p.add_run(rest_of_text)


    doc.add_heading("1.1. DE LOS PRODUCTOS", level=3)
    agregar_parrafo_con_texto(doc, "a. La presente licitación pública, se enfoca en la adquisición de los productos que se presentan en el cuadro siguiente, se evaluaran técnicamente cada producto. La adjudicación será por la totalidad.")
    agregar_parrafo_con_texto(doc, "b. La siguiente tabla presenta cantidades de consumo referenciales, la que se utilizara solo para términos de evaluación.")

    # Products Table
    products_data = [
        ["ITEM", "INSUMOS", "UD", "MONTO\nMÁXIMO\nA PAGAR"],
        ["1", "Recolector de contenido y exudado de herida con gel de 300 ml para presión negativa con conexión que mide 1,20 mt aprox., circuito cerrado, clamp integrado y filtro de carbón. Desechable.", "UD", "$150.000"],
        ["2", "Recolector de contenido y exudado de herida con gel de 500 ml para ser utilizado presión negativa, con conexión que mide 1,80 mt aprox., circuito cerrado, clamp integrado y filtro de carbón. Desechable.", "UD", "$180.000"],
        ["3", "Recolector de contenido y exudado de herida con gel de 1000 ml para ser utilizado con presión negativa con conexión que mide 1,80 mt aprox., circuito cerrado, clamp integrado y filtro de carbón. Desechable.", "UD", "$220.000"],
        ["4", "Kit de apósito espuma negra en forma ovalada 26 cm x15cm x3.2 cm aprox. tamaño LARGE, con láminas adhesivas transparentes, conector de succión de silicona flexible de 90 cm aprox., conector luer-lock, clamp y regla desechable.", "UD", "$100.000"],
        ["5", "Kit de apósito espuma negra en forma ovalada 60 cm x30cm x1.8 cm aprox. tamaño extra large, con láminas adhesivas transparentes, conector de succión de silicona flexible de 90 cm aprox., conector luer-lock, clamp, desechable.", "UD", "$370.000"],
        ["6", "Kit de apósito espuma negra precortada en forma de espiral 11.3 cm x 7.7cm x 1.75cm aprox. tamaño small con láminas adhesivas transparentes, conector de succión de silicona flexible de 90 cm, conector luer-lock, clamp, desechable.", "UD", "$100.000"],
        ["7", "Kit de apósito espuma negra precortada en forma de espiral 17.4cm x14.7cm x1.75 cm aprox. tamaño medium, con láminas adhesivas transparentes, conector de succión de silicona flexible de 90 cm, conector luer-lock, clamp, desechable.", "UD", "$120.000"],
        ["8", "Aposito de espuma blanca de alcohol polivinílico humedecido con agua estéril, 10CM ×7.5CM x 1 cm aprox. tamaño small", "UD", "$60.000"],
        ["9", "Apósito de espuma blanca de alcohol polivinílico, humedecido con agua estéril, hidrofílica, DE 10 CM x 15CM x 1 cm aprox. Tamaño Large (L)", "UD", "$70.000"],
        ["10", "Kit de apósito espuma hidrofóbica color gris de eter de poliuretano con plata metálica 10 cm x7.5 cm x3.2 cm aprox. tamaño small", "UD", "$120.000"],
        ["11", "Kit de apósito espuma hidrofóbica color gris de eter de poliuretano con plata metálica 18cm x12.5 cmx3.2 cm aprox. tamaño medio", "UD", "$130.000"],
        ["12", "Kit de apósito espuma hidrofóbica color gris de eter de poliuretano con plata metálica 26 cm x15cm x3.2 cm aprox. tamaño large.", "UD", "$230.000"],
        ["13", "kit de apósito abdominal para manejo de abdomen abierto con presión negativa, con lámina protectora visceral de poliuretano, láminas adhesivas transparentes, conector de succión de silicona flexible de 90 cm aprox., conector luer-lock y clamp, desechable", "UD", "$550.000"],
        ["14", "Kit de apósito para incisiones lineales 90 cm aprox. con espuma de poliuretano, laminas adhesivas transparentes y un conector de succión de silicona flexible de 90 cm aprox., conector luer-lock y clamp y desechable.", "UD", "$420.000"],
        ["15", "Kit de apósito para incisiones lineales 20 cm aprox. con espuma de poliuretano, laminas adhesivas transparentes y un conector de succión de silicona flexible de 90 cm aprox., conector luer-lock y clamp y desechable.", "UD", "$300.000"],
        ["16", "Kit de apósito para incisiones lineales 13 cm aprox. con espuma de poliuretano, laminas adhesivas transparentes y un conector de succión de silicona flexible de 90 cm aprox., conector luer-lock y clamp y desechable.", "UD", "$250.000"],
        ["17", "Kit de apósito para incisiones lineales 35 cm aprox. con espuma de poliuretano, laminas adhesivas transparentes y un conector de succión de silicona flexible de 90 cm aprox., conector luer-lock y clamp y desechable.", "UD", "$400.000"],
        ["18", "Lámina adhesiva transparente, hipoalergénica, semipermeable para realizar el sello de la terapia de presión negativa", "UD", "$50.000"],
        ["19", "Conector de succión de silicona flexible de 90 cm aprox. con sensores externos de luz de monitoreo continuo de la presión, sistema de detección de obstrucciones y sistema de ráfagas de aire cada 5 min para ayudar a reducir los bloqueos, conector luer-lock y clamp.", "UD", "$50.000"],
        ["20", "Kit de apósito para terapia de instilación tamaño médium, espuma de ester de poliuretano y reticulada con 3 capas: 1 capa en contacto con la herida que tiene orificios de 5 mm, una segunda capa fina de 8 mm y una tercera capa gruesa de 16 mm.", "UD", "$200.000"],
        ["21", "Cassete para conectar la solución para la terapia de instilación.", "UD", "$110.000"],
    ]
    tabla_products = crear_tabla(doc, products_data)
    # Center header text manually as create_tabla doesn't do it by default
    for j in range(len(tabla_products.rows[0].cells)):
        tabla_products.cell(0, j).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        tabla_products.cell(0, j).paragraphs[0].runs[0].bold = True
    # Center numerical/unit columns
    for i in range(1, len(products_data)):
        tabla_products.cell(i, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        tabla_products.cell(i, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        tabla_products.cell(i, 3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT # Price aligned right?


    adjudicacion_value_note = doc.add_paragraph()
    adjudicacion_value_note.add_run("c. ").bold = True
    adjudicacion_value_note.add_run("La adjudicación se realizará por ")
    adjudicacion_value_note.add_run("valor unitario").bold = True
    adjudicacion_value_note.add_run(" y tendrá una duración de 36 meses o hasta agotar el presupuesto, lo que ocurra primero, sin obligar al hospital a comprar una cantidad mínima establecida.").bold = True # Entire sentence bold in OCR

    admisibilidad_note = doc.add_paragraph()
    admisibilidad_note.add_run("d. ").bold = True
    admisibilidad_note.add_run("Se considera causal de ")
    admisibilidad_note.add_run("admisibilidad").bold = True
    admisibilidad_note.add_run(" que el proveedor adjunte ficha técnica en español de todos los productos solicitados al portal de Mercado Público.")


    doc.add_heading("1.2. ENTREGA DE MUESTRAS", level=3)
    entrega_muestras_list = [
        [("• Presentar muestras de los insumos y equipo en comodato para evaluación es de carácter ", ""), ("OBLIGATORIO.", "bold")],
        [("• Se debe presentar muestras de todos los insumos solicitados. En caso de no presentar muestras su propuesta podrá ser declarada ", ""), ("inadmisible", "bold"), (".", "")],
        ["• Cada muestra debe indicar nombre del proveedor, número de licitación y N° de línea del producto."],
        ["• Las muestras en ningún caso generarán costo para el Hospital y NO podrán ser devueltas a los oferentes, ya que serán utilizados para realizar pruebas de parte de los referentes técnicos para su evaluación."],
        ["• El oferente deberá permitir la apertura de cajas/bolsas de las muestras presentadas para una correcta recepción de estas."],
        [("• El oferente deberá entregar las muestras a la ", ""), ("UNIDAD DE ABASTECIMIENTO DEL HOSPITAL SAN JOSE DE MELIPILLA", "bold"), (", ubicada en Calle O'Higgins N.º 551, Comuna de Melipilla hasta el cierre de la licitación. ", ""), ("La no entrega de muestras en la forma y plazos establecidos en bases facultará al establecimiento a dejar inadmisible la oferta.", "bold")],
        ["• Toda muestra deberá ser acompañados de una guía de despacho o acta de recepción, la que será completada (firmada y timbrada) por Unidad de Abastecimiento. Este documento será el que respaldará la recepción de las muestras."],
    ]
    for item in entrega_muestras_list:
         agregar_parrafo_con_runs(doc, item, style='List Bullet')


    doc.add_heading("1.3. SOBRE LOS EQUIPOS SOLICITADOS EN COMODATO PARA EL USO DE LOS INSUMOS – CONDICION OBLIGATORIA", level=3)
    agregar_parrafo_con_texto(doc, "Para ejecutar el suministro, es obligatorio para el proveedor adjudicado, la entrega en comodato, a las unidades clínicas del hospital que lo soliciten, de los siguientes equipos médicos:")

    equipos_comodato_data = [
        ["NOMBRE EQUIPOS", "CANTIDAD ESTIMADA"],
        ["EQUIPO PARA TERAPIA PRESIÓN NEGATIVA\nINTRAHOSPITALARIO", "12"],
        ["EQUIPO PARA TERAPIA PRESIÓN NEGATIVA AMBULATORIO -\nDOMICILIARIO", "4"],
    ]
    tabla_equipos_comodato = crear_tabla(doc, equipos_comodato_data)
    for r_idx in range(len(equipos_comodato_data)):
         for c_idx in range(len(equipos_comodato_data[0])):
              tabla_equipos_comodato.cell(r_idx, c_idx).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
              if r_idx == 0:
                   tabla_equipos_comodato.cell(r_idx, c_idx).paragraphs[0].runs[0].bold = True


    agregar_parrafo_con_texto(doc, "La cantidad de equipos es estimada, tiene un carácter referencial. Debe adjuntar ficha técnica de los dispositivos. Estas cantidades podrán variar de acuerdo a la demanda interna, para lo cual el oferente adjudicado deberá dar respuesta las necesidades del hospital, máximo 10 días hábiles desde la adjudicación.")
    agregar_parrafo_con_texto(doc, "El proveedor que se adjudique, deberá entregar en comodato, a unidades clínicas del Hospital, Equipo para terapia presión negativa intrahospitalario y Equipo para terapia presión negativa ambulatorio compatibles con los insumos ofertados. Junto con la entrega en comodato, el proveedor se obliga a realizar la mantención preventiva y correctiva de los equipos. En caso de falla, se obliga a facilitar al Hospital, sin costo para este, equipos de similares características, mientras se entrega un nuevo equipo en comodato")

    doc.add_heading("CARACTERISTICAS GENERALES OBLIGATORIAS DE LOS EQUIPOS EN COMODATO", level=4)
    caracteristicas_comodato_list = [
        "Consolas con botón de encendido y apagado.",
        "Rangos de presión de -25 a -200 mmHg",
        "Terapias integradas en el mismo equipo .terapias de presión negativa estándar, manejo de abdomen abierto, prevención de dehiscencia de suturas e instilación",
        "Compatibilidad con contendor de 1000, 500 y 300 ml para uso intrahospitalario y contendor de 300 para uso ambulatorio-domiciliario",
        "Sistema de alarma de contenedor en su capacidad máxima y botón para liberación de contenedor",
        "Sistema de alarma de baja presión, terapia interrumpida u obstrucción.",
        "Sistema de alarmas para nivel crítico de batería.",
        "Cable para alimentación de corriente eléctrica o baterías en caso de equipo portátil.",
        "Memoria de uso de consola",
        "Autonomía de la batería de al menos 6 horas para equipo de uso hospitalario y 10 horas para equipo de uso domiciliario.",
    ]
    # Use simple paragraphs within a border or table to simulate the box
    # Easiest is just regular paragraphs, border is complex to add programmatically.
    # Or create a single cell table and put text inside
    table_caracteristicas = doc.add_table(1, 1, style='Table Grid')
    cell_caract = table_caracteristicas.cell(0, 0)
    for item in caracteristicas_comodato_list:
        cell_caract.add_paragraph(item)
    cell_caract.paragraphs[0].runs[0].bold = True # Make first item bold as it looks like a header in OCR

    doc.add_heading("1.4. ENTREGA Y RECEPCION", level=3)
    entrega_recepcion_list = [
        "La adquisición de estos productos será de forma parcializada según la cantidad y periocidad que el hospital considere necesario.",
        "El proveedor deberá despachar los productos señalando explícitamente Nombre del producto, Identificación del Proveedor y N° de Guía/Factura, Modelo (solo cuando corresponda), N° de Lote/Serie, Fecha de Vencimiento, de acuerdo a Norma Técnica de Minsal °226/22.",
        "Los productos deberán ser entregado en las dependencias del Hospital de Melipilla, considerando el traslado carga y descarga.",
        "La propuesta deberá contemplar todos los costos de trasporte para el despacho de los productos. El Hospital no cancelará ningún costo asociado a esta temática.",
        "Desde el requerimiento, el proveedor tendrá un máximo 7 días corridos para entregar los productos, siempre respetando los plazos ofertados según anexo Plazo de Entrega.",
        "El proveedor deberá realizar los cambios de los productos que no se ajusten a las bases técnicas y/o presenten deterioros en un plazo no mayor a 48 horas, con previo requerimiento del administrador del contrato.",
        "Los productos despachados que no se ajusten a la calidad ofertada serán rechazados e informado vía correo electrónico, para solicitar el cambio.",
        "El gasto que eventualmente se genere por artículos rechazados será de cargo de la empresa adjudicada.",
        "El embalaje deberá ser suficiente para soportar, sin límites, la manipulación brusca y descuidada durante el tránsito y la exposición a temperaturas extremas.",
        "El proveedor deberá permitir la apertura de cajas, bolsas, etc., para la correcta revisión de los productos entregados al Hospital por el personal de Bodega del Establecimiento para así dar una correcta recepción conforme.",
        "En los casos que los productos sean despachados por empresas de transporte estos deberán permitir la revisión de los productos, en caso contrario los productos serán rechazados.",
        "Los productos deberán ser entregados en Bodega de Farmacia del Hospital San José de Melipilla ubicada en calle O'Higgins #551, en los siguientes horarios: lunes a viernes de 8:00 a 14:00 horas",
    ]
    for item in entrega_recepcion_list:
         agregar_parrafo_con_texto(doc, item, style='List Bullet')

    doc.add_heading("a) MOTIVOS DE RECHAZO POR OBSERVACIÓN FÍSICA (ya iniciado en contrato):", level=4)
    agregar_parrafo_con_texto(doc, "Los artículos requeridos en la presente licitación podrán ser rechazados, al momento de la recepción en Bodega de Farmacia, por los siguientes motivos:")
    motivos_rechazo_list = [
         "Empaques deteriorados o visiblemente sucios, manchados, húmedos, etc.",
         "Cajas colectivas sin identificación de su contenido o leyendas ilegibles.",
         "Diferentes lotes no señalizados, Incluidos en un empaque colectivo.",
         "Textos o leyendas equivocadas, que puedan inducir a error.",
         "Envases con etiquetas e impresiones ilegibles o sin ellas.",
         "Discordancia entre envases ya sea colectivo, primarios o secundarios.",
         "Acondicionamiento inadecuado dentro de los envases primarios o secundarios.",
         "Envases vacíos o adulterados.",
         "Número de lote o fecha de vencimiento equivocada o ausente.",
         "Caja o etiqueta incorrecta, leyendas incompletas o ausentes.",
         "Sello violado o mal colocado.",
         "Contenido incorrecto, diferente o menor al etiquetado.",
         "Envases aplastados o deteriorados con motivo del traslado.",
         "Contaminación visible.",
         "Coloración no homogénea (intra o inter lote).",
         "Partículas extrañas observadas a simple vista o contraluz.",
         "Los insumos no deberán tener un vencimiento inferior a 2 año desde la fecha de recepción."
    ]
    for item in motivos_rechazo_list:
        agregar_parrafo_con_texto(doc, item, style='List Bullet')


    # ANEXO N°1
    doc.add_section(WD_SECTION_START.NEW_PAGE)
    doc.add_heading("ANEXO N° 1", level=1)
    doc.add_heading("IDENTIFICACIÓN DEL OFERENTE", level=2)
    doc.add_paragraph("PROPUESTA PÚBLICA: PARA EL SUMINISTRO DE INSUMOS Y ACCESORIOS PARA TERAPIA DE PRESIÓN NEGATIVA CON EQUIPOS EN COMODATO PARA EL HOSPITAL SAN JOSÉ DE MELIPILLA.").alignment = WD_ALIGN_PARAGRAPH.CENTER

    anexo1_data = [
        ["IDENTIFICACIÓN DEL OFERENTE", ""],
        ["R.U.T. DEL OFERENTE", ""],
        ["SIGLA PARA EL CASO DE EMPRESAS (Nombre de\nFantasía)", ""],
        ["DIRECCIÓN OFERENTE", ""],
        ["CIUDAD", ""],
        ["COMUNA", ""],
        ["TELÉFONOS", ""],
        ["NOMBRE DEL REPRESENTANTE LEGAL", ""],
        ["RUT DEL REPRESENTANTE LEGAL", ""],
        ["NOMBRE DE LA NOTARIA", ""],
        ["FECHA DONDE SE SEÑALA LA PERSONERIA DEL\nREPRESENTANTE LEGAL (adjuntar documento si no se\nencuentra actualizado en portal de Mercado Público)", ""],
        ["NOMBRE DEL CONTACTO COMERCIAL\n(ADMINISTRADOR EXTERNO DEL CONTRATO)", ""],
        ["CARGO DEL CONTACTO COMERCIAL", ""],
        ["RUBRO COMERCIAL", ""],
        ["E-MAIL", ""],
        ["NOMBRE DEL CONTACTO PARA EL SERVICIO", ""],
        ["TELEFONO", ""],
        ["E-MAIL", ""],
        ["CELULAR", ""],
        ["HORARIO DE ATENCION", ""],
    ]
    tabla_anexo1 = crear_tabla(doc, anexo1_data)
    # Optional: Adjust column widths

    doc.add_paragraph("\n") # Add some space
    agregar_parrafo_con_texto(doc, "FIRMA Y TIMBRE OFERENTE O REPRESENTANTE LEGAL", centrado=True)


    # ANEXO N°2
    doc.add_section(WD_SECTION_START.NEW_PAGE)
    agregar_parrafo_con_texto(doc, "Fecha: ______________", centrado=False)
    doc.add_heading("ANEXO N° 2", level=1)
    doc.add_heading("DECLARACION JURADA DE HABILIDAD", level=2)
    doc.add_paragraph("PROPUESTA PÚBLICA: PARA EL SUMINISTRO DE INSUMOS Y ACCESORIOS PARA TERAPIA DE PRESIÓN NEGATIVA CON EQUIPOS EN COMODATO PARA EL HOSPITAL SAN JOSÉ DE MELIPILLA.").alignment = WD_ALIGN_PARAGRAPH.CENTER

    anexo2_paras = [
        "Por la presente, el Oferente, <<NOMBRE PERSONA NATURAL O NOMBRE PERSONA JURIDICA>>, declara bajo juramento que no ha sido sancionado con la pena de prohibición perpetua o temporal (esta última vigente) para contratar con el Estado, por lavado de activos, financiamiento del terrorismo y cohecho, en virtud de lo dispuesto en los artículos 8° N°2 y N°10 de la Ley N°20.393.",
        'Asimismo, declara bajo juramento, sea persona natural o jurídica que no le afecta ninguna de las inhabilidades previstas en los incisos primero y sexto del artículo 4º de la Ley N°19.886, además, las previstas en el artículo 26 letra D del Decreto Ley N°211, que se transcriben en su parte pertinente:',
        '"(...) Quedaran excluidos quienes, dentro de los dos años anteriores al momento de la presentación de la oferta, de la formulación de la propuesta o de la suscripción de la convención, según se trate de licitaciones públicas, privadas o contratación directa, hayan sido condenados por prácticas antisindicales o infracción as los derechos fundamentales del trabajador, o por delitos concursales establecidos en el Código Penal" (inciso primero).',
        '“Ningún órgano de la Administración del Estado y de las empresas y corporaciones del Estado o en que este tenga participación, podrá suscribir contratos administrativos de provisión de bienes o prestación de servicios con los funcionarios directivos del mismo órgano o empresa, ni con personas unidas a ellos por los vínculos de parentesco descritos en la letra b) del artículos 54 de la ley N°18.575, ley Orgánica Constitucional de Bases Generales de la Administración del Estado, ni con sociedades de personas de las que aquellos o estas formen parte, mi con sociedades comanditas por acciones o anónima cerrada en que aquellos o estas sean accionistas, ni con sociedades anónima abiertas en que aquellos o estas sean dueños de acciones que representen el 10% o más del capital, ni con los gerentes, administradores, representantes o directores de cualquiera de las sociedades antedichas" (inciso sexto).',
        "En el caso de las conductas previstas en la letra a) del artículo 3°, podrá imponer, además, la prohibición de contratar a cualquier título con órganos de la administración centralizada o descentralizada del Estado, con organismos autónomos o con instituciones, organismos, empresas o servicios en los que el Estado efectúe aportes, con el Congreso Nacional y el Poder Judicial, así como la prohibición de adjudicarse cualquier concesión otorgada por el Estado, hasta el plazo de cinco años contado desde que la sentencia definitiva quede ejecutoriada.",
    ]
    for para in anexo2_paras:
         agregar_parrafo_con_texto(doc, para)

    doc.add_paragraph("\n")
    agregar_parrafo_con_texto(doc, "FIRMA Y TIMBRE OFERENTE O REPRESENTANTE LEGAL", centrado=True)


    # ANEXO N°3
    doc.add_section(WD_SECTION_START.NEW_PAGE)
    agregar_parrafo_con_texto(doc, "Fecha: ______________", centrado=False)
    doc.add_heading("ANEXO N° 3", level=1)
    doc.add_heading("DECLARACION JURADA DE CUMPLIMIENTO DE OBLIGACIONES LABORALES Y PREVISIONALES", level=2)
    doc.add_paragraph("PROPUESTA PÚBLICA: PARA EL SUMINISTRO DE INSUMOS Y ACCESORIOS PARA TERAPIA DE PRESIÓN NEGATIVA CON EQUIPOS EN COMODATO PARA EL HOSPITAL SAN JOSÉ DE MELIPILLA.").alignment = WD_ALIGN_PARAGRAPH.CENTER

    anexo3_paras = [
        "Por la presente, el Oferente, <<NOMBRE PERSONA NATURAL O PERSONA JURIDICA>>, declara bajo juramento:",
    ]
    for para in anexo3_paras:
         agregar_parrafo_con_texto(doc, para)

    anexo3_table_data = [
        ["¿POSEE\nDEUDA\nPREVISIONAL?\n(Marque el cuadro\nque corresponda)", "SI (Sí poseo deuda\nPrevisional)", "NO (No poseo deuda\nprevisional)", "MONTO"],
        ["", "█", "", "$__________"], # Simulate checkbox with block
    ]
    tabla_anexo3 = crear_tabla(doc, anexo3_table_data)
    centrar_verticalmente_tabla(tabla_anexo3)
    # Manually center headers and the block
    for j in range(len(anexo3_table_data[0])):
         tabla_anexo3.cell(0, j).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    tabla_anexo3.cell(1, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    tabla_anexo3.cell(1, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    tabla_anexo3.cell(1, 3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT


    anexo3_paras_cont = [
        "Comprometiéndose en caso de registrar saldos insolutos de remuneraciones o cotizaciones de seguridad social con sus actuales trabajadores o con aquellos que fueron contratados en los dos últimos años anteriores a la presente licitación, los primeros estados de pago producto del contrato licitado, deberán ser destinados al pago de dichas obligaciones, debiendo la empresa acreditar que la totalidad de las obligaciones se encuentran liquidadas al cumplirse la mitad del periodo de ejecución del contrato, con un máximo de seis meses.",
        "En virtud a lo establecido en el artículo 4º inciso 2º de la Ley N°19.886 de Bases sobre Contratos Administrativos de Suministro y Prestación de Servicios, en caso que el proveedor adjudicado registre saldos insolutos de remuneraciones o cotizaciones de seguridad social con sus actuales trabajadores contratados en los últimos 2 años, los primeros estados de pago producto del contrato licitado deberán ser destinados al pago de dichas obligaciones, debiendo la empresa acreditar que la totalidad de las obligaciones se encuentran liquidadas al cumplirse la mitad del periodo de ejecución del contrato, con un máximo de seis meses.", # Repeated in OCR
    ]
    for para in anexo3_paras_cont:
         agregar_parrafo_con_texto(doc, para)

    doc.add_paragraph("\n")
    agregar_parrafo_con_texto(doc, "FIRMA Y TIMBRE OFERENTE O REPRESENTANTE LEGAL", centrado=True)


    # ANEXO N°4
    doc.add_section(WD_SECTION_START.NEW_PAGE)
    agregar_parrafo_con_texto(doc, "Fecha: ______________", centrado=False)
    doc.add_heading("ANEXO N°4", level=1)
    doc.add_heading("DECLARACION JURADA SIMPLE", level=2)
    doc.add_paragraph("PROPUESTA PÚBLICA: PARA EL SUMINISTRO DE INSUMOS Y ACCESORIOS PARA TERAPIA DE PRESIÓN NEGATIVA CON EQUIPOS EN COMODATO PARA EL HOSPITAL SAN JOSÉ DE MELIPILLA.").alignment = WD_ALIGN_PARAGRAPH.CENTER

    anexo4_intro = doc.add_paragraph()
    anexo4_intro.add_run("En [ciudad/país], a [fecha] <<NOMBRE PERSONA NATURAL O PERSONA JURIDICA>>, integrante de la unión temporal declaro bajo juramento:")
    agregar_parrafo_con_texto(doc, "Que respecto no concurre ninguna de las prohibiciones descritas en el artículo 4º de la Ley 19.886, esto es:")

    anexo4_list_items = [
        "Ser funcionario directivo.",
        "Ser cónyuge, hijo, adoptado ni pariente hasta el tercer grado de consanguinidad ni segundo de afinidad de algún funcionario directivo;",
        "Que no tiene la calidad de gerente, administrador, representante o director de cualquiera de las sociedades referidas en las letras anteriores.",
        "Haber sido condenado(a) por prácticas antisindicales o infracción a los derechos del trabajador, en los últimos dos años anteriores a la presentación de la oferta y que, dentro en ese mismo lapso, no ha sido condenado por delitos concursales establecidos en el Código Penal.",
        "Que está en conocimiento que, en caso de ser adjudicada la licitación, y de existir saldos insolutos de remuneraciones o cotizaciones de seguridad social con los actuales trabajadores o trabajadores contratados en los dos últimos años, los pagos producto del contrato licitado deberán ser destinados primeramente a liquidar dichas deudas.",
        "Que no concurre la prohibición descrita en los artículos 8 y 10 de la Ley 20.393, esto es prohibición de celebrar actos y contratos con organismos del Estado.",
        "Que no concurre lo establecido en Decreto Ley N°211, que fija normas para la defensa de la libre competencia.",
    ]
    for i, item in enumerate(anexo4_list_items):
        p = doc.add_paragraph()
        p.add_run(item)
        aplicar_numeracion(p, num_id_anexo4, nivel=0) # Apply numbering 1, 2, 3...

    anexo4_note = doc.add_paragraph()
    anexo4_note.add_run("Nota:").bold = True
    anexo4_note.add_run(" Todos los representantes de la Unión Temporal de Proveedores deben presentar este anexo, en caso de que el oferente no sea una Unión Temporal de Proveedores no debe presentar este anexo.")


    doc.add_paragraph("\n")
    agregar_parrafo_con_texto(doc, "Fecha: ______________", centrado=False)
    agregar_parrafo_con_texto(doc, "FIRMA Y TIMBRE OFERENTE O REPRESENTANTE LEGAL", centrado=True)


    # ANEXO N°5
    doc.add_section(WD_SECTION_START.NEW_PAGE)
    doc.add_heading("ANEXO Nº5 (60%) OFERTA ECONÓMICA", level=1)
    doc.add_paragraph("PROPUESTA PÚBLICA: PARA EL SUMINISTRO DE INSUMOS Y ACCESORIOS PARA TERAPIA DE PRESIÓN NEGATIVA CON EQUIPOS EN COMODATO PARA EL HOSPITAL SAN JOSÉ DE MELIPILLA.").alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Economic Offer Table (Similar to Products Table, but with Unit Value column)
    anexo5_data = [
        ["ITEM", "INSUMOS", "CANTIDAD\nSOLICITADA", "VALOR\nUNITARIO NETO"],
        ["1", "Recolector de contenido y exudado de herida con gel de 300 ml para presión negativa con conexión que mide 1,20 mt aprox., circuito cerrado, clamp integrado y filtro de carbón. Desechable.", "1", "$_______"],
        ["2", "Recolector de contenido y exudado de herida con gel de 500 ml para ser utilizado presión negativa, con conexión que mide 1,80 mt aprox., circuito cerrado, clamp integrado y filtro de carbón. Desechable.", "1", "$_______"],
        ["3", "Recolector de contenido y exudado de herida con gel de 1000 ml para ser utilizado con presión negativa con conexión que mide 1,80 mt aprox., circuito cerrado, clamp integrado y filtro de carbón. Desechable.", "1", "$_______"],
        ["4", "Kit de apósito espuma negra en forma ovalada 26 cm x15cm x3.2 cm aprox. tamaño LARGE, con láminas adhesivas transparentes, conector de succión de silicona flexible de 90 cm aprox., conector luer-lock, clamp y regla desechable.", "1", "$_______"],
        ["5", "Kit de apósito espuma negra en forma ovalada 60 cm x30cm x1.8 cm aprox. tamaño extra large, con láminas adhesivas transparentes, conector de succión de silicona flexible de 90 cm aprox., conector luer-lock, clamp, desechable.", "1", "$_______"],
        ["6", "Kit de apósito espuma negra precortada en forma de espiral 11.3 cm x 7.7cm x 1.75cm aprox. tamaño small con láminas adhesivas transparentes, conector de succión de silicona flexible de 90 cm, conector luer-lock, clamp, desechable.", "1", "$_______"],
        ["7", "Kit de apósito espuma negra precortada en forma de espiral 17.4cm x14.7cm x1.75 cm aprox. tamaño medium, con láminas adhesivas transparentes, conector de succión de silicona flexible de 90 cm, conector luer-lock, clamp, desechable.", "1", "$_______"],
        ["8", "Aposito de espuma blanca de alcohol polivinílico humedecido con agua estéril, 10CM ×7.5CM x 1 cm aprox. tamaño small", "1", "$_______"],
        ["9", "Apósito de espuma blanca de alcohol polivinílico, humedecido con agua estéril, hidrofílica, DE 10 CM x 15CM x 1 cm aprox. Tamaño Large (L)", "1", "$_______"],
        ["10", "Kit de apósito espuma hidrofóbica color gris de eter de poliuretano con plata metálica 10 cm x7.5 cm x3.2 cm aprox. tamaño small", "1", "$_______"],
        ["11", "Kit de apósito espuma hidrofóbica color gris de eter de poliuretano con plata metálica 18cm x12.5 cmx3.2 cm aprox. tamaño medio", "1", "$_______"],
        ["12", "Kit de apósito espuma hidrofóbica color gris de eter de poliuretano con plata metálica 26 cm x15cm x3.2 cm aprox. tamaño large.", "1", "$_______"],
        ["13", "kit de apósito abdominal para manejo de abdomen abierto con presión negativa, con lámina protectora visceral de poliuretano, láminas adhesivas transparentes, conector de succión de silicona flexible de 90 cm aprox., conector luer-lock y clamp, desechable", "1", "$_______"],
        ["14", "Kit de apósito para incisiones lineales 90 cm aprox. con espuma de poliuretano, laminas adhesivas transparentes y un conector de succión de silicona flexible de 90 cm aprox., conector luer-lock y clamp y desechable.", "1", "$_______"],
        ["15", "Kit de apósito para incisiones lineales 20 cm aprox. con espuma de poliuretano, laminas adhesivas transparentes y un conector de succión de silicona flexible de 90 cm aprox., conector luer-lock y clamp y desechable.", "1", "$_______"],
        ["16", "Kit de apósito para incisiones lineales 13 cm aprox. con espuma de poliuretano, laminas adhesivas transparentes y un conector de succión de silicona flexible de 90 cm aprox., conector luer-lock y clamp y desechable.", "1", "$_______"],
        ["17", "Kit de apósito para incisiones lineales 35 cm aprox. con espuma de poliuretano, laminas adhesivas transparentes y un conector de succión de silicona flexible de 90 cm aprox., conector luer-lock y clamp y desechable.", "1", "$_______"],
        ["18", "Lámina adhesiva transparente, hipoalergénica, semipermeable para realizar el sello de la terapia de presión negativa", "1", "$_______"],
        ["19", "Conector de succión de silicona flexible de 90 cm aprox. con sensores externos de luz de monitoreo continuo de la presión, sistema de detección de obstrucciones y sistema de ráfagas de aire cada 5 min para ayudar a reducir los bloqueos, conector luer-lock y clamp.", "1", "$_______"],
        ["20", "Kit de apósito para terapia de instilación tamaño médium, espuma de ester de poliuretano y reticulada con 3 capas: 1 capa en contacto con la herida que tiene orificios de 5 mm, una segunda capa fina de 8 mm y una tercera capa gruesa de 16 mm.", "1", "$_______"],
        ["21", "Cassete para conectar la solución para la terapia de instilación.", "1", "$_______"],
        ["TOTAL VALOR NETO", "", "", "$_______"],
    ]
    tabla_anexo5 = crear_tabla(doc, anexo5_data)

    # Manually format headers
    for j in range(len(anexo5_data[0])):
        tabla_anexo5.cell(0, j).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        tabla_anexo5.cell(0, j).paragraphs[0].runs[0].bold = True

    # Center ITEM and CANTIDAD columns
    for i in range(1, len(anexo5_data) - 1): # Exclude header and total row
         tabla_anexo5.cell(i, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
         tabla_anexo5.cell(i, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Format Total row
    tabla_anexo5.cell(len(anexo5_data) - 1, 0).merge(tabla_anexo5.cell(len(anexo5_data) - 1, 2))
    tabla_anexo5.cell(len(anexo5_data) - 1, 0).text = "TOTAL VALOR NETO"
    tabla_anexo5.cell(len(anexo5_data) - 1, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    tabla_anexo5.cell(len(anexo5_data) - 1, 0).paragraphs[0].runs[0].bold = True

    # Add notes below the table
    anexo5_notes = [
        "* Al adjudicar, se considerará el valor unitario de los productos.",
        "* Este formulario deberá adjuntarlo obligatoriamente.",
        "* El valor ofertado deberá contemplar todos los costos asociados respecto a transporte o temáticas asociadas al despacho de los productos, el Hospital no pagará valores no contemplados en el presente anexo.",
        "* Predomina el valor unitario, el Hospital se reserva el derecho de modificar o corregir las sumatorias totales.",
    ]
    for note in anexo5_notes:
        agregar_parrafo_con_texto(doc, note)

    doc.add_paragraph("\n")
    agregar_parrafo_con_texto(doc, "FIRMA Y TIMBRE OFERENTE O REPRESENTANTE LEGAL", centrado=True)


    # ANEXO N°6
    doc.add_section(WD_SECTION_START.NEW_PAGE)
    doc.add_heading("ANEXO N°6 (20%)", level=1)
    doc.add_heading("PAUTA DE EVALUACIÓN TÉCNICA (SOLO PARA CONOCIMIENTO)", level=2)
    doc.add_paragraph("PROPUESTA PÚBLICA: PARA EL SUMINISTRO DE INSUMOS Y ACCESORIOS PARA TERAPIA DE PRESIÓN NEGATIVA CON EQUIPOS EN COMODATO PARA EL HOSPITAL SAN JOSÉ DE MELIPILLA.").alignment = WD_ALIGN_PARAGRAPH.CENTER

    anexo6_preamble = [
        "PRODUCTO:",
        "PROVEEDOR:",
        "NOMBRE EVALUADOR:",
    ]
    for text in anexo6_preamble:
         agregar_parrafo_con_texto(doc, text)

    agregar_parrafo_con_texto(doc, "A continuación, se presentan afirmaciones que debe calificar con nota de 1 a 7, donde 1 representa incumplimiento absoluto y 7 representa desempeño mayor al esperado. Solo debe marcar una alternativa.")

    # Evaluation Table
    anexo6_data = [
        ["N°", "ITEM I.-PARAMETROS A EVALUAR SOBRE CALIDAD Y\nFUNCIONALIDAD", "VALOR", "", "", "", "", "", "", "NOTA"],
        ["", "", "1", "2", "3", "4", "5", "6", "7", "N/A"],
        ["1", "El producto cumple con los objetivos esperados.", "","","","","","","",""],
        ["2", "El producto cumple con la calidad esperada.", "","","","","","","",""],
        ["3", "La funcionalidad del producto es la adecuada.", "","","","","","","",""],
        ["4", "El Equipo en comodato cumple con la funcionalidad y calidad esperada necesaria para su utilización.", "","","","","","","",""],
        ["5", "La manipulación del insumo en conjunto al equipamiento ofrece la seguridad óptima deseada.", "","","","","","","",""],
        ["NOTA FINAL * ITEM I.-", "", "", "", "", "", "", "", ""], # Merged cell
        ["PUNTAJE FINAL ITEM I.-", "", "", "", "", "", "", "", ""], # Merged cell
    ]
    tabla_anexo6_item1 = doc.add_table(rows=len(anexo6_data), cols=len(anexo6_data[0]), style='Table Grid')

    # Fill table content
    for r_idx, row_data in enumerate(anexo6_data):
        for c_idx, cell_content in enumerate(row_data):
            tabla_anexo6_item1.cell(r_idx, c_idx).text = cell_content

    # Merge cells and format headers
    tabla_anexo6_item1.cell(0, 1).merge(tabla_anexo6_item1.cell(0, 9))
    tabla_anexo6_item1.cell(0, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    tabla_anexo6_item1.cell(0, 0).paragraphs[0].runs[0].bold = True
    tabla_anexo6_item1.cell(0, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    tabla_anexo6_item1.cell(0, 1).paragraphs[0].runs[0].bold = True # ITEM I header
    tabla_anexo6_item1.cell(1, 2).merge(tabla_anexo6_item1.cell(1, 8)) # Merge VALOR 1-7
    tabla_anexo6_item1.cell(1, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    tabla_anexo6_item1.cell(1, 2).paragraphs[0].runs[0].bold = True # VALOR header

    # Bold headers for VALOR values (1 to 7 and N/A) - tricky with merged cell above
    # Let's recreate row 1 content more carefully
    tabla_anexo6_item1.cell(1, 2).text = 'VALOR' # Re-set merged cell text
    tabla_anexo6_item1.cell(1, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    tabla_anexo6_item1.cell(1, 2).paragraphs[0].runs[0].bold = True

    # Manually add 1-7 and N/A to separate cells below the merged 'VALOR'
    # This requires re-splitting the merged cell if it was merged programmatically first
    # A simpler approach is to create the row with individual cells and then merge the header
    # Let's create the table with correct rows/cols first and then set text/merge

    anexo6_data_structured = [
        ['N°', 'ITEM I.-PARAMETROS A EVALUAR SOBRE CALIDAD Y\nFUNCIONALIDAD', 'VALOR', '', '', '', '', '', '', 'NOTA'], # Row 0
        ['', '', '1', '2', '3', '4', '5', '6', '7', 'N/A'], # Row 1
        ['1', 'El producto cumple con los objetivos esperados.', '', '', '', '', '', '', '', ''], # Row 2
        ['2', 'El producto cumple con la calidad esperada.', '', '', '', '', '', '', '', ''], # Row 3
        ['3', 'La funcionalidad del producto es la adecuada.', '', '', '', '', '', '', '', ''], # Row 4
        ['4', 'El Equipo en comodato cumple con la funcionalidad y calidad esperada necesaria para su utilización.', '', '', '', '', '', '', '', ''], # Row 5
        ['5', 'La manipulación del insumo en conjunto al equipamiento ofrece la seguridad óptima deseada.', '', '', '', '', '', '', '', ''], # Row 6
        ['NOTA FINAL * ITEM I.-', '', '', '', '', '', '', '', '', ''], # Row 7
        ['PUNTAJE FINAL ITEM I.-', '', '', '', '', '', '', '', '', ''], # Row 8
    ]
    tabla_anexo6_item1 = doc.add_table(rows=len(anexo6_data_structured), cols=len(anexo6_data_structured[0]), style='Table Grid')
    for r_idx, row_data in enumerate(anexo6_data_structured):
        for c_idx, cell_content in enumerate(row_data):
            tabla_anexo6_item1.cell(r_idx, c_idx).text = cell_content

    # Merges
    tabla_anexo6_item1.cell(0, 1).merge(tabla_anexo6_item1.cell(0, 9)) # Main header merge
    tabla_anexo6_item1.cell(1, 2).merge(tabla_anexo6_item1.cell(1, 8)) # VALOR header merge
    tabla_anexo6_item1.cell(7, 0).merge(tabla_anexo6_item1.cell(7, 9)) # NOTA FINAL merge
    tabla_anexo6_item1.cell(8, 0).merge(tabla_anexo6_item1.cell(8, 9)) # PUNTAJE FINAL merge

    # Formatting (Center and Bold Headers)
    for r_idx in [0, 1]:
         for c_idx in range(len(anexo6_data_structured[0])):
              if tabla_anexo6_item1.cell(r_idx, c_idx).text: # Only format non-empty cells
                   tabla_anexo6_item1.cell(r_idx, c_idx).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                   tabla_anexo6_item1.cell(r_idx, c_idx).paragraphs[0].runs[0].bold = True
    # Bold first cell in data rows
    for r_idx in range(2, 7):
         tabla_anexo6_item1.cell(r_idx, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
         tabla_anexo6_item1.cell(r_idx, 0).paragraphs[0].runs[0].bold = True

    # Bold merged cells at the end
    tabla_anexo6_item1.cell(7, 0).paragraphs[0].runs[0].bold = True
    tabla_anexo6_item1.cell(8, 0).paragraphs[0].runs[0].bold = True


    anexo6_notes_item1 = [
         "*Notas finales del ITEM I.- Esta nota se pondera con un 15% del total de la Evaluación Técnica. Aquellas notas inferiores a 4 serán declaradas inadmisibles de evaluación.",
         "Puntaje final ITEM I.= Nota obtenida x 15%\nNota máxima (7)",
    ]
    for note in anexo6_notes_item1:
         agregar_parrafo_con_texto(doc, note)


    # ITEM II Evaluation Table
    anexo6_item2_data = [
        ["ITEM II.- PARAMETROS A EVALUAR SOBRE\nESPECIFICACIONES – FICHA TECNICA – OPORTUNIDAD", "EVALUACIÓN", "", "PUNTAJE (%)"],
        ["", "SI CUMPLE", "NO CUMPLE", ""],
        ["1", "El producto cumple con las especificaciones técnicas solicitadas.", "1%", "INADMISIBLE", ""],
        ["2", "La oferente entrega equipos en comodato para su evaluación", "2%", "INADMISIBLE", ""],
        ["3", "La oferente entrega muestras de los productos para su evaluación.", "2%", "INADMISIBLE", ""],
        ["SUMA DE PUNTAJE ITEM II.-", "", "", ""], # Merged cell
        ["PUNTAJE FINAL EVALUACION TECNICA (SUMA PUNTAJE FINAL ITEM I.- + PUNTAJE FINAL ITEM II.-", "", "", ""], # Merged cell
    ]
    tabla_anexo6_item2 = doc.add_table(rows=len(anexo6_item2_data), cols=len(anexo6_item2_data[0]), style='Table Grid')

    for r_idx, row_data in enumerate(anexo6_item2_data):
         for c_idx, cell_content in enumerate(row_data):
              tabla_anexo6_item2.cell(r_idx, c_idx).text = cell_content

    # Merges
    tabla_anexo6_item2.cell(0, 1).merge(tabla_anexo6_item2.cell(0, 3)) # EVALUACIÓN main header merge
    tabla_anexo6_item2.cell(5, 0).merge(tabla_anexo6_item2.cell(5, 3)) # SUMA merge
    tabla_anexo6_item2.cell(6, 0).merge(tabla_anexo6_item2.cell(6, 3)) # PUNTAJE FINAL merge

    # Formatting (Center and Bold Headers)
    for r_idx in [0, 1]:
         for c_idx in range(len(anexo6_item2_data[0])):
              if tabla_anexo6_item2.cell(r_idx, c_idx).text:
                   tabla_anexo6_item2.cell(r_idx, c_idx).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                   tabla_anexo6_item2.cell(r_idx, c_idx).paragraphs[0].runs[0].bold = True

    # Bold first cell in data rows (ITEM number and description)
    for r_idx in range(2, 5):
         tabla_anexo6_item2.cell(r_idx, 0).paragraphs[0].runs[0].bold = True

    # Bold merged cells at the end
    tabla_anexo6_item2.cell(5, 0).paragraphs[0].runs[0].bold = True
    tabla_anexo6_item2.cell(6, 0).paragraphs[0].runs[0].bold = True

    # Center relevant columns
    for r_idx in range(1, len(anexo6_item2_data)):
         for c_idx in range(1, len(anexo6_item2_data[0])):
              tabla_anexo6_item2.cell(r_idx, c_idx).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    centrar_verticalmente_tabla(tabla_anexo6_item2)


    doc.add_heading("OBSERVACIONES:", level=3)
    doc.add_paragraph("_________________________________________________________________________________________") # Placeholder lines
    doc.add_paragraph("_________________________________________________________________________________________")
    doc.add_paragraph("_________________________________________________________________________________________")
    doc.add_paragraph("_________________________________________________________________________________________")

    doc.add_paragraph("\n")
    agregar_parrafo_con_texto(doc, "FIRMA Y TIMBRE DE EVALUADOR", centrado=True)
    agregar_parrafo_con_texto(doc, "FECHA DE EVALUACION: _______/_______/________", centrado=False)


    # ANEXO N°7
    doc.add_section(WD_SECTION_START.NEW_PAGE)
    doc.add_heading("ANEXO N° 7", level=1)
    doc.add_heading("FICHA TECNICA DEL HSJM", level=2)
    doc.add_paragraph("PROPUESTA PÚBLICA: PARA EL SUMINISTRO DE INSUMOS Y ACCESORIOS PARA TERAPIA DE PRESIÓN NEGATIVA CON EQUIPOS EN COMODATO PARA EL HOSPITAL SAN JOSÉ DE MELIPILLA.").alignment = WD_ALIGN_PARAGRAPH.CENTER

    anexo7_data = [
        ["Nombre del Dispositivo Médico", ""],
        ["Proveedor", ""],
        ["ID", ""],
        ["Registro Sanitario", ""],
        ["Certificaciones", ""],
        ["Material", ""],
        ["Descripción de Dispositivo Médico", ""],
        ["Características Generales (Adjuntar Archivos de respaldo con fichas técnicas)", ""],
        ["Imagen o\nFotografía del\nProducto (Se\npermite señalar n°\nde página de\nfichas técnicas\nadjuntadas a sus\nofertas en portal)", ""],
        ["Otras Observaciones del Proveedor", ""],
    ]
    tabla_anexo7 = doc.add_table(rows=len(anexo7_data), cols=2, style='Table Grid')
    for r_idx, row_data in enumerate(anexo7_data):
         tabla_anexo7.cell(r_idx, 0).text = row_data[0]
         tabla_anexo7.cell(r_idx, 1).text = row_data[1]

    # Apply row spans manually where content goes across rows
    # Based on visual inspection, 'Descripción...', 'Características...', 'Imagen...', 'Otras Observaciones...' seem to span multiple conceptual rows in the form layout,
    # but representing this in a simple docx table structure can be tricky without explicit row merging in the OCR data.
    # The simplest representation is just a two-column table where the right column is left empty for input.
    # If specific rows need more height or multiple paragraphs, agregar_contenido_celda would be used, but for a form, leaving it blank is sufficient.
    # The 'Imagen o Fotografía...' cell content implies a larger area, but the table structure doesn't show merges.
    # We can leave the table as 2 columns as created above.

    doc.add_paragraph("\n")
    agregar_parrafo_con_texto(doc, "Fecha: _______/_______/________", centrado=False)
    doc.add_paragraph("\n")
    agregar_parrafo_con_texto(doc, "FIRMA Y TIMBRE OFERENTE O REPRESENTANTE LEGAL", centrado=True)
    doc.add_paragraph("\n")
    agregar_parrafo_con_texto(doc, "ESTE FORMULARIO DEBERA COMPLETARLO Y ADJUNTARLO A LA OFERTA, EN CASO CONTRARIO SU PROPUESTA PODRÁ SER DECLARADA INADMISIBLE.", centrado=True)


    # ANEXO N°8
    doc.add_section(WD_SECTION_START.NEW_PAGE)
    doc.add_heading("ANEXO Nº8 (10%)", level=1)
    doc.add_heading("PLAZO DE ENTREGA", level=2)
    doc.add_paragraph("PROPUESTA PÚBLICA: PARA EL SUMINISTRO DE INSUMOS Y ACCESORIOS PARA TERAPIA DE PRESIÓN NEGATIVA CON EQUIPOS EN COMODATO PARA EL HOSPITAL SAN JOSÉ DE MELIPILLA.").alignment = WD_ALIGN_PARAGRAPH.CENTER

    agregar_parrafo_con_texto(doc, 'Seleccionar Solo Una Opción Con Una "X" para el plazo de entrega de insumos:', centrado=False)

    anexo8_data = [
        ["PAUTA EVALUACION PLAZO DE ENTREGA PARA INSUMOS", "", ""], # Merged Header
        ["PLAZO DE ENTREGA", "PUNTAJE", "MARCAR CON\nUNA X"],
        ["Menor o igual a 2 días hábiles a partir del envío\nde la orden de compra.", "10 puntos", ""],
        ["3 a 4 días hábiles a partir del envío de la orden\nde compra.", "4 puntos", ""],
        ["5 a 6 días hábiles a partir del envío de la orden\nde compra.", "0 puntos", ""],
        [("Plazos Mayores a 6 días hábiles no será considerada su oferta por lo que será excluido automáticamente declarando \"", ""), ("inadmisible", "bold"), ("\" su propuesta.", ""), "Excluyente", ""], # Complex row
    ]
    tabla_anexo8 = doc.add_table(rows=len(anexo8_data), cols=len(anexo8_data[0]), style='Table Grid')

    # Fill content, handle special row 5
    for r_idx, row_data in enumerate(anexo8_data):
         if r_idx < 5:
              for c_idx, cell_content in enumerate(row_data):
                   tabla_anexo8.cell(r_idx, c_idx).text = cell_content
         else: # Row 5 handling
              # Text and bold for the first cell (index 0)
              cell0 = tabla_anexo8.cell(r_idx, 0)
              agregar_parrafo_con_runs(cell0, row_data[0])
              # Text for other cells (index 1, 2)
              tabla_anexo8.cell(r_idx, 1).text = row_data[1]
              tabla_anexo8.cell(r_idx, 2).text = row_data[2]


    # Merges
    tabla_anexo8.cell(0, 0).merge(tabla_anexo8.cell(0, 2)) # Main Header merge
    tabla_anexo8.cell(5, 0).merge(tabla_anexo8.cell(5, 2)) # Excluyente row merge (full width)

    # Formatting (Center and Bold Headers)
    for r_idx in [0, 1]:
         for c_idx in range(len(anexo8_data[0])):
              if tabla_anexo8.cell(r_idx, c_idx).text:
                   tabla_anexo8.cell(r_idx, c_idx).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                   tabla_anexo8.cell(r_idx, c_idx).paragraphs[0].runs[0].bold = True

    # Center Plazo and Puntaje cells in data rows
    for r_idx in range(2, 5):
         tabla_anexo8.cell(r_idx, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Center the "Excluyente" text in the merged cell
    tabla_anexo8.cell(5, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Bold the 'inadmisible' part within the merged cell runs (already handled by agregar_contenido_celda for row 5, col 0)


    doc.add_paragraph("\n")
    agregar_parrafo_con_texto(doc, "Fecha: _______/_______/________", centrado=False)
    doc.add_paragraph("\n")
    agregar_parrafo_con_texto(doc, "FIRMA Y TIMBRE OFERENTE O REPRESENTANTE LEGAL", centrado=True)
    doc.add_paragraph("\n")
    agregar_parrafo_con_texto(doc, "ESTE FORMULARIO DEBERA COMPLETARLO Y ADJUNTARLO A LA OFERTA, EN CASO CONTRARIO SU PROPUESTA PODRÁ SER DECLARADA INADMISIBLE.", centrado=True)


    # ANEXO N°9
    doc.add_section(WD_SECTION_START.NEW_PAGE)
    doc.add_heading("ANEXO N° 9 (10%)", level=1)
    doc.add_heading("SERVICIO POST VENTA", level=2)
    doc.add_paragraph("PROPUESTA PÚBLICA: PARA EL SUMINISTRO DE INSUMOS Y ACCESORIOS PARA TERAPIA DE PRESIÓN NEGATIVA CON EQUIPOS EN COMODATO PARA EL HOSPITAL SAN JOSÉ DE MELIPILLA.").alignment = WD_ALIGN_PARAGRAPH.CENTER

    anexo9_data = [
        ["N°", "Compromisos", "Seleccione la Alternativa Ofertada", ""], # Header
        ["", "", "SI", "NO"], # Sub-header
        ["1", "El Aseguramiento de stock para la totalidad de los productos adjudicados y durante todo el período establecido en bases de licitación", "1%", "Inadmisible"],
        ["2", "La entrega de 1 equipo Back up para uso hospitalario y 1 Equipo Back up para uso domiciliario", "1%", "Inadmisible"],
        ["3", "La entrega de asistencia Técnica vía Telefónica y/o correo Electrónico las 24 horas los 7 días de la semana.", "1%", "Inadmisible"],
        ["4", "Resolución de problemas técnicos en un periodo menor\na 72 horas", "3%", "1%\nMayor a\n96 horas\nInadmisible"], # Merged row conceptually in 3rd col
        ["5", "Asistencia en terreno menor a 48.", "3%", "1%\nMayor a\n72 horas\nInadmisible"], # Merged row conceptually in 3rd col
        ["6", "Capacitaciones teórico prácticas para el uso de los equipos según necesidad del Hospital.", "1%", "Inadmisible"],
        ["Puntaje Final", "", "", ""], # Merged cell
    ]
    tabla_anexo9 = doc.add_table(rows=len(anexo9_data), cols=4, style='Table Grid')

    # Fill content, handle complex cells
    for r_idx, row_data in enumerate(anexo9_data):
         if r_idx in [0, 1, 2, 3, 6, 8]: # Simple rows
              for c_idx, cell_content in enumerate(row_data):
                   tabla_anexo9.cell(r_idx, c_idx).text = cell_content
         elif r_idx in [4, 5]: # Rows with complex alternative options
              tabla_anexo9.cell(r_idx, 0).text = row_data[0]
              tabla_anexo9.cell(r_idx, 1).text = row_data[1]
              # Cell for SI/NO options (3 columns conceptually: SI, NO, Other Options)
              # Let's map them to the 2 columns (2 and 3) available
              tabla_anexo9.cell(r_idx, 2).text = row_data[2] # SI option
              tabla_anexo9.cell(r_idx, 3).text = row_data[3] # This cell contains multiple options in the OCR
              # Let's split the complex text in col 3 and put each option on a new line
              if "\n" in row_data[3]:
                   options_text = row_data[3].split("\n")
                   cell3 = tabla_anexo9.cell(r_idx, 3)
                   cell3.text = options_text[0] # First option
                   for opt in options_text[1:]:
                        cell3.add_paragraph(opt) # Subsequent options as new paragraphs in cell

    # Merges
    tabla_anexo9.cell(0, 2).merge(tabla_anexo9.cell(0, 3)) # Seleccione header merge
    tabla_anexo9.cell(8, 0).merge(tabla_anexo9.cell(8, 3)) # Puntaje Final merge

    # Formatting (Center and Bold Headers)
    for r_idx in [0, 1]:
         for c_idx in range(len(anexo9_data[0])):
              if tabla_anexo9.cell(r_idx, c_idx).text:
                   tabla_anexo9.cell(r_idx, c_idx).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                   tabla_anexo9.cell(r_idx, c_idx).paragraphs[0].runs[0].bold = True

    # Bold N° column in data rows
    for r_idx in range(2, 8):
         tabla_anexo9.cell(r_idx, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
         tabla_anexo9.cell(r_idx, 0).paragraphs[0].runs[0].bold = True

    # Bold Puntaje Final merged cell
    tabla_anexo9.cell(8, 0).paragraphs[0].runs[0].bold = True

    # Center SI/NO/Option cells in data rows
    for r_idx in range(2, 8):
         for c_idx in range(2, 4):
              tabla_anexo9.cell(r_idx, c_idx).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
              # If there are multiple paragraphs in cell 3 (for complex options)
              if r_idx in [4, 5]:
                  cell3 = tabla_anexo9.cell(r_idx, 3)
                  for p in cell3.paragraphs:
                      p.alignment = WD_ALIGN_PARAGRAPH.CENTER


    doc.add_paragraph("\n")
    agregar_parrafo_con_texto(doc, "Fecha: _______/_______/________", centrado=False)
    doc.add_paragraph("\n")
    agregar_parrafo_con_texto(doc, "FIRMA Y TIMBRE OFERENTE O REPRESENTANTE LEGAL", centrado=True)
    doc.add_paragraph("\n")
    agregar_parrafo_con_texto(doc, "NOTA: ESTE FORMULARIO DEBERÁ ADJUNTARLO OBLIGATORIAMENTE, EN CASO CONTRARIO SU PROPUESTA SERÁ DECLARADA INADMISIBLE.", centrado=True)


    # Final Resolution Text (Page 43)
    doc.add_section(WD_SECTION_START.NEW_PAGE)
    final_resolution_items = [
        (("4.- IMPÚTESE", "bold"), (" el gasto que se autoriza por la presente resolución, según clasificador presupuestario año 2024 al ITEM 22.04.005.003 \"Otros Insumos Clínicos\" del Hospital San José de Melipilla.", "")),
        (("5.- PUBLÍQUESE", "bold"), (" el presente llamado e ", ""), ("INFÓRMESE", "bold"), (" el proceso de licitación de la Adquisición requerida en el Sistema de Información establecido por la Dirección de Compras y Contratación Pública, de acuerdo a los dispuesto en el artículo 20 de la Ley N° 19.866 y su Reglamento con sus respectivas modificaciones.", "")),
        (("6.- DESÍGNASE", "bold"), (" como integrantes de la Comisión de Evaluación de la propuesta a los siguientes funcionarios: el Subdirector(a) Administrativo, subdirector(a) Médico Atención Abierta, subdirector(a) Médico Atención Cerrada, subdirector(a) de Gestión de Cuidado de Enfermería, subdirector(a) de Gestión y Desarrollo de las Personas, subdirector(a) de Matronería, subdirector(a) de Análisis de Información para la Gestión, subdirector de apoyo clínico o sus subrogantes. La misma Comisión estudiará los antecedentes de la Propuesta y elaborará un informe fundado para el director de este establecimiento, acerca de la conveniencia de resolver al respecto. Para los efectos del quórum para sesionar se requerirá un mínimo de tres miembros.", "")),
        (("7.- CORRESPONDERÁ", "bold"), (" al encargado(a) del proceso de Licitaciones coordinar y velar por el cumplimiento del procedimiento de la Propuesta Pública.", "")),
    ]
    for item in final_resolution_items:
         p = doc.add_paragraph()
         if isinstance(item, tuple):
             for part in item:
                 if isinstance(part, tuple):
                      p.add_run(part[0]).bold = True
                 else:
                      p.add_run(part)
         else:
             p.add_run(item)
         # Apply numbering (4, 5, 6, 7) - needs a new list ID
         aplicar_numeracion(p, num_id_final_resolution, nivel=0)


    agregar_parrafo_con_texto(doc, "ANÓTESE, PUBLÍQUESE Y ARCHÍVESE.", centrado=True)
    doc.add_paragraph("\n") # Space before signature

    # Signature block
    doc.add_paragraph("DR. OSCAR VARGAS DURANTI").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("DIRECTOR").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("HOSPITAL SAN JOSÉ DE MELIPILLA").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("\n") # Space

    # Distribution Block (Simulate alignment with spaces/table or just text)
    doc.add_paragraph("TRANSCRITO FIELMENTE").alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.add_paragraph("MARISOL ARAVENA REYYES").alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.add_paragraph("MINISTRO DE FE (S)").alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.add_paragraph("\n") # Space

    doc.add_paragraph("CRE/JHF/EYJ/ICL/MES/MMJ/FRA/MLG")
    doc.add_paragraph("Distribución Impresa:")
    distribucion_impresa_list = [
         "Unidad de Abastecimiento",
         "Oficina de Partes",
    ]
    for item in distribucion_impresa_list:
         agregar_parrafo_con_texto(doc, item, style='List Bullet')

    doc.add_paragraph("Distribución Digital:")
    distribucion_digital_list = [
         "Auditoria",
         "Unidad de Farmacia",
         "Bodega de Farmacia",
         "Servicio de Pabellón",
         "Unidad de Finanzas",
         "Unidad de Contabilidad",
    ]
    for item in distribucion_digital_list:
         agregar_parrafo_con_texto(doc, item, style='List Bullet')


    # Add page numbers to the footer
    # Note: This will add page numbers to *all* sections created so far.
    # If you need different headers/footers per section, you'd configure them
    # after adding each section break. Simple sequential numbering is applied here.
    for section in doc.sections:
        footer = section.footer
        if not footer.paragraphs:
            footer.add_paragraph()
        footer_paragraph = footer.paragraphs[0]
        footer_paragraph.text = "Page "
        run = footer_paragraph.add_run()
        run.add_field('PAGE')
        footer_paragraph.add_run(' of ')
        run = footer_paragraph.add_run()
        run.add_field('NUMPAGES')
        footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT


    # Save the document
    doc_path = 'Resolucion_Hospital_Melipilla.docx'
    doc.save(doc_path)
    print(f"Documento guardado como: {doc_path}")

if __name__ == "__main__":
    main()




"""import docx
import os
from Bases import configurar_directorio_trabajo

configurar_directorio_trabajo()

doc = docx.Document()

# Agregar título
doc.add_heading("Ejemplos de Estilos de Lista", 0)

# List
doc.add_paragraph("Este es un ejemplo de estilo 'List'", style="List")

# List 2
doc.add_paragraph("Este es un ejemplo de estilo 'List 2'", style="List 2")

# List 3
doc.add_paragraph("Este es un ejemplo de estilo 'List 3'", style="List 3")

# List Bullet
doc.add_paragraph("Este es un ejemplo de estilo 'List Bullet'", style="List Bullet")

# List Bullet 2
doc.add_paragraph("Este es un ejemplo de estilo 'List Bullet 2'", style="List Bullet 2")

# List Bullet 3
doc.add_paragraph("Este es un ejemplo de estilo 'List Bullet 3'", style="List Bullet 3")

# List Continue
doc.add_paragraph("Este es un ejemplo de estilo 'List Continue'", style="List Continue")

# List Continue 2
doc.add_paragraph("Este es un ejemplo de estilo 'List Continue 2'", style="List Continue 2")

# List Continue 3
doc.add_paragraph("Este es un ejemplo de estilo 'List Continue 3'", style="List Continue 3")

# List Number
doc.add_paragraph("Este es un ejemplo de estilo 'List Number'", style="List Number")

# List Number 2
doc.add_paragraph("Este es un ejemplo de estilo 'List Number 2'", style="List Number 2")

# List Number 3
doc.add_paragraph("Este es un ejemplo de estilo 'List Number 3'", style="List Number 3")

# List Paragraph
doc.add_paragraph("Este es un ejemplo de estilo 'List Paragraph'", style="List Paragraph")

# Guardar documento
doc.save("ejemplos_estilos_lista.docx")"""

