import os
import re
import random
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_BREAK
from docx.shared import Pt


def configurar_directorio_trabajo():
    """Configura el directorio de trabajo en la subcarpeta 'Files'."""
    cwd = os.getcwd()
    target_dir_name = "Files"
    wd = os.path.join(cwd, target_dir_name)
    pattern = r"Files\\Files"
    if re.search(pattern, wd):
        wd = wd.replace(r"\Files\Files", r"\Files")
    if os.path.isdir(wd):
        os.chdir(wd)
        print(f"Directorio de trabajo cambiado a: {wd}")
    else:
        print(f"Advertencia: El directorio '{wd}' no existe o no es válido.")

def agregar_bookmark(parrafo, bookmark_id, bookmark_name):
    """
    Agrega un bookmark a un párrafo específico.
    :param parrafo: El párrafo al que se le añadirá el bookmark.
    :param bookmark_id: ID único para el bookmark.
    :param bookmark_name: Nombre del bookmark.
    """
    p = parrafo._p
    # Crear elemento de inicio del bookmark
    bookmark_start = OxmlElement('w:bookmarkStart')
    bookmark_start.set(qn('w:id'), str(bookmark_id))
    bookmark_start.set(qn('w:name'), bookmark_name)
    p.append(bookmark_start)

    # Crear elemento de fin del bookmark (después del contenido del párrafo)
    bookmark_end = OxmlElement('w:bookmarkEnd')
    bookmark_end.set(qn('w:id'), str(bookmark_id))
    p.append(bookmark_end)

def crear_numeracion(doc):
    """Crea un formato de numeración y devuelve su ID."""
    part = doc._part
    if not hasattr(part, 'numbering_part'):
        part._add_numbering_part()
    return random.randint(1000, 9999)

def aplicar_numeracion(parrafo, num_id, nivel=0):
    """Aplica numeración a un párrafo con el ID y nivel especificados."""
    p = parrafo._p
    pPr = p.get_or_add_pPr()
    for child in pPr.iterchildren():
        if child.tag.endswith('numPr'):
            pPr.remove(child)
    numPr = OxmlElement('w:numPr')
    ilvl = OxmlElement('w:ilvl')
    ilvl.set(qn('w:val'), str(nivel))
    numPr.append(ilvl)
    numId = OxmlElement('w:numId')
    numId.set(qn('w:val'), str(num_id))
    numPr.append(numId)
    pPr.append(numPr)
    ind = OxmlElement('w:ind')
    ind.set(qn('w:left'), '720')
    ind.set(qn('w:hanging'), '360')
    pPr.append(ind)
    return parrafo

def centrar_verticalmente_tabla(tabla):
    """Aplica alineación vertical centrada a todas las celdas de una tabla."""
    for fila in tabla.rows:
        for celda in fila.cells:
            tc = celda._tc
            tcPr = tc.get_or_add_tcPr()
            for vAlign in tcPr.findall(qn('w:vAlign')):
                tcPr.remove(vAlign)
            vAlign = OxmlElement('w:vAlign')
            vAlign.set(qn('w:val'), 'center')
            tcPr.append(vAlign)

def agregar_parrafo_con_texto(doc, texto, estilo=None, negrita=False):
    """Agrega un párrafo con texto y aplica estilo o formato si se especifica."""
    p = doc.add_paragraph(texto, style=estilo)
    if negrita and p.runs:
        p.runs[0].bold = True
    return p

def agregar_parrafo_con_runs(doc, partes, estilo=None):
    """Agrega un párrafo con múltiples runs y aplica formatos específicos."""
    p = doc.add_paragraph(style=estilo)
    for texto, formato in partes:
        run = p.add_run(texto)
        if formato:
            for fmt in formato.split(','):
                if fmt == 'bold':
                    run.bold = True
                elif fmt == 'underline':
                    run.underline = True
    return p

def crear_tabla(doc, datos, estilo='Table Grid', centrar=True):
    """Crea una tabla con los datos proporcionados y aplica formato."""
    filas = len(datos)
    columnas = len(datos[0]) if filas > 0 else 0
    tabla = doc.add_table(filas, columnas, style=estilo)
    for i, fila in enumerate(datos):
        for j, texto in enumerate(fila):
            celda = tabla.cell(i, j)
            celda.text = texto
    if centrar:
        centrar_verticalmente_tabla(tabla)
    return tabla

def agregar_contenido_celda(tabla, fila, columna, contenidos):
    """Agrega contenido a una celda con múltiples párrafos o runs formateados."""
    celda = tabla.cell(fila, columna)
    if celda.paragraphs and not celda.paragraphs[0].text:
        p_element = celda.paragraphs[0]._p
        celda._element.remove(p_element)
    for contenido in contenidos:
        if isinstance(contenido, str):
            celda.add_paragraph(contenido)
        else:
            p = celda.add_paragraph()
            for item in contenido:
                if isinstance(item, tuple) and len(item) == 2:
                    texto, formato = item
                    run = p.add_run(texto)
                    if formato:
                        for fmt in formato.split(','):
                            if fmt == 'bold':
                                run.bold = True
                            elif fmt == 'underline':
                                run.underline = True
                else:
                    # If it's not a tuple or has unexpected length, treat it as plain text
                    run = p.add_run(str(item))

# Formated_Base_PEP8.py
from docx.enum.text import WD_ALIGN_PARAGRAPH

from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def aplicar_formato_global(doc):
    """Aplica Calibri Light 11 y justificado, excepto portada y título principal."""
    # 1) Detectar hasta dónde llega la portada
    limite_portada = 0
    for i, paragraph in enumerate(doc.paragraphs):
        # lógica original para fijar límite_portada
        ...

    # 2) Estilos a excluir del justificado
    excluded_styles = ['Title']

    # 3) Aplicar formato al resto
    for i, paragraph in enumerate(doc.paragraphs):
        if i <= limite_portada or paragraph.style.name in excluded_styles:
            continue
        paragraph.style.font.name = 'Calibri Light'
        paragraph.style.font.size = Pt(11)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # 4) Justificar también el contenido de tablas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if p.style.name in excluded_styles:
                        continue
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

#numero_base = str(140)

def main():
    configurar_directorio_trabajo()
    doc = Document("portada_melipilla_base.docx")
    list_style = 'List Number'
    if list_style not in doc.styles:
        doc.styles.add_style(list_style, WD_STYLE_TYPE.PARAGRAPH)

    # Crear IDs para numeración
    num_id_vistos = crear_numeracion(doc)
    num_id_resolucion = crear_numeracion(doc)
    num_id_bases_p1 = crear_numeracion(doc)
    num_id_requisitos = crear_numeracion(doc)
    num_id_consultas = crear_numeracion(doc)
    administrado_contrato_id_lista = crear_numeracion(doc)
    pacto_integridad_id = crear_numeracion(doc)

    # Títulos y secciones principales
    heading_paragraph = doc.add_heading('', level=0)

    # Add the first part of the text as a run
    run1 = heading_paragraph.add_run("RESOLUCIÓN EXENTA Nº1")
    run1.font.size = Pt(11)

    # Add a line break
    #run1.add_break(WD_BREAK.LINE)

    # Add the second part of the text as another run
    run2 = heading_paragraph.add_run("MELIPILLA")
    heading_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run2.font.size = Pt(11)

    doc.add_heading("VISTOS", level=2)
    agregar_parrafo_con_texto(doc, "Lo dispuesto en la Ley Nº 19.886 de Bases sobre Contratos Administrativos de Suministro y Prestación de Servicios; el Decreto Supremo Nº 250 /04 modificado por los Decretos Supremos Nº 1763/09, 1383/11 y 1410/14 todos del Ministerio de Hacienda; D. S. 38/2005, Reglamento Orgánico de los Establecimientos de Menor Complejidad y de los Establecimientos de Autogestión en Red; en uso de las atribuciones que me confieren el D.F.L. Nº 1/2.005, en virtud del cual se fija el texto refundido, coordinado y sistematizado del D.L. 2.763/79 y de las leyes 18.933 y 18.469; lo establecido en los Decretos Supremos Nos 140/04, Reglamento Orgánico de los Servicios de Salud; {{ director }}; lo dispuesto por las Resoluciones 10/2017, 7/2019 y 8/2019 ambas de la Contraloría General de la República, y,")
    # la Resolución Exenta RA 116395/343/2024 de fecha 12/08/2024 del SSMOCC., la cual nombra Director del Hospital San José de Melipilla al suscrito

    # Sección CONSIDERANDO
    doc.add_heading("CONSIDERANDO", level=2)
    vistos_items = [
        ("Que dada la alta complejidad que caracteriza al Hospital San José de Melipilla, obliga a efectuar mejoras constantes y permanentes a fin de brindar a toda nuestra comunidad el desarrollo de diversas funciones con alta calidad que el sistema público puede brindar.", list_style),
        ("Que, el Hospital de San José de Melipilla perteneciente a la red de salud del Servicio de Salud Metropolitano Occidente, tiene como misión otorgar una atención integral, oportuna y resolutiva a las personas y familias de la provincia de Melipilla y sus alrededores, con un equipo de salud competente, comprometido y solidario, entregando un servicio de calidad y seguridad, en coordinación con la red asistencial;", list_style),
        ("Que, dada la naturaleza del Establecimiento, la atención de los beneficiarios requiere una oportuna e inmediata resolución, que no puede en caso alguno diferirse en el tiempo, lo que nos compromete a disponer en forma constante, continua y permanente de los servicios necesarios para responder adecuadamente a la demanda asistencial y administrativa a su población beneficiaria.", list_style),
    ]
    for texto, estilo in vistos_items:
        p = agregar_parrafo_con_texto(doc, texto, estilo)
        aplicar_numeracion(p, num_id_vistos)

    vistos_p4 = doc.add_paragraph(style=list_style)
    vistos_p4.add_run("Que, existe la necesidad ")
    run_bold = vistos_p4.add_run("{{ nombre_adquisicion }}")
    run_bold.bold = True
    vistos_p4.add_run(", a fin de entregar una prestación de salud integral y oportuna a los usuarios del Hospital de San José de Melipilla, y de esta manera dar cumplimiento con el tratamiento de los pacientes.")
    aplicar_numeracion(vistos_p4, num_id_vistos)

    for texto in [
        "Que corresponde asegurar la transparencia en este proceso y conocer las condiciones de oferta imperantes en el mercado bajo la modalidad de la licitación pública en el sistema de compras y contratación públicas establecido en la Ley Nº 19.886 y su Reglamento.",
        "Que, considerando los montos de la contratación y en virtud de lo establecido en las resoluciones N°7/2019 y 16/2020 de la Contraloría General de la República, la presenta contratación no está sometida al trámite de toma de razón.",
        "Que revisado el catálogo de bienes y servicios ofrecidos en el sistema de información Mercado Público, se ha verificado la ausencia de contratos marcos vigentes para el servicio antes mencionado.",
        "Que, en consecuencia y en mérito de lo expuesto, para esta contratación se requiere llamar a licitación pública, debiendo esta regularse por la Bases Administrativas, Técnicas, Formularios y Anexos que se aprueban a través del presente acto administrativo.",
        "Que, en razón de lo expuesto y la normativa vigente;"
    ]:
        p = agregar_parrafo_con_texto(doc, texto, list_style)
        aplicar_numeracion(p, num_id_vistos)

    # Sección RESOLUCIÓN
    doc.add_heading("RESOLUCIÓN", level=2)
    resolucion_p1 = doc.add_paragraph(style=list_style)
    resolucion_p1.add_run("LLÁMASE ").bold = True
    resolucion_p1.add_run("a Licitación Pública Nacional a través del Portal Mercado Público, para la compra de ")
    resolucion_p1.add_run("{{ nombre_adquisicion }}").bold = True
    resolucion_p1.add_run("para el Hospital San José de Melipilla.")
    aplicar_numeracion(resolucion_p1, num_id_resolucion)

    resolucion_p2 = doc.add_paragraph(style=list_style)
    resolucion_p2.add_run("ACOGIENDOSE ").bold = True
    resolucion_p2.add_run("al Art.º 25 del decreto 250 que aprueba el reglamento de la ley Nº 19.886...")
    aplicar_numeracion(resolucion_p2, num_id_resolucion)

    resolucion_p3 = doc.add_paragraph(style=list_style)
    resolucion_p3.add_run("APRUÉBENSE las bases administrativas, técnicas y anexos N.º 1, 2, 3, 4, 5 {{ cantidad_anexos }}").bold = True
    resolucion_p3.add_run("desarrollados para efectuar el llamado a licitación, que se transcriben a continuación:")
    aplicar_numeracion(resolucion_p3, num_id_resolucion)

    # Sección BASES ADMINISTRATIVAS
    doc.add_section()
    main_header = doc.add_heading("",level=1)
    main_header.add_run("BASES ADMINISTRATIVAS PARA EL ")
    main_header.add_run("{{ nombre_adquisicion }}").upper = True
        #= doc.add_heading("BASES ADMINISTRATIVAS PARA EL {{ SUMINISTRO DE INSUMOS Y ACCESORIOS PARA TERAPIA DE PRESIÓN NEGATIVA CON EQUIPOS EN COMODATO PARA EL HOSPITAL SAN JOSÉ DE MELIPILLA }}", level=1)
    # SUMINISTRO DE INSUMOS Y ACCESORIOS PARA TERAPIA DE PRESIÓN NEGATIVA CON EQUIPOS EN COMODATO PARA EL HOSPITAL SAN JOSÉ DE MELIPILLA
    doc.add_heading("Antecedentes y Plazos", level=2)
    agregar_parrafo_con_texto(doc, "En Santiago, a 1 de enero de 2023, se resuelve lo siguiente:")

    bases_p1 = doc.add_paragraph(style=list_style)
    bases_p1.add_run("Antecedentes Básicos de la ENTIDAD LICITANTE").bold = True
    aplicar_numeracion(bases_p1, num_id_bases_p1)

    tabla_p1_datos = [
        ["Razón Social del organismo", "Hospital San José de Melipilla"],
        ["Unidad de Compra", "Unidad de Abastecimiento de Bienes y Servicios"],
        ["R.U.T. del organismo", "61.602.123-0"],
        ["Dirección", "O'Higgins #551"],
        ["Comuna", "Melipilla"],
        ["Región en que se genera la Adquisición", "Región Metropolitana"]
    ]
    tabla_p1 = crear_tabla(doc, tabla_p1_datos)

    bases_p2 = doc.add_paragraph(style=list_style)
    bases_p2.add_run("Antecedentes Administrativos").bold = True
    aplicar_numeracion(bases_p2, num_id_bases_p1)

    tabla_p2_datos = [
        ["Nombre Adquisición", "{{ nombre_adquisicion }}"],
        ["Descripción", "El Hospital requiere generar un convenio por el {{ nombre_adquisicion }}, en adelante “EL HOSPITAL”. El convenio tendrá una vigencia de {{ plazo_meses }} meses."],
        ["Tipo de Convocatoria", "Abierta"],
        ["Moneda o Unidad reajustable", "Pesos Chilenos"],
        ["Presupuesto Referencial", " {{ presupuesto_con_impuestos }}.- (Impuestos incluidos)"],
        ["Etapas del Proceso de Apertura", "Una Etapa (Etapa de Apertura Técnica y Etapa de Apertura Económica en una misma instancia)."],
        ["Opciones de pago", "Transferencia Electrónica"],
        ["Tipo de Adjudicación", "{{ tipo_adjudicacion }}"]
    ]
    tabla_p2 = crear_tabla(doc, tabla_p2_datos)

    bases_p2_runs = doc.add_paragraph()
    bases_p2_runs.add_run("* Presupuesto referencial:").underline = True
    bases_p2_runs.add_run(" El Hospital se reserva el derecho de aumentar, previo acuerdo entre las partes, hasta un 30% el presupuesto referencial estipulado en las presentes bases de licitación.")

    # Definiciones
    bases_p2_definiciones = doc.add_paragraph()
    bases_p2_definiciones.add_run("Definiciones").bold = True

    definiciones = [
        ("Proponente u oferente:", "El proveedor o prestador que participa en el proceso de licitación mediante la presentación de una propuesta, en la forma y condiciones establecidas en las Bases.", "List Number 3"),
        ("Administrador o coordinador Externo del Contrato", "Persona designada por el oferente adjudicado, quien actuará como contraparte ante el Hospital.", "List Number 3"),
        ("Días Hábiles:", "Son todos los días de la semana, excepto los sábados, domingos y festivos.", "List Number 3"),
        ("Días Corridos:", "Son los días de la semana que se computan uno a uno en forma correlativa. Salvo que se exprese lo contrario, los plazos de días señalados en las presentes bases de licitación son días corridos. En caso que el plazo expire en días sábados, domingos o festivos se entenderá prorrogados para el día hábil siguiente.", "List Number 3"),
        ("Administrador del Contrato y/o Referente Técnico:", "Es el funcionario designado por el Hospital para supervisar la correcta ejecución del contrato, solicitar órdenes de compra, validar prefacturas, gestionar multas y/o toda otra labor que guarde relación con la ejecución del contrato.", "List Number 3"),
        ("Gestor de Contrato:", "Es el funcionario a cargo de la ejecución del presente proceso de Licitación, desde la publicación de las Bases hasta la generación del contrato en formato documental, como la elaboración de ficha en la plataforma “gestor de contrato” en el portal de mercado público, además de ser el responsable de dar seguimiento y cumplimiento a los procesos y plazos establecidos.", "List Number 3")
    ]
    for titulo, descripcion, estilo in definiciones:
        p = doc.add_paragraph(style=estilo)
        p.add_run(f"{titulo}").bold = True
        p.add_run(f" {descripcion}")

    # Tabla de tipos de licitación
    tabla_licitaciones_datos = [
        ["RANGO (en UTM)", "TIPO LICITACION PUBLICA", "PLAZO PUBLICACION EN DIAS CORRIDOS"],
        ["<100", "L1", "5"],
        ["<=100 y <1000", "LE", "10, rebajable a 5"],
        ["<=1000 y <2000", "LP", "20, rebajable a 10"],
        ["<=2000 y <5000", "LQ", "20, rebajable a 10"],
        ["<=5000", "LR", "30"]
    ]
    crear_tabla(doc, tabla_licitaciones_datos)

    bases_p3 = doc.add_paragraph(style=list_style)
    bases_p3.add_run("Etapas y plazos:").bold = True
    aplicar_numeracion(bases_p3, num_id_bases_p1)

    tabla_plazos_datos = [
        ["VIGENCIA DE LA PUBLICACION {{ dias_vigencia_publicacion }} DIAS CORRIDOS", "", ""],
        ["Consultas", "Hasta las 15:00 Horas de {{ plazo_consultas }} día corrido de publicada la Licitación.", "Deben ser ingresadas al portal www.mercadopublico.cl"],
        ["Respuestas a Consultas", "Hasta las 17:00 Horas de {{ plazo_respuesta }} día corrido de publicada la Licitación.", "Deben ser ingresadas al portal www.mercadopublico.cl"],
        ["Aclaratorias", "Hasta 1 días corrido antes del cierre de recepción de ofertas.", "Deben ser ingresadas al portal www.mercadopublico.cl"],
        ["Recepción de ofertas", "Hasta las 17:00 Horas de {{ plazo_recepcion_ofertas }} día corrido de publicada la Licitación.", "Deben ser ingresadas al portal www.mercadopublico.cl"],
        ["Evaluación de las Ofertas", "Máximo 40 días corridos a partir del cierre de la Licitación.", ""],
        ["Plazo Adjudicaciones", "Máximo 20 días {{ corridos }} a partir de la fecha del acta de evaluación de las ofertas.", ""],
        ["Suscripción de Contrato", "Máximo de {{ plazo_suscripcion }} desde la Adjudicación de la Licitación.", ""],
        ["Consideración", "Los plazos de días establecidos en la cláusula 3, Etapas y Plazos, son de días corridos, excepto el plazo para emitir la orden de compra, el que se considerará en días hábiles, entendiéndose que son inhábiles los sábados, domingos y festivos en Chile, sin considerar los feriados regionales.", ""]
    ]
    tabla_plazos = crear_tabla(doc, tabla_plazos_datos)
    tabla_plazos.cell(0, 0).merge(tabla_plazos.cell(0, 2))
    tabla_plazos.cell(0, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    tabla_plazos.cell(0, 0).paragraphs[0].runs[0].bold = True

    # Sección Consultas, Aclaraciones y Modificaciones
    doc.add_heading("Consultas, Aclaraciones y modificaciones a las bases.", level=2)
    consultas_items = [
        [("Las consultas de los participantes se deberán realizar únicamente a través del portal ", ""),
         ("www.mercadopublico.cl", "bold"),
         (" conforme el cronograma de actividades de esta licitación señalado en el punto 3 precedente. A su vez, las respuestas y aclaraciones estarán disponibles a través del portal de Mercado Público, en los plazos indicados en el cronograma señalado precedentemente, información que se entenderá conocida por todos los interesados desde el momento de su publicación.", "")],
        [("No serán admitidas las consultas formuladas fuera de plazo o por un conducto diferente al señalado.", "bold")],
        [("“EL HOSPITAL” realizará las aclaraciones a las Bases comunicando las respuestas a través del Portal Web de Mercado Público, sitio", ""),
         (" www.mercadopublico.cl", "bold")],
        [("Las aclaraciones, derivadas de este proceso de consultas, formarán parte integrante de las Bases, teniéndose por conocidas y aceptadas por todos los participantes aun cuando el oferente no las hubiere solicitado, por lo que los proponentes no podrán alegar desconocimiento de las mismas.", "")],
        [("“EL HOSPITAL” podrá modificar las presentes bases y sus anexos previa autorización por acto administrativo, durante el periodo de presentación de las ofertas, hasta antes de fecha de cierre de recepción de ofertas. Estas modificaciones, que se lleven a cabo, serán informadas a través del portal", ""),
         (" www.mercadopublico.cl", "bold")],
        [("Estas consultas, aclaratorias y modificaciones formaran parte integra de las bases y estarán vigentes desde la total tramitación del acto administrativo que las apruebe. Junto con aprobar las modificaciones, deberá establecer un nuevo plazo prudencial cuando lo amerite para el cierre o recepción de las propuestas, a fin de que los potenciales oferentes puedan adecuar sus ofertas.", "")],
        [("No se aceptarán consultas realizadas por otros medios, tales como correos electrónicos, fax u otros.", "")]
    ]
    for item in consultas_items:
        p = agregar_parrafo_con_runs(doc, item)

    # Requisitos Mínimos para Participar
    doc.add_heading("Requisitos Mínimos para Participar.", level=2)
    requisitos_items = [
        "No haber sido condenado por prácticas antisindicales, infracción a los derechos fundamentales del trabajador o por delitos concursales establecidos en el Código Penal dentro de los dos últimos años anteriores a la fecha de presentación de la oferta, de conformidad con lo dispuesto en el artículo 4° de la ley N° 19.886.",
        "No haber sido condenado por el Tribunal de Defensa de la Libre Competencia a la medida dispuesta en la letra d) del artículo 26 del Decreto con Fuerza de Ley N°1, de 2004, del Ministerio de Economía, Fomento y Reconstrucción, que Fija el texto refundido, coordinado y sistematizado del Decreto Ley N° 211, de 1973, que fija normas para la defensa de la libre competencia, hasta por el plazo de cinco años contado desde que la sentencia definitiva quede ejecutoriada.",
        "No ser funcionario directivo de la respectiva entidad compradora; o una persona unida a aquél por los vínculos de parentesco descritos en la letra b) del artículo 54 de la ley N° 18.575; o una sociedad de personas de las que aquél o esta formen parte; o una sociedad comandita por acciones o anónima cerrada en que aquélla o esta sea accionista; o una sociedad anónima abierta en que aquél o esta sean dueños de acciones que representen el 10% o más del capital; o un gerente, administrador, representante o director de cualquiera de las sociedades antedichas.",
        "Tratándose exclusivamente de una persona jurídica, no haber sido condenada conforme a la ley N° 20.393 a la pena de prohibición de celebrar actos y contratos con el Estado, mientras esta pena esté vigente.",
        "A fin de acreditar el cumplimiento de dichos requisitos, los oferentes deberán presentar una “Declaración jurada de requisitos para ofertar”, la cual será generada completamente en línea a través de www.mercadopublico.cl en el módulo de presentación de las ofertas. Sin perjuicio de lo anterior, la entidad licitante podrá verificar la veracidad de la información entregada en la declaración, en cualquier momento, a través de los medios oficiales disponibles."
    ]
    for texto in requisitos_items:
        p = agregar_parrafo_con_texto(doc, texto)
        aplicar_numeracion(p, num_id_requisitos)

    req_p6 = doc.add_paragraph()
    req_p6.add_run("En caso de que los antecedentes administrativos solicitados en esta sección no sean entregados y/o completados en forma correcta y oportuna, se desestimará la propuesta, no será evaluada y será declarada ")
    req_p6.add_run("inadmisible").bold = True
    req_p6.add_run(".")
    aplicar_numeracion(req_p6, num_id_requisitos)

    # Instrucciones para la Presentación de Ofertas
    doc.add_heading("Instrucciones para la Presentación de Ofertas.", level=3)
    tabla_ofertas = doc.add_table(rows=4, cols=2, style='Table Grid')
    tabla_ofertas.cell(0, 0).text = "Presentar Ofertas por Sistema."
    tabla_ofertas.cell(0, 1).text = "Obligatorio."

    agregar_contenido_celda(tabla_ofertas, 1, 0, ["Anexos Administrativos."])
    agregar_contenido_celda(tabla_ofertas, 1, 1, [
        [("Anexo N° 1 Identificación del Oferente.", "bold")],
        [("Anexo N° 2 Declaración Jurada de Habilidad.", "bold")],
        [("Anexo N° 3 Declaración Jurada de Cumplimiento de Obligaciones Laborales y Previsionales.", "bold")],
        [("Declaración jurada online:", "bold"), (" Los oferentes deberán presentar una ", ""), ("Declaración jurada de requisitos para ofertar", "bold"), (", la cual será generada completamente en línea a través de www.mercadopublico.cl en el módulo de presentación de las ofertas.", "")],
        [("Unión Temporal de Proveedores (UTP):", "bold"), (" Solo en el caso de que la oferta sea presentada por una unión temporal de proveedores deberán presentar obligatoriamente la siguiente documentación en su totalidad, en caso contrario, ésta no será sujeta a aclaración y la oferta será declarada ", ""), ("inadmisible", "bold"), (".", "")],
        [("Anexo N°4. Declaración para Uniones Temporales de Proveedores:", "bold"), (" Debe ser presentado por el miembro de la UTP que presente la oferta en el Sistema de Información y quien realiza la declaración a través de la “Declaración jurada de requisitos para ofertar” electrónica presentada junto a la oferta.", "")],
        ["Las ofertas presentadas por una Unión Temporal de Proveedores (UTP) deberán contar con un apoderado, el cual debe corresponder a un integrante de la misma, ya sea persona natural o jurídica. En el caso que el apoderado sea una persona jurídica, ésta deberá actuar a través de su representante legal para ejercer sus facultades."],
        [("En caso de no presentarse debidamente la declaración jurada online constatando la ausencia de conflictos de interés e inhabilidades por condenas, o no presentarse el Anexo N°4, la oferta será declarada ", ""), ("inadmisible", "bold"), (".", "")]
    ])

    agregar_contenido_celda(tabla_ofertas, 2, 0, ["Anexos\nEconómicos."])
    agregar_contenido_celda(tabla_ofertas, 2, 1, [
        [("Anexo N°5: Oferta económica", "bold")],
        ["El anexo referido debe ser ingresado a través del sistema www.mercadopublico.cl , en la sección Anexos Económicos."],
        [("En caso de que no se presente debidamente el Anexo N°5 “Oferta económica”, la oferta será declarada ", ""), ("inadmisible", "bold")]
    ])

    agregar_contenido_celda(tabla_ofertas, 3, 0, ["Anexos Técnicos."])
    agregar_contenido_celda(tabla_ofertas, 3, 1, [
        [("{{anexo_6}}", "bold")],
        [("{{anexo_7}}", "bold")],
        [("{{anexo_8}}", "bold")],
        [("{{anexo_9}}", "bold")],
        ["Los anexos referidos deben ser ingresados a través del sistema www.mercadopublico.cl. en la sección Anexos Técnicos."],
        [("{{ausencia_para_inadmisible}} ", ""), ("inadmisible", "bold")]
    ])
    centrar_verticalmente_tabla(tabla_ofertas)

    # Observaciones
    doc.add_heading("Observaciones", level=3)
    agregar_parrafo_con_runs(doc, [
        ("Los oferentes deberán presentar su oferta a través de su cuenta en el Sistema de Información www.mercadopublico.cl. De existir discordancia entre el oferente o los antecedentes de su oferta y la cuenta a través de la cual la presenta, esta no será evaluada, siendo desestimada del proceso y declarada como", ""),
        (" inadmisible", "bold")
    ])

    agregar_parrafo_con_runs(doc, [
        ("Las únicas ofertas válidas serán las presentadas a través del portal", ""),
        (" www.mercadopublico.cl", "bold"),
        (", en la forma en que se solicita en estas bases. No se aceptarán ofertas que se presenten por un medio distinto al establecido en estas Bases, a menos que se acredite la indisponibilidad técnica del sistema, de conformidad con el artículo 62 del Reglamento de la Ley de Compras. Será responsabilidad de los oferentes adoptar las precauciones necesarias para ingresar oportuna y adecuadamente sus ofertas.", "")
    ])

    agregar_parrafo_con_texto(doc, "Los oferentes deben constatar que el envío de su oferta a través del portal electrónico de compras públicas haya sido realizado con éxito, incluyendo el previo ingreso de todos los formularios y anexos requeridos completados de acuerdo con lo establecido en las presentes bases. Debe verificar que los archivos que se ingresen contengan efectivamente los anexos solicitados.")

    agregar_parrafo_con_texto(doc, "Asimismo, se debe comprobar siempre, luego de que se finalice la última etapa de ingreso de la oferta respectiva, que se produzca el despliegue automático del “Comprobante de Envío de Oferta” que se entrega en dicho Sistema, el cual puede ser impreso por el proponente para su resguardo. En dicho comprobante será posible visualizar los anexos adjuntos, cuyo contenido es de responsabilidad del oferente.")

    agregar_parrafo_con_texto(doc, "El hecho de que el oferente haya obtenido el “Comprobante de envío de ofertas” señalado, únicamente acreditará el envío de ésta a través del Sistema, pero en ningún caso certificará la integridad o la completitud de ésta, lo cual será evaluado por la comisión evaluadora. En caso de que, antes de la fecha de cierre de la licitación, un proponente edite una oferta ya enviada, deberá asegurarse de enviar nuevamente la oferta una vez haya realizado los ajustes que estime, debiendo descargar un nuevo Comprobante.")

    agregar_parrafo_con_runs(doc, [
        ("Si la propuesta económica subida al portal, presenta diferencias entre el valor del anexo económico solicitado y el valor indicado en la línea de la plataforma", ""),
        (" www.mercadopublico.cl", "bold"),
        (", prevalecerá la oferta del anexo económico solicitado en bases. Sin embargo, el Hospital San José de Melipilla, podrá solicitar aclaraciones de las ofertas realizadas a través del portal.", "")
    ])

    # Antecedentes legales para poder ser contratado
    doc.add_heading("Antecedentes legales para poder ser contratado.", level=3)
    tabla_legal = doc.add_table(rows=7, cols=3, style='Table Grid')

    start_natural_row = 0
    end_natural_row = 3
    tabla_legal.cell(start_natural_row, 0).text = "Si el oferente\nes Persona\nNatural"
    tabla_legal.cell(start_natural_row, 2).text = "Acreditar en\nel Registro de\nProveedores"
    tabla_legal.cell(start_natural_row, 1).add_paragraph().add_run("Inscripción (en estado hábil) en el Registro electrónico oficial de contratistas de la Administración, en adelante “").bold = True
    tabla_legal.cell(start_natural_row, 1).paragraphs[0].add_run("Registro de Proveedores").bold = True
    tabla_legal.cell(start_natural_row, 1).paragraphs[0].add_run("”.").bold = True
    tabla_legal.cell(start_natural_row + 1, 1).add_paragraph().add_run("Anexo N°3. Declaración Jurada de Cumplimiento de Obligaciones Laborales y Previsionales.").bold = True
    tabla_legal.cell(start_natural_row + 2, 1).text = "Todos los Anexos deben ser firmados por la persona natural respectiva."
    tabla_legal.cell(start_natural_row + 3, 1).text = "Fotocopia de su cédula de identidad."
    tabla_legal.cell(start_natural_row, 0).merge(tabla_legal.cell(end_natural_row, 0))
    tabla_legal.cell(start_natural_row, 2).merge(tabla_legal.cell(end_natural_row, 2))

    start_nonatural_row = end_natural_row + 1
    end_nonatural_row = start_nonatural_row + 2
    tabla_legal.cell(start_nonatural_row, 0).text = "Si el oferente\nno es\nPersona\nNatural"
    tabla_legal.cell(start_nonatural_row, 2).text = "Acreditar en\nel Registro de\nProveedores"
    tabla_legal.cell(start_nonatural_row, 1).add_paragraph().add_run("Inscripción (en estado hábil) en el Registro de Proveedores.").bold = True
    tabla_legal.cell(start_nonatural_row + 1, 1).text = "Certificado de Vigencia del poder del representante legal, con una antigüedad no superior a 60 días corridos, contados desde la fecha de notificación de la adjudicación, otorgado por el Conservador de Bienes Raíces correspondiente o, en los casos que resulte procedente, cualquier otro antecedente que acredite la vigencia del poder del representante legal del oferente, a la época de presentación de la oferta."
    tabla_legal.cell(start_nonatural_row + 2, 1).text = "Certificado de Vigencia de la Sociedad con una antigüedad no superior a 60 días corridos, contados desde la fecha de notificación de la adjudicación"
    tabla_legal.cell(start_nonatural_row, 0).merge(tabla_legal.cell(end_nonatural_row, 0))
    tabla_legal.cell(start_nonatural_row, 2).merge(tabla_legal.cell(end_nonatural_row, 2))
    centrar_verticalmente_tabla(tabla_legal)

    # Observaciones (continuación)
    doc.add_heading("Observaciones", level=4)
    for texto in [
        "Los antecedentes legales para poder ser contratado, sólo se requerirán respecto del adjudicatario y deberán estar disponibles en el Registro de Proveedores.",
        "Lo señalado en el párrafo precedente no resultará aplicable a la garantía de fiel cumplimiento de contrato, la cual podrá ser entregada físicamente en los términos que indican las presentes bases en aquellos casos que aplique su entrega.",
        "En los casos en que se otorgue de manera electrónica, deberá ajustarse a la ley N° 19.799 sobre documentos electrónicos, firma electrónica y servicios de certificación de dicha firma, y remitirse en la forma señalada en la cláusula 8.2 de estas bases."
    ]:
        agregar_parrafo_con_texto(doc, texto)

    observ_parafo_2 = doc.add_paragraph()
    observ_parafo_2.add_run("Si el respectivo proveedor no entrega la totalidad de los antecedentes requeridos para ser contratado, dentro del plazo fatal de 10 días hábiles contados desde la notificación de la resolución de adjudicación o no suscribe el contrato en los plazos establecidos en estas bases, la entidad licitante podrá readjudicar de conformidad a lo establecido en la")
    observ_parafo_2.add_run(" cláusula 9 letra i")
    observ_parafo_2.add_run(" de las presentes bases. Además, tales incumplimientos darán origen al cobro de la garantía de seriedad de la oferta, si la hubiere.")

    # Inscripción en el Registro de Proveedores
    doc.add_heading("Inscripción en el registro de proveedores", level=4)
    for texto in [
        "En caso de que el proveedor que resulte adjudicado no se encuentre inscrito en el Registro Electrónico Oficial de Contratistas de la Administración (Registro de Proveedores), deberá inscribirse dentro del plazo de 15 días hábiles, contados desde la notificación de la resolución de adjudicación.",
        "Tratándose de los adjudicatarios de una Unión Temporal de Proveedores, cada integrante de ésta deberá inscribirse en el Registro de Proveedores, dentro del plazo de 15 días hábiles, contados desde la notificación de la resolución de adjudicación."
    ]:
        agregar_parrafo_con_texto(doc, texto)

    # Naturaleza y monto de las garantías
    doc.add_heading("Naturaleza y monto de las garantías", level=2)
    doc.add_heading("Garantía de Seriedad de la Oferta", level = 3)
    agregar_parrafo_con_texto(doc, "{{garantia_seriedad_oferta_p1}}")
    agregar_parrafo_con_texto(doc, "{{garantia_seriedad_oferta_p2}}")
    agregar_parrafo_con_texto(doc, "{{garantia_seriedad_oferta_p3}}")

    runs = [("{{run_1_garantia}}", ""),
            ("“{{run_2_garantia}} {{ nombre_adquisicion }}", "bold"),
            ("{{run_3_garantia}}", "")
            ]
    parrafo_loop = doc.add_paragraph()
    for text, style in runs:
        # It's usually better practice to add the paragraph first, then the run
        run = parrafo_loop.add_run(text)
        if style == "bold":
            run.bold = True


    parrafos_garantia = ["La(s) garantía(s) deberá(n) tener como vigencia mínima 120 días corridos desde el cierre y apertura de la oferta. ",
                         "Toda oferta que no acompañe la garantía de seriedad, en la forma y términos expresados, será rechazada por el Hospital San José de Melipilla.",
                         "Será responsabilidad del oferente mantener vigente la(s) garantía(s), debiendo reemplazarla si por razones sobrevinientes a su presentación, deja de cubrir la vigencia mínima exigida en esta cláusula, como por ejemplo ampliación de fecha de cierre de la licitación o del proceso de evaluación.",
                         "Como beneficiario del instrumento debe figurar la razón social y RUT de la entidad licitante, indicadas en la presente licitación, numeral N°1.",
                         "Si el instrumento que se presenta expresa su monto en unidades de fomento (UF), se considerará para determinar su equivalente en pesos chilenos (CLP), el valor de la UF a la fecha en que se realice la apertura de la oferta, considerando las variaciones en el mercado monto que debe ser detallado en peso en el mismo documento de garantía. ",
                         "Esta(s) garantía(s) se otorgará(n) para caucionar la seriedad de la oferta, pudiendo ser ejecutada unilateralmente por vía administrativa por la entidad licitante, siempre que los incumplimientos sean imputables al proveedor, en los siguientes casos:"
                         ]

    for texto in parrafos_garantia:
        doc.add_paragraph(texto)

    elementos = ["Por no suscripción del contrato definitivo o se rechace la orden de compra por parte del proveedor adjudicado, si corresponde;",
                 "Por la no entrega de los antecedentes requeridos para la elaboración del contrato, de acuerdo con las presentes bases, si corresponde;",
                 "Por el desistimiento de la oferta dentro de su plazo de validez establecido en las presentes bases;",
                 "Por la presentación de una oferta no fidedigna, manifiestamente errónea o conducente a error, y que así se justifique mediante resolución fundada del órgano comprador.",
                 "Por la no inscripción en el Registro de Proveedores dentro de los plazos establecidos en las presentes bases;",
                 "Por la no presentación oportuna de la garantía de fiel cumplimiento del contrato, en el caso del proveedor adjudicado."]
    for texto in elementos:
        doc.add_paragraph(texto, style = "List Number")

    doc.add_heading("Forma y oportunidad de restitución de la garantía de la seriedad de la oferta", level = 3)
    agregar_parrafo_con_texto(doc, "{{forma_garantia_seriedad_oferta_p1}}")
    agregar_parrafo_con_texto(doc, "{{forma_garantia_seriedad_oferta_p2}}")
    agregar_parrafo_con_texto(doc, "{{forma_garantia_seriedad_oferta_p3}}")

    tipo = "contrato"
    if tipo == "contrato":
        doc.add_heading("{{ Septimo_DeLaGarantíaFielCumplimiento }}{{ espacio }}Garantía de Fiel Cumplimiento de Contrato.", level = 3)

        par_garantia_fiel = doc.add_paragraph()
        par_garantia_fiel.add_run("Para garantizar el fiel y cabal cumplimiento de las obligaciones que impone el contrato, el adjudicatario entrega una Garantía pagadera a la vista y de carácter irrevocable, por un monto  equivalente al 5% del valor que involucre este convenio {{texto_gar_1}}{{monto_contrato_garantia}}{{ texto_gar_2 }}")
        par_garantia_fiel.add_run("{{ nombre_adquisicion }}").bold = True
        par_garantia_fiel.add_run(" ID" + "{{ID_licitacion}}")
        par_garantia_fiel.add_run("y/o de las obligaciones laborales y sociales del adjudicatario”. La garantía se hará efectiva ante cualquier incumplimiento a las condiciones y exigencias expuestas en las bases.")

        doc.add_paragraph("Se deja constancia que el proveedor adjudicatario entrega certificado de fianza de institución FINFAST como garantía de fiel cumplimiento del contrato con los siguientes datos:")
        doc.add_paragraph("[[TABLE_PLACEHOLDER]]")
        doc.add_paragraph("[[TABLE_PLACEHOLDER]]")

        doc.add_paragraph("Será responsabilidad del proveedor adjudicado mantener vigente la garantía de fiel cumplimiento, al menos hasta 120 días corridos después de culminado el contrato. Mientras se encuentre vigente el contrato, las renovaciones de esta garantía serán de exclusiva responsabilidad del proveedor adjudicado.")
        doc.add_paragraph("En caso de cobro de esta garantía, derivado del incumplimiento de las obligaciones contractuales del adjudicatario indicadas en las bases de licitación, éste deberá reponer la garantía por igual monto y por el mismo plazo de vigencia que la que reemplaza en un plazo de 15 días hábiles, contados desde la notificación de cobro.")
        par2_garantia_fiel = doc.add_paragraph()
        par2_garantia_fiel.add_run("La restitución de esta garantía será realizada una vez que se haya cumplido su fecha de vencimiento, en los términos indicados en el presente contrato, y su retiro será obligación y responsabilidad exclusiva del contratado previa solicitud por correo electrónico a: ")
        par2_garantia_fiel.add_run("garantias.hsjm@hospitaldemelipilla.cl").bold = True
        par2_garantia_fiel.add_run(" , con copia a ")
        par2_garantia_fiel.add_run("Manuel.lara@hospitaldemelipilla.cl").bold = True
        par2_garantia_fiel.add_run(",  en el siguiente horario: de lunes a viernes de 09:00 horas hasta las 16:00 horas. Lo anterior previa confirmación por parte del Establecimiento. Para el retiro de la garantía deberá presentar un poder simple timbrado por la empresa, fotocopia de la cédula de identidad de la persona que retira y el Rut de la empresa, siempre que no existan observaciones pendientes.")
        doc.add_paragraph("Cabe señalar que toda clase de garantías o cauciones que se constituyan en el contexto de esta cláusula, se enmarcan de acuerdo a lo dispuesto por el artículo 11 de la Ley N°19.886, a partir de lo cual se asegurará el fiel y oportuno cumplimiento del contrato, el pago de las obligaciones laborales y sociales con los trabajadores de los contratantes, y permanecerán vigentes hasta 120 días corridos después de culminado el contrato. Asimismo, con cargo a estas mismas cauciones podrán hacerse efectivas las multas y demás sanciones que afecten a los contratistas adjudicados.")



    else :
        loop_fiel = ["Para garantizar el fiel y oportuno cumplimiento del contrato, el adjudicado debe presentar una o más garantías de la misma naturaleza, equivalentes en total al porcentaje del 5% del valor total del contrato adjudicado ",
                     #{{ monto_contrato_garantia }}.
                     "La(s) garantía(s) debe(n) ser entregada(s) en la dirección de la entidad licitante indicada: Oficina de Partes del Hospital San José de Melipilla, ubicado en calle O’Higgins Nº 551 comuna de Melipilla, Región Metropolitana, dentro de los 10 días hábiles contados desde la notificación de la adjudicación en horario de 8:00 a 14:00 horas.",
                     "Si la(s) garantía(s) fuera(n) en soporte electrónico (garantía emitida por las instituciones de Garantía recíproca (IGR), Internacionalmente conocidas como SGR), se deberá enviar al correo electrónico garantias.hsjm@hospitaldemelipilla.cl, si no se presenta esta garantía en tiempo y forma, el Hospital San José de Melipilla podrá hacer efectiva la garantía de seriedad de la oferta y dejar sin efecto administrativamente la adjudicación, sin perjuicio de otros derechos."
                     ]
        for texto in loop_fiel:
            doc.add_paragraph(texto)
        parrafo_loop_fiel = doc.add_paragraph()
        lista_loop_fiel = [("Se aceptará cualquier tipo de instrumento de garantía que asegure su cobro de manera rápida y efectiva, pagadera a la vista y con el carácter de irrevocable, y siempre que cumpla con los requisitos dispuestos por el artículo 68 del reglamento de la ley N°19.886. El instrumento deberá incluir la glosa: Para garantizar el fiel cumplimiento del contrato denominado: ", ""),
                           ("{{nombre_adquisicion}} ID _________________", "bold"),
                           ("y/o de las obligaciones laborales y sociales del adjudicatario”. En caso de que el instrumento no permita la inclusión de la glosa señalada, el oferente deberá dar cumplimiento a la incorporación de ésta en forma manuscrita en el mismo instrumento, o bien, mediante un documento anexo a la garantía. Como ejemplos de garantías se pueden mencionar los siguientes instrumentos: Boleta de Garantía, Certificado de Fianza a la Vista, Vale Vista o Póliza de Seguro, entre otros. ", "")
                           ]
        for texto, style in lista_loop_fiel:
            run = parrafo_loop_fiel.add_run(texto)
            if style == "bold":
                run.bold = True




    doc.add_heading("Evaluación y adjudicación de las ofertas", level=2)

    comis_eval_p1 = doc.add_paragraph()
    comis_eval_p1.add_run("Comisión Evaluadora: ").bold = True
    comis_eval_p1.add_run("La Dirección del Hospital San José de Melipilla designa como integrantes de la Comisión de Evaluación de la propuesta a los siguientes funcionarios: el Subdirector(a) Administrativo, Subdirector(a) Médico de Atención Abierta, Subdirector(a) Médico de Atención Cerrada, Subdirector(a) de Gestión del Cuidado de Enfermería, Subdirector(a) de Gestión y Desarrollo de las Personas, Subdirector(a) de Matronería, Subdirector(a) de Análisis de Información para la Gestión, Subdirector(a) de Apoyo Clínico o sus subrogantes. Para los efectos del quórum para sesionar se requerirá un mínimo de tres miembros. Lo anterior en conformidad con lo dispuesto en el artículo 37 del Decreto Nº 250 que establece el Reglamento de la Ley Nº 19.886. Los miembros de la Comisión Evaluadora no podrán:")

    for texto, estilo in [
        ("Tener contactos con los oferentes, salvo en cuanto proceda alguno de mecanismos regulados por los artículos 27, 39 y 40 del reglamento de la ley N° 19.886.", 'List Bullet'),
        ("Aceptar solicitudes de reunión, de parte de terceros, sobre asuntos vinculados directa o indirectamente con esta licitación, mientras integren la Comisión Evaluadora.", 'List Bullet'),
        ("Aceptar ningún donativo de parte de terceros. Entiéndase como terceros, entre otros, a las empresas que prestan servicios de asesoría, o bien, sociedades consultoras, asociaciones, gremios o corporaciones. La misma Comisión estudiará los antecedentes de la Propuesta y elaborará un informe fundado para el Director de este Establecimiento, quien podrá declarar, mediante resolución fundada, admisible aquellas ofertas que cumplan con los requisitos establecidos en las bases de licitación, como también podrá declarar, mediante resolución fundada, inadmisible aquellas ofertas que no cumplan los requisitos establecidos en las bases. En caso de no presentarse oferentes o cuando las ofertas no resulten convenientes para los intereses del Establecimiento, podrá declarar desierta la licitación, fundándose en razones objetivas y no discriminatorias. Esta Comisión Evaluadora podrá invitar a profesionales técnicos para colaborar en el proceso de adjudicación.", "List Bullet")
    ]:
        agregar_parrafo_con_texto(doc, texto, estilo)

    agregar_parrafo_con_texto(doc, "La misma Comisión estudiará los antecedentes de la Propuesta y elaborará un informe fundado para el Director de este Establecimiento, quien podrá declarar, mediante resolución fundada, admisible aquellas ofertas que cumplan con los requisitos establecidos en las bases de licitación, como también podrá declarar, mediante resolución fundada, inadmisible aquellas ofertas que no cumplan los requisitos establecidos en las bases. En caso de no presentarse oferentes o cuando las ofertas no resulten convenientes para los intereses del Establecimiento, podrá declarar desierta la licitación, fundándose en razones objetivas y no discriminatorias.Esta Comisión Evaluadora podrá invitar a profesionales técnicos para colaborar en el proceso de adjudicación")

    consider_general_p1 = doc.add_paragraph()
    consider_general_p1.add_run("Consideraciones Generales: ").bold = True
    consider_general_p1.add_run("Se exigirá el cumplimiento de los requerimientos establecidos en la cláusula 6, “Instrucciones para Presentación de Ofertas”, de las presentes Bases de Licitación. Aquellas ofertas que no fueran presentadas a través del portal, en los términos solicitados, se declararán como propuestas inadmisibles, por tanto, no serán consideradas en la evaluación. Lo anterior, sin perjuicio de que concurra y se acredite algunas de las causales de excepción establecidas en el artículo 62 del Reglamento de la Ley de Compras.")

    for texto in [
        "La entidad licitante declarará inadmisible cualquiera de las ofertas presentadas que no cumplan los requisitos o condiciones establecidos en las presentes bases, sin perjuicio de la facultad de la entidad licitante de solicitar a los oferentes que salven errores u omisiones formales, de acuerdo con lo establecido en el artículo 40 del Reglamento de la Ley N°19.886 y en las presentes bases.",
        "Los documentos solicitados por la entidad licitante deben estar vigentes a la fecha de cierre de la presentación de las ofertas indicado en la cláusula 3 de las presentes bases y ser presentados como copias simples, legibles y firmadas por el representante legal de la empresa o persona natural. Sin perjuicio de ello, la entidad licitante podrá verificar la veracidad de la información entregada por el proveedor. En el caso en que el proveedor esté inscrito y habilitado por el Registro de Proveedores, serán suficientes los antecedentes que se encuentren en dicho Registro, en la medida que se haya dado cumplimiento a las normas de actualización de documentos que establece el Registro de Proveedores."
    ]:
        agregar_parrafo_con_texto(doc, texto)

    subsan_err_p1 = doc.add_paragraph()
    subsan_err_p1.add_run("Subsanación de errores u omisiones formales: ").bold = True
    subsan_err_p1.add_run("Una vez realizada la apertura electrónica de las ofertas, la entidad licitante podrá solicitar a los oferentes que salven errores u omisiones formales, siempre y cuando las rectificaciones de dichos vicios u omisiones no les confieran a esos oferentes una situación de privilegio respecto de los demás competidores, esto es, en tanto no se afecten los principios de estricta sujeción a las bases y de igualdad de los oferentes, y se informe de dicha solicitud al resto de los oferentes, a través del Sistema de Información")
    subsan_err_p1.add_run("www.mercadopublico.cl.").bold = True

    agregar_parrafo_con_texto(doc, "El plazo que tendrán los oferentes, en este caso para dar cumplimiento a lo solicitado por el mandante, no será inferior a las 24 horas, contadas desde la fecha de publicación de la solicitud por parte del Hospital, la que se informará a través del Sistema de información www.mercadopublico.cl. La responsabilidad de revisar oportunamente dicho sistema durante el período de evaluación recae exclusivamente en los respectivos oferentes.")

    inadmisibilidad_p1 = doc.add_paragraph()
    inadmisibilidad_p1.add_run("Inadmisibilidad de las ofertas y declaración de desierta de la licitación: ").bold = True
    inadmisibilidad_p1.add_run("La entidad licitante declarará inadmisible las ofertas presentadas que no cumplan los requisitos mínimos establecidos en los Anexos N°5, N°6, N°7, N°8 y N°9 y/o las condiciones establecidas en las presentes bases de licitación, sin perjuicio de la facultad para solicitar a los oferentes que salven errores u omisiones formales de acuerdo con lo establecido en las presentes bases.")

    agregar_parrafo_con_texto(doc, "La entidad licitante podrá, además, declarar desierta la licitación cuando no se presenten ofertas o cuando éstas no resulten convenientes a sus intereses. Dichas declaraciones deberán materializarse a través de la dictación de una resolución fundada y no darán derecho a indemnización alguna a los oferentes.")

    criterios_eval_p1 = doc.add_paragraph()
    criterios_eval_p1.add_run("Criterios de evaluación y procedimientos de las ofertas: ").bold = True
    criterios_eval_p1.add_run("La evaluación de las ofertas se realizará en una etapa, utilizando criterios técnicos, económicos y administrativos.")

    criterios_eval_p2 = doc.add_paragraph()
    criterios_eval_p2.add_run("La evaluación de las ofertas presentadas para el ")
    criterios_eval_p2.add_run("{{ nombre_adquisicion }}").bold = True
    criterios_eval_p2.add_run("se regirá por las siguientes ponderaciones y criterios a evaluar:")

    # Tabla de criterios de evaluación
    tabla_criterios_datos = [
        ["CRITERIOS", "", "PONDERACIÓN", "EVALUADO\nSEGÚN ANEXO"],
        ["ECONÓMICO", "OFERTA ECONÓMICA", "60%", "ANEXO N°5"],
        ["TÉCNICOS", "{{anexo_6}}".upper(), "{{20%}}", "ANEXO N°6"],
        ["", "{{anexo_7}}".upper(), "{{10%}}", "ANEXO N°8"],
        ["", "{{anexo_8}}".upper(), "{{10%}}", "ANEXO N°9"]
    ]

    tabla_criterios = crear_tabla(doc, tabla_criterios_datos)
    tabla_criterios.cell(0, 0).merge(tabla_criterios.cell(0, 1))
    tabla_criterios.cell(0, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    tabla_criterios.cell(0, 0).paragraphs[0].runs[0].bold = True
    tabla_criterios.cell(0, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    tabla_criterios.cell(0, 2).paragraphs[0].runs[0].bold = True
    tabla_criterios.cell(0, 3).paragraphs[0].text = ""
    tabla_criterios.cell(0, 3).paragraphs[0].add_run("EVALUADO").bold = True
    tabla_criterios.cell(0, 3).paragraphs[0].add_run("\n")
    tabla_criterios.cell(0, 3).paragraphs[0].add_run("SEGÚN ANEXO").bold = True
    tabla_criterios.cell(0, 3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    tabla_criterios.cell(2, 0).merge(tabla_criterios.cell(4, 0))
    tabla_criterios.cell(2, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    tabla_criterios.cell(2, 0).paragraphs[0].runs[0].bold = True
    for i in range(1, 5):
        tabla_criterios.cell(i, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        tabla_criterios.cell(i, 3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    tabla_criterios.cell(1, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    tabla_criterios.cell(1, 0).paragraphs[0].runs[0].bold = True

    # Cálculo del puntaje
    calculo_puntaje = doc.add_paragraph()
    calculo_puntaje.add_run("Cálculo del puntaje total de evaluación: ").bold = True

    for texto in [
        "El Puntaje de la Evaluación Final estará dado por el siguiente polinomio:",
        "Puntaje Evaluación Técnica + Puntaje Evaluación Económica",
        "Donde el Puntaje Evaluación Técnica = Evaluación Técnica + Plazo de Entrega + Servicio Post-Venta.",
        "Donde Puntaje Evaluación Económica = Precio."
    ]:
        agregar_parrafo_con_texto(doc, texto)

    # Criterios de evaluación
    doc.add_heading("Criterios de evaluación", level=3)
    doc.add_heading("Criterios Económicos", level=4)
    crit_econ = doc.add_paragraph()
    crit_econ.add_run("OFERTA ECONÓMICA 60%").bold = True
    agregar_parrafo_con_texto(doc, "Valor ítem ofertado. Para calcular el puntaje correspondiente al precio se utilizará la siguiente fórmula: A este puntaje se le aplicará la ponderación del 60 %. El oferente deberá declarar en Anexo N°5, los valores ofertados considerando todos los gastos involucrados e impuestos que apliquen.")

    doc.add_heading("Criterios Técnicos", level=4)
    crit_tecn_para = doc.add_paragraph(style="List Bullet")
    crit_tecn_run = crit_tecn_para.add_run("EVALUACIÓN TÉCNICA 20%: ")
    crit_tecn_run.bold = True
    crit_tecn_para.add_run("Se evaluará según información presentada para el Anexo N°7 que deberá ser adjuntada en su oferta en el Portal de Mercado Público, junto con la pauta de evaluación del Anexo N°6. Se evaluará por producto ofertado, donde el puntaje total será el promedio de la evaluación de todos los insumos ofertados.")

    plazo_entrega_para = doc.add_paragraph(style="List Bullet")
    plazo_entrega_run = plazo_entrega_para.add_run("PLAZO DE ENTREGA 10%: ")
    plazo_entrega_run.bold = True
    plazo_entrega_para.add_run("Se evaluará según información presentada en el Anexo N°8 de la presente base de licitación.")

    serv_post_venta_para = doc.add_paragraph(style="List Bullet")
    serv_post_venta_run = serv_post_venta_para.add_run("Servicio Post-Venta 10%: ")
    serv_post_venta_run.bold = True
    serv_post_venta_para.add_run("Se evaluará según información presentada en el Anexo N°9 de la presente base de licitación.")

    # Adjudicación y otras secciones
    doc.add_heading("Adjudicación", level=3)
    agregar_parrafo_con_texto(doc, "Se adjudicará al oferente que obtenga el mayor puntaje, en los términos descritos en las presentes bases. La presente licitación se adjudicará a través de una resolución dictada por la autoridad competente, la que será publicada en www.mercadopublico.cl, una vez que se encuentre totalmente tramitada.")

    doc.add_heading("Mecanismo de Resolución de empates.", level=3)
    agregar_parrafo_con_texto(doc, "En el evento de que, una vez culminado el proceso de evaluación de ofertas, hubiese dos o más proponentes que hayan obtenido el mismo puntaje en la evaluación final, quedando más de uno en condición de resultar adjudicado, se optará por aquella oferta que cuente con un mayor puntaje de acuerdo con la secuencia de los criterios que resulten aplicables, de acuerdo al siguiente orden: {{ resolucion_desempates}}. Finalmente, de mantenerse la igualdad, se adjudicará a aquel oferente que haya ingresado primero su propuesta en el portal Mercado Público considerándose la hora en que aquello se efectúe.")

    doc.add_heading("Resolución de consultas respecto de la Adjudicación.", level=3)
    resolucion_consultas_par = doc.add_paragraph()
    resolucion_consultas_par.add_run("Las consultas sobre la adjudicación deberán realizarse dentro del plazo fatal de 5 días hábiles contados desde la publicación de la resolución en el Sistema de Información ")
    resolucion_consultas_par.add_run("www.mercadopublico.cl").bold = True
    resolucion_consultas_par.add_run("a través del siguiente enlace: ")
    resolucion_consultas_par.add_run("http://ayuda.mercadopublico.cl ").bold = True

    doc.add_heading("Readjudicación", level=3)
    agregar_parrafo_con_texto(doc, "Si el adjudicatario se desistiere de firmar el contrato o de aceptar la orden de compra, o no cumpliese con las demás condiciones y requisitos establecidos en las presentes bases para la suscripción o aceptación de los referidos documentos, la entidad licitante podrá, junto con dejar sin efecto la adjudicación original, adjudicar la licitación al oferente que le seguía en puntaje, o a los que le sigan sucesivamente, dentro del plazo de 60 días corridos contados desde la publicación de la adjudicación original.")

    doc.add_section()
    doc.add_heading("Condiciones Contractuales, Vigencia de las Condiciones Comerciales, Operatoria de la Licitación y Otras Cláusulas:", level=1)
    doc.add_heading("{{ Documentos_Integrantes }}{{ espacio }}Documentos Integrantes", level=2)
    agregar_parrafo_con_texto(doc, "La relación contractual que se genere entre la entidad licitante y el adjudicatario se ceñirá a los siguientes documentos:")
    for texto in [
        "Bases de licitación y sus anexos.",
        "Aclaraciones, respuestas y modificaciones a las Bases, si las hubiere.",
        "Oferta.",
        "El presente contrato",
        "Orden de compra"
    ]:
        agregar_parrafo_con_texto(doc, texto, estilo="List Bullet")

    agregar_parrafo_con_texto(doc, "Todos los documentos antes mencionados forman un todo integrado y se complementan recíprocamente, especialmente respecto de las obligaciones que aparezcan en uno u otro de los documentos señalados. Se deja constancia que se considerará el principio de preeminencia de las Bases.")

    doc.add_heading("Validez de la Oferta", level=2)
    for texto in [
        "La oferta tendrá validez de ciento veinte días (120) días corridos, contados desde la fecha de apertura de la propuesta. La oferta cuyo periodo de validez sea menor que el requerido, será rechazada de inmediato.",
        "Si vencido el plazo señalado precedentemente, el Hospital San José de Melipilla no ha realizado la adjudicación, podrá solicitar a los Proponentes la prórroga de sus ofertas y garantías. Los proponentes podrán ratificar sus ofertas o desistir de ellas, formalizando su decisión mediante comunicación escrita dirigida al Hospital. Se devolverá la garantía a aquellos que no accedan a la prórroga."
    ]:
        agregar_parrafo_con_texto(doc, texto)

    doc.add_heading("Suscripción del Contrato", level=2)
    for texto in [
        "Para suscribir el contrato o aceptar la orden de compra contemplada en el artículo 63 del reglamento de la Ley de Compras, el adjudicado deberá estar inscrito en el Registro de Proveedores.",
        "Para formalizar las adquisiciones de bienes y servicios regidas por la ley Nº 19.886, se requerirá la suscripción de un contrato, la que en este caso se verá reflejada por la sola aceptación de la respectiva Orden de Compras.",
        "El respectivo contrato deberá suscribirse dentro de los {{ plazo_suscripcion }} siguientes a la notificación de la resolución de adjudicación totalmente tramitada. Asimismo, cuando corresponda, la orden de compra que formaliza la adquisición deberá ser aceptada por el adjudicatario dentro de ese mismo plazo.",
        "Si por cualquier causa que no sea imputable a la entidad licitante, el contrato no se suscribe dentro de dicho plazo, o no se acepta la orden de compra que formaliza la adquisición dentro de ese mismo término, se entenderá desistimiento de la oferta, pudiendo readjudicar la licitación al oferente que le seguía en puntaje, o a los que le sigan sucesivamente."
    ]:
        agregar_parrafo_con_texto(doc, texto)

    doc.add_heading("{{ Cuarto_ModificacionDelContrato }}{{espacio}}Modificaciones Del Contrato", level=2)
    agregar_parrafo_con_texto(doc, "Las partes de común acuerdo podrán modificar el contrato aumentando o disminuyendo los Bienes o servicios licitados, como también se podrán pactar nuevos bienes o servicios que no alteren la naturaleza del contrato. Estas modificaciones podrán ser hasta un 30% el presupuesto disponible estipulado en las presentes bases de licitación.")
    agregar_parrafo_con_texto(doc ,"En el caso de aumentar los bienes o servicios contratados, la garantía fiel cumplimiento de contrato también podrá readecuarse en proporción al monto de la modificación que se suscriba según aquellos casos que apliquen. En caso de aumentar o disminuir los bienes o servicios contratados, los valores a considerar, serán aquellos ofertados en el anexo oferta económica. ")
    agregar_parrafo_con_texto(doc ,"Con todo, las eventuales modificaciones que se pacten no producirán efecto alguno sino desde la total tramitación del acto administrativo que las apruebe.")

    doc.add_heading("{{ Quinto_GastoseImpuestos }}{{ espacio }}Gastos e Impuestos", level=2)
    agregar_parrafo_con_texto(doc, "Todos los gastos e impuestos que se generen o produzcan por causa o con ocasión de este Contrato, tales como los gastos notariales de celebración de contratos y/o cualesquiera otros que se originen en el cumplimiento de obligaciones que, según las Bases, ha contraído el oferente adjudicado, serán de cargo exclusivo de éste.")

    doc.add_heading("{{ Sexto_EfectosDerivadosDeIncumplimiento }}{{ espacio }}Efectos derivados de Incumplimiento del proveedor", level=2)
    agregar_parrafo_con_texto(doc, "En función de la gravedad de la infracción cometida por el adjudicatario, se le aplicarán las siguientes sanciones:")

    clasificacion_par1 = doc.add_paragraph()
    clasificacion_par1.add_run("Amonestación: ").bold = True
    clasificacion_par1.add_run("Corresponde a un registro escrito, que dejará de manifiesto cualquier falta menor cometida por el adjudicado. Se entenderá por falta menor aquella que no ponga en riesgo de forma alguna la prestación del servicio o la vida e integridad psíquica y física de los pacientes, que se vinculen a temas administrativos y técnicos y que no sea constitutiva de multa. La amonestación no estará afecta a sanción pecuniaria.")

    clasificacion_par2 = doc.add_paragraph()
    clasificacion_par2.add_run("Multa: ").bold = True
    clasificacion_par2.add_run("Corresponde a la sanción de cualquier falta, de gravedad leve, moderada o grave en que incurra el adjudicado, cada vez que éste no dé cumplimiento a cualquiera de las obligaciones contempladas en las presentes bases. Se expresará en Unidades Tributarias Mensuales (UTM).El monto de cada multa, dependerá de la gravedad de la infracción cometida, en este sentido las multas se clasifican en:")

    p_leve = doc.add_paragraph()
    p_leve.add_run("Multa Leve: ").bold = True
    p_leve.add_run("Sera considerada LEVE aquella situación originada por una falta de carácter menor, que no origina riesgos a las personas, ni daños a los bienes de la Institución o a su imagen. Su importe será de 3 Unidades Tributarias Mensuales (UTM). Las conductas que pueden estar afectas a multa leve son:")

    for texto in [
        "Entrega de productos con atraso de hasta dos (2) días hábiles, contados desde la fecha de entrega estipulada en el contrato, sin que estos se hubiesen re-pactado en su plazo de entrega.",
        "Conducta o trato irrespetuoso de parte del personal del Oferente adjudicado o su cadena de distribución.",
        "La acumulación de dos amonestaciones.",
        "Incumplimiento del contrato que no origine riesgos a las personas o daño a los bienes del establecimiento o a su imagen."
    ]:
        agregar_parrafo_con_texto(doc, texto, estilo='List Bullet')

    p_moderada = doc.add_paragraph()
    p_moderada.add_run("Multa Moderada: ").bold = True
    p_moderada.add_run("Sera considerada MODERADA, aquella situación originada por una falta que afecte o ponga en riesgo, directa o indirectamente a personas o a la Institución o que limite significativamente la atención y calidad del servicio, pero que es factible de ser corregida. Su importe será de 6 Unidades Tributarias Mensuales (UTM). Las conductas que puedan estar afectas a multa moderada son:")

    for texto in [
        "No aceptar la Orden de Compra dentro de los dos (4) días hábiles siguientes al envío de la orden a través del portal de Mercado Publico.",
        "Entrega de los productos con atraso de entre tres (3) y seis (6) días hábiles inclusive, contados desde la fecha de entrega estipulada en el contrato, sin que estos se hubiesen re-pactado en su plazo de entrega.",
        "Despacho de productos en lugares no autorizados por el Hospital.",
        "La acumulación de dos multas leves trimestres móviles.",
        "cualquier falta que afecte o ponga en riesgo, directa o indirectamente, a personas o a la institución, o que limite significativamente la atención y calidad del servicio, pero que es factible de ser corregida."
    ]:
        agregar_parrafo_con_texto(doc, texto, estilo='List Bullet')

    p_grave = doc.add_paragraph()
    p_grave.add_run("Multa Grave: ").bold = True
    p_grave.add_run("Sera considerada GRAVE, aquella situación originada por una falta que atente, directa o indirectamente con la atención y calidad del servicio. Su importe será de 10 Unidades Tributarias Mensuales (UTM). Las conductas que pueden estar afectas a multa grave son:")

    for texto in [
        "Incumplimiento de la totalidad de lo requerido en la orden de compra.",
        "Entrega de productos con atraso de más de  {{ atraso_para_multa_grave }}, contados desde la fecha de entrega estipulada en el contrato, sin que estos se hubiesen re-pactado en su plazo de entrega.",
        "Rechazo total de los productos por no ajustarse a las especificaciones técnicas.",
        "El incumplimiento en el recambio de los productos que presenten problemas de estabilidad, empaque, envases en mal estado, conservación inadecuada, vencida, defectuosa o dañada en un periodo máximo de cuarenta y ocho (48) horas.",
        "La acumulación de dos multas moderadas en 2 trimestres móviles",
        "Si el adjudicatario no inicia sus labores en la fecha acordada",
        "Si se acreditaren acciones y/u omisiones maliciosas y/o negligentes que comprometan la eficiencia y eficacia del servicio o la seguridad y/o bienes dispuestos por el Hospital para la correcta ejecución del convenio.",
        "No cumplir con lo señalado en los requerimientos técnicos.",
        "Vulnerar normas referidas al uso de información reservada o confidencial al que se tenga acceso en razón del servicio especializado contratado.",
        "No cumplir con los horarios establecidos en las bases",
        "No cumplir con los tiempos de respuestas establecidos en bases de licitación según sea el caso.",
        "Incumplimiento de las normas y procedimientos internos vigentes, tanto técnicas como administrativas impartidas por el Hospital, a través del administrador interno del contrato",
        "cualquier falta que atente, directa o indirectamente, contra la integridad física de los pacientes o funcionarios, que implique u obstruye la atención, calidad del servicio y/o según lo establecido en las bases."
    ]:
        agregar_parrafo_con_texto(doc, texto, estilo='List Bullet')

    for texto in [
        "Las referidas multas, en total, no podrán sobrepasar el 20% del valor total neto del contrato. Igualmente, no se le podrán cursar más de 6 multas totalmente tramitadas en un período de 6 meses consecutivos. En ambos casos, superado cada límite, se configurará una causal de término anticipado del contrato.",
        "Las multas deberán ser pagadas en el plazo máximo de 5 días hábiles contados desde la notificación de la resolución que aplica la multa.",
        "Cuando el cálculo del monto de la respectiva multa, convertido a pesos chilenos, resulte un número con decimales, éste se redondeará al número entero más cercano. La fecha de conversión de la UTM será la del día de emisión del respectivo acto administrativo que origina el cobro de la multa",
        "Las multas se aplicarán sin perjuicio del derecho de la entidad licitante de recurrir ante los Tribunales Ordinarios de Justicia ubicados en la ciudad de Melipilla, a fin de hacer efectiva la responsabilidad del contratante incumplidor.",
        "No procederá el cobro de las multas señaladas en este punto, si el incumplimiento se debe a un caso fortuito o fuerza mayor, de acuerdo con los artículos 45 y 1547 del Código Civil o una causa enteramente ajena a la voluntad de las partes, el cual será calificado como tal por la Entidad Licitante, en base al estudio de los antecedentes por los cuales el oferente adjudicado acredite el hecho que le impide cumplir."
    ]:
        agregar_parrafo_con_texto(doc, texto)

    doc.add_heading("{{ Octavo_CobroDeLaGarantiaFielCumplimiento }}{{ espacio }}Cobro de la Garantía de Fiel Cumplimiento de Contrato", level=2)
    agregar_parrafo_con_texto(doc, "Al Adjudicatario le podrá ser aplicada la medida de cobro de la Garantía por Fiel Cumplimiento del Contrato por la entidad licitante, en los siguientes casos:")
    for texto in [
        "No pago de multas dentro de los plazos establecidos en las presentes bases y/o el respectivo contrato.",
        "Incumplimientos de las exigencias técnicas de los bienes y servicios (en caso de que hayan sido requeridos) adjudicados establecidos en el Contrato.",
        "Cualquiera de las causales señaladas en el N°10.6.3 sobre “Término Anticipado del Contrato”, a excepción del numeral 3) y numeral 16), en todas estas causales señaladas, se procederá al cobro de la garantía de fiel cumplimiento del contrato, si se hubiere exigido dicha caución en las Bases."
    ]:
        agregar_parrafo_con_texto(doc, texto, estilo='List Bullet')

    doc.add_heading("{{ Noveno_TerminoAnticipadoDelContrato }}{{ espacio }}Término anticipado del contrato", level=2)
    agregar_parrafo_con_texto(doc,
                              "El hospital está facultado para declarar administrativamente mediante resolución fundada el término anticipado del contrato, en cualquier momento, sin derecho a indemnización alguna para el adjudicado, si concurre alguna de las causales que se señalan a continuación:")

    # Items principales con numeración
    num_id_termino = crear_numeracion(doc)

    # Primer ítem con viñetas internas
    termino_p1 = doc.add_paragraph(style=list_style)
    termino_p1.add_run(
        "Por incumplimiento grave de las obligaciones contraídas por el proveedor adjudicado, cuando sea imputable a éste. Se entenderá por incumplimiento grave la no ejecución o la ejecución parcial por parte del adjudicatario de las obligaciones contractuales, descritas en las presentes Bases, sin que exista alguna causal que le exima de responsabilidad, y cuando dicho incumplimiento le genere al hospital un perjuicio en el cumplimiento de sus funciones. Alguno de estos motivos puede ser:")
    aplicar_numeracion(termino_p1, num_id_termino)

    # Subitems para el primer ítem (viñetas)
    subitems_1 = [
        "Entrega injustificada fuera de los plazos convenidos.",
        "La imposibilidad fundada de entregar los productos en los plazos comprometidos, en más de 2 oportunidades.",
        "Por entregar productos no solicitados."
    ]

    for subitem in subitems_1:
        p_sub = doc.add_paragraph(subitem, style="List Bullet")
        # Aumentar sangría para mostrar que pertenece al ítem numerado
        p_sub.paragraph_format.left_indent = Pt(72)

    # Ítems 2 a 6 sin viñetas internas
    for i, texto in enumerate([
        "Si el adjudicado se encuentra en estado de notoria insolvencia o fuere declarado deudor en un procedimiento concursal de liquidación. En el caso de una UTP, aplica para cualquiera de sus integrantes. En este caso no procederá el término anticipado si se mejoran las cauciones entregadas o las existentes sean suficientes para garantizar el cumplimiento del contrato.",
        "Por exigirlo la necesidad del servicio, el interés público o la seguridad nacional.",
        "Registrar, a la mitad del período de ejecución contractual, con un máximo de seis meses, saldos insolutos de remuneraciones o cotizaciones de seguridad social con sus actuales trabajadores o con trabajadores contratados en los últimos 2 años.",
        "Si se disuelve la sociedad o empresa adjudicada, o en caso de fallecimiento del contratante, si se trata de una persona natural.",
        "Incumplimiento de uno o más de los compromisos asumidos por los adjudicatarios, en virtud del \"Pacto de integridad\" contenido en estas bases. Cabe señalar que en el caso que los antecedentes den cuenta de una posible afectación a la libre competencia, el organismo licitante pondrá dichos antecedentes en conocimiento de la Fiscalía Nacional Económica."
    ]):
        p = doc.add_paragraph(style=list_style)
        p.add_run(texto)
        aplicar_numeracion(p, num_id_termino)

    # Séptimo ítem con subitems
    termino_p7 = doc.add_paragraph(style=list_style)
    termino_p7.add_run("Sin perjuicio de lo señalado en el Pacto de integridad, si el adjudicatario, sus representantes, o el personal dependiente de aquél, no observaren el más alto estándar ético exigible, durante la ejecución de la licitación, y propiciaren prácticas corruptas, tales como:")
    aplicar_numeracion(termino_p7, num_id_termino)

    # Subitems para el séptimo ítem
    subitems_7 = [
        "Dar u ofrecer obsequios, regalías u ofertas especiales al personal de la entidad licitante, que pudiere implicar un conflicto de intereses, presente o futuro, entre el respectivo adjudicatario y la entidad licitante.",
        "Efectuar reuniones en dependencias del organismo comprador, con el objeto de solicitar beneficios económicos respecto de la presente licitación.",
        "Efectuar contactos con funcionarios de la entidad compradora, fuera de la plataforma www.mercadopublico.cl, con el objeto de participar u obtener información de la presente licitación."
    ]

    for subitem in subitems_7:
        p_sub = doc.add_paragraph(subitem, style="List Bullet")
        p_sub.paragraph_format.left_indent = Pt(72)

    # Items 8 y 9 sin viñetas internas
    for texto in [
        "No renovación oportuna de la Garantía de Fiel Cumplimiento, según lo establecido en la cláusula 8.2 de las bases de licitación cuando aplique.",
        "La comprobación de la falta de idoneidad, de fidelidad o de completitud de los antecedentes aportados por el proveedor adjudicado, para efecto de ser adjudicado o contratado.",
        "La comprobación de que el adjudicatario, al momento de presentar su oferta contaba con información o antecedentes relacionados con el proceso de diseño de las bases, encontrándose a consecuencia de ello en una posición de privilegio en relación al resto de los oferentes, ya sea que dicha información hubiese sido conocida por el proveedor en razón de un vínculo laboral o profesional entre éste y las entidades compradoras, o bien, como resultado de prácticas contrarias al ordenamiento jurídico."
    ]:
        p = doc.add_paragraph(style=list_style)
        p.add_run(texto)
        aplicar_numeracion(p, num_id_termino)

    # Item 11 con subitems
    termino_p11 = doc.add_paragraph(style=list_style)
    termino_p11.add_run("En caso de ser el adjudicatario de una Unión Temporal de Proveedores (UTP):")
    aplicar_numeracion(termino_p11, num_id_termino)

    # Subitems para el ítem 11
    subitems_11 = [
        "Concurra alguna de las causales de término respecto de cualquiera de sus integrantes.",
        "La UTP no ha cumplido con su obligación de informar a la entidad licitante sobre cambios de sus integrantes."
    ]

    for subitem in subitems_11:
        p_sub = doc.add_paragraph(subitem, style="List Bullet")
        p_sub.paragraph_format.left_indent = Pt(72)

    # Items restantes sin viñetas internas
    for texto in [
        "En caso de infracción de lo dispuesto en la cláusula sobre Cesión de contrato y Subcontratación",
        "En caso de que las multas cursadas, en total, sobrepasen el 20 % del valor total contratado con impuestos incluidos o se apliquen más de 6 multas totalmente tramitadas en un periodo de 6 meses consecutivos.",
        "Por el no pago de las multas aplicadas.",
        "Por la aplicación de dos multas graves en que incurra el adjudicatario en virtud del incumplimiento de las obligaciones reguladas en las bases y del presente contrato.",
        "Si el Hospital San José de Melipilla cesara su funcionamiento en lugar de origen por cambio de ubicación de sus dependencias.",
        "Por la comprobación de la inhabilidad del adjudicatario para contratar con la Administración del Estado en portal de mercado público, durante la ejecución del presente contrato. Solo en el caso que el proveedor desde la notificación de esta situación no regularice su registro en un plazo superior a 15 días hábiles.",
        "Por incumplimiento de obligaciones de confidencialidad establecidas en las respectivas Bases."
    ]:
        p = doc.add_paragraph(style=list_style)
        p.add_run(texto)
        aplicar_numeracion(p, num_id_termino)

    for texto in [
        "De concurrir cualquiera de las causales anteriormente señaladas como término anticipado del contrato, exceptuando las causales número 3 y número 16, se procederá al cobro de la garantía de fiel cumplimiento del contrato, siempre y cuando se hubiere exigido dicha caución en las Bases.",
        "El término anticipado por incumplimientos se aplicará siguiendo el procedimiento establecido en la cláusula “sobre aplicación de Medidas derivadas de incumplimientos.”"
    ]:
        agregar_parrafo_con_texto(doc, texto)

    doc.add_heading("{{ Decimo_ResciliacionMutuoAcuerdo }}{{ espacio }}Resciliación de Mutuo Acuerdo", level=2)
    agregar_parrafo_con_texto(doc, "Sin perjuicio de lo anterior, la entidad licitante y el respectivo adjudicatario podrán poner término al contrato en cualquier momento, de común acuerdo, sin constituir una medida por incumplimiento.")

    doc.add_heading("{{ DecimoPrimero_ProcedimientoIncumplimiento }}{{ espacio }}Procedimiento para Aplicación de Medidas derivadas de incumplimientos", level=2)
    for texto in [
        "Detectada una situación que amerite la aplicación de una multa u otra medida derivada de incumplimientos contemplada en las presentes bases, o que constituya una causal de término anticipado, con excepción de la resciliación, el referente técnico o administrador del contrato notificará de ello al oferente adjudicado, informándole sobre la medida a aplicar y sobre los hechos que la fundamentan.",
        "A contar de la notificación singularizada en el párrafo anterior, el proveedor adjudicado tendrá un plazo de 5 días hábiles para efectuar sus descargos por escrito, acompañando todos los antecedentes que lo fundamenten. Vencido el plazo indicado sin que se hayan presentados descargos, la Dirección del Hospital resolverá según la naturaleza de la infracción, notificando al proveedor la resolución del caso por parte del Hospital.",
        "Si el proveedor adjudicado ha presentado sus descargos dentro del plazo establecido para estos efectos, el Hospital tendrá un plazo de 30 días hábiles, contados desde la recepción de los descargos del proveedor, para rechazarlos o acogerlos, total o parcialmente. Al respecto, el rechazo total o parcial de los descargos del respectivo proveedor deberá formalizarse a través de la dictación de una resolución fundada del hospital, en la cual deberá detallarse el contenido y las características de la medida. La indicada resolución será notificada al proveedor adjudicado.",
        "Con todo, el adjudicatario solo será responsable por hechos imputables a su incumplimiento directo y no por indisponibilidades de servicio ocasionadas por fallas ajenas a su gestión y control, lo que deberá, en todo caso, acreditarse debidamente. Sin perjuicio de lo anterior, el adjudicatario deberá adoptar medidas que ofrezcan continuidad operativa a los servicios materia de la respectiva licitación.",
        "Una vez finalizados los trámites administrativos señalados precedentemente y para el evento de que esta conlleve la aplicación de una multa o sanción, el Hospital San José de Melipilla podrá realizar el cobro de la multa o sanción que será debidamente notificado junto con el acto administrativo que lo autoriza. El monto de las multas podría ser rebajado del pago, que el Hospital deba efectuar al proveedor, en el estado de pago más próximo a la notificación del acto administrativo, pudiéndose aplicar tanto en la emisión de la orden de compra, como también en la aplicación del descuento en el pago de facturas. De no ser suficiente este monto o en caso de no existir pagos pendientes, el proveedor deberá pagar directamente al Hospital San José de Melipilla, el monto indicado en el acto administrativo previamente notificado, este pago no podrá ser superior a los 5 días hábiles desde su notificación. Si el proveedor no paga dentro de dicho plazo, se hará efectivo el cobro de la garantía de fiel cumplimiento del contrato, debiendo reponer una nueva boleta de garantía por un monto igual al original, en un plazo no superior a 5 días hábiles en caso que aplique la solicitud de dicha caución.",
        "En el caso de no reponer la boleta de garantía, el hospital podrá proceder a tramitar el término anticipado del contrato en aquellos casos que aplique con la solicitud de dicha caución.",
        "El valor de la UTM a considerar será el equivalente a su valor en pesos del mes en el cual se aplicó la multa."
    ]:
        agregar_parrafo_con_texto(doc, texto)

    doc.add_heading("{{ DecimoSegundo_EmisionOC }}{{ espacio }}Emisión de la Orden de Compra", level=2)
    for texto in [
        "Las órdenes de compra se emitirán previa solicitud del administrador del contrato, quien, en función de la necesidad y demanda del servicio, realizara los pedidos correspondientes.",
        "La orden de compra sólo se emitirá en los casos que el proveedor este en estado hábil para ser contratado por el Estado de Chile y sólo se emitirá el documento a nombre del proveedor adjudicado por el Hospital.",
        "Al inicio del convenio, por registros en la plataforma y tramites del “gestor de contratos” se emitirá una orden de compras por un monto mínimo, la que solo debe ser aceptada por el proveedor, sin tramitar dicho servicio. Todo cambio respecto a este punto, será informado con la respectiva anticipación."
    ]:
        agregar_parrafo_con_texto(doc, texto)

    doc.add_heading("{{ DecimoTercero_DelPago }}{{espacio}}Del Pago", level=2)
    for texto in [
        "El pago se efectuará una vez que el “Hospital” haya recibido oportunamente y a su entera satisfacción dichos bienes o servicios y desde la recepción conforme de la factura u otro instrumento de cobro.",
        "El pago será efectuado dentro de los 30 días corridos siguientes, contados desde la recepción de la factura respectiva, salvo las excepciones indicadas en el artículo 79 bis del Reglamento de la Ley N°19.886.",
        "El proveedor solo podrá facturar los bienes o servicios efectivamente entregados y recibidos conforme por este organismo comprador, una vez que el administrador del contrato por parte del organismo comprador autorice la facturación en virtud de la recepción conforme de los bienes o servicios. “El Hospital” rechazará todas las facturas que hayan sido emitidas sin contar con la recepción conforme de los bienes o servicios y la autorización expresa de facturar por parte de éste.",
        "Para efectos del pago, el proveedor adjudicado deberá indicar en la factura el número de orden de compra, además, no podrá superar el monto de la orden de compra, de lo contrario, se cancelará la factura por “forma”."
    ]:
        agregar_parrafo_con_texto(doc, texto)

    del_pago_correo = doc.add_paragraph()
    del_pago_correo.add_run("La factura electrónica deberá ser enviada al correo: ")
    del_pago_correo.add_run("facturas.hjsm@redsalud.gov.cl").bold = True
    del_pago_correo.add_run(" con copia al correo ")
    del_pago_correo.add_run("dipresrecepcion@custodium.com").bold = True
    del_pago_correo.add_run("(En formato PDF y XML)")

    for texto in [
        "El valor del convenio se reajustará anualmente de acuerdo con la variación que haya experimentado el Índice de Precios al Consumidor IPC, obtenido del promedio de la sumatoria de los IPC de los doce meses inmediatamente anteriores al mes en que se efectúa su cálculo. Este reajuste es de exclusiva responsabilidad de la empresa adjudicada; si por alguna razón no lo aplicare, no se permitirá su cobro en forma retroactiva. Su precio se pagará conforme a lo establecido.",
        "En ningún caso procederán cobros adicionales por bienes o servicios no convenidos previamente, ni por tiempos en que el proveedor no preste los servicios.",
        "Cabe señalar que, cuando el resultado del monto a facturar resulte un número con decimales, éste se redondeará al número entero siguiente en caso de que la primera cifra decimal sea igual o superior a 5. En caso contrario el monto deberá ser redondeado al número entero anterior."
    ]:
        agregar_parrafo_con_texto(doc, texto)

    doc.add_heading("{{ DecimoCuarto_VigenciaContrato }}{{espacio}}Vigencia del Contrato", level=2)
    agregar_parrafo_con_texto(doc, "El contrato tendrá una duración {{ plazo_meses }} meses contados desde la total tramitación del acto administrativo que aprueba la adjudicación o hasta que se cumpla con el monto estipulado en las presentes bases, lo que suceda primero y sin perjuicio, que por razones de buen servicio las prestaciones materia de la licitación podrían iniciarse desde el momento de la suscripción del mismo, sin que proceda pago alguno en el tiempo intermedio.")

    doc.add_heading("{{ DecimoQuinto_AdministradorContrato}}{{espacio}}Administrador del Contrato y/o Referente Técnico.", level=2)
    agregar_parrafo_con_texto(doc, "Con el objeto de supervisar y verificar el cumplimiento materia de la presente licitación, El Hospital designará a {{ opciones_referente_tecnico_adm }}, para coordinar y fiscalizar la efectiva ejecución del contrato en términos administrativos.")

    administrado_contrato = doc.add_paragraph()
    administrado_contrato.add_run("El adjudicatario ").bold = True
    administrado_contrato.add_run("{{ coordinador }}{{ nombre_coordinador }}")
    # deberá nombrar un coordinador del contrato, cuya identidad deberá ser informada al Hospital.
    # El adjudicatario nombra coordinador del contrato a doña MARIA GABRIELA CARDENAS en el desempeño de su cometido, el coordinador del contrato deberá, a lo menos:
    agregar_parrafo_con_texto(doc, "En el desempeño de su cometido, el coordinador del contrato deberá, a lo menos:")

    for texto in [
        "Informar oportunamente al órgano comprador de todo hecho relevante que pueda afectar el cumplimiento del contrato.",
        "Representar al proveedor en la discusión de las materias relacionadas con la ejecución del contrato.",
        "Coordinar las acciones que sean pertinentes para la operación y cumplimiento de este contrato."
    ]:
        p = agregar_parrafo_con_texto(doc, texto)
        aplicar_numeracion(p, administrado_contrato_id_lista)

    agregar_parrafo_con_texto(doc, "La designación del coordinador y todo cambio posterior deberá ser informado por el adjudicatario al responsable de administrar el contrato y/o referente técnico por parte del órgano comprador, a más tardar dentro de las 24 horas siguientes de efectuada la designación o el cambio, por medio del correo electrónico institucional del funcionario.")

    doc.add_heading("{{ DecimoSexto_PactoDeIntegrida }}{{ espacio }}Pacto de Integridad", level=2)
    agregar_parrafo_con_texto(doc, "El oferente declara que, por el sólo hecho de participar en la presente licitación, acepta expresamente el presente pacto de integridad, obligándose a cumplir con todas y cada una de las estipulaciones contenidas en el mismo, sin perjuicio de las que se señalen en el resto de las bases de licitación y demás documentos integrantes. Especialmente, el oferente acepta el suministrar toda la información y documentación que sea considerada necesaria y exigida de acuerdo con las presentes bases de licitación, asumiendo expresamente los siguientes compromisos:")

    pacto_items = [
        "El oferente se compromete a respetar los derechos fundamentales de sus trabajadores, entendiéndose por éstos los consagrados en la Constitución Política de la República en su artículo 19, números 1º, 4º, 5º, 6º, 12º, y 16º, en conformidad al artículo 485 del Código del Trabajo. Asimismo, el oferente se compromete a respetar los derechos humanos, lo que significa que debe evitar dar lugar o contribuir a efectos adversos en los derechos humanos mediante sus actividades, bienes o servicios, y subsanar esos efectos cuando se produzcan, de acuerdo con los Principios Rectores de Derechos Humanos y Empresas de Naciones Unidas.",
        "El oferente se obliga a no ofrecer ni conceder, ni intentar ofrecer o conceder, sobornos, regalos, premios, dádivas o pagos, cualquiera fuese su tipo, naturaleza y/o monto, a ningún funcionario público en relación con su oferta, con el proceso de licitación pública, ni con la ejecución de el o los contratos que eventualmente se deriven de la misma, ni tampoco a ofrecerlas o concederlas a terceras personas que pudiesen influir directa o indirectamente en el proceso licitatorio, en su toma de decisiones o en la posterior adjudicación y ejecución del o los contratos que de ello se deriven.",
        "El oferente se obliga a no intentar ni efectuar acuerdos o realizar negociaciones, actos o conductas que tengan por objeto influir o afectar de cualquier forma la libre competencia, cualquiera fuese la conducta o acto específico, y especialmente, aquellos acuerdos, negociaciones, actos o conductas de tipo o naturaleza colusiva, en cualquiera de sus tipos o formas.",
        "El oferente se obliga a revisar y verificar toda la información y documentación, que deba presentar para efectos del presente proceso licitatorio, tomando todas las medidas que sean necesarias para asegurar su veracidad, integridad, legalidad, consistencia, precisión y vigencia.",
        "El oferente se obliga a ajustar su actuar y cumplir con los principios de legalidad, probidad y transparencia en el presente proceso licitatorio.",
        "El oferente manifiesta, garantiza y acepta que conoce y respetará las reglas y condiciones establecidas en las bases de licitación, sus documentos integrantes y él o los contratos que de ellos se derivase.",
        "El oferente reconoce y declara que la oferta presentada en el proceso licitatorio es una propuesta seria, con información fidedigna y en términos técnicos y económicos ajustados a la realidad, que aseguren la posibilidad de cumplir con la misma en las condiciones y oportunidad ofertadas.",
        "El oferente se obliga a tomar todas las medidas que fuesen necesarias para que las obligaciones anteriormente señaladas sean asumidas y cabalmente cumplidas por sus empleados, dependientes, asesores y/o agentes y, en general, todas las personas con que éste o éstos se relacionen directa o indirectamente en virtud o como efecto de la presente licitación, incluidos sus subcontratistas, haciéndose plenamente responsable de las consecuencias de su infracción, sin perjuicio de las responsabilidades individuales que también procediesen y/o fuesen determinadas por los organismos correspondientes."
    ]
    for texto in pacto_items:
        p = agregar_parrafo_con_texto(doc, texto)
        aplicar_numeracion(p, pacto_integridad_id)

    doc.add_heading("{{ DecimoSeptimo_ComportamientoEticoAdjudic }}{{ espacio }}Comportamiento ético del Adjudicatario.", level=2)
    agregar_parrafo_con_texto(doc, "El adjudicatario que preste los servicios deberá observar, durante toda la época de ejecución del contrato, el más alto estándar ético exigible a los funcionarios públicos. Tales estándares de probidad deben entenderse equiparados a aquellos exigidos a los funcionarios de la Administración Pública, en conformidad con el Título III de la ley N°18.575, Orgánica Constitucional de Bases Generales de la Administración del Estado.")

    doc.add_heading("{{ DecimoOctavo_Auditorias }}{{espacio}}Auditorías", level=2)
    agregar_parrafo_con_texto(doc, "El adjudicatario podrá ser sometido a auditorías externas, contratadas por la entidad licitante a empresas auditoras independientes, con la finalidad de velar por el cumplimiento de las obligaciones contractuales y de las medidas de seguridad comprometidas por el adjudicatario en su oferta. Si el resultado de estas auditorías evidencia incumplimientos contractuales por parte del adjudicatario, el proveedor quedará sujeto a las medidas que corresponda aplicar la entidad licitante, según las presentes bases.")

    doc.add_heading("{{ DecimoNoveno_Confidencialidad }}{{ espacio }}Confidencialidad", level=2)
    for texto in [
        "El adjudicatario no podrá utilizar para ninguna finalidad ajena a la ejecución del contrato, la documentación, los antecedentes y, en general, cualquier información, que haya conocido o a la que haya accedido, en virtud de cualquier actividad relacionada con el contrato.",
        "El adjudicatario, así como su personal dependiente que se haya vinculado a la ejecución del contrato, en cualquiera de sus etapas, deben guardar confidencialidad sobre los antecedentes relacionados con el proceso licitatorio y el respectivo contrato.",
        "El adjudicatario debe adoptar medidas para el resguardo de la confidencialidad de la información, reservándose el órgano comprador el derecho de ejercer las acciones legales que correspondan, de acuerdo con las normas legales vigentes, en caso de divulgación no autorizada, por cualquier medio, de la totalidad o parte de la información referida.",
        "La divulgación, por cualquier medio, de la totalidad o parte de la información referida en los párrafos anteriores, por parte del proveedor, durante la vigencia del contrato o dentro de los 5 años siguientes después de finalizado éste, podrá dar pie a que la Entidad entable en su contra las acciones judiciales que correspondan. Con todo, tratándose de bases de datos de carácter personal, la obligación de confidencialidad dura indefinidamente, de acuerdo con la Ley N°19.628, sobre Protección de la Vida Privada."
    ]:
        agregar_parrafo_con_texto(doc, texto)

    doc.add_heading("{{ Vigesimo_PropiedadDeLaInformacion }}{{espacio}}Propiedad de la información", level=2)
    agregar_parrafo_con_texto(doc, "La entidad licitante será la titular de todos los datos de transacciones, bitácoras (logs), parámetros, documentos electrónicos y archivos adjuntos y, en general, de las bases de datos y de toda información contenida en la infraestructura física y tecnológica que le suministre el proveedor contratado y que se genere en virtud de la ejecución de los servicios objeto de la presente licitación. El proveedor no podrá utilizar la información indicada en el párrafo anterior, durante la ejecución del contrato ni con posterioridad al término de su vigencia, sin autorización escrita de la entidad licitante. Por tal motivo, una vez que el proveedor entregue dicha información a la entidad o al finalizar la relación contractual, deberá borrarla de sus registros lógicos y físicos.")

    doc.add_heading("{{ VigesimoPrimero_SaldosInsolutos }}{{ espacio }}Saldos insolutos de remuneraciones o cotizaciones de seguridad social.", level=2)
    for texto in [
        "Durante la vigencia del respectivo contrato el adjudicatario deberá acreditar que no registra saldos insolutos de obligaciones laborales y sociales con sus actuales trabajadores o con trabajadores contratados en los últimos dos años.",
        "El órgano comprador podrá requerir al adjudicatario, en cualquier momento, los antecedentes que estime necesarios para acreditar el cumplimiento de las obligaciones laborales y sociales antes señaladas.",
        "En caso de que la empresa adjudicada registre saldos insolutos de remuneraciones o cotizaciones de seguridad social con sus actuales trabajadores o con trabajadores contratados en los últimos dos años, los primeros estados de pago de los bienes y servicios de esta licitación deberán ser destinados al pago de dichas obligaciones, debiendo la empresa acreditar que la totalidad de las obligaciones se encuentran liquidadas al cumplirse la mitad del período de ejecución de las prestaciones, con un máximo de seis meses.",
        "La entidad licitante deberá exigir que la empresa adjudicada proceda a dichos pagos y le presente los comprobantes y planillas que demuestren el total cumplimiento de la obligación. El incumplimiento de estas obligaciones por parte de la empresa adjudicataria dará derecho a terminar la relación contractual, pudiendo llamarse a una nueva licitación en la que la empresa referida no podrá participar."
    ]:
        agregar_parrafo_con_texto(doc, texto)

    doc.add_heading("{{ VigesimoSegundo_NormasLaboralesAplicable }}{{ espacio }}Normas Laborales Aplicables", level=2)
    for texto in [
        "El adjudicatario, en su calidad de empleador, será responsable exclusivo del cumplimiento íntegro y oportuno de las normas del Código del Trabajo y leyes complementarias, leyes sociales, de previsión, de seguros, de enfermedades profesionales, de accidentes del trabajo y demás pertinentes respecto de sus trabajadores y/o integrantes de sus respectivos equipos de trabajo.",
        "En consecuencia, el adjudicatario será responsable, en forma exclusiva, y sin que la enumeración sea taxativa, del pago oportuno de las remuneraciones, honorarios, indemnizaciones, desahucios, gratificaciones, gastos de movilización, beneficios y, en general, de toda suma de dinero que, por cualquier concepto, deba pagarse a sus trabajadores y/o integrantes de sus respectivos equipos de trabajo.",
        "El Hospital se reserva el derecho a exigir al contratista, a simple requerimiento de la contraparte técnica, y sin perjuicio de lo dispuesto en el artículo 4° de la Ley de Compras y el artículo 183-C del Código del Trabajo, un certificado que acredite el monto y estado de cumplimiento de las obligaciones laborales y previsionales emitido por la Inspección del Trabajo respectiva, o bien, por medios idóneos que garanticen la veracidad de dicho monto y estado de cumplimiento, respecto de sus trabajadores. Ello, con el propósito de hacer efectivo por parte del órgano comprador, su derecho a ser informado y el derecho de retención, consagrados en los incisos segundo y tercero del artículo 183-C del Código del Trabajo, en el marco de la responsabilidad subsidiaria derivada de dichas obligaciones laborales y previsionales, a la que alude el artículo 183-D del mismo Código.",
        "Por otra parte, se deja expresa constancia que la suscripción del contrato respectivo no significará en caso alguno que el adjudicatario, sus trabajadores o integrantes de los equipos presentados por éstos, adquieran la calidad de funcionarios públicos, no existiendo vínculo alguno de subordinación o dependencia de ellos con el órgano comprador."
    ]:
        agregar_parrafo_con_texto(doc, texto)

    doc.add_heading("{{ VigesimoTercero_CambioPersonalProveedor }}{{ espacio }}Cambio de personal del proveedor adjudicado.", level=2)
    for texto in [
        "El Hospital San José de Melipilla podrá, por razones de buen servicio, solicitar el cambio de trabajadores, expresando la causa del derecho a cambiar al personal del proveedor, entendiéndose como el derecho a prohibir unilateralmente la continuidad de funciones de un trabajador que implique un potencial riesgo a los pacientes, funcionarios, bienes e imagen de la organización.",
        "El Proveedor adjudicado deberá reemplazar al personal, dentro del plazo que se le indique. La decisión del Hospital San José de Melipilla se comunicará por escrito al Proveedor precisando las causas que motivan la solicitud, con a lo menos 5 días de anticipación a la fecha en que se solicita deje de prestar servicios en sus dependencias, el trabajador que se indique."
    ]:
        agregar_parrafo_con_texto(doc, texto)

    doc.add_heading("{{ VigesimoCuarto_CesionySubcontratacion }}{{ espacio }}Cesión y subcontratación.", level=2)
    for texto in [
        "El proveedor adjudicado no podrá ceder ni transferir en forma alguna, total ni parcialmente, los derechos y obligaciones que nacen del desarrollo de esta licitación, y, en especial, los establecidos en los respectivos contratos que se celebren con los órganos públicos mandantes.",
        "La infracción de esta prohibición será causal inmediata de término del contrato, sin perjuicio de las acciones legales que procedan ante esta situación.",
        "Durante la ejecución del contrato, y previa autorización por escrito del Hospital, el adjudicatario sólo podrá efectuar aquellas subcontrataciones que sean indispensables para la realización de tareas específicas, todo lo cual será calificado por el coordinador del contrato. En todo caso, el adjudicatario seguirá siendo el único responsable de las obligaciones contraídas en virtud del respectivo contrato suscrito con el Hospital.",
        "Así mismo, el subcontratista debe encontrarse hábil en el registro de Proveedores del Estado y tratándose de servicios, acreditar el cumplimiento de obligaciones laborales, conforme lo establece el artículo 4° inciso 2° de la Ley N°19.886.",
        "En todos los casos es el oferente y eventual adjudicatario el único responsable del pleno cumplimiento de lo señalado en estas bases (Art. N° 76, Reglamento de la Ley N°19.886)."
    ]:
        agregar_parrafo_con_texto(doc, texto)

    doc.add_heading("{{ VigesimoQuinto_Discrepancias }}{{espacio}}Discrepancias", level=2)
    for texto in [
        "Si con motivo de la ejecución del contrato se presentaran denuncias, querellas o demandas ante el Ministerio Público o los Tribunales Ordinarios de Justicia; o reclamos ante el Consejo de Defensa del Estado por el cuestionamiento en la prestación otorgada y que corresponda al objeto del contrato celebrado, será el proveedor el único responsable por tales actos, por lo que, sí el Hospital fuese condenado a pagar una multa o indemnización, en razón de los actos precedentemente enunciados o el Hospital tuviera que pagar alguna transacción judicial o extrajudicial que deba celebrarse en razón de las situaciones antes enunciadas, el proveedor deberá reembolsar al Hospital el total del monto resultante de un fallo ejecutoriado o de una transacción judicial o extrajudicial o de un procedimiento de medición de acuerdo a la Ley Nº 19.966.",
        "Asimismo, serán responsables de todos los daños, pérdidas, deterioros o perjuicios de bienes muebles e inmuebles del Hospital, producto del mal uso ocasionado en virtud de la prestación de servicio, debiendo restituir al Hospital los costos en que deba incurrir para reparar los daños producidos por este motivo. Esta obligación se mantendrá aun cuando el presente contrato que al efecto se suscriba se dé por terminado ya sea por expiración del plazo establecido o por decisión del Hospital."
    ]:
        agregar_parrafo_con_texto(doc, texto)

    doc.add_heading("Constancia", level=2)
    agregar_parrafo_con_texto(doc, "Se deja expresa constancia que todas y cada una de las cláusulas contenidas en las presentes Bases, Anexos y aclaratorias, se entienden incorporadas sin necesidad de mención expresa en el correspondiente contrato que se materialice con el adjudicado y éste se hace responsable del cumplimiento de las obligaciones de tales documentos, Bases Administrativas y Contrato que se deriven.")

    # Ahora la sección Bases Ténicas para BASES TECNICAS PARA EL SUMINISTRO DE INSUMOS Y ACCESORIOS PARA TERAPIA DE PRESIÓN NEGATIVA CON EQUIPOS EN COMODATO PARA EL HOSPITAL SAN JOSÉ DE MELIPILLA

    doc.add_section()
    doc.add_heading("BASES TECNICAS PARA EL {{ nombre_adquisicion }}", level = 1)
    doc.add_heading("Requisitos para adjudicarse disposiciones generales", level=2)

    doc.add_paragraph("Para adjudicarse el presente proceso de licitación, los oferentes participantes deberán cumplir con lo siguiente:")
    for texto in ["Haber llenado y presentado los Anexos Administrativos N°1, N°2, N°3, N°4, el presentar estos anexos habilita al proveedor a participar en la presente licitación.",
                  "Deben presentar el Económico N°5 {{ anexos_tecnicos }} con toda la información requerida, debidamente firmados por el representante legal de la empresa o la persona natural, según corresponda.",
                  "Deberán entregar toda la información necesaria para poder evaluar a la empresa en cada uno de los ítems de los Criterios de Evaluación.",
                  "Deberán dar respuesta a los requisitos generados por foro inverso en los plazos y/o periodos establecidos en las presentes Bases de Licitación.",
                  "Presentar ficha técnica y certificados de los productos ofertados.",
                  "Entregar muestras de los productos solicitados y comodato ofertado.",
                  "Entregar garantías de la oferta."]:
        agregar_parrafo_con_texto(doc, texto, estilo="List Bullet")

    parrafo_nuevo = doc.add_paragraph()
    parrafo_nuevo.add_run("Nota:").bold = True
    parrafo_nuevo.add_run(" Los oferentes que no cumplan con estos requisitos no serán evaluados, declarándose inadmisible su oferta.")

    doc.add_heading("Disposiciones de la Licitación", level=2)
    doc.add_paragraph("Determinar las directrices y características técnicas necesarias para el suministro de insumos y accesorios para terapia de presión negativa con equipos en comodato.")

    doc.add_heading("GENERALIDADES:", level = 3)

    generalidades_lista = doc.add_paragraph(style="List Bullet")
    generalidades_lista.add_run("La Licitación será adjudicada por {{ metodo_adjudicacion }}").bold = True
    generalidades_lista.add_run(
        " y se podrá aumentar o disminuir hasta un 30% por cada línea adjudicada sin superar el monto total presupuestado para la Licitación.")

    gen_p1 = doc.add_paragraph(
        "La propuesta deberá contemplar todos los costos de trasporte para el despacho de los productos. El Hospital no cancelará ningún costo asociado a esta temática.",
        style="List Bullet")

    gen_p2 = doc.add_paragraph(
        "El proveedor deberá mantener la calidad de los productos ofertados para cada solicitud de compra bajo esta Licitación, situación que será constantemente evaluada por el administrador del contrato.",
        style="List Bullet")

    gen_p3 = doc.add_paragraph(
        "El proveedor deberá permitir la revisión de los productos entregados al Hospital por el personal que se disponga por parte del Establecimiento, para así dar una correcta recepción conforme.",
        style="List Bullet")

    gen_p4 = doc.add_paragraph(
        "En los casos que los productos sean despachados por empresas de transporte estos deberán permitir la revisión de los productos, en caso contrario los productos serán rechazados.",
        style="List Bullet")

    gen_p5 = doc.add_paragraph(
        "Toda entrega deberá adjuntar un documento que acredite la compra (Guía de despacho, Factura u Orden de Compra).",
        style="List Bullet")

    gen_p6 = doc.add_paragraph(style="List Bullet")
    gen_p6.add_run("La adquisición de estos productos será de forma parcializada durante un periodo máximo de ")
    gen_p6.add_run("{{ plazo_meses }} meses").bold = True
    gen_p6.add_run(" o hasta la duración del monto estipulado en base.")

    gen_p7 = doc.add_paragraph(
        "Deberán entregar toda la información necesaria para poder evaluar a la empresa en cada uno de los ítems de los Criterios de Evaluación.",
        style="List Bullet")

    gen_p8 = doc.add_paragraph(
        "El administrador técnico del contrato será {{ administrador_tecnico_administrativo }}",
        style="List Bullet")

    # Los productos
    doc.add_heading("De los Productos", level = 4)
    num_id_productos = crear_numeracion(doc)

    de_productos_p1 = doc.add_paragraph("La presente licitación pública, se enfoca en la adquisición de los productos que se presentan en el cuadro siguiente, se evaluaran técnicamente cada producto. La adjudicación será por {{ metodo_adjudicacion }}.", style = "List Number")
    de_productos_p2 = doc.add_paragraph("La siguiente tabla presenta cantidades de consumo referenciales, la que se utilizara solo para términos de evaluación.", style = "List Number")
    aplicar_numeracion(de_productos_p1, num_id_productos)
    aplicar_numeracion(de_productos_p2, num_id_productos)

    # Data para la tabla de Insumos (extraída de la imagen original + nuevos ítems)
    insumos_header = ["ITEM", "INSUMOS", "UD", "MONTO MÁXIMO A PAGAR"]

    insumos_data = [
        ["1",
         "Recolector de contenido y exudado de herida con gel de 300 ml para presión negativa con conexión que mide 1,20 mt aprox., circuito cerrado, clamp integrado y filtro de carbón. Desechable.",
         "UD", "$150.000"],
        ["2",
         "Recolector de contenido y exudado de herida con gel de 500 ml para ser utilizado con presión negativa, con conexión que mide 1,80 mt aprox., circuito cerrado, clamp integrado y filtro de carbón. Desechable.",
         "UD", "$180.000"],
        ["3",
         "Recolector de contenido y exudado de herida con gel de 1000 ml para ser utilizado con presión negativa con conexión que mide 1,80 mt aprox., circuito cerrado, clamp integrado y filtro de carbón. Desechable.",
         "UD", "$220.000"],
        ["4",
         "Kit de apósito espuma negra en forma ovalada 26 cm x15cm x3.2 cm aprox. tamaño LARGE, con láminas adhesivas transparentes, conector de succión de silicona flexible de 90 cm aprox., conector luer-lock, clamp y regla desechable.",
         "UD", "$100.000"],
        ["5",
         "Kit de apósito espuma negra pre cortada en forma ovalada 60 cm x30cm x1.8 cm aprox. tamaño extra large, con láminas adhesivas transparentes, conector de succi��n de silicona flexible de 90 cm aprox., conector luer-lock, clamp, desechable.",
         "UD", "$370.000"],
        ["6",
         "Kit de apósito espuma negra pre cortada en forma de espiral 11.3 cm x 7.7cm x 1.75cm aprox. tamaño small con láminas adhesivas transparentes, conector de succión de silicona flexible de 90 cm, conector luer-lock, clamp, desechable.",
         "UD", "$100.000"],
        ["7",
         "Kit de apósito espuma negra pre cortada en forma de espiral 17.4cm x14.7cm x1.75 cm aprox. tamaño medium, con láminas adhesivas transparentes, conector de succión de silicona flexible de 90 cm, conector luer-lock, clamp, desechable.",
         "UD", "$120.000"],
        ["8",
         "Apósito de espuma blanca de alcohol polivinílico humedecido con agua estéril, hidrofílica, DE 10 CM X 10CM X 1 cm aprox. tamaño small",
         "UD", "$60.000"],
        ["9",
         "Apósito de espuma blanca de alcohol polivinílico humedecido con agua estéril, hidrofílica, DE 10 CM X 15CM X 1 cm aprox. Tamaño Large (L)",
         "UD", "$70.000"],
        ["10",
         "Kit de apósito espuma hidrofóbica color gris de eter de poliuretano con plata metálica 10 cm x7.5 cm x3.2 cm aprox. tamaño small.",
         "UD", "$120.000"],
        ["11",
         "Kit de apósito espuma hidrofóbica color gris de eter de poliuretano con plata metálica 18cm x12.5 cmx3.2 cm aprox. tamaño medio",
         "UD", "$130.000"],
        ["12",
         "Kit de apósito espuma hidrofóbica color gris de eter de poliuretano con plata metálica 26 cm x15cm x3.2 cm aprox. tamaño large.",
         "UD", "$230.000"],
        ["13",
         "Kit de apósito abdominal para manejo de abdomen abierto con presión negativa, con lámina protectora visceral de poliuretano, láminas adhesivas transparentes, conector de succión de silicona flexible de 90 cm aprox., conector luer-lock y clamp, desechable",
         "UD", "$550.000"],
        ["14",
         "Kit de apósito para incisiones lineales 90 cm aprox. con espuma de poliuretano, láminas adhesivas transparentes y un conector de succión de silicona flexible de 90 cm aprox., conector luer-lock y clamp y desechable.",
         "UD", "$420.000"],
        ["15",
         "Kit de apósito para incisiones lineales 20 cm aprox. con espuma de poliuretano, láminas adhesivas transparentes y un conector de succión de silicona flexible de 90 cm aprox., conector luer-lock y clamp y desechable.",
         "UD", "$300.000"],
        # Nuevos ítems añadidos
        ["16",
         "Kit de apósito para incisiones lineales 13 cm aprox. con espuma de poliuretano, laminas adhesivas transparentes y un conector de succión de silicona flexible de 90 cm aprox., conector luer-lock y clamp y desechable.",
         "UD", "$250.000"],
        ["17",
         "Kit de apósito para incisiones lineales 35 cm aprox. con espuma de poliuretano, laminas adhesivas transparentes y un conector de succión de silicona flexible de 90 cm aprox., conector luer-lock y clamp y desechable.",
         "UD", "$400.000"],
        ["18",
         "Lámina adhesiva transparente, hipoalergénica, semipermeable para realizar el sello de la terapia de presión negativa",
         "UD", "$50.000"],
        ["19",
         "Conector de succión de silicona flexible de 90 cm aprox. con sensores externos de luz de monitoreo continuo de la presión, sistema de detección de obstrucciones y sistema de ráfagas de aire cada 5 min para ayudar a reducir los bloqueos, conector luer-lock y clamp.",
         "UD", "$50.000"],
        ["20",
         "Kit de apósito para terapia de instilación tamaño médium, espuma de ester de poliuretano y reticulada con 3 capas: 1 capa en contacto con la herida que tiene orificios de 5 mm, una segunda capa fina de 8 mm y una tercera capa gruesa de 16 mm.",
         "UD", "$200.000"],
        ["21", "Cassete para conectar la solución para la terapia de instilación.", "UD", "$110.000"]
    ]

    # Resto del código para crear y formatear la tabla (sin cambios)
    # ... (código para crear la tabla usando crear_tabla o directamente)
    # ... (código para formatear encabezado y celdas de datos)
    # ... (código para centrar verticalmente la tabla si es necesario)
    # Añadir una tabla para los insumos
    # Número de filas = 1 (encabezado) + número de filas de datos
    num_rows_insumos = len(insumos_data) + 1
    num_cols_insumos = len(insumos_header)
    # Usa tu función existente 'crear_tabla' o añade la tabla directamente
    # Aquí usamos la función crear_tabla que ya definiste
    tabla_insumos = crear_tabla(doc, [insumos_header] + insumos_data, estilo='Table Grid',
                                centrar=False)  # Centrar verticalmente después si es necesario

    de_productos_p3 = doc.add_paragraph("La adjudicación se realizará por valor unitario y tendrá una duración de {{ plazo_meses }} meses o hasta agotar el presupuesto, lo que ocurra primero, sin obligar al hospital a comprar una cantidad mínima establecida.")
    aplicar_numeracion(de_productos_p3, num_id_productos)


    parte_d_generalidades = doc.add_paragraph(style = "List Number")
    parte_d_generalidades.add_run("se considera causal")
    parte_d_generalidades.add_run(" admisibilidad").bold = True
    parte_d_generalidades.add_run("que el proveedor adjunte ficha técnica en español de todos los productos solicitados al portal de Mercado Público.")
    aplicar_numeracion(parte_d_generalidades, num_id_productos)

    # Entrega de Muestras
    doc.add_heading("Entrega de Muestras", level = 4)
    entrega_muestras_p1 = doc.add_paragraph(style = "List Bullet")
    entrega_muestras_p1.add_run("Presentar muestras de los insumos y equipo en comodato para evaluación es de carácter ")
    entrega_muestras_p1.add_run("OBLIGATORIO").bold = True

    entrega_muestras_p2 = doc.add_paragraph(style = "List Bullet")
    entrega_muestras_p2.add_run("Se debe presentar muestras de todos los insumos solicitados. En caso de no presentar muestras su propuesta podrá ser declarada ")
    entrega_muestras_p2.add_run("inadmisible").bold = True

    entrega_muestras_p3 = doc.add_paragraph(style = "List Bullet")
    entrega_muestras_p3.add_run("Cada muestra debe indicar nombre del proveedor, número de licitación y N° de línea del producto. ")

    entrega_muestras_p4 = doc.add_paragraph(style = "List Bullet")
    entrega_muestras_p4.add_run("Las muestras en ningún caso generarán costo para el Hospital y NO podrán ser devueltas a los oferentes, ya que serán utilizados para realizar pruebas de parte de los referentes técnicos para su evaluación.")

    entrega_muestras_p5 = doc.add_paragraph(style = "List Bullet")
    entrega_muestras_p5.add_run("El oferente deberá permitir la apertura de cajas/bolsas de las muestras presentadas para una correcta recepción de estas.")

    entrega_muestras_p6 = doc.add_paragraph(style = "List Bullet")
    entrega_muestras_p6.add_run("El oferente deberá entregar las muestras a la ")
    entrega_muestras_p6.add_run("UNIDAD DE ABASTECIMIENTO DEL HOSPITAL SAN JOSE DE MELIPILLA, ").bold = True
    entrega_muestras_p6.add_run("ubicada en Calle O’Higgins N.º 551, Comuna de Melipilla hasta el cierre de la licitación ")
    entrega_muestras_p6.add_run("La no entrega de muestras en la forma y plazos establecidos en bases facultará al establecimiento a dejar inadmisible la oferta.").bold = True

    entrega_muestras_p7 = doc.add_paragraph(style = "List Bullet")
    entrega_muestras_p7.add_run("Toda muestra deberá ser acompañados de una guía de despacho o acta de recepción, la que será completada (firmada y timbrada) por Unidad de Abastecimiento. Este documento será el que respaldará la recepción de las muestras.")


    doc.add_heading("Sobre los equipos solicitados en comodato para el uso de insumos condición obligatoria", level = 4)
    doc.add_paragraph("Para ejecutar el suministro, es obligatorio para el proveedor adjudicado, la entrega en comodato, a las unidades clínicas del hospital que lo soliciten, de los siguientes equipos médicos:")

    datos_tabla_equipos = [
        ["NOMBRE EQUIPOS", "CANTIDAD ESTIMADA"],
        ["EQUIPO PARA TERAPIA PRESIÓN NEGATIVA INTRAHOSPITALARIO", "12"],
        ["EQUIPO PARA TERAPIA PRESIÓN NEGATIVA AMBULATORIO - DOMICILIARIO", "4"]
    ]

    # Crear la tabla de equipos
    tabla_equipos = crear_tabla(doc, datos_tabla_equipos, estilo='Table Grid', centrar=True)

    doc.add_paragraph("La cantidad de equipos es estimada, tiene un carácter referencial. Debe adjuntar ficha técnica de los dispositivos. Estas cantidades podrán variar de acuerdo a la demanda interna, para lo cual el oferente adjudicado deberá dar respuesta las necesidades del hospital, máximo 10 días hábiles desde la adjudicación.")

    doc.add_paragraph("El proveedor que se adjudique, deberá entregar en comodato, a unidades clínicas del Hospital, Equipo para terapia presión negativa intrahospitalario y Equipo para terapia presión negativa ambulatorio compatibles con los insumos ofertados. Junto con la entrega en comodato, el proveedor se obliga a realizar la mantención preventiva y correctiva de los equipos. En caso de falla, se obliga a facilitar al Hospital, sin costo para este, equipos de similares características, mientras se entrega un nuevo equipo en comodato")

    datos_tabla_caracteristicas = [
        ["CARACTERISTICAS GENERALES OBLIGATORIAS DE LOS EQUIPOS EN COMODATO"],
        ["Consolas con botón de encendido y apagado."],
        ["Rangos de presión de -25 a -200 mmHg"],
        [
            "Terapias integradas en el mismo equipo. Terapias de presión negativa estándar, manejo de abdomen abierto, prevención de dehiscencia de suturas e instilación"],
        [
            "Compatibilidad con contendor de 1000, 500 y 300 ml para uso intrahospitalario y contendor de 300 para uso ambulatorio-domiciliario"],
        ["Sistema de alarma de contenedor en su capacidad máxima y botón para liberación de contenedor"],
        ["Sistema de alarma de baja presión, terapia interrumpida u obstrucción."],
        ["Sistema de alarmas para nivel crítico de batería."],
        ["Cable para alimentación de corriente eléctrica o baterías en caso de equipo portátil."],
        ["Memoria de uso de consola"],
        [
            "Autonomía de la batería de al menos 6 horas para equipo de uso hospitalario y 10 horas para equipo de uso domiciliario."]
    ]

    # Crear la tabla de características de equipos
    tabla_caracteristicas = crear_tabla(doc, datos_tabla_caracteristicas, estilo='Table Grid',
                                        centrar=False)  # Centrar=False o True según preferencia

    doc.add_heading("Entrega y Recepción", level=4)  # Asumiendo level 4, ajustar si es necesario

    entrega_recepcion_items = [
        "La adquisición de estos productos será de forma parcializada según la cantidad y periocidad que el hospital considere necesario.",
        "El proveedor deberá despachar los productos señalando explícitamente Nombre del producto, Identificación del Proveedor y N° de Guía/Factura, Modelo (solo cuando corresponda), N° de Lote/Serie, Fecha de Vencimiento, de acuerdo a Norma Técnica de Minsal °226/22.",
        "Los productos deberán ser entregado en las dependencias del Hospital de Melipilla, considerando el traslado carga y descarga.",
        "La propuesta deberá contemplar todos los costos de trasporte para el despacho de los productos. El Hospital no cancelará ningún costo asociado a esta temática.",
        "Desde el requerimiento, el proveedor tendrá un máximo 7 días corridos para entregar los productos, siempre respetando los plazos ofertados según anexo Plazo de Entrega.",
        "El proveedor deberá realizar los cambios de los productos que no se ajusten a las bases técnicas y/o presenten deterioros en un plazo no mayor a 48 horas, con previo requerimiento del administrador del contrato.",
        "Los productos despachados que no se ajusten a la calidad ofertada serán rechazados e informado vía correo electrónico, para solicitar el cambio.",
        "El gasto que eventualmente se genere por artículos rechazados será de cargo a la empresa adjudicada.",
        "El embalaje deberá ser suficiente para soportar, sin límites, la manipulación brusca y descuidada durante el tránsito y la exposición a temperaturas extremas.",
        "El proveedor deberá permitir la apertura de cajas, bolsas, etc., para la correcta revisión de los productos entregados al Hospital por el personal de Bodega del Establecimiento para así dar una correcta recepción conforme.",
        "En los casos que los productos sean despachados por empresas de transporte estos deberán permitir la revisión de los productos, en caso contrario los productos serán rechazados."
    ]

    for item_text in entrega_recepcion_items:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(item_text)

    # Último ítem con parte en negrita
    p_ultimo = doc.add_paragraph(style='List Bullet')
    p_ultimo.add_run("Los productos deberán ser entregados en ")
    p_ultimo.add_run("Bodega de Farmacia").bold = True
    p_ultimo.add_run(
        " del Hospital San José de Melipilla ubicada en calle O’Higgins #551, en los siguientes horarios: lunes a viernes de 8:00 a 14:00 horas")


    doc.add_heading("MOTIVOS DE RECHAZO POR OBSERVACIÓN FÍSICA (ya iniciado en contrato):", level=4)

    rechazo_p = doc.add_paragraph()
    rechazo_p.add_run("Los artículos requeridos en la presente licitación podrán ser rechazados, al momento de la recepción en ")
    rechazo_p.add_run("Bodega de Farmacia").bold = True
    rechazo_p.add_run(", por los siguientes motivos:")

    motivos_rechazo_items = [
        "Empaques deteriorados o visiblemente sucios, manchados, húmedos, etc.",
        "Cajas colectivas sin identificación de su contenido o leyendas ilegibles.",
        "Diferentes lotes no señalizados, Incluidos en un empaque colectivo.",
        "Textos o leyendas equivocadas, que puedan inducir a error.",
        "Envases con etiquetas e impresiones ilegibles o sin ellas.",
        "Discordancia entre envases ya sea colectivo, primarios o secundarios.",
        "Acondicionamiento inadecuado dentro de los envases primarios o secundarios.",
        "Envases vacíos o adulterados.",
        "Número de lote o fecha de vencimiento equivocada o ausente.",
        "Caja o etiqueta incorrecta, leyendas incompletas o ausentes.",
        "Sello violado o mal colocado.",
        "Contenido incorrecto, diferente o menor al etiquetado.",
        "Envases aplastados o deteriorados con motivo del traslado.",
        "Contaminación visible.",
        "Coloración no homogénea (intra o inter lote).",
        "Partículas extrañas observadas a simple vista o contraluz.",
        "Los insumos no deberán tener un vencimiento inferior a 2 año desde la fecha de recepción."
    ]

    for item_text in motivos_rechazo_items:
        p = doc.add_paragraph(style='List Bullet')
        # El carácter '' se reemplaza por una viñeta estándar de Word.
        # Si necesitas específicamente ese carácter, tendrías que agregarlo manualmente al inicio de cada string.
        p.add_run(item_text)

    doc.add_section()

    p1 = doc.add_heading("Anexos", level=1)
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p1_nexo1 = doc.add_heading("Anexo N°1", level=2)
    p1_nexo1.allignment = WD_ALIGN_PARAGRAPH.CENTER

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.add_run("IDENTIFICACIÓN DEL OFERENTE")

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.add_run("PROPUESTA PÚBLICA: {{ nombre_adquisicion }}")

    # REMOVED: doc.add_paragraph()  # Removed the extra space before the table

    # --- Table Data ---
    # Structure: (Text for cell 1, Text for cell 2, Does cell 1 span both columns?)
    table_data = [
        ("IDENTIFICACIÓN DEL OFERENTE", "", True),
        ("R.U.T. DEL OFERENTE", "", False),
        ("SIGLA PARA EL CASO DE EMPRESAS (Nombre de Fantasía)", "", False),
        ("DIRECCIÓN OFERENTE", "", False),
        ("CIUDAD", "", False),
        ("COMUNA", "", False),
        ("TELÉFONOS", "", False),
        ("", "", True),  # This was the blue separator row, now just a spanned empty row
        ("NOMBRE DEL REPRESENTANTE LEGAL", "", False),
        ("RUT DEL REPRESENTANTE LEGAL", "", False),
        ("NOMBRE DE LA NOTARIA", "", False),
        (
            "FECHA DONDE SE SEÑALA LA PERSONERIA DEL REPRESENTANTE LEGAL (adjuntar documento si no se encuentra actualizado en portal de Mercado Público)",
            "", False),
        ("NOMBRE DEL CONTACTO COMERCIAL (ADMINISTRADOR EXTERNO DEL CONTRATO)", "", False),
        ("CARGO DEL CONTACTO COMERCIAL", "", False),
        ("RUBRO COMERCIAL", "", False),
        ("E-MAIL", "", False),
        ("NOMBRE DEL CONTACTO PARA EL SERVICIO", "", False),
        ("TELEFONO", "", False),
        ("E-MAIL", "", False),
        ("CELULAR", "", False),
        ("HORARIO DE ATENCION", "", False),
    ]

    num_rows = len(table_data)
    num_cols = 2
    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.style = 'TableGrid'  # This will add basic borders

    # Populate table
    for i, (text1, text2, span) in enumerate(table_data):
        row_cells = table.rows[i].cells
        cell1 = row_cells[0]
        cell2 = row_cells[1]

        if span:
            cell1.merge(cell2)
            cell1.text = text1
            if text1 == "IDENTIFICACIÓN DEL OFERENTE":  # Center this specific header
                cell1.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            cell1.text = text1
            cell2.text = text2

    # --- Footer Text ---
    doc.add_paragraph()  # Keep this space after the table

    # Add the underline for the signature line BEFORE the signature text
    pf_signature_line = doc.add_paragraph()
    pf_signature_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Adjust the number of underscores as needed for desired length
    pf_signature_line.add_run("____________________________________")

    pf1 = doc.add_paragraph()
    pf1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf1.add_run("FIRMA Y TIMBRE OFERENTE O REPRESENTANTE LEGAL")

    doc.add_paragraph()

    pf2 = doc.add_paragraph()
    pf2.add_run("Fecha: ____________________________________")
    # Guardar el documento

    anexo2_header = doc.add_heading("ANEXO N° 2", level=2 )

    # Add DECLARACION JURADA Header as a paragraph
    declaracion_header = doc.add_paragraph()
    declaracion_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    declaracion_header.add_run("DECLARACION JURADA DE HABILIDAD")

    # Add the Proposal Title again as a paragraph
    title_anexo2 = doc.add_paragraph()
    title_anexo2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_anexo2.add_run("PROPUESTA PÚBLICA: PARA EL SUMINISTRO DE INSUMOS Y ACCESORIOS PARA TERAPIA DE "
                         "PRESIÓN NEGATIVA CON EQUIPOS EN COMODATO PARA EL HOSPITAL SAN JOSÉ DE MELIPILLA.")

    # Add the declaration paragraphs
    para1 = doc.add_paragraph()
    para1.add_run(
        "Por la presente, el Oferente, <<NOMBRE PERSONA NATURAL O NOMBRE PERSONA JURIDICA>>, declara bajo juramento que no ha sido sancionado con la pena de prohibición perpetua o temporal (esta última vigente) para contratar con el Estado, por lavado de activos, financiamiento del terrorismo y cohecho, en virtud de lo dispuesto en los artículos 8° N°2 y N°10 de la Ley N°20.393.")


    para2 = doc.add_paragraph()
    para2.add_run(
        "Asimismo, declara bajo juramento, sea persona natural o jurídica que no le afecta ninguna de las inhabilidades previstas en los incisos primero y sexto del artículo 4° de la Ley N°19.886, además, las previstas en el artículo 26 letra D del Decreto Ley N°211, que se transcriben en su parte pertinente:")

    para3 = doc.add_paragraph()
    para3.add_run(
        "“(...) Quedarán excluidos quienes, dentro de los dos años anteriores al momento de la presentación de la oferta, de la formulación de la propuesta o de la suscripción de la convención, según se trate de licitaciones públicas, privadas o contratación directa, hayan sido condenados por prácticas antisindicales o infracción a los derechos fundamentales del trabajador, o por delitos concursales establecidos en el Código Penal” (inciso primero).")

    para4 = doc.add_paragraph()
    para4.add_run(
        "“Ningún órgano de la Administración del Estado y de las empresas y corporaciones del Estado o en que éste tenga participación, podrá suscribir contratos administrativos de provisión de bienes o prestación de servicios con los funcionarios directivos del mismo órgano o empresa, ni con personas unidas a ellos por los vínculos o parentescos descritos en la letra b) del artículo 54 de la ley N°18.575, ley Orgánica Constitucional de Bases Generales de la Administración del Estado, ni con sociedades de personas de las que aquéllos o éstas formen parte, ni con sociedades comandita por acciones o anónimas cerradas en que aquéllos o estas sean accionistas, ni con sociedades anónimas abiertas en que aquéllos o estas sean dueños de acciones que representen el 10% o más del capital, ni con los gerentes, administradores, representantes o directores de cualquiera de las sociedades antedichas” (inciso sexto).")

    para5 = doc.add_paragraph()
    para5.add_run(
        "En el caso de las conductas previstas en la letra a) del artículo 3°, podrá imponer, además, la prohibición de contratar a cualquier título con órganos de la administración centralizada o descentralizada del Estado, con organismos autónomos o con instituciones, organismos, empresas o servicios en los que el Estado efectúe aportes, con el Congreso Nacional y el Poder Judicial, así como la prohibición de adjudicarse cualquier concesión otorgada por el Estado, hasta el plazo de cinco años contado desde que la sentencia definitiva quede ejecutoriada.")

    # Add the underline for the signature line BEFORE the signature text
    pf_signature_line = doc.add_paragraph()
    pf_signature_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Adjust the number of underscores as needed for desired length
    pf_signature_line.add_run("____________________________________")

    pf1 = doc.add_paragraph()
    pf1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf1.add_run("FIRMA Y TIMBRE OFERENTE O REPRESENTANTE LEGAL")


    pf2 = doc.add_paragraph()
    pf2.add_run("Fecha: ____________________________________")

    # Aplicar formato global
    aplicar_formato_global(doc)

    doc_path = 'base_automatizada.docx'
    doc.save(doc_path)
    print(f"Documento guardado como: {doc_path}")

if __name__ == "__main__":
    main()